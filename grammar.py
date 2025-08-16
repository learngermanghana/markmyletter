# ==== Standard Library ====
import atexit, base64, difflib, hashlib
import html as html_stdlib
import io, json, os, random, math, re, sqlite3, tempfile, time
import urllib.parse as _urllib
import calendar
from datetime import date, datetime, timedelta, timezone
from uuid import uuid4
from typing import Optional

# ==== Third-Party Packages ====
import bcrypt
import firebase_admin
import matplotlib.pyplot as plt
import pandas as pd
import requests
import streamlit as st
import streamlit.components.v1 as components
from bs4 import BeautifulSoup
from docx import Document
from firebase_admin import credentials, firestore
from fpdf import FPDF
from gtts import gTTS
from openai import OpenAI
from streamlit.components.v1 import html as st_html
from streamlit_cookies_manager import EncryptedCookieManager
from streamlit_quill import st_quill

# ---- Streamlit page config MUST be first Streamlit call ----
st.set_page_config(
    page_title="Falowen – Your German Conversation Partner",
    page_icon="👋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Top spacing + chrome (tighter)
st.markdown("""
<style>
/* Remove Streamlit's top padding */
[data-testid="stAppViewContainer"] > .main .block-container {
  padding-top: 0 !important;
}

/* First rendered block (often a head-inject) — keep a small gap only */
[data-testid="stAppViewContainer"] .main .block-container > div:first-child {
  margin-top: 0 !important;
  margin-bottom: 8px !important;   /* was 24px */
  padding-top: 0 !important;
  padding-bottom: 0 !important;
}

/* If that first block is an iframe, collapse it completely */
[data-testid="stAppViewContainer"] .main .block-container > div:first-child [data-testid="stIFrame"] {
  display: block;
  height: 0 !important;
  min-height: 0 !important;
  margin: 0 !important;
  padding: 0 !important;
  border: 0 !important;
  overflow: hidden !important;
}

/* Keep hero flush and compact */
  .hero {
    margin-top: 2px !important;      /* was 0/12 — pulls hero up */
    margin-bottom: 4px !important;   /* tighter space before tabs */
    padding-top: 6px !important;
    display: flow-root;
  }
.hero h1:first-child { margin-top: 0 !important; }
/* Trim default gap above Streamlit tabs */
[data-testid="stTabs"] {
  margin-top: 8px !important;
}

/* Hide default Streamlit chrome */
#MainMenu { visibility: hidden; }
footer { visibility: hidden; }
</style>
""", unsafe_allow_html=True)

# Compatibility alias
html = st_html

# ---- PWA head helper (define BEFORE you call it) ----
BASE = st.secrets.get("PUBLIC_BASE_URL", "")
_manifest = f'{BASE}/static/manifest.webmanifest' if BASE else "/static/manifest.webmanifest"
_icon180  = f'{BASE}/static/icons/falowen-180.png' if BASE else "/static/icons/falowen-180.png"

def _inject_meta_tags():
    components.html(f"""
      <link rel="manifest" href="{_manifest}">
      <link rel="apple-touch-icon" href="{_icon180}">
      <meta name="apple-mobile-web-app-capable" content="yes">
      <meta name="apple-mobile-web-app-title" content="Falowen">
      <meta name="apple-mobile-web-app-status-bar-style" content="black">
      <meta name="theme-color" content="#000000">
      <meta name="viewport" content="width=device-width, initial-scale=1, viewport-fit=cover">
    """, height=0)

# --- State bootstrap ---
def _bootstrap_state():
    defaults = {
        "logged_in": False,
        "student_row": None,
        "student_code": "",
        "student_name": "",
        "session_token": "",
        "cookie_synced": False,
        "__last_refresh": 0.0,
        "__ua_hash": "",
        "_oauth_state": "",
        "_oauth_code_redeemed": "",
    }
    for k, v in defaults.items():
        st.session_state.setdefault(k, v)
_bootstrap_state()

# ==== Hide Streamlit chrome ====
st.markdown("""
<style>
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
</style>
""", unsafe_allow_html=True)

# ==== FIREBASE ADMIN INIT (Firestore only; no Firebase Auth in login) ====
try:
    if not firebase_admin._apps:
        cred_dict = dict(st.secrets["firebase"])
        cred = credentials.Certificate(cred_dict)
        firebase_admin.initialize_app(cred)
    db = firestore.client()
except Exception as e:
    st.error(f"Firebase init failed: {e}")
    st.stop()

# ---- Firestore sessions (server-side auth state) ----
SESSIONS_COL = "sessions"
SESSION_TTL_MIN = 60 * 24 * 14         # 14 days
SESSION_ROTATE_AFTER_MIN = 60 * 24 * 7 # 7 days

def _rand_token(nbytes: int = 48) -> str:
    return base64.urlsafe_b64encode(os.urandom(nbytes)).rstrip(b"=").decode("ascii")

def create_session_token(student_code: str, name: str, ua_hash: str = "") -> str:
    now = time.time()
    token = _rand_token()
    db.collection(SESSIONS_COL).document(token).set({
        "student_code": (student_code or "").strip().lower(),
        "name": name or "",
        "issued_at": now,
        "expires_at": now + (SESSION_TTL_MIN * 60),
        "ua_hash": ua_hash or "",
    })
    return token

def validate_session_token(token: str, ua_hash: str = "") -> dict | None:
    if not token:
        return None
    try:
        snap = db.collection(SESSIONS_COL).document(token).get()
        if not snap.exists:
            return None
        data = snap.to_dict() or {}
        if float(data.get("expires_at", 0)) < time.time():
            return None
        if data.get("ua_hash") and ua_hash and data["ua_hash"] != ua_hash:
            return None
        return data
    except Exception:
        return None

def refresh_or_rotate_session_token(token: str) -> str:
    try:
        ref = db.collection(SESSIONS_COL).document(token)
        snap = ref.get()
        if not snap.exists:
            return token
        data = snap.to_dict() or {}
        now = time.time()
        # Extend TTL
        ref.update({"expires_at": now + (SESSION_TTL_MIN * 60)})
        # Rotate if old
        if now - float(data.get("issued_at", now)) > (SESSION_ROTATE_AFTER_MIN * 60):
            new_token = _rand_token()
            db.collection(SESSIONS_COL).document(new_token).set({
                **data,
                "issued_at": now,
                "expires_at": now + (SESSION_TTL_MIN * 60),
            })
            try:
                ref.delete()
            except Exception:
                pass
            return new_token
    except Exception:
        pass
    return token

def destroy_session_token(token: str) -> None:
    try:
        db.collection(SESSIONS_COL).document(token).delete()
    except Exception:
        pass

# ==== OPENAI CLIENT SETUP ====
OPENAI_API_KEY = st.secrets.get("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    st.error("Missing OpenAI API key. Please add OPENAI_API_KEY in Streamlit secrets.")
    st.stop()
os.environ["OPENAI_API_KEY"] = OPENAI_API_KEY
client = OpenAI(api_key=OPENAI_API_KEY)

# ==== DB CONNECTION & INITIALIZATION ====
def get_connection():
    if "conn" not in st.session_state:
        st.session_state["conn"] = sqlite3.connect(
            "vocab_progress.db", check_same_thread=False
        )
        atexit.register(st.session_state["conn"].close)
    return st.session_state["conn"]

def init_db():
    conn = get_connection()
    c = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS vocab_progress (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            student_code TEXT,
            name TEXT,
            level TEXT,
            word TEXT,
            student_answer TEXT,
            is_correct INTEGER,
            date TEXT
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS schreiben_progress (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            student_code TEXT,
            name TEXT,
            level TEXT,
            essay TEXT,
            score INTEGER,
            feedback TEXT,
            date TEXT
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS sprechen_progress (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            student_code TEXT,
            name TEXT,
            level TEXT,
            teil TEXT,
            message TEXT,
            score INTEGER,
            feedback TEXT,
            date TEXT
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS exam_progress (
            student_code TEXT,
            level TEXT,
            teil TEXT,
            remaining TEXT,
            used TEXT,
            PRIMARY KEY (student_code, level, teil)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS my_vocab (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            student_code TEXT,
            level TEXT,
            word TEXT,
            translation TEXT,
            date_added TEXT
        )
    """)
    for tbl in ["sprechen_usage", "letter_coach_usage", "schreiben_usage"]:
        c.execute(f"""
            CREATE TABLE IF NOT EXISTS {tbl} (
                student_code TEXT,
                date TEXT,
                count INTEGER,
                PRIMARY KEY (student_code, date)
            )
        """)
    conn.commit()
init_db()

# ==== CONSTANTS ====
FALOWEN_DAILY_LIMIT = 20
VOCAB_DAILY_LIMIT = 20
SCHREIBEN_DAILY_LIMIT = 5

def get_sprechen_usage(student_code):
    today = str(date.today())
    conn = get_connection()
    c = conn.cursor()
    c.execute(
        "SELECT count FROM sprechen_usage WHERE student_code=? AND date=?",
        (student_code, today)
    )
    row = c.fetchone()
    return row[0] if row else 0

def inc_sprechen_usage(student_code):
    today = str(date.today())
    conn = get_connection()
    c = conn.cursor()
    c.execute(
        """
        INSERT INTO sprechen_usage (student_code, date, count)
        VALUES (?, ?, 1)
        ON CONFLICT(student_code, date)
        DO UPDATE SET count = count + 1
        """,
        (student_code, today)
    )
    conn.commit()

def has_sprechen_quota(student_code, limit=FALOWEN_DAILY_LIMIT):
    return get_sprechen_usage(student_code) < limit

def has_sprechen_quota(student_code, limit=FALOWEN_DAILY_LIMIT):
    return get_sprechen_usage(student_code) < limit

# ==== YOUTUBE PLAYLIST HELPERS ====
YOUTUBE_API_KEY = st.secrets.get("YOUTUBE_API_KEY", "AIzaSyBA3nJi6dh6-rmOLkA4Bb0d7h0tLAp7xE4")

YOUTUBE_PLAYLIST_IDS = {
    "A1": ["PL5vnwpT4NVTdwFarD9kwm1HONsqQ11l-b"],
    "A2": ["PLs7zUO7VPyJ7YxTq_g2Rcl3Jthd5bpTdY", "PLquImyRfMt6dVHL4MxFXMILrFh86H_HAc", "PLs7zUO7VPyJ5Eg0NOtF9g-RhqA25v385c"],
    "B1": ["PLs7zUO7VPyJ5razSfhOUVbTv9q6SAuPx-", "PLB92CD6B288E5DB61"],
    "B2": ["PLs7zUO7VPyJ5XMfT7pLvweRx6kHVgP_9C", "PLs7zUO7VPyJ6jZP-s6dlkINuEjFPvKMG0", "PLs7zUO7VPyJ4SMosRdB-35Q07brhnVToY"],
}


@st.cache_data(ttl=43200)
def fetch_youtube_playlist_videos(playlist_id, api_key=YOUTUBE_API_KEY):
    base_url = "https://www.googleapis.com/youtube/v3/playlistItems"
    params = {"part": "snippet", "playlistId": playlist_id, "maxResults": 50, "key": api_key}
    videos, next_page = [], ""
    while True:
        if next_page:
            params["pageToken"] = next_page
        response = requests.get(base_url, params=params, timeout=12)
        data = response.json()
        for item in data.get("items", []):
            vid = item["snippet"]["resourceId"]["videoId"]
            videos.append({"title": item["snippet"]["title"], "url": f"https://www.youtube.com/watch?v={vid}"})
        next_page = data.get("nextPageToken")
        if not next_page:
            break
    return videos

# ==== STUDENT SHEET LOADING ====
GOOGLE_SHEET_CSV = "https://docs.google.com/spreadsheets/d/12NXf5FeVHr7JJT47mRHh7Jp-TC1yhPS7ZG6nzZVTt1U/gviz/tq?tqx=out:csv&sheet=Sheet1"

@st.cache_data(ttl=300)
def load_student_data():
    try:
        resp = requests.get(GOOGLE_SHEET_CSV, timeout=12)
        resp.raise_for_status()
        # guard: ensure CSV not HTML
        txt = resp.text
        if "<html" in txt[:512].lower():
            raise RuntimeError("Expected CSV, got HTML (check sheet privacy).")
        df = pd.read_csv(io.StringIO(txt), dtype=str, keep_default_na=True, na_values=["", " ", "nan", "NaN", "None"])
    except Exception as e:
        st.error(f"❌ Could not load student data. {e}")
        st.stop()

    # Normalize headers and trim cells while preserving NaN
    df.columns = df.columns.str.strip().str.replace(" ", "")
    for col in df.columns:
        s = df[col]
        df[col] = s.where(s.isna(), s.astype(str).str.strip())

    # Keep only rows with a ContractEnd value
    df = df[df["ContractEnd"].notna() & (df["ContractEnd"].str.len() > 0)]

    # Robust parse
    def _parse_contract_end(s: str):
        for fmt in ("%m/%d/%Y", "%d/%m/%Y", "%Y-%m-%d"):
            try:
                return pd.to_datetime(s, format=fmt, errors="raise")
            except Exception:
                continue
        return pd.to_datetime(s, errors="coerce")

    df["ContractEnd_dt"] = df["ContractEnd"].apply(_parse_contract_end)
    df = df[df["ContractEnd_dt"].notna()]

    # Normalize identifiers
    if "StudentCode" in df.columns:
        df["StudentCode"] = df["StudentCode"].str.lower().str.strip()
    if "Email" in df.columns:
        df["Email"] = df["Email"].str.lower().str.strip()

    # Keep most recent per student
    df = (df.sort_values("ContractEnd_dt", ascending=False)
            .drop_duplicates(subset=["StudentCode"], keep="first")
            .drop(columns=["ContractEnd_dt"]))
    return df

def is_contract_expired(row):
    expiry_str = str(row.get("ContractEnd", "") or "").strip()
    if not expiry_str or expiry_str.lower() == "nan":
        return True
    expiry_date = None
    for fmt in ("%m/%d/%Y", "%d/%m/%Y", "%Y-%m-%d"):
        try:
            expiry_date = datetime.strptime(expiry_str, fmt); break
        except ValueError:
            continue
    if expiry_date is None:
        parsed = pd.to_datetime(expiry_str, errors="coerce")
        if pd.isnull(parsed): return True
        expiry_date = parsed.to_pydatetime()
    return expiry_date.date() < datetime.utcnow().date()

# ==== Query param helpers (stable) ====
def qp_get():
    # returns a dict-like object
    return st.query_params

def qp_clear():
    # clears all query params from the URL
    st.query_params.clear()

def qp_clear_keys(*keys):
    # remove only the specified keys
    for k in keys:
        try:
            del st.query_params[k]
        except KeyError:
            pass

# ==== Cookie helpers (normal cookies) ====
def _expire_str(dt: datetime) -> str:
    return dt.strftime("%a, %d %b %Y %H:%M:%S GMT")

def _js_set_cookie(name: str, value: str, max_age_sec: int, expires_gmt: str, secure: bool, domain: Optional[str] = None):
    base = (
        f'var c = {json.dumps(name)} + "=" + {json.dumps(_urllib.quote(value, safe=""))} + '
        f'"; Path=/; Max-Age={max_age_sec}; Expires={json.dumps(expires_gmt)}; SameSite=Lax";\n'
        f'if ({str(bool(secure)).lower()}) c += "; Secure";\n'
    )
    if domain:
        base += f'c += "; Domain=" + {domain};\n'
    base += "document.cookie = c;\n"
    return base

def set_student_code_cookie(cookie_manager, value: str, expires: datetime):
    key = "student_code"
    norm = (value or "").strip().lower()
    use_secure = (os.getenv("ENV", "prod") != "dev")
    max_age = 60 * 60 * 24 * 180  # 180 days
    exp_str = _expire_str(expires)
    # Library cookie (encrypted; host-only)
    try:
        cookie_manager.set(key, norm, expires=expires, secure=use_secure, samesite="Lax", path="/")
        cookie_manager.save()
    except Exception:
        try:
            cookie_manager[key] = norm; cookie_manager.save()
        except Exception:
            pass
    # JS host-only + base-domain (guard invalid hosts)
    host_cookie_name = (getattr(cookie_manager, 'prefix', '') or '') + key
    host_js = _js_set_cookie(host_cookie_name, norm, max_age, exp_str, use_secure, domain=None)
    script = f"""
    <script>
      (function(){{
        try {{
          {host_js}
          try {{
            var h = (window.location.hostname||'').split('.').filter(Boolean);
            if (h.length >= 2) {{
              var base = '.' + h.slice(-2).join('.');
              {_js_set_cookie(host_cookie_name, norm, max_age, exp_str, use_secure, "base")}
            }}
          }} catch(e) {{}}
          try {{ localStorage.setItem('student_code', {json.dumps(norm)}); }} catch(e) {{}}
        }} catch(e) {{}}
      }})();
    </script>
    """
    components.html(script, height=0)

def set_session_token_cookie(cookie_manager, token: str, expires: datetime):
    key = "session_token"
    val = (token or "").strip()
    use_secure = (os.getenv("ENV", "prod") != "dev")
    max_age = 60 * 60 * 24 * 30  # 30 days
    exp_str = _expire_str(expires)
    try:
        cookie_manager.set(key, val, expires=expires, secure=use_secure, samesite="Lax", path="/")
        cookie_manager.save()
    except Exception:
        try:
            cookie_manager[key] = val; cookie_manager.save()
        except Exception:
            pass
    host_cookie_name = (getattr(cookie_manager, 'prefix', '') or '') + key
    host_js = _js_set_cookie(host_cookie_name, val, max_age, exp_str, use_secure, domain=None)
    script = f"""
    <script>
      (function(){{
        try {{
          {host_js}
          try {{
            var h = (window.location.hostname||'').split('.').filter(Boolean);
            if (h.length >= 2) {{
              var base = '.' + h.slice(-2).join('.');
              {_js_set_cookie(host_cookie_name, val, max_age, exp_str, use_secure, "base")}
            }}
          }} catch(e) {{}}
          try {{ localStorage.setItem('session_token', {json.dumps(val)}); }} catch(e) {{}}
        }} catch(e) {{}}
      }})();
    </script>
    """
    components.html(script, height=0)

def _persist_session_client(token: str, student_code: str = "") -> None:
    components.html(f"""
    <script>
      try {{
        localStorage.setItem('session_token', {json.dumps(token)});
        if ({json.dumps(student_code)} !== "") {{
          localStorage.setItem('student_code', {json.dumps(student_code)});
        }}
        const u = new URL(window.location);
        ['code','state'].forEach(k => u.searchParams.delete(k));
        window.history.replaceState({{}}, '', u);
      }} catch(e) {{}}
    </script>
    """, height=0)

# ==== Cookie manager init ====
COOKIE_SECRET = os.getenv("COOKIE_SECRET") or st.secrets.get("COOKIE_SECRET")
if not COOKIE_SECRET:
    st.error("Cookie secret missing. Add COOKIE_SECRET to your Streamlit secrets.")
    st.stop()
cookie_manager = EncryptedCookieManager(prefix="falowen_", password=COOKIE_SECRET)
if not cookie_manager.ready():
    st.warning("Cookies not ready; please refresh.")
    st.stop()

# ---- Restore from existing session token (cookie) ----
restored = False
if not st.session_state.get("logged_in", False):
    cookie_tok = (cookie_manager.get("session_token") or "").strip()
    if cookie_tok:
        data = validate_session_token(cookie_tok, st.session_state.get("__ua_hash", ""))
        if data:
            # Validate the student still exists and contract active
            try:
                df_students = load_student_data()
                found = df_students[df_students["StudentCode"] == data.get("student_code","")]
            except Exception:
                found = pd.DataFrame()
            if not found.empty and not is_contract_expired(found.iloc[0]):
                row = found.iloc[0]
                st.session_state.update({
                    "logged_in": True,
                    "student_row": row.to_dict(),
                    "student_code": row["StudentCode"],
                    "student_name": row["Name"],
                    "session_token": cookie_tok,
                })
                new_tok = refresh_or_rotate_session_token(cookie_tok) or cookie_tok
                st.session_state["session_token"] = new_tok
                set_session_token_cookie(cookie_manager, new_tok, expires=datetime.utcnow() + timedelta(days=30))
                restored = True


# --- 2) Global CSS (tightened spacing) ---
st.markdown("""
<style>
  .hero {
    background: #fff; border-radius: 12px; padding: 24px; margin: 12px auto; max-width: 800px;
    box-shadow: 0 4px 16px rgba(0,0,0,0.05);
  }
  .help-contact-box {
    background: #fff; border-radius: 14px; padding: 20px; margin: 8px auto; max-width: 500px;
    box-shadow: 0 2px 10px rgba(0,0,0,0.04); border:1px solid #ebebf2; text-align:center;
  }
  .quick-links { display: flex; flex-wrap: wrap; gap:12px; justify-content:center; }
  .quick-links a {
    background: #e2e8f0; padding: 8px 16px; border-radius: 8px; font-weight:600; text-decoration:none;
    color:#0f172a; border:1px solid #cbd5e1;
  }
  .quick-links a:hover { background:#cbd5e1; }
  .stButton > button { background:#2563eb; color:#ffffff; font-weight:700; border-radius:8px; border:2px solid #1d4ed8; }
  .stButton > button:hover { background:#1d4ed8; }
  a:focus-visible, button:focus-visible, input:focus-visible, textarea:focus-visible, [role="button"]:focus-visible {
    outline:3px solid #f59e0b; outline-offset:2px; box-shadow:none !important;
  }
  input, textarea { color:#0f172a !important; }
  .page-wrap { max-width: 1100px; margin: 0 auto; }
  @media (max-width:600px){ .hero, .help-contact-box { padding:16px 4vw; } }
</style>
""", unsafe_allow_html=True)

GOOGLE_CLIENT_ID     = st.secrets.get("GOOGLE_CLIENT_ID", "180240695202-3v682khdfarmq9io9mp0169skl79hr8c.apps.googleusercontent.com")
GOOGLE_CLIENT_SECRET = st.secrets.get("GOOGLE_CLIENT_SECRET", "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")
REDIRECT_URI         = st.secrets.get("GOOGLE_REDIRECT_URI", "https://www.falowen.app/")


def _handle_google_oauth(code: str, state: str) -> None:
    df = load_student_data()
    df["Email"] = df["Email"].str.lower().str.strip()
    try:
        if st.session_state.get("_oauth_state") and state != st.session_state["_oauth_state"]:
            st.error("OAuth state mismatch. Please try again."); return
        if st.session_state.get("_oauth_code_redeemed") == code:
            return
        token_url = "https://oauth2.googleapis.com/token"
        data = {
            "code": code,
            "client_id": GOOGLE_CLIENT_ID,
            "client_secret": GOOGLE_CLIENT_SECRET,
            "redirect_uri": REDIRECT_URI,
            "grant_type": "authorization_code",
        }
        resp = requests.post(token_url, data=data, timeout=10)
        if not resp.ok:
            st.error(f"Google login failed: {resp.status_code} {resp.text}"); return
        access_token = resp.json().get("access_token")
        if not access_token:
            st.error("Google login failed: no access token."); return
        st.session_state["_oauth_code_redeemed"] = code
        userinfo = requests.get(
            "https://www.googleapis.com/oauth2/v2/userinfo",
            headers={"Authorization": f"Bearer {access_token}"},
            timeout=10,
        ).json()
        email = (userinfo.get("email") or "").lower().strip()
        match = df[df["Email"] == email]
        if match.empty:
            st.error("No student account found for that Google email."); return
        student_row = match.iloc[0]
        if is_contract_expired(student_row):
            st.error("Your contract has expired. Contact the office."); return
        ua_hash = st.session_state.get("__ua_hash", "")
        sess_token = create_session_token(student_row["StudentCode"], student_row["Name"], ua_hash=ua_hash)
        st.session_state.update({
            "logged_in": True,
            "student_row": student_row.to_dict(),
            "student_code": student_row["StudentCode"],
            "student_name": student_row["Name"],
            "session_token": sess_token,
        })
        set_student_code_cookie(cookie_manager, student_row["StudentCode"], expires=datetime.utcnow() + timedelta(days=180))
        _persist_session_client(sess_token, student_row["StudentCode"])
        set_session_token_cookie(cookie_manager, sess_token, expires=datetime.utcnow() + timedelta(days=30))
        qp_clear()
        st.success(f"Welcome, {student_row['Name']}!")
        st.rerun()
    except Exception as e:
        st.error(f"Google OAuth error: {e}")


def render_google_oauth():
    import secrets, urllib.parse

    def _qp_first(val):
        return val[0] if isinstance(val, list) else val

    qp = qp_get()
    code = _qp_first(qp.get("code")) if hasattr(qp, "get") else None
    state = _qp_first(qp.get("state")) if hasattr(qp, "get") else None
    if code:
        _handle_google_oauth(code, state)
        return
    st.session_state["_oauth_state"] = secrets.token_urlsafe(24)
    params = {
        "client_id": GOOGLE_CLIENT_ID,
        "redirect_uri": REDIRECT_URI,
        "response_type": "code",
        "scope": "openid email profile",
        "prompt": "select_account",
        "state": st.session_state["_oauth_state"],
        "include_granted_scopes": "true",
        "access_type": "online",
    }
    auth_url = "https://accounts.google.com/o/oauth2/v2/auth?" + urllib.parse.urlencode(params)
    st.markdown(
        """<div class="page-wrap" style='text-align:center;margin:12px 0;'>
                <a href="{url}">
                    <button aria-label="Sign in with Google"
                            style="background:#4285f4;color:white;padding:8px 24px;border:none;border-radius:6px;cursor:pointer;">
                        Sign in with Google
                    </button>
                </a>
           </div>""".replace("{url}", auth_url),
        unsafe_allow_html=True,
    )


def render_login_form():
    with st.form("login_form", clear_on_submit=False):
        login_id = st.text_input("Student Code or Email", help="Use your school email or Falowen code (e.g., felixa2)." ).strip().lower()
        login_pass = st.text_input("Password", type="password")
        login_btn = st.form_submit_button("Log In")
    if not login_btn:
        return
    df = load_student_data()
    df["StudentCode"] = df["StudentCode"].str.lower().str.strip()
    df["Email"] = df["Email"].str.lower().str.strip()
    lookup = df[(df["StudentCode"] == login_id) | (df["Email"] == login_id)]
    if lookup.empty:
        st.error("No matching student code or email found."); return
    student_row = lookup.iloc[0]
    if is_contract_expired(student_row):
        st.error("Your contract has expired. Contact the office."); return
    doc_ref = db.collection("students").document(student_row["StudentCode"])
    doc = doc_ref.get()
    if not doc.exists:
        st.error("Account not found. Please use 'Sign Up (Approved)' first."); return
    data = doc.to_dict() or {}
    stored_pw = data.get("password", "")
    is_hash = stored_pw.startswith(("$2a$", "$2b$", "$2y$")) and len(stored_pw) >= 60
    try:
        ok = bcrypt.checkpw(login_pass.encode("utf-8"), stored_pw.encode("utf-8")) if is_hash else stored_pw == login_pass
        if ok and not is_hash:
            new_hash = bcrypt.hashpw(login_pass.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")
            doc_ref.update({"password": new_hash})
    except Exception:
        ok = False
    if not ok:
        st.error("Incorrect password."); return
    ua_hash = st.session_state.get("__ua_hash", "")
    sess_token = create_session_token(student_row["StudentCode"], student_row["Name"], ua_hash=ua_hash)
    st.session_state.update({
        "logged_in": True,
        "student_row": dict(student_row),
        "student_code": student_row["StudentCode"],
        "student_name": student_row["Name"],
        "session_token": sess_token,
    })
    set_student_code_cookie(cookie_manager, student_row["StudentCode"], expires=datetime.utcnow() + timedelta(days=180))
    _persist_session_client(sess_token, student_row["StudentCode"])
    set_session_token_cookie(cookie_manager, sess_token, expires=datetime.utcnow() + timedelta(days=30))
    st.success(f"Welcome, {student_row['Name']}!")
    st.rerun()


def render_signup_form():
    with st.form("signup_form", clear_on_submit=False):
        new_name = st.text_input("Full Name", key="ca_name")
        new_email = st.text_input(
            "Email (must match teacher’s record)",
            help="Use the school email your tutor added to the roster.",
            key="ca_email",
        ).strip().lower()
        new_code = st.text_input("Student Code (from teacher)", help="Example: felixa2", key="ca_code").strip().lower()
        new_password = st.text_input("Choose a Password", type="password", key="ca_pass")
        signup_btn = st.form_submit_button("Create Account")
    if not signup_btn:
        return
    if not (new_name and new_email and new_code and new_password):
        st.error("Please fill in all fields."); return
    if len(new_password) < 8:
        st.error("Password must be at least 8 characters."); return
    df = load_student_data()
    df["StudentCode"] = df["StudentCode"].str.lower().str.strip()
    df["Email"] = df["Email"].str.lower().str.strip()
    valid = df[(df["StudentCode"] == new_code) & (df["Email"] == new_email)]
    if valid.empty:
        st.error("Your code/email aren’t registered. Use 'Request Access' first."); return
    doc_ref = db.collection("students").document(new_code)
    if doc_ref.get().exists:
        st.error("An account with this student code already exists. Please log in instead."); return
    hashed_pw = bcrypt.hashpw(new_password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")
    doc_ref.set({"name": new_name, "email": new_email, "password": hashed_pw})
    st.success("Account created! Please log in on the Returning tab.")


def render_reviews():
    # Richer, clearer data: goal, time, features used, outcome
    REVIEWS = [
        {
            "quote": "Falowen helped me pass A2 in 8 weeks. The assignments and feedback were spot on.",
            "author": "Ama",
            "location": "Accra, Ghana 🇬🇭",
            "level": "A2",
            "time": "20 weeks",
            "used": ["Course Book", "Assignments", "Results emails"],
            "outcome": "Passed Goethe A2"
        },
        {
            "quote": "The Course Book and Results emails keep me consistent. The vocab trainer is brilliant.",
            "author": "Tunde",
            "location": "Lagos, Nigeria 🇳🇬",
            "level": "B1",
            "time": "30 weeks",
            "used": ["Vocab Trainer", "Results emails", "Course Book"],
            "outcome": "Completed B1 modules"
        },
        {
            "quote": "Clear lessons, easy submissions, and I get notified quickly when marked.",
            "author": "Mariama",
            "location": "Freetown, Sierra Leone 🇸🇱",
            "level": "A1",
            "time": "10 weeks",
            "used": ["Assignments", "Course Book"],
            "outcome": "A1 basics completed"
        },
        {
            "quote": "I like the locked submissions and the clean Results tab.",
            "author": "Kwaku",
            "location": "Kumasi, Ghana 🇬🇭",
            "level": "B2",
            "time": "40 weeks",
            "used": ["Results tab", "Assignments"],
            "outcome": "B2 writing improved"
        },
    ]

    _html = """
    <div class="page-wrap" style="max-width:900px;margin-top:8px;">
      <section id="reviews" aria-label="Student stories" class="rev-wrap" tabindex="-1">
        <header class="rev-head">
          <h3 class="rev-title">Student stories</h3>
          <div class="rev-cta">
            <button class="rev-btn" id="rev_prev" aria-label="Previous review" title="Previous">◀</button>
            <button class="rev-btn" id="rev_next" aria-label="Next review" title="Next">▶</button>
          </div>
        </header>

        <article class="rev-card" aria-live="polite" aria-atomic="true">
          <blockquote id="rev_quote" class="rev-quote"></blockquote>
          <div class="rev-meta">
            <div class="rev-name" id="rev_author"></div>
            <div class="rev-sub"  id="rev_location"></div>
          </div>

          <div class="rev-badges">
            <span class="badge" id="rev_level"></span>
            <span class="badge" id="rev_time"></span>
            <span class="badge badge-ok" id="rev_outcome"></span>
          </div>

          <div class="rev-used" id="rev_used" aria-label="Features used"></div>
        </article>

        <nav class="rev-dots" aria-label="Slide indicators" id="rev_dots"></nav>
      </section>
    </div>

    <style>
      .rev-wrap{
        background:#fff; border:1px solid #e5e7eb; border-radius:12px; padding:14px; 
        box-shadow:0 4px 16px rgba(0,0,0,.05);
      }
      .rev-head{ display:flex; align-items:center; justify-content:space-between; margin-bottom:8px; }
      .rev-title{ margin:0; font-size:1.05rem; color:#25317e; }
      .rev-cta{ display:flex; gap:6px; }
      .rev-btn{
        background:#eef3fc; border:1px solid #cbd5e1; border-radius:8px; padding:4px 10px; cursor:pointer; 
        font-weight:700;
      }
      .rev-btn:hover{ background:#e2e8f0; }

      .rev-card{ position:relative; min-height:190px; }
      .rev-quote{ font-size:1.06rem; line-height:1.45; margin:4px 0 10px 0; color:#0f172a; }
      .rev-meta{ display:flex; gap:10px; flex-wrap:wrap; margin-bottom:8px; }
      .rev-name{ font-weight:700; color:#1e293b; }
      .rev-sub{ color:#475569; }

      .rev-badges{ display:flex; gap:6px; flex-wrap:wrap; margin:6px 0 8px; }
      .badge{
        display:inline-block; background:#f1f5f9; border:1px solid #e2e8f0; color:#0f172a;
        padding:4px 8px; border-radius:999px; font-size:.86rem; font-weight:600;
      }
      .badge-ok{ background:#ecfdf5; border-color:#bbf7d0; color:#065f46; }

      .rev-used{ display:flex; gap:6px; flex-wrap:wrap; }
      .rev-used .chip{
        background:#eef2ff; border:1px solid #c7d2fe; color:#3730a3; 
        padding:3px 8px; border-radius:999px; font-size:.82rem; font-weight:600;
      }

      .rev-dots{ display:flex; gap:6px; justify-content:center; margin-top:10px; }
      .rev-dot{
        width:8px; height:8px; border-radius:999px; background:#cbd5e1; border:none; padding:0; cursor:pointer;
      }
      .rev-dot[aria-current="true"]{ background:#25317e; }

      /* Motion awareness */
      .fade{ opacity:0; transform:translateY(4px); transition:opacity .28s ease, transform .28s ease; }
      .fade.show{ opacity:1; transform:none; }
      @media (prefers-reduced-motion: reduce){
        .fade{ transition:none; opacity:1; transform:none; }
      }
    </style>

    <script>
      const DATA = __DATA__;
      const q  = (id) => document.getElementById(id);
      const qs = (sel) => document.querySelector(sel);
      const wrap = qs("#reviews");
      const quote = q("rev_quote");
      const author = q("rev_author");
      const locationEl = q("rev_location");
      const level = q("rev_level");
      const time  = q("rev_time");
      const outcome = q("rev_outcome");
      const used = q("rev_used");
      const dots = q("rev_dots");
      const prevBtn = q("rev_prev");
      const nextBtn = q("rev_next");

      let i = 0, timer = null, hovered = false;
      const reduced = window.matchMedia('(prefers-reduced-motion: reduce)').matches;

      function setUsedChips(items){
        used.innerHTML = "";
        (items || []).forEach(t => {
          const s = document.createElement("span");
          s.className = "chip";
          s.textContent = t;
          used.appendChild(s);
        });
      }

      function setDots(){
        dots.innerHTML = "";
        DATA.forEach((_, idx) => {
          const b = document.createElement("button");
          b.className = "rev-dot";
          b.setAttribute("aria-label", "Go to review " + (idx+1));
          if(idx === i) b.setAttribute("aria-current","true");
          b.addEventListener("click", () => { i = idx; show(true); restart(); });
          dots.appendChild(b);
        });
      }

      function show(animate){
        const c = DATA[i];
        quote.textContent = '"' + (c.quote || '') + '"';
        author.textContent = c.author ? c.author + ' — ' : '';
        locationEl.textContent = c.location || '';
        level.textContent = 'Level: ' + (c.level || '—');
        time.textContent  = 'Time: ' + (c.time  || '—');
        outcome.textContent = c.outcome || '';

        setUsedChips(c.used);
        setDots();

        const card = wrap.querySelector(".rev-card");
        if(animate && !reduced){
          card.classList.remove("show");
          card.classList.add("fade");
          requestAnimationFrame(() => {
            requestAnimationFrame(() => card.classList.add("show"));
          });
        }
      }

      function next(){ i = (i + 1) % DATA.length; show(true); }
      function prev(){ i = (i - 1 + DATA.length) % DATA.length; show(true); }

      function start(){
        if(reduced) return;
        timer = setInterval(() => { if(!hovered) next(); }, 6000);
      }
      function stop(){ if(timer){ clearInterval(timer); timer = null; } }
      function restart(){ stop(); start(); }

      // Events
      nextBtn.addEventListener("click", () => { next(); restart(); });
      prevBtn.addEventListener("click", () => { prev(); restart(); });
      wrap.addEventListener("mouseenter", () => { hovered = true; });
      wrap.addEventListener("mouseleave", () => { hovered = false; });

      // Keyboard nav
      wrap.addEventListener("keydown", (e) => {
        if(e.key === "ArrowRight"){ next(); restart(); }
        if(e.key === "ArrowLeft"){  prev(); restart(); }
      });

      // Init
      show(false);
      start();
    </script>
    """
    # NOTE: height tuned; no scrollbars; fixed a padding typo from previous HTML
    _json = json.dumps(REVIEWS)
    components.html(_html.replace("__DATA__", _json), height=300, scrolling=False)

def login_page():

    # Optional container width helper (safe if you already defined it in global CSS)
    st.markdown('<style>.page-wrap{max-width:1100px;margin:0 auto;}</style>', unsafe_allow_html=True)

    # HERO FIRST — this is the first visible element on the page
    st.markdown("""
    <div class="page-wrap">
      <div class="hero" aria-label="Falowen app introduction">
        <h1 style="text-align:center; color:#25317e;">👋 Welcome to <strong>Falowen</strong></h1>
        <p style="text-align:center; font-size:1.1em; color:#555;">
          Falowen is your all-in-one German learning platform, powered by
          <b>Learn Language Education Academy</b>, with courses and vocabulary from
          <b>A1 to C1</b> levels and live tutor support.
        </p>
        <ul style="max-width:700px; margin:16px auto; color:#444; font-size:1em; line-height:1.5;">
          <li>📊 <b>Dashboard</b>: Track your learning streaks, assignment progress, active contracts, and more.</li>
          <li>📚 <b>Course Book</b>: Access lecture videos, grammar modules, and submit assignments for levels A1–C1 in one place.</li>
          <li>📝 <b>Exams & Quizzes</b>: Take practice tests and official exam prep right in the app.</li>
          <li>💬 <b>Custom Chat</b>: Sprechen & expression trainer for live feedback on your speaking.</li>
          <li>🏆 <b>Results Tab</b>: View your grades, feedback, and historical performance at a glance.</li>
          <li>🔤 <b>Vocab Trainer</b>: Practice and master A1–C1 vocabulary with spaced-repetition quizzes.</li>
          <li>✍️ <b>Schreiben Trainer</b>: Improve your writing with guided exercises and instant corrections.</li>
        </ul>
      </div>
    </div>
    """, unsafe_allow_html=True)

    
    # ===== Compact stats strip =====
    st.markdown("""
      <style>
        .stats-strip { display:flex; flex-wrap:wrap; gap:10px; justify-content:center; margin:10px auto 4px auto; max-width:820px; }
        .stat { background:#0ea5e9; color:#ffffff; border-radius:12px; padding:12px 14px; min-width:150px; text-align:center;
                box-shadow:0 2px 10px rgba(2,132,199,0.15); outline: none; }
        .stat:focus-visible { outline:3px solid #1f2937; outline-offset:2px; }
        .stat .num { font-size:1.25rem; font-weight:800; line-height:1; }
        .stat .label { font-size:.92rem; opacity:.98; }
        @media (max-width:560px){ .stat { min-width:46%; } }
      </style>
      <div class="stats-strip" role="list" aria-label="Falowen highlights">
        <div class="stat" role="listitem" tabindex="0" aria-label="Active learners: over 300">
          <div class="num">300+</div>
          <div class="label">Active learners</div>
        </div>
        <div class="stat" role="listitem" tabindex="0" aria-label="Assignments submitted">
          <div class="num">1,200+</div>
          <div class="label">Assignments submitted</div>
        </div>
        <div class="stat" role="listitem" tabindex="0" aria-label="Levels covered: A1 to C1">
          <div class="num">A1–C1</div>
          <div class="label">Full course coverage</div>
        </div>
        <div class="stat" role="listitem" tabindex="0" aria-label="Average student feedback">
          <div class="num">4.8/5</div>
          <div class="label">Avg. feedback</div>
        </div>
      </div>
    """, unsafe_allow_html=True)

    with st.expander("📌 Which option should I choose?", expanded=True):
        st.markdown("""
        <div class="option-box">
          <div class="option-item">
            <div class="option-icon">👋</div>
            <div><b>Returning Student</b>: You already created a password — simply log in to continue your learning.</div>
          </div>
          <div class="option-item">
            <div class="option-icon">🧾</div>
            <div><b>Sign Up (Approved)</b>: You’ve paid and your email + code are already on our roster, but you don’t have an account yet — create one here.</div>
          </div>
          <div class="option-item">
            <div class="option-icon">📝</div>
            <div><b>Request Access</b>: New to Falowen? Fill out our form and we’ll get in touch to guide you through the next steps.</div>
          </div>
        </div>
        """, unsafe_allow_html=True)


    tab1, tab2, tab3 = st.tabs(["👋 Returning", "🧾 Sign Up (Approved)", "📝 Request Access"])
#

    with tab1:
        render_google_oauth()
        st.markdown("<div class='page-wrap' style='text-align:center; margin:8px 0;'>⎯⎯⎯ or ⎯⎯⎯</div>", unsafe_allow_html=True)
        render_login_form()

    with tab2:
        render_signup_form()

    # --- Request Access ---
    with tab3:
        st.markdown(
            """
            <div class="page-wrap" style="text-align:center; margin-top:20px;">
                <p style="font-size:1.1em; color:#444;">
                    If you don't have an account yet, please request access by filling out this form.
                </p>
                <a href="https://docs.google.com/forms/d/e/1FAIpQLSenGQa9RnK9IgHbAn1I9rSbWfxnztEUcSjV0H-VFLT-jkoZHA/viewform?usp=header" 
                   target="_blank" rel="noopener">
                    <button style="background:#25317e; color:white; padding:10px 20px; border:none; border-radius:6px; cursor:pointer;">
                        📝 Open Request Access Form
                    </button>
                </a>
            </div>
            """,
            unsafe_allow_html=True
        )

    
    st.markdown("""
    <div class="page-wrap">
      <div class="help-contact-box" aria-label="Help and contact options">
        <b>❓ Need help or access?</b><br>
        <a href="https://api.whatsapp.com/send?phone=233205706589" target="_blank" rel="noopener">📱 WhatsApp us</a>
        &nbsp;|&nbsp;
        <a href="mailto:learngermanghana@gmail.com" target="_blank" rel="noopener">✉️ Email</a>
      </div>
    </div>
    """, unsafe_allow_html=True)


    # --- Centered Video (pick a frame style by changing the class) ---
    st.markdown("""
    <div class="page-wrap">
      <div class="video-wrap">
        <div class="video-shell style-gradient">
          <video
            width="360"
            autoplay
            muted
            loop
            playsinline
            tabindex="-1"
            oncontextmenu="return false;"
            draggable="false"
            style="pointer-events:none; user-select:none; -webkit-user-select:none; -webkit-touch-callout:none;">
            <source src="https://raw.githubusercontent.com/learngermanghana/a1spreche/main/falowen.mp4" type="video/mp4">
            Sorry, your browser doesn't support embedded videos.
          </video>
        </div>
      </div>
    </div>

    <style>
      /* Layout */
      .video-wrap{
        display:flex; justify-content:center; align-items:center;
        margin: 12px 0 24px;
      }
      .video-shell{
        position:relative; border-radius:16px; padding:4px;
      }
      .video-shell > video{
        display:block; width:min(360px, 92vw); border-radius:12px; margin:0;
        box-shadow: 0 4px 12px rgba(0,0,0,.08);
      }

      /* 1) Soft gradient frame (default) */
      .video-shell.style-gradient{
        background: linear-gradient(135deg,#e8eeff,#f6f9ff);
        box-shadow: 0 8px 24px rgba(0,0,0,.08);
      }

      /* 2) Glow pulse */
      .video-shell.style-glow{
        background:#0b1220;
        box-shadow: 0 0 0 2px #1d4ed8, 0 0 18px #1d4ed8;
        animation: glowPulse 3.8s ease-in-out infinite;
      }
      @keyframes glowPulse{
        0%,100%{ box-shadow:0 0 0 2px #1d4ed8, 0 0 12px #1d4ed8; }
        50%{    box-shadow:0 0 0 2px #06b6d4, 0 0 22px #06b6d4; }
      }

      /* 3) Glassmorphism */
      .video-shell.style-glass{
        background: rgba(255,255,255,.25);
        border: 1px solid rgba(255,255,255,.35);
        backdrop-filter: blur(6px);
        -webkit-backdrop-filter: blur(6px);
        box-shadow: 0 10px 30px rgba(0,0,0,.10);
      }

      /* 4) Animated dashes */
      .video-shell.style-dash{
        padding:6px; border-radius:18px;
        background:
          repeating-linear-gradient(90deg,#1d4ed8 0 24px,#93c5fd 24px 48px);
        background-size: 48px 100%;
        animation: dashMove 6s linear infinite;
      }
      @keyframes dashMove { to { background-position: 48px 0; } }

      /* 5) Shimmer frame */
      .video-shell.style-shimmer{
        background: linear-gradient(120deg,#e5e7eb, #f8fafc, #e5e7eb);
        background-size: 200% 200%;
        animation: shimmer 6s linear infinite;
        box-shadow: 0 8px 24px rgba(0,0,0,.08);
      }
      @keyframes shimmer{ 0%{background-position:0% 50%;} 100%{background-position:100% 50%;} }

      /* Mobile nudge */
      @media (max-width:600px){
        .video-wrap{ margin: 8px 0 16px; }
      }
    </style>
    """, unsafe_allow_html=True)
    #
#

    # Quick Links
    st.markdown("""
    <div class="page-wrap">
      <div class="quick-links" aria-label="Useful links">
        <a href="https://www.learngermanghana.com/tutors"           target="_blank" rel="noopener">👩‍🏫 Tutors</a>
        <a href="https://www.learngermanghana.com/upcoming-classes" target="_blank" rel="noopener">🗓️ Upcoming Classes</a>
        <a href="https://www.learngermanghana.com/accreditation"    target="_blank" rel="noopener">✅ Accreditation</a>
        <a href="https://www.learngermanghana.com/privacy-policy"   target="_blank" rel="noopener">🔒 Privacy</a>
        <a href="https://www.learngermanghana.com/terms-of-service" target="_blank" rel="noopener">📜 Terms</a>
        <a href="https://www.learngermanghana.com/contact-us"       target="_blank" rel="noopener">✉️ Contact</a>
      </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")

    LOGIN_IMG_URL      = "https://i.imgur.com/pFQ5BIn.png"
    COURSEBOOK_IMG_URL = "https://i.imgur.com/pqXoqSC.png"
    RESULTS_IMG_URL    = "https://i.imgur.com/uiIPKUT.png"

    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown(f"""
        <img src="{LOGIN_IMG_URL}" alt="Login screenshot"
             style="width:100%; height:220px; object-fit:cover; border-radius:12px; pointer-events:none; user-select:none;">
        <div style="height:8px;"></div>
        <h3 style="margin:0 0 4px 0;">1️⃣ Sign in</h3>
        <p style="margin:0;">Use your <b>student code or email</b> and start your level (A1–C1).</p>
        """, unsafe_allow_html=True)
    with c2:
        st.markdown(f"""
        <img src="{COURSEBOOK_IMG_URL}" alt="Course Book screenshot"
             style="width:100%; height:220px; object-fit:cover; border-radius:12px; pointer-events:none; user-select:none;">
        <div style="height:8px;"></div>
        <h3 style="margin:0 0 4px 0;">2️⃣ Learn & submit</h3>
        <p style="margin:0;">Watch lessons, practice vocab, and <b>submit assignments</b> in the Course Book.</p>
        """, unsafe_allow_html=True)
    with c3:
        st.markdown(f"""
        <img src="{RESULTS_IMG_URL}" alt="Results screenshot"
             style="width:100%; height:220px; object-fit:cover; border-radius:12px; pointer-events:none; user-select:none;">
        <div style="height:8px;"></div>
        <h3 style="margin:0 0 4px 0;">3️⃣ Get results</h3>
        <p style="margin:0;">You’ll get an <b>email when marked</b>. Check <b>Results & Resources</b> for feedback.</p>
        """, unsafe_allow_html=True)

        # --- Student Stories Section ---
    st.markdown("""
    <style>
      .section-title {
        font-weight:700;
        font-size:1.15rem;
        padding-left:12px;
        border-left:5px solid #2563eb;
        margin: 12px 0 12px 0;
      }
      @media (prefers-color-scheme: dark){
        .section-title { border-left-color:#3b82f6; color:#f1f5f9; }
      }
    </style>
    <div class="page-wrap">
      <div class="section-title">💬 Student Stories</div>
    </div>
    """, unsafe_allow_html=True)

    def render_reviews():
        REVIEWS = [
            {"quote": "Falowen helped me pass A2 in 8 weeks. The assignments and feedback were spot on.", "author": "Ama — Accra, Ghana 🇬🇭", "level": "A2"},
            {"quote": "The Course Book and Results emails keep me consistent. The vocab trainer is brilliant.", "author": "Tunde — Lagos, Nigeria 🇳🇬", "level": "B1"},
            {"quote": "Clear lessons, easy submissions, and I get notified quickly when marked.", "author": "Mariama — Freetown, Sierra Leone 🇸🇱", "level": "A1"},
            {"quote": "I like the locked submissions and the clean Results tab.", "author": "Kwaku — Kumasi, Ghana 🇬🇭", "level": "B2"},
        ]

        _reviews_html = """
        <style>
          :root{
            --bg: #0b1220;
            --card:#ffffffcc;
            --text:#0f172a;
            --muted:#475569;
            --brand:#2563eb;
            --chip:#e0f2fe;
            --chip-text:#0369a1;
            --ring:#93c5fd;
          }
          @media (prefers-color-scheme: dark){
            :root{
              --card:#0b1220cc;
              --text:#e2e8f0;
              --muted:#94a3b8;
              --chip:#1e293b;
              --chip-text:#e2e8f0;
              --ring:#334155;
            }
          }
          .page-wrap{max-width:900px;margin:8px auto;}
          .rev-shell{
            position:relative; isolation:isolate;
            border-radius:16px; padding:18px 16px 20px 16px;
            background: radial-gradient(1200px 300px at 10% -10%, #e0f2fe55, transparent),
                        radial-gradient(1200px 300px at 90% 110%, #c7d2fe44, transparent);
            border:1px solid rgba(148,163,184,.25);
            box-shadow: 0 10px 30px rgba(2,6,23,.08);
            overflow:hidden;
          }
          .rev-card{
            background: var(--card);
            backdrop-filter: blur(8px);
            border:1px solid rgba(148,163,184,.25);
            border-radius:16px; padding:20px 18px; min-height:170px;
          }
          .rev-quote{
            font-size:1.06rem; line-height:1.55; color:var(--text); margin:0;
          }
          .rev-meta{
            display:flex; align-items:center; gap:10px; margin-top:14px; color:var(--muted);
          }
          .rev-chip{
            font-size:.78rem; font-weight:700;
            background:var(--chip); color:var(--chip-text);
            border-radius:999px; padding:6px 10px;
          }
          .rev-author{ font-weight:700; color:var(--text); }
          .rev-dots{
            display:flex; gap:6px; justify-content:center; margin-top:14px;
          }
          .rev-dot{
            width:8px; height:8px; border-radius:999px;
            background:#cbd5e1; opacity:.8; transform:scale(.9);
            transition: all .25s ease;
          }
          .rev-dot[aria-current="true"]{
            background:var(--brand); opacity:1; transform:scale(1.15);
            box-shadow:0 0 0 4px var(--ring);
          }
        </style>
        <div class="page-wrap">
          <div id="reviews" class="rev-shell">
            <div class="rev-card" id="rev_card">
              <p id="rev_quote" class="rev-quote"></p>
              <div class="rev-meta">
                <span id="rev_level" class="rev-chip"></span>
                <span id="rev_author" class="rev-author"></span>
              </div>
              <div class="rev-dots" id="rev_dots"></div>
            </div>
          </div>
        </div>
        <script>
          const data = __DATA__;
          const q = document.getElementById('rev_quote');
          const a = document.getElementById('rev_author');
          const l = document.getElementById('rev_level');
          const dotsWrap = document.getElementById('rev_dots');
          let i = 0;
          function setActiveDot(idx){
            [...dotsWrap.children].forEach((d, j) => d.setAttribute('aria-current', j === idx ? 'true' : 'false'));
          }
          function render(idx){
            const c = data[idx];
            q.textContent = c.quote;
            a.textContent = c.author;
            l.textContent = "Level " + c.level;
            setActiveDot(idx);
          }
          function next(){
            i = (i + 1) % data.length;
            render(i);
          }
          data.forEach((_, idx) => {
            const dot = document.createElement('button');
            dot.className = 'rev-dot';
            dot.type = 'button';
            dot.addEventListener('click', () => { i = idx; render(i); });
            dotsWrap.appendChild(dot);
          });
          setInterval(next, 6000);
          render(i);
        </script>
        """
        _reviews_json = json.dumps(REVIEWS, ensure_ascii=False)
        components.html(_reviews_html.replace("__DATA__", _reviews_json), height=300, scrolling=False)

    # --- Render reviews below Quick Links + Steps ---
    render_reviews()
#


    st.markdown("---")

    with st.expander("How do I log in?"):
        st.write("Use your school email **or** Falowen code (e.g., `felixa2`). If you’re new, request access first.")
    with st.expander("Where do I see my scores?"):
        st.write("Scores are emailed to you and live in **Results & Resources** inside the app.")
    with st.expander("How do assignments work?"):
        st.write("Type your answer, confirm, and **submit**. The box locks. Your tutor is notified automatically.")
    with st.expander("What if I open the wrong lesson?"):
        st.write("Check the blue banner at the top (Level • Day • Chapter). Use the dropdown to switch to the correct page.")

    st.markdown("""
    <div class="page-wrap" style="text-align:center; margin:24px 0;">
      <a href="https://www.youtube.com/YourChannel" target="_blank" rel="noopener">📺 YouTube</a>
      &nbsp;|&nbsp;
      <a href="https://api.whatsapp.com/send?phone=233205706589" target="_blank" rel="noopener">📱 WhatsApp</a>
    </div>
    """, unsafe_allow_html=True)

    st.markdown(f"""
    <div class="page-wrap" style="text-align:center;color:#64748b; margin-bottom:16px;">
      © {datetime.utcnow().year} Learn Language Education Academy • Accra, Ghana<br>
      Need help? <a href="mailto:learngermanghana@gmail.com">Email</a> • 
      <a href="https://api.whatsapp.com/send?phone=233205706589" target="_blank" rel="noopener">WhatsApp</a>
    </div>
    """, unsafe_allow_html=True)

    st.stop()

# =========================
# Logged-in header + Logout (no callback; rerun works)
# =========================


# --- helper for query params ---
def qp_clear_keys(*keys):
    for k in keys:
        try:
            del st.query_params[k]
        except KeyError:
            pass

# --- run once right after a logout to clean client storage & URL ---
if st.session_state.pop("_inject_logout_js", False):
    components.html("""
      <script>
        try {
          localStorage.removeItem('student_code');
          localStorage.removeItem('session_token');
          const u = new URL(window.location);
          ['code','state','token'].forEach(k => u.searchParams.delete(k));
          window.history.replaceState({}, '', u);
        } catch(e) {}
      </script>
    """, height=0)

# ===== AUTH GUARD =====
if not st.session_state.get("logged_in", False):
    login_page()
    st.stop()

# ===== Header + plain button (no on_click) =====
st.markdown("""
<style>
  .post-login-header { margin-top:0; margin-bottom:4px; }
  .block-container { padding-top: 0.6rem !important; }
  div[data-testid="stExpander"] { margin-top: 6px !important; margin-bottom: 6px !important; }
  .your-notifs { margin: 4px 0 !important; }
</style>
""", unsafe_allow_html=True)

st.markdown("<div class='post-login-header'>", unsafe_allow_html=True)
col1, col2 = st.columns([0.85, 0.15])
with col1:
    st.write(f"👋 Welcome, **{st.session_state.get('student_name','Student')}**")
with col2:
    st.markdown("<div style='display:flex;justify-content:flex-end;align-items:center;'>", unsafe_allow_html=True)
    _logout_clicked = st.button("Log out", key="logout_btn")  # <-- no on_click
    st.markdown("</div>", unsafe_allow_html=True)
st.markdown("</div>", unsafe_allow_html=True)

# ===== Logout handling (works in all versions) =====
if _logout_clicked:
    # 1) Revoke server token if available
    try:
        tok = st.session_state.get("session_token", "")
        if tok and "destroy_session_token" in globals():
            destroy_session_token(tok)
    except Exception as e:
        st.warning(f"Logout warning (revoke): {e}")

    # 2) Expire cookies
    try:
        expires_past = datetime.utcnow() - timedelta(seconds=1)
        if "set_student_code_cookie" in globals():
            set_student_code_cookie(cookie_manager, "", expires=expires_past)
        if "set_session_token_cookie" in globals():
            set_session_token_cookie(cookie_manager, "", expires=expires_past)
    except Exception as e:
        st.warning(f"Logout warning (expire cookies): {e}")

    try:
        cookie_manager.delete("student_code")
        cookie_manager.delete("session_token")
        cookie_manager.save()
    except Exception:
        pass

    # 3) Clean server-side URL params
    qp_clear_keys("code", "state", "token")

    # 4) Reset session state
    st.session_state.update({
        "logged_in": False,
        "student_row": None,
        "student_code": "",
        "student_name": "",
        "session_token": "",
        "cookie_synced": False,
        "__last_refresh": 0.0,
        "__ua_hash": "",
        "_oauth_state": "",
        "_oauth_code_redeemed": "",
    })

    # 5) On next run, clear localStorage & URL on the client
    st.session_state["_inject_logout_js"] = True

    # 6) Now safe to rerun (not in a callback)
    st.rerun()



# =========================================================
# ============= Announcements (mobile-friendly) ===========
# =========================================================
def render_announcements(ANNOUNCEMENTS: list):
    """Responsive rotating announcement board with mobile-first, light card on phones."""
    if not ANNOUNCEMENTS:
        st.info("📣 No announcements to show.")
        return

    _html = """
    <style>
      /* ---------- THEME TOKENS ---------- */
      :root{
        /* brand */
        --brand:#1d4ed8;      /* primary */
        --ring:#93c5fd;

        /* light defaults */
        --text:#0b1220;
        --muted:#475569;
        --card:#ffffff;       /* <- light card by default */
        --chip-bg:#eaf2ff;
        --chip-fg:#1e3a8a;
        --link:#1d4ed8;
        --shell-border: rgba(2,6,23,.08);
      }

      /* Dark scheme (desktop/tablet). We will still force light card on phones below. */
      @media (prefers-color-scheme: dark){
        :root{
          --text:#e5e7eb;
          --muted:#cbd5e1;
          --card:#111827;
          --chip-bg:#1f2937;
          --chip-fg:#e5e7eb;
          --link:#93c5fd;
          --shell-border: rgba(148,163,184,.25);
        }
      }

      /* ---------- LAYOUT ---------- */
      .page-wrap{max-width:1100px;margin:0 auto;padding:0 10px;}
      .ann-title{
        font-weight:800; font-size:1.05rem; line-height:1.2;
        padding-left:12px; border-left:5px solid var(--brand);
        margin: 0 0 6px 0; color: var(--text);
        letter-spacing: .2px;
      }
      .ann-shell{
        border-radius:14px;
        border:1px solid var(--shell-border);
        background:var(--card);
        box-shadow:0 6px 18px rgba(2,6,23,.12);
        padding:12px 14px; isolation:isolate; overflow:hidden;
      }
      .ann-heading{
        display:flex; align-items:center; gap:10px; margin:0 0 6px 0;
        font-weight:800; color:var(--text); letter-spacing:.2px;
      }
      .ann-chip{
        font-size:.78rem; font-weight:800; text-transform:uppercase;
        background:var(--chip-bg); color:var(--chip-fg);
        padding:4px 9px; border-radius:999px; border:1px solid var(--shell-border);
      }
      .ann-body{ color:var(--muted); margin:0; line-height:1.55; font-size:1rem }
      .ann-actions{ margin-top:8px }
      .ann-actions a{ color:var(--link); text-decoration:none; font-weight:700 }

      .ann-dots{
        display:flex; gap:12px; justify-content:center; margin-top:12px
      }
      .ann-dot{
        width:11px; height:11px; border-radius:999px; background:#9ca3af;
        opacity:.9; transform:scale(.95);
        transition:transform .2s, background .2s, opacity .2s;
        border:none; cursor:pointer;
      }
      .ann-dot[aria-current="true"]{
        background:var(--brand); opacity:1; transform:scale(1.22);
        box-shadow:0 0 0 4px var(--ring)
      }

      @keyframes fadeInUp{from{opacity:0;transform:translateY(6px)}to{opacity:1;transform:translateY(0)}}
      .ann-anim{animation:fadeInUp .25s ease both}
      @media (prefers-reduced-motion: reduce){ .ann-anim{animation:none} .ann-dot{transition:none} }

      /* ---------- MOBILE OVERRIDES ---------- */
      @media (max-width: 640px){
        /* Force a light look on phones, regardless of system dark mode */
        :root{
          --card:#ffffff !important;
          --text:#0b1220 !important;
          --muted:#334155 !important;
          --link:#1d4ed8 !important;
          --chip-bg:#eaf2ff !important;
          --chip-fg:#1e3a8a !important;
          --shell-border: rgba(2,6,23,.10) !important;
        }
        .page-wrap{ padding:0 8px; }
        .ann-shell{ padding:10px 12px; border-radius:12px; }
        .ann-title{ font-size:1rem; margin:0 0 4px 0; }
        .ann-heading{ gap:8px; }
        .ann-chip{ font-size:.72rem; padding:3px 8px; }
        .ann-body{ font-size:1.02rem; line-height:1.6; }
        .ann-dots{ gap:10px; margin-top:10px; }
        .ann-dot{ width:12px; height:12px; }
      }

      /* Tight spacer utility for Streamlit blocks around this widget */
      .tight-section{ margin:6px 0 !important; }
    </style>

    <div class="page-wrap tight-section">
      <div class="ann-title">📣 Announcements</div>
      <div class="ann-shell" id="ann_shell" aria-live="polite">
        <div class="ann-anim" id="ann_card">
          <div class="ann-heading">
            <span class="ann-chip" id="ann_tag" style="display:none;"></span>
            <span id="ann_title"></span>
          </div>
          <p class="ann-body" id="ann_body">loading…</p>
          <div class="ann-actions" id="ann_action" style="display:none;"></div>
        </div>
        <div class="ann-dots" id="ann_dots" role="tablist" aria-label="Announcement selector"></div>
      </div>
    </div>

    <script>
      const data = __DATA__;
      const titleEl = document.getElementById('ann_title');
      const bodyEl  = document.getElementById('ann_body');
      const tagEl   = document.getElementById('ann_tag');
      const actionEl= document.getElementById('ann_action');
      const dotsWrap= document.getElementById('ann_dots');
      const card    = document.getElementById('ann_card');
      const shell   = document.getElementById('ann_shell');
      const reduced = window.matchMedia('(prefers-reduced-motion: reduce)').matches;

      let i = 0, timer = null;
      const INTERVAL = 6500;

      function setActiveDot(idx){
        [...dotsWrap.children].forEach((d, j)=> d.setAttribute('aria-current', j===idx ? 'true' : 'false'));
      }
      function render(idx){
        const c = data[idx] || {};
        card.classList.remove('ann-anim'); void card.offsetWidth; card.classList.add('ann-anim');

        titleEl.textContent = c.title || '';
        bodyEl.textContent  = c.body  || '';

        if (c.tag){
          tagEl.textContent = c.tag;
          tagEl.style.display='';
        } else {
          tagEl.style.display='none';
        }

        if (c.href){
          const link = document.createElement('a');
          link.href = c.href; link.target = '_blank'; link.rel = 'noopener';
          link.textContent = 'Open';
          actionEl.textContent = '';
          actionEl.appendChild(link);
          actionEl.style.display='';
        } else {
          actionEl.style.display='none';
          actionEl.textContent = '';
        }
        setActiveDot(idx);
      }
      function next(){ i = (i+1) % data.length; render(i); }
      function start(){ if (!reduced && data.length > 1) timer = setInterval(next, INTERVAL); }
      function stop(){ if (timer) clearInterval(timer); timer = null; }
      function restart(){ stop(); start(); }

      data.forEach((_, idx)=>{
        const dot = document.createElement('button');
        dot.className='ann-dot'; dot.type='button'; dot.setAttribute('role','tab');
        dot.setAttribute('aria-label','Show announcement '+(idx+1));
        dot.addEventListener('click', ()=>{ i=idx; render(i); restart(); });
        dotsWrap.appendChild(dot);
      });

      shell.addEventListener('mouseenter', stop);
      shell.addEventListener('mouseleave', start);
      shell.addEventListener('focusin', stop);
      shell.addEventListener('focusout', start);

      render(i); start();
    </script>
    """
    data_json = json.dumps(ANNOUNCEMENTS, ensure_ascii=False)
    components.html(_html.replace("__DATA__", data_json), height=220, scrolling=False)


# Optional: extra style injector for status chips & mini-cards if you want to reuse elsewhere
def inject_notice_css():
    st.markdown("""
    <style>
      :root{
        --chip-border: rgba(148,163,184,.35);
      }
      @media (prefers-color-scheme: dark){
        :root{
          --chip-border: rgba(148,163,184,.28);
        }
      }
      .statusbar { display:flex; flex-wrap:wrap; gap:8px; margin:8px 0 6px 0; }
      .chip { display:inline-flex; align-items:center; gap:8px;
              padding:8px 12px; border-radius:999px; font-weight:700; font-size:.98rem;
              border:1px solid var(--chip-border); mix-blend-mode: normal; }
      .chip-red   { background:#fef2f2; color:#991b1b; border-color:#fecaca; }
      .chip-amber { background:#fff7ed; color:#7c2d12; border-color:#fed7aa; }
      .chip-blue  { background:#eef4ff; color:#2541b2; border-color:#c7d2fe; }
      .chip-gray  { background:#f1f5f9; color:#334155; border-color:#cbd5e1; }

      .minirow { display:flex; flex-wrap:wrap; gap:10px; margin:6px 0 2px 0; }
      .minicard { flex:1 1 280px; border:1px solid var(--chip-border); border-radius:12px; padding:12px;
                  background: #ffffff; isolation:isolate; mix-blend-mode: normal; }
      .minicard h4 { margin:0 0 6px 0; font-size:1.02rem; color:#0f172a; }
      .minicard .sub { color:#475569; font-size:.92rem; }

      .pill { display:inline-block; padding:3px 9px; border-radius:999px; font-weight:700; font-size:.92rem; }
      .pill-green { background:#e6ffed; color:#0a7f33; }
      .pill-purple { background:#efe9ff; color:#5b21b6; }
      .pill-amber { background:#fff7ed; color:#7c2d12; }

      @media (max-width: 640px){
        .chip{ padding:7px 10px; font-size:.95rem; }
        .minicard{ padding:11px; }
      }
    </style>
    """, unsafe_allow_html=True)


# =========================================================
# ================== Demo data for announcements ==========
# =========================================================
announcements = [
    {
        "title": "New! Dictionary tab in Vocab Trainer",
        "body":  "Look up quick definitions and example sentences. Open: Vocab Trainer → Dictionary.",
        "tag":   "Update"
    },
    {
        "title": "Calendar & Zoom in My Course → Classroom (no WhatsApp)",
        "body":  "All class reminders and updates are now inside the app. Go to My Course → Classroom to add the calendar and join Zoom.",
        "tag":   "Action"
    },
    {
        "title": "Answers by Email + Weekly Summary",
        "body":  "Your marked answers will arrive in your email. You’ll also get a weekly summary of your submissions.",
        "tag":   "Email"
    },
    {
        "title": "Before You Submit Assignments",
        "body":  "Always make sure your assignment is fully complete before submitting — add your Schreiben, Lesen & Hören work if required, type your final answers in My Course → Classroom (answer box), then submit.",
        "tag":   "Reminder"
    }
]


# =========================================================
# ============== Data loaders & helpers ===================
# =========================================================
@st.cache_data
def load_assignment_scores():
    SHEET_ID = "1BRb8p3Rq0VpFCLSwL4eS9tSgXBo9hSWzfW_J_7W36NQ"
    url = f"https://docs.google.com/spreadsheets/d/{SHEET_ID}/gviz/tq?tqx=out:csv&sheet=Sheet1"
    df = pd.read_csv(url, dtype=str)
    df.columns = df.columns.str.strip().str.lower()
    for col in df.columns:
        df[col] = df[col].astype(str).str.strip()
    return df

@st.cache_data(ttl=43200)
def load_full_vocab_sheet():
    SHEET_ID = "1I1yAnqzSh3DPjwWRh9cdRSfzNSPsi7o4r5Taj9Y36NU"
    csv_url = f"https://docs.google.com/spreadsheets/d/{SHEET_ID}/export?format=csv&gid=0"
    try:
        df = pd.read_csv(csv_url, dtype=str)
    except Exception:
        st.error("Could not load vocab sheet.")
        return pd.DataFrame(columns=["level", "german", "english", "example"])
    df.columns = df.columns.str.strip().str.lower()

    def _match(colnames, *cands):
        s = set(colnames)
        for c in cands:
            if c in s: return c
        for c in colnames:
            if any(c.startswith(x) for x in cands): return c
        return None

    col_level   = _match(df.columns, "level")
    col_german  = _match(df.columns, "german", "de", "word", "wort")
    col_english = _match(df.columns, "english", "en", "meaning", "translation")
    col_example = _match(df.columns, "example", "sentence", "usage")
    if not (col_level and col_german and col_english):
        return pd.DataFrame(columns=["level", "german", "english", "example"])

    rename = {col_level:"level", col_german:"german", col_english:"english"}
    if col_example: rename[col_example] = "example"
    df = df.rename(columns=rename)
    if "example" not in df.columns: df["example"] = ""
    for c in ["level","german","english","example"]:
        df[c] = df[c].astype(str).str.strip()
    df = df[df["level"].notna() & (df["level"] != "")]
    df["level"] = df["level"].str.upper()
    return df[["level","german","english","example"]]

def get_vocab_of_the_day(df: pd.DataFrame, level: str):
    if df is None or df.empty: return None
    if not {"level","german","english","example"}.issubset(df.columns): return None
    lvl = (level or "").upper().strip()
    subset = df[df["level"] == lvl]
    if subset.empty: return None
    idx = date.today().toordinal() % len(subset)
    row = subset.reset_index(drop=True).iloc[idx]
    return {"german": row.get("german",""), "english": row.get("english",""), "example": row.get("example","")}

def parse_contract_end(date_str):
    if not date_str or str(date_str).strip().lower() in ("nan","none",""): return None
    for fmt in ("%Y-%m-%d","%m/%d/%Y","%d.%m.%y","%d/%m/%Y","%d-%m-%Y"):
        try: return datetime.strptime(date_str, fmt)
        except ValueError: continue
    return None


@st.cache_data
def load_reviews():
    SHEET_ID = "137HANmV9jmMWJEdcA1klqGiP8nYihkDugcIbA-2V1Wc"
    url = f"https://docs.google.com/spreadsheets/d/{SHEET_ID}/gviz/tq?tqx=out:csv&sheet=Sheet1"
    df = pd.read_csv(url)
    df.columns = df.columns.str.strip().str.lower()
    return df

def parse_contract_start(date_str: str):
    return parse_contract_end(date_str)

def add_months(dt: datetime, n: int) -> datetime:
    """
    Add n calendar months to dt, clamping the day to the last day of the target month.
    """
    y = dt.year + (dt.month - 1 + n) // 12
    m = (dt.month - 1 + n) % 12 + 1
    last_day = calendar.monthrange(y, m)[1]
    d = min(dt.day, last_day)
    return dt.replace(year=y, month=m, day=d)

def months_between(start_dt: datetime, end_dt: datetime) -> int:
    months = (end_dt.year - start_dt.year) * 12 + (end_dt.month - start_dt.month)
    if end_dt.day < start_dt.day: months -= 1
    return months

# =========================================================
# ===================== Tabs UI ===========================
# =========================================================
def render_dropdown_nav():
    """
    Mobile-friendly dropdown nav with a clear banner that says:
    '🧭 Main Menu — use the selector below to switch sections.'
    Also keeps URL (?tab=...) and st.session_state in sync.
    """
    tabs = [
        "Dashboard",
        "My Course",
        "My Results and Resources",
        "Exams Mode & Custom Chat",
        "Vocab Trainer",
        "Schreiben Trainer",
    ]
    icons = {
        "Dashboard": "🏠",
        "My Course": "📚",
        "My Results and Resources": "📊",
        "Exams Mode & Custom Chat": "🤖",
        "Vocab Trainer": "🗣️",
        "Schreiben Trainer": "✍️",
    }

    # --- Clean, simple banner: always visible, right above the selector ---
    st.markdown(
        """
        <div style="
            padding:12px 14px;
            background:#ecfeff;
            border:1px solid #67e8f9;
            border-radius:12px;
            margin: 4px 0 10px 0;
            display:flex;align-items:center;gap:10px;justify-content:space-between;">
          <div style="font-weight:800;color:#0f172a;font-size:1.05rem;">
            🧭 Main Menu
          </div>
          <div style="color:#0c4a6e;font-size:0.95rem;">
            Use the selector <b>below</b> to switch sections
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # --- Default from URL (?tab=...) or session ---
    default = st.query_params.get(
        "tab",
        [st.session_state.get("main_tab_select", "Dashboard")]
    )[0]
    if default not in tabs:
        default = "Dashboard"

    # --- Selectbox with icons in labels ---
    def _fmt(x: str) -> str:
        return f"{icons.get(x,'•')}  {x}"

    sel = st.selectbox(
        "🧭 Main menu (tap ▾)",
        tabs,
        index=tabs.index(default),
        key="nav_dd",
        format_func=_fmt,
        help="This is the main selector. Tap the arrow ▾ to view all sections.",
    )

    # --- Persist selection to URL + session ---
    if sel != default:
        st.query_params["tab"] = sel
    st.session_state["main_tab_select"] = sel

    # Small “you are here” chip (helps on mobile)
    st.markdown(
        f"""
        <div style="margin-top:6px;">
          <span style="background:#e0f2fe;border:1px solid #7dd3fc;color:#075985;
                       padding:4px 10px;border-radius:999px;font-size:0.92rem;">
            You’re viewing: {icons.get(sel,'•')} <b>{sel}</b>
          </span>
        </div>
        """,
        unsafe_allow_html=True,
    )

    return sel

# usage:
tab = render_dropdown_nav()


# =========================================================
# ===================== Dashboard =========================
# =========================================================
if tab == "Dashboard":
    # ---------- Helpers ----------
    def safe_get(row, key, default=""):
        try: return row.get(key, default)
        except Exception: pass
        try: return getattr(row, key, default)
        except Exception: pass
        try: return row[key]
        except Exception: return default

    # Fallback parsers if globals not present
    def _fallback_parse_date(s):
        fmts = ("%Y-%m-%d", "%m/%d/%Y", "%d.%m.%y", "%d/%m/%Y", "%d-%m-%Y")
        for f in fmts:
            try: return datetime.strptime(str(s).strip(), f)
            except Exception: pass
        return None

    def _fallback_add_months(dt, n):
        y = dt.year + (dt.month - 1 + n) // 12
        m = (dt.month - 1 + n) % 12 + 1
        d = min(dt.day, calendar.monthrange(y, m)[1])
        return dt.replace(year=y, month=m, day=d)

    parse_contract_start_fn = globals().get("parse_contract_start", _fallback_parse_date)
    parse_contract_end_fn   = globals().get("parse_contract_end",   _fallback_parse_date)
    add_months_fn           = globals().get("add_months",           _fallback_add_months)

    # Global styles for chips & mini-cards
    inject_notice_css()

    # ---------- Ensure we have a student row ----------
    load_student_data_fn = globals().get("load_student_data")
    if load_student_data_fn is None:
        def load_student_data_fn():
            return pd.DataFrame(columns=["StudentCode"])

    df_students = load_student_data_fn()
    student_code = (st.session_state.get("student_code", "") or "").strip().lower()

    student_row = {}
    if student_code and not df_students.empty and "StudentCode" in df_students.columns:
        try:
            matches = df_students[df_students["StudentCode"].astype(str).str.lower() == student_code]
            if not matches.empty:
                student_row = matches.iloc[0].to_dict()
        except Exception:
            pass

    if (not student_row) and isinstance(st.session_state.get("student_row"), dict) and st.session_state["student_row"]:
        student_row = st.session_state["student_row"]

    st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)

    if not student_row:
        st.info("🚩 No student selected.")
        st.stop()
        
    st.divider()
    # ---------- 1) Announcements (top) ----------
    render_announcements(announcements)
    st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)

    st.divider()
    # ---------- 3) Motivation mini-cards (streak / vocab / leaderboard) ----------
    _student_code = (st.session_state.get("student_code", "") or "").strip().lower()
    _df_assign = load_assignment_scores()
    _df_assign["date"] = pd.to_datetime(_df_assign["date"], errors="coerce").dt.date
    _mask_student = _df_assign["studentcode"].str.lower().str.strip() == _student_code

    _dates = sorted(_df_assign[_mask_student]["date"].dropna().unique(), reverse=True)
    _streak = 1 if _dates else 0
    for i in range(1, len(_dates)):
        if (_dates[i - 1] - _dates[i]).days == 1:
            _streak += 1
        else:
            break

    _monday = date.today() - timedelta(days=date.today().weekday())
    _weekly_goal = 3
    _submitted_this_week = _df_assign[_mask_student & (_df_assign["date"] >= _monday)].shape[0]
    _goal_left = max(0, _weekly_goal - _submitted_this_week)

    _level = (safe_get(student_row, "Level", "A1") or "A1").upper().strip()
    _vocab_df = load_full_vocab_sheet()
    _vocab_item = get_vocab_of_the_day(_vocab_df, _level)

    _df_assign['level'] = _df_assign['level'].astype(str).str.upper().str.strip()
    _df_assign['score'] = pd.to_numeric(_df_assign['score'], errors='coerce')
    _min_assignments = 3
    _df_level = (
        _df_assign[_df_assign['level'] == _level]
        .groupby(['studentcode', 'name'], as_index=False)
        .agg(total_score=('score', 'sum'), completed=('assignment', 'nunique'))
    )
    _df_level = _df_level[_df_level['completed'] >= _min_assignments]
    _df_level = _df_level.sort_values(['total_score', 'completed'], ascending=[False, False]).reset_index(drop=True)
    _df_level['Rank'] = _df_level.index + 1
    _your_row = _df_level[_df_level['studentcode'].str.lower() == _student_code.lower()]
    _total_students = len(_df_level)

    _streak_line = (
        f"<span class='pill pill-green'>{_streak} day{'s' if _streak != 1 else ''} streak</span>"
        if _streak > 0 else
        "<span class='pill pill-amber'>Start your streak today</span>"
    )
    _goal_line = (
        f"Submitted {_submitted_this_week}/{_weekly_goal} this week"
        + (f" — {_goal_left} to go" if _goal_left else " — goal met 🎉")
    )

    if _vocab_item:
        _vocab_chip = f"<span class='pill pill-purple'>{_vocab_item.get('german','')}</span>"
        _vocab_sub = f"{_vocab_item.get('english','')} · Level {_level}"
    else:
        _vocab_chip = "<span class='pill pill-amber'>No vocab available</span>"
        _vocab_sub = f"Level {_level}"

    if not _your_row.empty:
        _rank = int(_your_row.iloc[0]["Rank"])
        _rank_text = f"Rank #{_rank} of {_total_students}"
        _lead_chip = "<span class='pill pill-purple'>On the board</span>"
    else:
        _rank_text = "Complete 3+ assignments to be ranked"
        _lead_chip = "<span class='pill pill-amber'>Not ranked yet</span>"

    st.markdown(
        f"""
        <div class="minirow">
          <div class="minicard">
            <h4>🏅 Assignment Streak</h4>
            <div>{_streak_line}</div>
            <div class="sub">{_goal_line}</div>
          </div>
          <div class="minicard">
            <h4>🗣️ Vocab of the Day</h4>
            <div>{_vocab_chip}</div>
            <div class="sub">{_vocab_sub}</div>
          </div>
          <div class="minicard">
            <h4>🏆 Leaderboard</h4>
            <div>{_lead_chip}</div>
            <div class="sub">{_rank_text}</div>
          </div>
        </div>
        """,
        unsafe_allow_html=True
    )
    st.divider()
    # ---------- Student header (compact) + details (expander) ----------
    name = safe_get(student_row, "Name")
    level = safe_get(student_row, "Level", "")
    code  = safe_get(student_row, "StudentCode", "")
    try:
        bal_val = float(str(safe_get(student_row, "Balance", 0)).replace(",", "").strip() or 0)
    except Exception:
        bal_val = 0.0

    # Always-visible compact header (one line)
    st.markdown(
        f"<div style='display:flex;flex-wrap:wrap;gap:10px;align-items:center;"
        f"padding:8px 10px;border:1px solid rgba(148,163,184,.35);border-radius:10px;"
        f"background:#ffffff;'>"
        f"<b>👤 {name}</b>"
        f"<span style='background:#eef4ff;color:#2541b2;padding:2px 8px;border-radius:999px;'>Level: {level}</span>"
        f"<span style='background:#f1f5f9;color:#334155;padding:2px 8px;border-radius:999px;'>Code: <code>{code}</code></span>"
        + (f"<span style='background:#fff7ed;color:#7c2d12;padding:2px 8px;border-radius:999px;'>Balance: ₵{bal_val:,.2f}</span>"
           if bal_val > 0 else
           "<span style='background:#ecfdf5;color:#065f46;padding:2px 8px;border-radius:999px;'>Balance: ₵0.00</span>")
        + "</div>",
        unsafe_allow_html=True
    )

    # Full details inside an expander
    with st.expander("👤 Student details", expanded=False):
        info_html = f"""
        <div style='
            background:#f8fbff;
            border:1.6px solid #cfe3ff;
            border-radius:12px;
            padding:12px 14px;
            margin-top:8px;
            box-shadow:0 2px 8px rgba(44,106,221,0.04);
            font-size:1.04em;
            color:#17325e;
            font-family:"Segoe UI","Arial",sans-serif;
            letter-spacing:.01em;'>
            <div style="font-weight:700;font-size:1.12em;margin-bottom:6px;">
                👤 {name}
            </div>
            <div style="font-size:1em; margin-bottom:4px;">
                <b>Level:</b> {safe_get(student_row, 'Level', '')} &nbsp;|&nbsp; 
                <b>Code:</b> <code>{safe_get(student_row, 'StudentCode', '')}</code> &nbsp;|&nbsp;
                <b>Status:</b> {safe_get(student_row, 'Status', '')}
            </div>
            <div style="font-size:1em; margin-bottom:4px;">
                <b>Email:</b> {safe_get(student_row, 'Email', '')} &nbsp;|&nbsp;
                <b>Phone:</b> {safe_get(student_row, 'Phone', '')} &nbsp;|&nbsp;
                <b>Location:</b> {safe_get(student_row, 'Location', '')}
            </div>
            <div style="font-size:1em;">
                <b>Contract:</b> {safe_get(student_row, 'ContractStart', '')} ➔ {safe_get(student_row, 'ContractEnd', '')} &nbsp;|&nbsp;
                <b>Enroll Date:</b> {safe_get(student_row, 'EnrollDate', '')}
            </div>
        </div>
        """
        st.markdown(info_html, unsafe_allow_html=True)

    # ---------- Payments & Renewal (policy-aligned, all inside one expander) ----------
    from datetime import datetime as _dt
    import calendar as _cal

    # Safe money reader (fallback if not provided elsewhere)
    _read_money = globals().get("_read_money")
    if _read_money is None:
        def _read_money(x):
            try:
                s = str(x).replace(",", "").strip()
                return float(s) if s not in ("", "nan", "None") else 0.0
            except Exception:
                return 0.0

    # Fallbacks for date parsing / month add
    def _fallback_parse_date(s):
        for f in ("%Y-%m-%d", "%m/%d/%Y", "%d.%m.%y", "%d/%m/%Y", "%d-%m-%Y"):
            try:
                return _dt.strptime(str(s).strip(), f)
            except Exception:
                pass
        return None

    def _fallback_add_months(dt, n):
        y = dt.year + (dt.month - 1 + n) // 12
        m = (dt.month - 1 + n) % 12 + 1
        d = min(dt.day, _cal.monthrange(y, m)[1])
        return dt.replace(year=y, month=m, day=d)

    # Use app-provided helpers if available, otherwise fallbacks
    _parse_start = (
        globals().get("parse_contract_start_fn")
        or globals().get("parse_contract_start")
        or _fallback_parse_date
    )
    _parse_end = (
        globals().get("parse_contract_end_fn")
        or globals().get("parse_contract_end")
        or _fallback_parse_date
    )
    _add_months = (
        globals().get("add_months_fn")
        or globals().get("add_months")
        or _fallback_add_months
    )

    # Normalize "today" to a date
    _today = _dt.today().date()

    # Contract start -> first payment due (start + 1 month)
    _cs = None
    for _k in ["ContractStart", "StartDate", "ContractBegin", "Start", "Begin"]:
        _s = str(safe_get(student_row, _k, "") or "").strip()
        if _s:
            _cs = _parse_start(_s)
            break
    _first_due_dt = _add_months(_cs, 1) if _cs else None
    _first_due = _first_due_dt.date() if _first_due_dt and hasattr(_first_due_dt, "date") else _first_due_dt

    # Read balance and compute status
    _balance = _read_money(safe_get(student_row, "Balance", 0))

    # Build expander title/body according to policy
    _exp_title = "💳 Payments (info)"
    _severity = "info"
    if _balance > 0 and _first_due:
        if _today > _first_due:
            _days_over = (_today - _first_due).days
            _exp_title = f"💳 Payments • overdue {_days_over}d"
            _severity = "error"
            _msg = (
                f"💸 **Overdue by {_days_over} day{'s' if _days_over != 1 else ''}.** "
                f"Amount due: **₵{_balance:,.2f}**. First due: {_first_due:%d %b %Y}."
            )
        elif _today == _first_due:
            _exp_title = "💳 Payments • due today"
            _severity = "warning"
            _msg = f"⏳ **Payment due today** ({_first_due:%d %b %Y}). Amount due: **₵{_balance:,.2f}**."
        else:
            # Balance positive but still before first due → not expected to pay yet
            _exp_title = "💳 Payments (info)"
            _severity = "info"
            _days_left = (_first_due - _today).days
            _msg = (
                f"No payment expected yet. Your first payment date is **{_first_due:%d %b %Y}** "
                f"(in {_days_left} day{'s' if _days_left != 1 else ''}). Current balance: **₵{_balance:,.2f}**."
            )
    elif _balance > 0 and not _first_due:
        _exp_title = "💳 Payments • schedule unknown"
        _severity = "info"
        _msg = (
            "ℹ️ You have a positive balance, but I couldn’t read your contract start date "
            "to compute the first payment date. Please contact the office."
        )
    else:
        # balance <= 0 → not expected to pay anything now
        _exp_title = "💳 Payments (info)"
        _severity = "info"
        if _first_due:
            _msg = (
                "No outstanding balance. You’re not expected to pay anything now. "
                f"Your first payment date (if applicable) is **{_first_due:%d %b %Y}**."
            )
        else:
            _msg = (
                "No outstanding balance. You’re not expected to pay anything now. "
                "We’ll compute your first payment date after your contract start is on file."
            )

    with st.expander(_exp_title, expanded=False):
        if _severity == "error":
            st.error(_msg)
        elif _severity == "warning":
            st.warning(_msg)
        else:
            st.info(_msg)

        # Always show raw details
        _cs_str = _cs.strftime("%d %b %Y") if _cs else "—"
        _fd_str = _first_due.strftime("%d %b %Y") if _first_due else "—"
        st.markdown(
            f"""
            **Details**
            - Contract start: **{_cs_str}**
            - First payment due (start + 1 month): **{_fd_str}**
            - Current balance: **₵{_balance:,.2f}**
            """
        )

        # ---- Renewal (contract end → extension policy) ----
        EXT_FEE = 1000
        _ce = _parse_end(safe_get(student_row, "ContractEnd", ""))
        _ce_date = _ce.date() if hasattr(_ce, "date") else _ce
        if _ce_date:
            _days_left = (_ce_date - _today).days
            if _days_left < 0:
                st.error(
                    f"⚠️ Your contract ended on **{_ce_date:%d %b %Y}**. "
                    f"If you need more time, extension costs **₵{EXT_FEE:,}/month**."
                )
            elif _days_left <= 14:
                st.warning(
                    f"⏰ Your contract ends in **{_days_left} day{'s' if _days_left != 1 else ''}** "
                    f"(**{_ce_date:%d %b %Y}**). Extension costs **₵{EXT_FEE:,}/month**."
                )
        # If contract end is further out, we stay silent per policy.

    # ---------- Always-visible Contract Alert (cannot be missed) ----------
    from datetime import datetime as _dt

    # Fallback date parser if app helpers aren’t injected
    def _fallback_parse_date(_s):
        for _f in ("%Y-%m-%d", "%m/%d/%Y", "%d.%m.%y", "%d/%m/%Y", "%d-%m-%Y"):
            try:
                return _dt.strptime(str(_s).strip(), _f)
            except Exception:
                pass
        return None

    _parse_end = (
        globals().get("parse_contract_end_fn")
        or globals().get("parse_contract_end")
        or _fallback_parse_date
    )

    _today = _dt.today().date()
    _ce_raw = _parse_end(safe_get(student_row, "ContractEnd", ""))
    _ce_date = _ce_raw.date() if hasattr(_ce_raw, "date") else _ce_raw

    # Mobile-friendly, readable alert styles
    st.markdown("""
    <style>
      .contract-alert { border-radius:12px; padding:12px 14px; margin:8px 0 10px 0; font-weight:600; }
      .ca-warn { background:#fff7ed; color:#7c2d12; border:1px solid #fed7aa; }
      .ca-err  { background:#fef2f2; color:#991b1b; border:1px solid #fecaca; }
      .ca-text { font-size:1rem; line-height:1.55; }
      .ca-cta  { margin-top:6px; font-size:.95rem; }
      @media (max-width:640px){
        .contract-alert{ padding:10px 12px; }
        .ca-text{ font-size:1.02rem; }
      }
    </style>
    """, unsafe_allow_html=True)

    if _ce_date:
        _days_left = (_ce_date - _today).days
        _student_code = str(safe_get(student_row, "StudentCode", "") or "").strip().lower()
        _alert_key = f"hide_contract_alert:{_student_code}:{_ce_date.isoformat()}:{_today.isoformat()}"
        _ext_fee = 1000

        if not st.session_state.get(_alert_key, False):
            if _days_left < 0:
                _msg = (
                    f"⚠️ <b>Your contract ended on {_ce_date:%d %b %Y}.</b> "
                    f"To continue, extension costs <b>₵{_ext_fee:,}/month</b>."
                )
                _cls = "ca-err"
            elif _days_left <= 14:
                _msg = (
                    f"⏰ <b>Your contract ends in {_days_left} day{'s' if _days_left != 1 else ''} "
                    f"({_ce_date:%d %b %Y}).</b> Extension costs <b>₵{_ext_fee:,}/month</b>."
                )
                _cls = "ca-warn"
            else:
                _msg = ""
                _cls = ""

            if _msg:
                st.markdown(
                    f"<div class='contract-alert {_cls}'><div class='ca-text'>{_msg}</div></div>",
                    unsafe_allow_html=True
                )
                # Dismiss for today (so students can acknowledge but can't claim they never saw it)
                if st.button("Got it — hide this notice for today", key=f"btn_contract_alert_{_student_code}"):
                    st.session_state[_alert_key] = True
                    st.rerun()
#


     # ---------- Class schedules ----------
    with st.expander("🗓️ Class Schedule & Upcoming Sessions", expanded=False):
        GROUP_SCHEDULES = {
            "A1 Munich Klasse": {
                "days": ["Monday", "Tuesday", "Wednesday"],
                "time": "6:00pm–7:00pm",
                "start_date": "2025-07-08",
                "end_date": "2025-09-02",
                "doc_url": "https://drive.google.com/file/d/1en_YG8up4C4r36v4r7E714ARcZyvNFD6/view?usp=sharing"
            },
            "A1 Berlin Klasse": {
                "days": ["Thursday", "Friday", "Saturday"],
                "time": "Thu/Fri: 6:00pm–7:00pm, Sat: 8:00am–9:00am",
                "start_date": "2025-06-14",
                "end_date": "2025-08-09",
                "doc_url": "https://drive.google.com/file/d/1foK6MPoT_dc2sCxEhTJbtuK5ZzP-ERzt/view?usp=sharing"
            },
            "A1 Koln Klasse": {
                "days": ["Thursday", "Friday", "Saturday"],
                "time": "Thu/Fri: 6:00pm–7:00pm, Sat: 8:00am–9:00am",
                "start_date": "2025-08-15",
                "end_date": "2025-10-11",
                "doc_url": "https://drive.google.com/file/d/1d1Ord557jGRn5NxYsmCJVmwUn1HtrqI3/view?usp=sharing"
            },
            "A2 Munich Klasse": {
                "days": ["Monday", "Tuesday", "Wednesday"],
                "time": "7:30pm–9:00pm",
                "start_date": "2025-06-24",
                "end_date": "2025-08-26",
                "doc_url": "https://drive.google.com/file/d/1Zr3iN6hkAnuoEBvRELuSDlT7kHY8s2LP/view?usp=sharing"
            },
            "A2 Berlin Klasse": {
                "days": ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"],
                "time": "Mon–Wed: 11:00am–12:00pm, Thu/Fri: 11:00am–12:00pm, Wed: 2:00pm–3:00pm",
                "start_date": "",
                "end_date": "",
                "doc_url": ""
            },
            "A2 Koln Klasse": {
                "days": ["Wednesday", "Thursday", "Friday"],
                "time": "11:00am–12:00pm",
                "start_date": "2025-08-06",
                "end_date": "2025-10-08",
                "doc_url": "https://drive.google.com/file/d/19cptfdlmBDYe9o84b8ZCwujmxuMCKXAD/view?usp=sharing"
            },
            "B1 Munich Klasse": {
                "days": ["Thursday", "Friday"],
                "time": "7:30pm–9:00pm",
                "start_date": "2025-08-07",
                "end_date": "2025-11-07",
                "doc_url": "https://drive.google.com/file/d/1CaLw9RO6H8JOr5HmwWOZA2O7T-bVByi7/view?usp=sharing"
            },
            "B2 Munich Klasse": {
                "days": ["Friday", "Saturday"],
                "time": "Fri: 2pm-3:30pm, Sat: 9:30am-10am",
                "start_date": "2025-08-08",
                "end_date": "2025-10-08",
                "doc_url": "https://drive.google.com/file/d/1gn6vYBbRyHSvKgqvpj5rr8OfUOYRL09W/view?usp=sharing"
            },
        }

        from datetime import datetime as _dt_local, timedelta as _td_local
        class_name = str(safe_get(student_row, "ClassName", "")).strip()
        class_schedule = GROUP_SCHEDULES.get(class_name)
        week_days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]

        if not class_name or not class_schedule:
            st.info("🚩 Your class is not set yet. Please contact your teacher or the office.")
        else:
            days = class_schedule.get("days", [])
            time_str = class_schedule.get("time", "")
            start_dt = class_schedule.get("start_date", "")
            end_dt = class_schedule.get("end_date", "")
            doc_url = class_schedule.get("doc_url", "")

            today = _dt_local.today().date()
            start_date_obj = None
            end_date_obj = None
            try:
                if start_dt:
                    start_date_obj = _dt_local.strptime(start_dt, "%Y-%m-%d").date()
            except Exception:
                start_date_obj = None
            try:
                if end_dt:
                    end_date_obj = _dt_local.strptime(end_dt, "%Y-%m-%d").date()
            except Exception:
                end_date_obj = None

            before_start = bool(start_date_obj and today < start_date_obj)
            after_end = bool(end_date_obj and today > end_date_obj)
            day_indices = [week_days.index(d) for d in days if d in week_days] if isinstance(days, list) else []

            def get_next_sessions(from_date, weekday_indices, limit=3, end_date=None):
                results = []
                if not weekday_indices:
                    return results
                check_date = from_date
                while len(results) < limit:
                    if end_date and check_date > end_date:
                        break
                    if check_date.weekday() in weekday_indices:
                        results.append(check_date)
                    check_date += _td_local(days=1)
                return results

            if before_start and start_date_obj:
                upcoming_sessions = get_next_sessions(start_date_obj, day_indices, limit=3, end_date=end_date_obj)
            elif after_end:
                upcoming_sessions = []
            else:
                upcoming_sessions = get_next_sessions(today, day_indices, limit=3, end_date=end_date_obj)

            if after_end:
                end_str = end_date_obj.strftime('%d %b %Y') if end_date_obj else end_dt
                st.error(f"❌ Your class ({class_name}) ended on {end_str}. Please contact the office for next steps.")
            else:
                if upcoming_sessions:
                    items = []
                    for session_date in upcoming_sessions:
                        weekday_name = week_days[session_date.weekday()]
                        display_date = session_date.strftime("%d %b")
                        items.append(
                            f"<li style='margin-bottom:6px;'><b>{weekday_name}</b> "
                            f"<span style='color:#1976d2;'>{display_date}</span> "
                            f"<span style='color:#333;'>{time_str}</span></li>"
                        )
                    session_items_html = "<ul style='padding-left:16px; margin:9px 0 0 0;'>" + "".join(items) + "</ul>"
                else:
                    session_items_html = "<span style='color:#c62828;'>No upcoming sessions in the visible window.</span>"

                if before_start and start_date_obj:
                    days_until = (start_date_obj - today).days
                    label = f"Starts in {days_until} day{'s' if days_until != 1 else ''} (on {start_date_obj.strftime('%d %b %Y')})"
                    bar_html = f"""
        <div style="margin-top:8px; font-size:0.85em;">
          <div style="margin-bottom:4px;">{label}</div>
          <div style="background:#ddd; border-radius:6px; overflow:hidden; height:12px; width:100%;">
            <div style="width:3%; background:#1976d2; height:100%;"></div>
          </div>
        </div>"""
                elif start_date_obj and end_date_obj:
                    total_days = (end_date_obj - start_date_obj).days + 1
                    elapsed = max(0, (today - start_date_obj).days + 1) if today >= start_date_obj else 0
                    remaining = max(0, (end_date_obj - today).days)
                    percent = int((elapsed / total_days) * 100) if total_days > 0 else 100
                    percent = min(100, max(0, percent))
                    label = f"{remaining} day{'s' if remaining != 1 else ''} remaining in course"
                    bar_html = f"""
        <div style="margin-top:8px; font-size:0.85em;">
          <div style="margin-bottom:4px;">{label}</div>
          <div style="background:#ddd; border-radius:6px; overflow:hidden; height:12px; width:100%;">
            <div style="width:{percent}%; background: linear-gradient(90deg,#1976d2,#4da6ff); height:100%;"></div>
          </div>
          <div style="margin-top:2px; font-size:0.75em;">
            Progress: {percent}% (started {elapsed} of {total_days} days)
          </div>
        </div>"""
                else:
                    bar_html = f"""
        <div style="margin-top:8px; font-size:0.85em;">
          <b>Course period:</b> {start_dt or '[not set]'} to {end_dt or '[not set]'}
        </div>"""

                period_str = f"{start_dt or '[not set]'} to {end_dt or '[not set]'}"
                st.markdown(
                    f"""
        <div style='border:2px solid #17617a; border-radius:14px;
                    padding:13px 11px; margin-bottom:13px;
                    background:#eaf6fb; font-size:1.15em;
                    line-height:1.65; color:#232323;'>
          <b style="font-size:1.09em;">🗓️ Your Next Classes ({class_name}):</b><br>
          {session_items_html}
          {bar_html}
          <div style="font-size:0.98em; margin-top:6px;">
            <b>Course period:</b> {period_str}
          </div>
          {f'<a href="{doc_url}" target="_blank" '
            f'style="font-size:1em;color:#17617a;text-decoration:underline;margin-top:6px;display:inline-block;">📄 View/download full class schedule</a>'
            if doc_url else ''}
        </div>""",
                    unsafe_allow_html=True,
                )

    # ---------- Goethe exam & video ----------
    with st.expander("⏳ Goethe Exam Countdown & Video of the Day", expanded=False):
        from datetime import date
        GOETHE_EXAM_DATES = {
            "A1": (date(2025, 10, 13), 2850, None),
            "A2": (date(2025, 10, 14), 2400, None),
            "B1": (date(2025, 10, 15), 2750, 880),
            "B2": (date(2025, 10, 16), 2500, 840),
            "C1": (date(2025, 10, 17), 2450, 700),
        }
        level = (safe_get(student_row, "Level", "") or "").upper().replace(" ", "")
        exam_info = GOETHE_EXAM_DATES.get(level)

        if exam_info:
            exam_date, fee, module_fee = exam_info
            days_to_exam = (exam_date - date.today()).days
            fee_text = f"**Fee:** ₵{fee:,}"
            if module_fee:
                fee_text += f" &nbsp; | &nbsp; **Per Module:** ₵{module_fee:,}"
            if days_to_exam > 0:
                st.info(
                    f"Your {level} exam is in {days_to_exam} days ({exam_date:%d %b %Y}).  \n"
                    f"{fee_text}  \n"
                    "[Register online here](https://www.goethe.de/ins/gh/en/spr/prf.html)"
                )
            elif days_to_exam == 0:
                st.success("🚀 Exam is today! Good luck!")
            else:
                st.error(
                    f"❌ Your {level} exam was on {exam_date:%d %b %Y}, {abs(days_to_exam)} days ago.  \n"
                    f"{fee_text}"
                )

            playlist_id = (globals().get("YOUTUBE_PLAYLIST_IDS") or {}).get(level)
            fetch_videos = globals().get("fetch_youtube_playlist_videos")
            api_key = globals().get("YOUTUBE_API_KEY")
            if playlist_id and fetch_videos and api_key:
                try:
                    video_list = fetch_videos(playlist_id, api_key)
                except Exception:
                    video_list = []
                if video_list:
                    pick = date.today().toordinal() % len(video_list)
                    video = video_list[pick]
                    st.markdown(f"**🎬 Video of the Day for {level}: {video.get('title','')}**")
                    st.video(video.get('url',''))
                else:
                    st.info("No videos found for your level’s playlist. Check back soon!")
            else:
                st.info("No playlist found for your level yet. Stay tuned!")
        else:
            st.warning("No exam date configured for your level.")
    
    # ---------- Reviews ----------
    with st.expander("🗣️ What Our Students Say", expanded=False):
        import datetime as _pydt
        reviews = load_reviews()
        if reviews.empty:
            st.info("No reviews yet. Be the first to share your experience!")
        else:
            rev_list = reviews.to_dict("records")
            pick = _pydt.date.today().toordinal() % len(rev_list)
            r = rev_list[pick]
            try:
                rating = int(r.get("rating", 5))
            except Exception:
                rating = 5
            rating = max(0, min(5, rating))
            stars = "★" * rating + "☆" * (5 - rating)
            st.markdown(
                f"> {r.get('review_text','')}\n"
                f"> — **{r.get('student_name','')}**  \n"
                f"> {stars}"
            )
   
    st.divider()
  
#



def get_a1_schedule():
    return [
        # DAY 1
        {
            "day": 1,
            "topic": "Lesen & Hören 0.1",
            "chapter": "0.1",
            "goal": "You will learn to introduce yourself, greet others in German, and ask about people's well-being.",
            "instruction": "Watch the video, review grammar, do the workbook, submit assignment.",
            "grammar_topic": "Formal and Informal Greetings",
            "assignment": True,
            "lesen_hören": {
                "video": "https://youtu.be/bK1HEZEaTVM",
                "youtube_link": "https://youtu.be/bK1HEZEaTVM",
                "grammarbook_link": "https://drive.google.com/file/d/1D9Pwg29qZ89xh6caAPBcLJ1K671VUc0_/view?usp=sharing",
                "workbook_link": "https://drive.google.com/file/d/1wjtEyPphP0N7jLbF3AWb5wN_FuJZ5jUQ/view?usp=sharing"
            }
        },
        # DAY 2 – Multi chapter
        {
            "day": 2,
            "topic": "Lesen & Hören 0.2 and 1.1 ",
            "chapter": "0.2_1.1",
            "goal": "Understand the German alphabets, personal pronouns and verb conjugation in German.",
            "instruction": "You are doing Lesen and Hören chapter 0.2 and 1.1. Make sure to follow up attentively.",
            "grammar_topic": "German Alphabets and Personal Pronouns",
            "lesen_hören": [
                {
                    "chapter": "0.2",
                    "video": "https://youtu.be/S7n6TlAQRLQ",
                    "youtube_link": "https://youtu.be/S7n6TlAQRLQ",
                    "grammarbook_link": "https://drive.google.com/file/d/1KtJCF15Ng4cLU88wdUCX5iumOLY7ZA0a/view?usp=sharing",
                    "assignment": True,
                    "workbook_link": "https://drive.google.com/file/d/1R6PqzgsPm9f5iVn7JZXSNVa_NttoPU9Q/view?usp=sharing",
                },
                {
                    "chapter": "1.1",
                    "video": "https://youtu.be/AjsnO1hxDs4",
                    "youtube_link": "https://youtu.be/AjsnO1hxDs4",
                    "grammarbook_link": "https://drive.google.com/file/d/1DKhyi-43HX1TNs8fxA9bgRvhylubilBf/view?usp=sharing",
                    "assignment": True,
                    "workbook_link": "https://drive.google.com/file/d/1A1D1pAssnoncF1JY0v54XT2npPb6mQZv/view?usp=sharing",
                }
            ]
        },
        # DAY 3
        {
            "day": 3,
            "topic": "Schreiben & Sprechen 1.1 and Lesen & Hören 1.2",
            "chapter": "1.1_1.2",
            "goal": "Recap what we have learned so far: be able to introduce yourself in German and know all the pronouns.",
            "instruction": (
                "Begin with the practicals at **Schreiben & Sprechen** (writing & speaking). "
                "Then, move to **Lesen & Hören** (reading & listening). "
                "**Do assignments only at Lesen & Hören.**\n\n"
                "Schreiben & Sprechen activities are for self-practice and have answers provided for self-check. "
                "Main assignment to be marked is under Lesen & Hören below."
            ),
            "grammar_topic": "German Pronouns",
            "schreiben_sprechen": {
                "video": "https://youtu.be/hEe6rs0lkRg",
                "youtube_link": "https://youtu.be/hEe6rs0lkRg",
                "workbook_link": "https://drive.google.com/file/d/1GXWzy3cvbl_goP4-ymFuYDtX4X23D70j/view?usp=sharing",
                "assignment": False,
            },
            "lesen_hören": [
                {
                    "chapter": "1.2",
                    "video": "https://youtu.be/NVCN4fZXEk0",
                    "youtube_link": "https://youtu.be/NVCN4fZXEk0",
                    "grammarbook_link": "https://drive.google.com/file/d/1OUJT9aSU1XABi3cdZlstUvfBIndyEOwb/view?usp=sharing",
                    "workbook_link": "https://drive.google.com/file/d/1Lubevhd7zMlbvPcvHHC1D0GzW7xqa4Mp/view?usp=sharing",
                    "assignment": True
                }
            ]
        },
        # DAY 4
        {
            "day": 4,
            "topic": "Lesen & Hören 2",
            "chapter": "2",
            "goal": "Learn numbers from one to 10 thousand. Also know the difference between city and street",
            "instruction": "Watch the video, study the grammar, complete the workbook, and send your answers.",
            "grammar_topic": "German Numbers",
            "assignment": True,
            "lesen_hören": {
                "video": "https://youtu.be/BzI2n4A8Oak",
                "youtube_link": "https://youtu.be/BzI2n4A8Oak",
                "grammarbook_link": "https://drive.google.com/file/d/1f2CJ492liO8ccudCadxHIISwGJkHP6st/view?usp=sharing",
                "workbook_link": "https://drive.google.com/file/d/1C4VZDUj7VT27Qrn9vS5MNc3QfRqpmDGE/view?usp=sharing",
                "assignment": True
            }
        },
        # DAY 5
        {
            "day": 5,
            "topic": "Schreiben & Sprechen 1.2 (Recap)",
            "chapter": "1.2",
            "goal": "Consolidate your understanding of introductions.",
            "instruction": "Use self-practice workbook and review answers for self-check.",
            "assignment": False,
            "schreiben_sprechen": {
                "video": "",
                "youtube_link": "",
                "workbook_link": "https://drive.google.com/file/d/1ojXvizvJz_qGes7I39pjdhnmlul7xhxB/view?usp=sharing"
            }
        },
        # DAY 6
        {
            "day": 6,
            "topic": "Schreiben & Sprechen 2.3",
            "chapter": "2.3",
            "goal": "Learn about family and expressing your hobby",
            "assignment": False,
            "instruction": "Use self-practice workbook and review answers for self-check.",
            "schreiben_sprechen": {
                "video": "https://youtu.be/JrYSpnZN6P0",
                "youtube_link": "https://youtu.be/JrYSpnZN6P0",
                "workbook_link": "https://drive.google.com/file/d/1xellIzaxzoBTFOUdaCEHu_OiiuEnFeWT/view?usp=sharing"
            }
        },
        # DAY 7
        {
            "day": 7,
            "topic": "Lesen & Hören 3",
            "chapter": "3",
            "goal": "Know how to ask for a price and also the use of mogen and gern to express your hobby",
            "instruction": "Watch the video, study the grammar, complete the workbook, and send your answers.Do schreiben and sprechen 2.3 before this chapter for better understanding",
            "grammar_topic": "Fragen nach dem Preis; gern/lieber/mögen (Talking about price and preferences)",
            "assignment": True,
            "lesen_hören": {
                "video": "https://youtu.be/dGIj1GbK4sI",
                "youtube_link": "https://youtu.be/dGIj1GbK4sI",
                "grammarbook_link": "https://drive.google.com/file/d/1sCE5y8FVctySejSVNm9lrTG3slIucxqY/view?usp=sharing",
                "workbook_link": "https://drive.google.com/file/d/1lL4yrZLMtKLnNuVTC2Sg_ayfkUZfIuak/view?usp=sharing"
            }
        },
        # DAY 8
        {
            "day": 8,
            "topic": "Lesen & Hören 4",
            "chapter": "4",
            "goal": "Learn about schon mal and noch nie, irregular verbs and all the personal pronouns",
            "instruction": "Watch the video, study the grammar, complete the workbook, and send your answers.",
            "grammar_topic": "schon mal, noch nie; irregular verbs; personal pronouns",
            "assignment": True,
            "lesen_hören": {
                "video": "https://youtu.be/JfTc1G9mubs",
                "youtube_link": "https://youtu.be/JfTc1G9mubs",
                "grammarbook_link": "https://drive.google.com/file/d/1obsYT3dP3qT-i06SjXmqRzCT2pNoJJZp/view?usp=sharing",
                "workbook_link": "https://drive.google.com/file/d/1woXksV9sTZ_8huXa8yf6QUQ8aUXPxVug/view?usp=sharing"
            }
        },
        # DAY 9
        {
            "day": 9,
            "topic": "Lesen & Hören 5",
            "chapter": "5",
            "goal": "Learn about the German articles and cases",
            "instruction": "Watch the video, study the grammar, complete the workbook, and send your answers.",
            "grammar_topic": "Nominative & Akkusative, Definite & Indefinite Articles",
            "assignment": True,
            "lesen_hören": {
                "video": "https://youtu.be/Yi5ZA-XD-GY?si=nCX_pceEYgAL-FU0",
                "youtube_link": "https://youtu.be/Yi5ZA-XD-GY?si=nCX_pceEYgAL-FU0",
                "grammarbook_link": "https://drive.google.com/file/d/17y5fGW8nAbfeVgolV7tEW4BLiLXZDoO6/view?usp=sharing",
                "workbook_link": "https://drive.google.com/file/d/1zjAqvQqNb7iKknuhJ79bUclimEaTg-mt/view?usp=sharing"
            }
        },
        # DAY 10
        {
            "day": 10,
            "topic": "Lesen & Hören 6 and Schreiben & Sprechen 2.4",
            "chapter": "6_2.4",
            "goal": "Understand Possessive Determiners and its usage in connection with nouns",
            "instruction": "The assignment is the lesen and horen chapter 6 but you must also go through schreiben and sprechnen 2.4 for full understanding",         
            "lesen_hören": {
                "video": "https://youtu.be/SXwDqcwrR3k",
                "youtube_link": "https://youtu.be/SXwDqcwrR3k",
                "grammarbook_link": "https://drive.google.com/file/d/1Fy4bKhaHHb4ahS2xIumrLtuqdQ0YAFB4/view?usp=sharing",
                "assignment": True,
                "workbook_link": "https://drive.google.com/file/d/1Da1iw54oAqoaY-UIw6oyIn8tsDmIi1YR/view?usp=sharing"
            },
            "schreiben_sprechen": {
                "video": "https://youtu.be/lw9SsojpKf8",
                "youtube_link": "https://youtu.be/lw9SsojpKf8",
                "workbook_link": "https://drive.google.com/file/d/1GbIc44ToWh2upnHv6eX3ZjFrvnf4fcEM/view?usp=sharing",
                "assignment": False,
            }
        },
        # DAY 11
        {
            "day": 11,
            "topic": "Lesen & Hören 7",
            "chapter": "7",
            "goal": "Understand the 12 hour clock system",
            "instruction": "Watch the video, study the grammar, complete the workbook, and send your answers.",
            "assignment": True,
            "lesen_hören": {
                "video": "https://youtu.be/uyvXoCoqjiE",
                "youtube_link": "https://youtu.be/uyvXoCoqjiE",
                "grammarbook_link": "https://drive.google.com/file/d/1pSaloRhfh8eTKK_r9mzwp6xkbfdkCVox/view?usp=sharing",
                "workbook_link": "https://drive.google.com/file/d/1QyDdRae_1qv_umRb15dCJZTPdXi7zPWd/view?usp=sharing"
            }
        },
        # DAY 12
        {
            "day": 12,
            "topic": "Lesen & Hören 8",
            "chapter": "8",
            "goal": "Understand the 24 hour clock and date system in German",
            "instruction": "Watch the video, study the grammar, complete the workbook, and send your answers.",
            "assignment": True,
            "lesen_hören": {
                "video": "https://youtu.be/hLpPFOthVkU",
                "youtube_link": "https://youtu.be/hLpPFOthVkU",
                "grammarbook_link": "https://drive.google.com/file/d/1fW2ChjnDKW_5SEr65ZgE1ylJy1To46_p/view?usp=sharing",
                "workbook_link": "https://drive.google.com/file/d/1onzokN8kQualNO6MSsPndFXiRwsnsVM9/view?usp=sharing"
            }
        },
        # DAY 13
        {
            "day": 13,
            "topic": "Schreiben & Sprechen 3.5",
            "chapter": "3.5",
            "goal": "Recap from the lesen and horen. Understand numbers, time, asking of price and how to formulate statements in German",
            "instruction": "Use the statement rule to talk about your weekly routine using the activities listed. Share with your tutor when done",
            "schreiben_sprechen": {
                "video": "https://youtu.be/PwDLGmfBUDw",
                "youtube_link": "https://youtu.be/PwDLGmfBUDw",
                "assignment": False,
                "workbook_link": "https://drive.google.com/file/d/12oFKrKrHBwSpSnzxLX_e-cjPSiYtCFVs/view?usp=sharing"
            }
        },
        # DAY 14
        {
            "day": 14,
            "topic": "Schreiben & Sprechen 3.6",
            "chapter": "3.6",
            "goal": "Understand how to use modal verbs with main verbs and separable verbs",
            "assignment": False,
            "instruction": "This is a practical exercise. All the answers are included in the document except for the last paragraph. You can send a screenshot of that to your tutor",
            "grammar_topic": "Modal Verbs",
            "schreiben_sprechen": {
                "video": "https://youtu.be/XwFPjLjvDog",
                "youtube_link": "https://youtu.be/XwFPjLjvDog",
                "workbook_link": "https://drive.google.com/file/d/1wnZehLNfkjgKMFw1V3BX8V399rZg6XLv/view?usp=sharing"
            }
        },
        # DAY 15
        {
            "day": 15,
            "topic": "Schreiben & Sprechen 4.7",
            "chapter": "4.7",
            "assignment": False,
            "goal": "Understand imperative statements and learn how to use them in your Sprechen exams, especially in Teil 3.",
            "instruction": "After completing this chapter, go to the Falowen Exam Chat Mode, select A1 Teil 3, and start practicing",
            "grammar_topic": "Imperative",
            "schreiben_sprechen": {
                "video": "https://youtu.be/IVtUc9T3o0Y",
                "youtube_link": "https://youtu.be/IVtUc9T3o0Y",
                "workbook_link": "https://drive.google.com/file/d/1953B01hB9Ex7LXXU0qIaGU8xgCDjpSm4/view?usp=sharing"
            }
        },
        # DAY 16
        {
            "day": 16,
            "topic": "Lesen & Hören 9 and 10",
            "chapter": "9_10",
            "goal": "Understand how to negate statements using nicht,kein and nein",
            "instruction": "This chapter has two assignments. Do the assignments for chapter 9 and after chapter 10. Chapter 10 has no grammar",
            "grammar_topic": "Negation",
            "lesen_hören": [
                {
                    "chapter": "9",
                    "video": "https://youtu.be/MrB3BPtQN6A",
                    "youtube_link": "https://youtu.be/MrB3BPtQN6A",
                    "assignment": True,
                    "grammarbook_link": "https://drive.google.com/file/d/1g-qLEH1ZDnFZCT83TW-MPLxNt2nO7UAv/view?usp=sharing",
                    "workbook_link": "https://drive.google.com/file/d/1hKtQdXg5y3yJyFBQsCMr7fZ11cYbuG7D/view?usp=sharing"
                },
                {
                    "chapter": "10",
                    "video": "",
                    "youtube_link": "",
                    "grammarbook_link": "",
                    "assignment": True,
                    "workbook_link": "https://drive.google.com/file/d/1rJXshXQSS5Or4ipv1VmUMsoB0V1Vx4VK/view?usp=sharing"
                }
            ]
        },
        # DAY 17
        {
            "day": 17,
            "topic": "Lesen & Hören 11",
            "chapter": "11",
            "goal": "Understand instructions and request in German using the Imperative rule",
            "grammar_topic": "Direction",
            "instruction": "",
            "lesen_hören": {
                "video": "https://youtu.be/k2ZC3rXPe1k",
                "youtube_link": "https://youtu.be/k2ZC3rXPe1k",
                "assignment": True,
                "grammarbook_link": "https://drive.google.com/file/d/1lMzZrM4aAItO8bBmehODvT6gG7dz8I9s/view?usp=sharing",
                "workbook_link": "https://drive.google.com/file/d/17FNSfHBxyga9sKxzicT_qkP7PA4vB5-A/view?usp=sharing"
            }
        },
        # DAY 18
        {
            "day": 18,
            "topic": "Lesen & Hören 12.1 and 12.2",
            "chapter": "12.1_12.2",
            "goal": "Learn about German professions and how to use two-way prepositions",
            "instruction": "Do assignments for 12.1 and 12.2 and use the schreiben and sprechen below for practicals for full understanding",
            "grammar_topic": "Two Case Preposition",
            "lesen_hören": [
                {
                    "chapter": "12.1",
                    "video": "https://youtu.be/-vTEvx9a8Ts",
                    "youtube_link": "https://youtu.be/-vTEvx9a8Ts",
                    "assignment": True,
                    "grammarbook_link": "https://drive.google.com/file/d/1wdWYVxBhu4QtRoETDpDww-LjjzsGDYva/view?usp=sharing",
                    "workbook_link": "https://drive.google.com/file/d/1A0NkFl1AG68jHeqSytI3ygJ0k7H74AEX/view?usp=sharing"
                },
                {
                    "chapter": "12.2",
                    "video": "",
                    "youtube_link": "",
                    "assignment": True,
                    "grammarbook_link": "",
                    "workbook_link": "https://drive.google.com/file/d/1xojH7Tgb5LeJj3nzNSATUVppWnJgJLEF/view?usp=sharing"
                }
            ],
            "schreiben_sprechen": {
                "video": "https://youtu.be/xVyYo7upDGo",
                "youtube_link": "https://youtu.be/xVyYo7upDGo",
                "assignment": False,
                "workbook_link": "https://drive.google.com/file/d/1iyYBuxu3bBEovxz0j9QeSu_1URX92fvN/view?usp=sharing"
            }
        },
        # DAY 19
        {
            "day": 19,
            "topic": "Schreiben & Sprechen 5.9",
            "chapter": "5.9",
            "goal": "Understand the difference between Erlaubt and Verboten and how to use it in the exams hall",
            "instruction": "Review the workbook and do the practicals in it. Answers are attached",
            "grammar_topic": "Erlaubt and Verboten",
            "schreiben_sprechen": {
                "video": "https://youtu.be/MqAp84GthAo",
                "youtube_link": "https://youtu.be/MqAp84GthAo",
                "assignment": False,
                "workbook_link": "https://drive.google.com/file/d/1CkoYa_qeqsGju0kTS6ElurCAlEW6pVFL/view?usp=sharing"
            }
        },
        # DAY 20
        {
            "day": 20,
            "topic": "Introduction to Letter Writing 12.3 ",
            "chapter": "12.3",
            "goal": "Practice how to write both formal and informal letters",
            "assignment": True,
            "instruction": "Write all the two letters in this document and send to your tutor for corrections",
            "grammar_topic": "Formal and Informal Letter",
            "schreiben_sprechen": {
                "video": "https://youtu.be/sHRHE1soH6I",
                "youtube_link": "https://youtu.be/sHRHE1soH6I",
                "workbook_link": "https://drive.google.com/file/d/1SjaDH1bYR7O-BnIbM2N82XOEjeLCfPFb/view?usp=sharing"
            }
        },
        # DAY 21
        {
            "day": 21,
            "topic": "Lesen & Hören 13",
            "chapter": "13",
            "assignment": True,
            "goal": "",
            "instruction": "Watch the video, study the grammar, complete the workbook, and send your answers.",
            "grammar_topic": "Weather and Past Tense. How to form Perfekt statement in German",
            "lesen_hören": {
                "video": "https://youtu.be/6cBs3Qfvdk4",
                "youtube_link": "https://youtu.be/6cBs3Qfvdk4",
                "assignment": True,
                "grammarbook_link": "https://drive.google.com/file/d/1PCXsTIg9iNlaAUkwH8BYekw_3v1HJjGq/view?usp=sharing",
                "workbook_link": "https://drive.google.com/file/d/1GZeUi5p6ayDGnPcebFVFfaNavmoWyoVM/view?usp=sharing"
            }
        },
        # DAY 22
        {
            "day": 22,
            "topic": "Lesen & Hören 14.1",
            "chapter": "14.1",
            "goal": "Understand health and talking about body parts in German",
            "instruction": "Watch the video, study the grammar, complete the workbook, and send your answers.",
            "grammar_topic": "Health and Body Parts",
            "lesen_hören": {
                "video": "https://youtu.be/Zx_TFF9FNGo",
                "youtube_link": "https://youtu.be/Zx_TFF9FNGo",
                "assignment": True,
                "grammarbook_link": "https://drive.google.com/file/d/1QoG4mNxA1w8AeTMPfLtMQ_rAHrmC1DdO/view?usp=sharing",
                "workbook_link": "https://drive.google.com/file/d/1LkDUU7r78E_pzeFnHKw9vfD9QgUAAacu/view?usp=sharing"
            }
        },
        # DAY 23
        {
            "day": 23,
            "topic": "Lesen & Hören 14.2",
            "chapter": "14.2",
            "goal": "Understand adjective declension and dative verbs",
            "instruction": " This chapter has no assignment. Only grammar",
            "grammar_topic": "Adjective Declension and Dative Verbs",
            "lesen_hören": {
                "video": "",
                "youtube_link": "",
                "assignment": False,
                "grammarbook_link": "https://drive.google.com/file/d/16h-yS0gkB2_FL1zxCC4MaqRBbKne7GI1/view?usp=sharing",
                "workbook_link": ""
            }
        },
        # DAY 24
        {
            "day": 24,
            "topic": "Schreiben & Sprechen 5.10",
            "chapter": "5.10",
            "goal": "Learn about conjunctions and how to apply them in your exams",
            "instruction": "This chapter has no assignments. It gives you ideas to progress for A2 and how to use conjunctions",
            "grammar_topic": "German Conjunctions",
            "assignment": False,
            "schreiben_sprechen": {
                "video": "https://youtu.be/WVq9x69dCeE",
                "youtube_link": "https://youtu.be/WVq9x69dCeE",
                "workbook_link": "https://drive.google.com/file/d/1LE1b9ilkLLobE5Uw0TVLG0RIVpLK5k1t/view?usp=sharing"
            }
        },
        # DAY 25
        {
            "day": 25,
            "topic": "Goethe Mock Test 15",
            "chapter": "15",
            "assignment": True,
            "goal": "This test should help the student have an idea about how the lesen and horen will look like",
            "instruction": "Open the link and answer the questions using the link. After submit and alert your tutor.",
            "schreiben_sprechen": {
                "video": "",
                "youtube_link": "",
                "workbook_link": "https://forms.gle/FP8ZPNhwxcAZsTfY6"
            }
        }
    ]


def get_a2_schedule():
    return [
        # DAY 1
        {
            "day": 1,
            "topic": "Small Talk 1.1 (Exercise)",
            "chapter": "1.1",
            "goal": "Practice basic greetings and small talk.",
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "assignment": True,
            "video": "https://youtu.be/siF0jWZdIwk",
            "youtube_link": "https://youtu.be/siF0jWZdIwk",
            "grammarbook_link": "https://drive.google.com/file/d/1NsCKO4K7MWI-queLWCeBuclmaqPN04YQ/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1LXDI1yyJ4aT4LhX5eGDbKnkCkJZ2EE2T/view?usp=sharing"
        },
        # DAY 2
        {
            "day": 2,
            "topic": "Personen Beschreiben 1.2 (Exercise)",
            "chapter": "1.2",
            "goal": "Describe people and their appearance.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "grammar_topic": "Subordinate Clauses (Nebensätze) with dass and weil",
            "video": "https://youtu.be/FYaXSvZsEDM?si=0e_sHxslHQL7FGDk",
            "youtube_link": "https://youtu.be/FYaXSvZsEDM?si=0e_sHxslHQL7FGDk",
            "grammarbook_link": "https://drive.google.com/file/d/1xMpEAPD8C0HtIFsmgqYO-wZaKDrQtiYp/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/128lWaKgCZ2V-3tActM-dwNy6igLLlzH3/view?usp=sharing"
        },
        # DAY 3
        {
            "day": 3,
            "topic": "Dinge und Personen vergleichen 1.3",
            "chapter": "1.3",
            "goal": "Learn to compare things and people.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "grammar_topic": "Positive, Comparative, and Superlative in German",
            "video": "https://youtu.be/oo3pUo5OSDE",
            "youtube_link": "https://youtu.be/oo3pUo5OSDE",
            "grammarbook_link": "https://drive.google.com/file/d/1Z3sSDCxPQz27TDSpN9r8lQUpHhBVfhYZ/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/18YXe9mxyyKTars1gL5cgFsXrbM25kiN8/view?usp=sharing"
        },
        # DAY 4
        {
            "day": 4,
            "topic": "Wo möchten wir uns treffen? 2.4",
            "chapter": "2.4",
            "goal": "Arrange and discuss meeting places.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "grammar_topic": "Nominalization of Verbs",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "https://drive.google.com/file/d/14qE_XJr3mTNr6PF5aa0aCqauh9ngYTJ8/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1RaXTZQ9jHaJYwKrP728zevDSQHFKeR0E/view?usp=sharing"
        },
        # DAY 5
        {
            "day": 5,
            "topic": "Was machst du in deiner Freizeit? 2.5 ",
            "chapter": "2.5",
            "goal": "Talk about free time activities.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "grammar_topic": "Dative Preposition",
            "video": "https://youtu.be/8dX40NXG_gI",
            "youtube_link": "https://youtu.be/8dX40NXG_gI",
            "grammarbook_link": "https://drive.google.com/file/d/11yEcMioSB9x1ZD-x5_67ApFzP53iau-N/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1dIsFg7wNaqyyOHm95h7xv4Ssll5Fm0V1/view?usp=sharing"
        },
        # DAY 6
        {
            "day": 6,
            "topic": "Möbel und Räume kennenlernen 3.6",
            "chapter": "3.6",
            "goal": "Identify furniture and rooms.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "grammar_topic": "Two Case Preposition",
            "video": "https://youtu.be/am3WqQaCibE",
            "youtube_link": "https://youtu.be/am3WqQaCibE",
            "grammarbook_link": "https://drive.google.com/file/d/1MSahBEyElIiLnitWoJb5xkvRlB21yo0y/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/16UfBIrL0jxCqWtqqZaLhKWflosNQkwF4/view?usp=sharing"
        },
        # DAY 7
        {
            "day": 7,
            "topic": "Eine Wohnung suchen (Übung) 3.7",
            "chapter": "3.7",
            "goal": "Practice searching for an apartment.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "grammar_topic": "Identifying German Nouns and their Gender",
            "video": "https://youtu.be/ScU6w8VQgNg", 
            "youtube_link": "https://youtu.be/ScU6w8VQgNg",
            "grammarbook_link": "https://drive.google.com/file/d/1clWbDAvLlXpgWx7pKc71Oq3H2p0_GZnV/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1EF87TdHa6Y-qgLFUx8S6GAom9g5EBQNP/view?usp=sharing"
        },
        # DAY 8
        {
            "day": 8,
            "topic": "Rezepte und Essen (Exercise) 3.8",
            "chapter": "3.8",
            "assignment": True,
            "goal": "Learn about recipes and food. Practice using sequence words like zuerst', 'nachdem', and 'außerdem' to organize your letter.",
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "grammar_topic": "Zuerst, Nachdem, and Talking About Sequence in German",
            "video": "https://youtu.be/_xQMNp3qcDQ",
            "youtube_link": "https://youtu.be/_xQMNp3qcDQ",
            "grammarbook_link": "https://drive.google.com/file/d/16lh8sPl_IDZ3dLwYNvL73PqOFCixidrI/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1c8JJyVlKYI2mz6xLZZ6RkRHLnH3Dtv0c/view?usp=sharing"
        },
        # DAY 9
        {
            "day": 9,
            "topic": "Urlaub 4.9",
            "chapter": "4.9",
            "goal": "Discuss vacation plans.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "grammar_topic": "Understanding Präteritum and Perfekt",
            "video": "https://youtu.be/NxoQH-BY9Js",
            "youtube_link": "https://youtu.be/NxoQH-BY9Js",
            "grammarbook_link": "https://drive.google.com/file/d/1kOb7c08Pkxf21OQE_xIGEaif7Xq7k-ty/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1NzRxbGUe306Vq0mq9kKsc3y3HYqkMhuA/view?usp=sharing"
        },
        # DAY 10
        {
            "day": 10,
            "topic": "Tourismus und Traditionelle Feste 4.10",
            "chapter": "4.10",
            "assignment": True,
            "goal": "Learn about tourism and festivals.",
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "grammar_topic": "Präteritum",
            "video": "https://youtu.be/XFxV3GSSm8E",
            "youtube_link": "https://youtu.be/XFxV3GSSm8E",
            "grammarbook_link": "https://drive.google.com/file/d/1snFsDYBK8RrPRq2n3PtWvcIctSph-zvN/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1vijZn-ryhT46cTzGmetuF0c4zys0yGlB/view?usp=sharing"
        },
        # DAY 11
        {
            "day": 11,
            "topic": "Unterwegs: Verkehrsmittel vergleichen 4.11",
            "chapter": "4.11",
            "assignment": True,
            "goal": "Compare means of transportation.",
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "grammar_topic": "Prepositions in and naxh",
            "video": "https://youtu.be/RkvfRiPCZI4",
            "youtube_link": "https://youtu.be/RkvfRiPCZI4",
            "grammarbook_link": "https://drive.google.com/file/d/19I7oOHX8r4daxXmx38mNMaZO10AXHEFu/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1c7ITea0iVbCaPO0piark9RnqJgZS-DOi/view?usp=sharing"
        },
        # DAY 12
        {
            "day": 12,
            "topic": "Mein Traumberuf (Übung) 5.12",
            "chapter": "5.12",
            "assignment": True,
            "goal": "Learn how to talk about a dream job and future goals.",
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "grammar_topic": "Konjunktiv II",
            "video": "https://youtu.be/w81bsmssGXQ",
            "youtube_link": "https://youtu.be/w81bsmssGXQ",
            "grammarbook_link": "https://drive.google.com/file/d/1dyGB5q92EePy8q60eWWYA91LXnsWQFb1/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/18u6FnHpd2nAh1Ev_2mVk5aV3GdVC6Add/view?usp=sharing"
        },
        # DAY 13
        {
            "day": 13,
            "topic": "Ein Vorstellungsgespräch (Exercise) 5.13",
            "chapter": "5.13",
            "assignment": True,
            "goal": "Prepare for a job interview.",
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "grammar_topic": "Konjunktive II with modal verbs",
            "video": "https://youtu.be/urKBrX5VAYU",
            "youtube_link": "https://youtu.be/urKBrX5VAYU",
            "grammarbook_link": "https://drive.google.com/file/d/1tv2tYzn9mIG57hwWr_ilxV1My7kt-RKQ/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1sW2yKZptnYWPhS7ciYdi0hN5HV-ycsF0/view?usp=sharing"
        },
        # DAY 14
        {
            "day": 14,
            "topic": "Beruf und Karriere (Exercise) 5.14",
            "chapter": "5.14",
            "assignment": True,
            "goal": "Discuss jobs and careers.",
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "grammar_topic": "Modal Verbs",
            "video": "https://youtu.be/IyBvx-yVT-0",
            "youtube_link": "https://youtu.be/IyBvx-yVT-0",
            "grammarbook_link": "https://drive.google.com/file/d/13mVpVGfhY1NQn-BEb7xYUivnaZbhXJsK/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1rlZoo49bYBRjt7mu3Ydktzgfdq4IyK2q/view?usp=sharing"
        },
        # DAY 15
        {
            "day": 15,
            "topic": "Mein Lieblingssport 6.15",
            "chapter": "6.15",
            "assignment": True,
            "goal": "Talk about your favorite sport.",
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "grammar_topic": "Reflexive Pronouns",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "https://drive.google.com/file/d/1dGZjcHhdN1xAdK2APL54RykGH7_msUyr/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1iiExhUj66r5p0SJZfV7PsmCWOyaF360s/view?usp=sharing"
        },
        # DAY 16
        {
            "day": 16,
            "topic": "Wohlbefinden und Entspannung 6.16",
            "chapter": "6.16",
            "goal": "Express well-being and relaxation.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "grammar_topic": "Verbs and Adjectives with Prepositions",
            "video": "https://youtu.be/r4se8KuS8cA",
            "youtube_link": "https://youtu.be/r4se8KuS8cA",
            "grammarbook_link": "https://drive.google.com/file/d/1BiAyDazBR3lTplP7D2yjaYmEm2btUT1D/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1G_sRFKG9Qt5nc0Zyfnax-0WXSMmbWB70/view?usp=sharing"
        },
        # DAY 17
        {
            "day": 17,
            "topic": "In die Apotheke gehen 6.17",
            "chapter": "6.17",
            "goal": "Learn phrases for the pharmacy.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "grammar_topic": "Notes on German Indefinite Pronouns",
            "video": "https://youtu.be/Xjp2A1hU1ag",
            "youtube_link": "https://youtu.be/Xjp2A1hU1ag",
            "grammarbook_link": "https://drive.google.com/file/d/1O040UoSuBdy4llTK7MbGIsib63uNNcrV/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1vsdVR_ubbu5gbXnm70vZS5xGFivjBYoA/view?usp=sharing"
        },
        # DAY 18
        {
            "day": 18,
            "topic": "Die Bank anrufen 7.18",
            "chapter": "7.18",
            "goal": "Practice calling the bank.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "grammar_topic": "Notes on Opening a Bank Account in Germany",
            "video": "https://youtu.be/ahIUVAbsuxU",
            "youtube_link": "https://youtu.be/ahIUVAbsuxU",
            "grammarbook_link": "https://drive.google.com/file/d/1qNHtY8MYOXjtBxf6wHi6T_P_X1DGFtPm/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1GD7cCPU8ZFykcwsFQZuQMi2fiNrvrCPg/view?usp=sharing"
        },
        # DAY 19
        {
            "day": 19,
            "topic": "Einkaufen? Wo und wie? (Exercise) 7.19",
            "chapter": "7.19",
            "goal": "Shop and ask about locations.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "video": "https://youtu.be/TOTK1yohCTg",
            "youtube_link": "https://youtu.be/TOTK1yohCTg",
            "grammarbook_link": "https://drive.google.com/file/d/1Qt9oxn-74t8dFdsk-NjSc0G5OT7MQ-qq/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1CEFn14eYeomtf6CpZJhyW00CA2f_6VRc/view?usp=sharing"
        },
        # DAY 20
        {
            "day": 20,
            "topic": "Typische Reklamationssituationen üben 7.20",
            "chapter": "7.20",
            "goal": "Handle typical complaints.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "video": "https://youtu.be/utAO9hvGF18",
            "youtube_link": "https://youtu.be/utAO9hvGF18",
            "grammarbook_link": "https://drive.google.com/file/d/1-72wZuNJE4Y92Luy0h5ygWooDnBd9PQW/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1_GTumT1II0E1PRoh6hMDwWsTPEInGeed/view?usp=sharing"
        },
        # DAY 21
        {
            "day": 21,
            "topic": "Ein Wochenende planen 8.21",
            "chapter": "8.21",
            "goal": "Plan a weekend.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "https://drive.google.com/file/d/1FcCg7orEizna4rAkX3_FCyd3lh_Bb3IT/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1mMtZza34QoJO_lfUiEX3kwTa-vsTN_RK/view?usp=sharing"
        },
        # DAY 22
        {
            "day": 22,
            "topic": "Die Woche Planung 8.22",
            "chapter": "8.22",
            "goal": "Make a weekly plan.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "video": "https://youtu.be/rBuEEFfee1c?si=YJpKuM0St2gWN67H",
            "youtube_link": "https://youtu.be/rBuEEFfee1c?si=YJpKuM0St2gWN67H",
            "grammarbook_link": "https://drive.google.com/file/d/1AvLYxZKq1Ae6_4ACJ20il1LqCOv2jQbb/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1mg_2ytNAYF00_j-TFQelajAxgQpmgrhW/view?usp=sharing"
        },
        # DAY 23
        {
            "day": 23,
            "topic": "Wie kommst du zur Schule / zur Arbeit? 9.23",
            "chapter": "9.23",
            "goal": "Talk about your route to school or work.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "video": "https://youtu.be/c4TpUe3teBE",
            "youtube_link": "https://youtu.be/c4TpUe3teBE",
            "grammarbook_link": "https://drive.google.com/file/d/1XbWKmc5P7ZAR-OqFce744xqCe7PQguXo/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1Ialg19GIE_KKHiLBDMm1aHbrzfNdb7L_/view?usp=sharing"
        },
        # DAY 24
        {
            "day": 24,
            "topic": "Einen Urlaub planen 9.24",
            "chapter": "9.24",
            "goal": "Plan a vacation.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "https://drive.google.com/file/d/1tFXs-DNKvt97Q4dsyXsYvKVQvT5Qqt0y/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1t3xqddDJp3-1XeJ6SesnsYsTO5xSm9vG/view?usp=sharing"
        },
        # DAY 25
        {
            "day": 25,
            "topic": "Tagesablauf (Exercise) 9.25",
            "chapter": "9.25",
            "goal": "Describe a daily routine.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "video": "",
            "youtube_link": "",
            "workbook_link": "https://drive.google.com/file/d/1jfWDzGfXrzhfGZ1bQe1u5MXVQkR5Et43/view?usp=sharing"
        },
        # DAY 26
        {
            "day": 26,
            "topic": "Gefühle in verschiedenen Situationen beschreiben 10.26",
            "chapter": "10.26",
            "goal": "Express feelings in various situations.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "video": "",
            "youtube_link": "",
            "workbook_link": "https://drive.google.com/file/d/126MQiti-lpcovP1TdyUKQAK6KjqBaoTx/view?usp=sharing"
        },
        # DAY 27
        {
            "day": 27,
            "topic": "Digitale Kommunikation 10.27",
            "chapter": "10.27",
            "goal": "Talk about digital communication.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "video": "",
            "youtube_link": "",
            "workbook_link": "https://drive.google.com/file/d/1UdBu6O2AMQ2g6Ot_abTsFwLvT87LHHwY/view?usp=sharing"
        },
        # DAY 28
        {
            "day": 28,
            "topic": "Über die Zukunft sprechen 10.28",
            "chapter": "10.28",
            "goal": "Discuss the future.",
            "assignment": True,
            "instruction": "Watch the video, review grammar, and complete your workbook.",
            "video": "",
            "youtube_link": "",
            "workbook_link": "https://drive.google.com/file/d/1164aJFtkZM1AMb87s1-K59wuobD7q34U/view?usp=sharing"
        },
        # DAY 29
        {
            "day": 29,
            "topic": "Goethe Mock Test 10.29",
            "chapter": "10.29",
            "goal": "Practice how the final exams for the lesen will look like",
            "assignment": True,
            "instruction": "Answer everything on the phone and dont write in your book. The answers will be sent to your email",
            "video": "",
            "youtube_link": "",
            "workbook_link": "https://forms.gle/YqCEMXTF5d3N9Q7C7"
        },
    ]
#
def get_b1_schedule():
    return [
        # TAG 1
        {
            "day": 1,
            "topic": "Traumwelten (Übung) 1.1",
            "chapter": "1.1",
            "goal": "Über Traumwelten und Fantasie sprechen.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "grammar_topic": "Präsens & Perfekt",
            "video": "https://youtu.be/wMrdW2DhD5o",
            "youtube_link": "https://youtu.be/wMrdW2DhD5o",
            "grammarbook_link": "https://drive.google.com/file/d/17dO2pWXKQ3V3kWZIgLHXpLJ-ozKHKxu5/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1gTcOHHGW2bXKkhxAC38jdl6OikgHCT9g/view?usp=sharing"
        },
        # TAG 2
        {
            "day": 2,
            "topic": "Freunde fürs Leben (Übung) 1.2",
            "chapter": "1.2",
            "goal": "Freundschaften und wichtige Eigenschaften beschreiben.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "grammar_topic": "Präteritum – Vergangene Erlebnisse erzählen",
            "video": "https://youtu.be/piJE4ucYFuc",
            "youtube_link": "https://youtu.be/piJE4ucYFuc",
            "grammarbook_link": "https://drive.google.com/file/d/1St8MpH616FiJmJjTYI9b6hEpNCQd5V0T/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1AgjhFYw07JYvsgVP1MBKYEMFBjeAwQ1e/view?usp=sharing"
        },
        # TAG 3
        {
            "day": 3,
            "topic": "Erfolgsgeschichten (Übung) 1.3",
            "chapter": "1.3",
            "goal": "Über Erfolge und persönliche Erlebnisse berichten.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "grammar_topic": "Adjektivdeklination mit unbestimmten Artikeln",
            "video": "https://youtu.be/8k0Iaw_-o8c",
            "youtube_link": "https://youtu.be/8k0Iaw_-o8c",
            "grammarbook_link": "https://drive.google.com/file/d/1kUtriLOZfJXUxj2IVU2VHZZkghIWDWKv/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1qVANqTLg4FOU40_WfLZyVTu5KBluzYrh/view?usp=sharing"
        },
        # TAG 4
        {
            "day": 4,
            "topic": "Wohnung suchen (Übung) 2.4",
            "chapter": "2.4",
            "goal": "Über Wohnungssuche und Wohnformen sprechen.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "grammar_topic": "Wechselpräpositionen – In der Stadt, auf dem Land",
            "video": "https://youtu.be/kR8SmSY99c8",
            "youtube_link": "https://youtu.be/kR8SmSY99c8",
            "grammarbook_link": "https://drive.google.com/file/d/1NW5F0R5zj6nn2SqDjhpQlkGcfK-UBUqk/view?usp=drive_link",
            "workbook_link": "https://drive.google.com/file/d/12r_HE51QtpknXSSU0R75ur-EDFpTjzXU/view?usp=sharing"
        },
        # TAG 5
        {
            "day": 5,
            "topic": "Der Besichtigungstermin (Übung) 2.5",
            "chapter": "2.5",
            "goal": "Einen Besichtigungstermin beschreiben.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 6
        {
            "day": 6,
            "topic": "Leben in der Stadt oder auf dem Land? 2.6",
            "chapter": "2.6",
            "goal": "Stadtleben und Landleben vergleichen.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 7
        {
            "day": 7,
            "topic": "Fast Food vs. Hausmannskost 3.7",
            "chapter": "3.7",
            "goal": "Fast Food und Hausmannskost vergleichen.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 8
        {
            "day": 8,
            "topic": "Alles für die Gesundheit 3.8",
            "chapter": "3.8",
            "goal": "Tipps für Gesundheit geben und Arztbesuche besprechen.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 9
        {
            "day": 9,
            "topic": "Work-Life-Balance im modernen Arbeitsumfeld 3.9",
            "chapter": "3.9",
            "goal": "Über Work-Life-Balance und Stress sprechen.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 10
        {
            "day": 10,
            "topic": "Digitale Auszeit und Selbstfürsorge 4.10",
            "chapter": "4.10",
            "goal": "Über digitale Auszeiten und Selbstfürsorge sprechen.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 11
        {
            "day": 11,
            "topic": "Teamspiele und Kooperative Aktivitäten 4.11",
            "chapter": "4.11",
            "goal": "Über Teamarbeit und kooperative Aktivitäten sprechen.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 12
        {
            "day": 12,
            "topic": "Abenteuer in der Natur 4.12",
            "chapter": "4.12",
            "goal": "Abenteuer und Erlebnisse in der Natur beschreiben.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 13
        {
            "day": 13,
            "topic": "Eigene Filmkritik schreiben 4.13",
            "chapter": "4.13",
            "goal": "Eine Filmkritik schreiben und Filme bewerten.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 14
        {
            "day": 14,
            "topic": "Traditionelles vs. digitales Lernen 5.14",
            "chapter": "5.14",
            "goal": "Traditionelles und digitales Lernen vergleichen.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 15
        {
            "day": 15,
            "topic": "Medien und Arbeiten im Homeoffice 5.15",
            "chapter": "5.15",
            "goal": "Über Mediennutzung und Homeoffice sprechen.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 16
        {
            "day": 16,
            "topic": "Prüfungsangst und Stressbewältigung 5.16",
            "chapter": "5.16",
            "goal": "Prüfungsangst und Strategien zur Stressbewältigung besprechen.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 17
        {
            "day": 17,
            "topic": "Wie lernt man am besten? 5.17",
            "chapter": "5.17",
            "goal": "Lerntipps geben und Lernstrategien vorstellen.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 18
        {
            "day": 18,
            "topic": "Wege zum Wunschberuf 6.18",
            "chapter": "6.18",
            "goal": "Über Wege zum Wunschberuf sprechen.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 19
        {
            "day": 19,
            "topic": "Das Vorstellungsgespräch 6.19",
            "chapter": "6.19",
            "goal": "Über Vorstellungsgespräche berichten und Tipps geben.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 20
        {
            "day": 20,
            "topic": "Wie wird man …? (Ausbildung und Qu) 6.20",
            "chapter": "6.20",
            "goal": "Über Ausbildung und Qualifikationen sprechen.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 21
        {
            "day": 21,
            "topic": "Lebensformen heute – Familie, Wohnge 7.21",
            "chapter": "7.21",
            "goal": "Lebensformen, Familie und Wohngemeinschaften beschreiben.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 22
        {
            "day": 22,
            "topic": "Was ist dir in einer Beziehung wichtig? 7.22",
            "chapter": "7.22",
            "goal": "Über Werte in Beziehungen sprechen.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 23
        {
            "day": 23,
            "topic": "Erstes Date – Typische Situationen 7.23",
            "chapter": "7.23",
            "goal": "Typische Situationen beim ersten Date beschreiben.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        # TAG 24
        {
            "day": 24,
            "topic": "Konsum und Nachhaltigkeit 8.24",
            "chapter": "8.24",
            "goal": "Nachhaltigen Konsum und Umweltschutz diskutieren.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": "https://drive.google.com/file/d/1x8IM6xcjR2hv3jbnnNudjyxLWPiT0-VL/view?usp=sharing"
        },
        # TAG 25
        {
            "day": 25,
            "topic": "Online einkaufen – Rechte und Risiken 8.25",
            "chapter": "8.25",
            "goal": "Rechte und Risiken beim Online-Shopping besprechen.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": "https://drive.google.com/file/d/1If0R3cIT8KwjeXjouWlQ-VT03QGYOSZz/view?usp=sharing"
        },
        # TAG 26
        {
            "day": 26,
            "topic": "Reiseprobleme und Lösungen 9.26",
            "chapter": "9.26",
            "goal": "Reiseprobleme und Lösungen beschreiben.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": "https://drive.google.com/file/d/1BMwDDkfPJVEhL3wHNYqGMAvjOts9tv24/view?usp=sharing"
        },
        # TAG 27
        {
            "day": 27,
            "topic": "Umweltfreundlich im Alltag 10.27",
            "chapter": "10.27",
            "goal": "Umweltfreundliches Verhalten im Alltag beschreiben.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": "https://drive.google.com/file/d/15fjOKp_u75GfcbvRJVbR8UbHg-cgrgWL/view?usp=sharing"
        },
        # TAG 28
        {
            "day": 28,
            "topic": "Klimafreundlich leben 10.28",
            "chapter": "10.28",
            "goal": "Klimafreundliche Lebensweisen vorstellen.",
            "assignment": True,
            "instruction": "Schau das Video, wiederhole die Grammatik und mache die Aufgabe.",
            "video": "",
            "youtube_link": "",
            "grammarbook_link": "",
            "workbook_link": "https://drive.google.com/file/d/1iBeZHMDq_FnusY4kkRwRQvyOfm51-COU/view?usp=sharing"
        },
    ]



def get_b2_schedule():
    return [
        {
            "day": 1,
            "topic": "Persönliche Identität und Selbstverständnis",
            "chapter": "1.1",
            "goal": "Drücken Sie Ihre persönliche Identität und Ihre Werte aus.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "https://youtu.be/a9LxkxNdnEg",
            "youtube_link": "https://youtu.be/a9LxkxNdnEg",
            "grammarbook_link": "https://drive.google.com/file/d/17pVc0VfLm32z4zmkaaa_cdshKJEQQxYa/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1D1eb-iwfl_WA2sXPOSPD_66NCiTB4o2w/view?usp=sharing",
            "grammar_topic": "Adjektivdeklination (Adjektivendungen nach bestimmten/unbestimmten Artikeln)"
        },
        {
            "day": 2,
            "topic": "Beziehungen und Kommunikation",
            "chapter": "1.2",
            "goal": "Diskutieren Sie über Beziehungstypen und Kommunikationsstrategien.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "https://youtu.be/gCzZnddwC_c",
            "youtube_link": "https://youtu.be/gCzZnddwC_c",
            "grammarbook_link": "https://drive.google.com/file/d/1Mlt-cK6YqPuJe9iCWfqT9DOG9oKhJBdK/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1XCLW0y-MMyIu_bNO3EkKIgp-8QLKgEek/view?usp=sharing",
            "grammar_topic": "Konjunktiv II (höfliche Bitten & hypothetische Situationen)"
        },
        {
            "day": 3,
            "topic": "Öffentliches vs. Privates Leben",
            "chapter": "1.3",
            "goal": "Vergleichen Sie das öffentliche und private Leben in Deutschland und Ihrem Land.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "https://drive.google.com/file/d/1R0sQc4uSWQNUxPa0_Gdz7PiQaiCyQrrL/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1VteR5sVx_uiKdhSVMBosMxiXe1lfnQnW/view?usp=sharing",
            "grammar_topic": "Passiv (Präsens und Vergangenheit)"
        },
        {
            "day": 4,
            "topic": "Beruf und Karriere",
            "chapter": "1.4",
            "goal": "Sprechen Sie über Berufe, Lebensläufe und Vorstellungsgespräche.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "https://drive.google.com/file/d/1_xVoBqbwCSCs0Xps2Rlx92Ho43Pcbreu/view?usp=sharing",
            "workbook_link": "https://drive.google.com/file/d/1tEKd5Umb-imLpPYrmFfNQyjf4oe2weBp/view?usp=sharing",
            "grammar_topic": "Konjunktiv I"
        },
        {
            "day": 5,
            "topic": "Bildung und Lernen",
            "chapter": "1.5",
            "goal": "Diskutieren Sie das Bildungssystem und lebenslanges Lernen.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Nominalisierung von Verben"
        },
        {
            "day": 6,
            "topic": "Migration und Integration",
            "chapter": "2.1",
            "goal": "Erforschen Sie Migration, Integration und kulturelle Identität.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Temporale Nebensätze (als, wenn, nachdem, während, bevor)"
        },
        {
            "day": 7,
            "topic": "Gesellschaftliche Vielfalt",
            "chapter": "2.2",
            "goal": "Untersuchen Sie Vielfalt und Inklusion in modernen Gesellschaften.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Relativsätze mit Präpositionen"
        },
        {
            "day": 8,
            "topic": "Politik und Engagement",
            "chapter": "2.3",
            "goal": "Lernen Sie politische Systeme und bürgerschaftliches Engagement kennen.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Finale und kausale Nebensätze (damit, um...zu, weil, da)"
        },
        {
            "day": 9,
            "topic": "Technologie und Digitalisierung",
            "chapter": "2.4",
            "goal": "Diskutieren Sie die digitale Transformation und deren Auswirkungen.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Infinitivkonstruktionen mit zu (ohne zu, anstatt zu, um zu, etc.)"
        },
        {
            "day": 10,
            "topic": "Umwelt und Nachhaltigkeit",
            "chapter": "2.5",
            "goal": "Sprechen Sie über Umweltschutz und Nachhaltigkeit.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Konjunktiv II Vergangenheit (hypothetische Vergangenheit)"
        },
        {
            "day": 11,
            "topic": "Gesundheit und Wohlbefinden",
            "chapter": "3.1",
            "goal": "Beschreiben Sie Gesundheit, Wohlbefinden und Lebensstil.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Reflexive Verben und Pronomen"
        },
        {
            "day": 12,
            "topic": "Konsum und Medien",
            "chapter": "3.2",
            "goal": "Analysieren Sie Medieneinfluss und Konsumgewohnheiten.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Modalverben im Passiv"
        },
        {
            "day": 13,
            "topic": "Reisen und Mobilität",
            "chapter": "3.3",
            "goal": "Planen Sie Reisen und diskutieren Sie Transportmöglichkeiten.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Präpositionen mit Genitiv"
        },
        {
            "day": 14,
            "topic": "Wohnen und Zusammenleben",
            "chapter": "3.4",
            "goal": "Vergleichen Sie verschiedene Wohnformen und Gemeinschaften.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Steigerung der Adjektive (Komparativ & Superlativ)"
        },
        {
            "day": 15,
            "topic": "Kunst und Kultur",
            "chapter": "3.5",
            "goal": "Entdecken Sie Kunst, Literatur und kulturelle Veranstaltungen.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Satzbau und Satzstellung"
        },
        {
            "day": 16,
            "topic": "Wissenschaft und Forschung",
            "chapter": "4.1",
            "goal": "Diskutieren Sie wissenschaftliche Entdeckungen und Forschung.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Partizipialkonstruktionen"
        },
        {
            "day": 17,
            "topic": "Feste und Traditionen",
            "chapter": "4.2",
            "goal": "Beschreiben Sie traditionelle Feste und Bräuche.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": ""
        },
        {
            "day": 18,
            "topic": "Freizeit und Hobbys",
            "chapter": "4.3",
            "goal": "Sprechen Sie über Freizeit und Hobbys.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Pronominaladverbien (darauf, worüber, etc.)"
        },
        {
            "day": 19,
            "topic": "Ernährung und Esskultur",
            "chapter": "4.4",
            "goal": "Diskutieren Sie über Essen, Ernährung und Essgewohnheiten.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Indirekte Rede"
        },
        {
            "day": 20,
            "topic": "Mode und Lebensstil",
            "chapter": "4.5",
            "goal": "Untersuchen Sie Mode- und Lebensstiltrends.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": ""
        },
        {
            "day": 21,
            "topic": "Werte und Normen",
            "chapter": "5.1",
            "goal": "Analysieren Sie Werte, Normen und deren Auswirkungen.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Negation: kein-, nicht, ohne, weder...noch"
        },
        {
            "day": 22,
            "topic": "Sprache und Kommunikation",
            "chapter": "5.2",
            "goal": "Diskutieren Sie Sprachenlernen und Kommunikationsstrategien.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Nominalstil vs. Verbalstil"
        },
        {
            "day": 23,
            "topic": "Innovation und Zukunft",
            "chapter": "5.3",
            "goal": "Spekulieren Sie über die Zukunft und Innovationen.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Futur I und II"
        },
        {
            "day": 24,
            "topic": "Gesellschaftliche Herausforderungen",
            "chapter": "5.4",
            "goal": "Diskutieren Sie gesellschaftliche Herausforderungen und mögliche Lösungen.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Subjekt- und Objektive Sätze"
        },
        {
            "day": 25,
            "topic": "Globalisierung und internationale Beziehungen",
            "chapter": "5.5",
            "goal": "Erforschen Sie Globalisierung und deren Auswirkungen.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": "Partizipialattribute"
        },
        {
            "day": 26,
            "topic": "Kreatives Schreiben & Projekte",
            "chapter": "6.1",
            "goal": "Entwickeln Sie kreative Schreibfähigkeiten.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": ""
        },
        {
            "day": 27,
            "topic": "Prüfungstraining & Wiederholung",
            "chapter": "6.2",
            "goal": "Wiederholen Sie B2-Themen und üben Sie Prüfungsformate.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": ""
        },
        {
            "day": 28,
            "topic": "Abschlusspräsentation & Feedback",
            "chapter": "6.3",
            "goal": "Fassen Sie die Kursthemen zusammen und reflektieren Sie Ihren Fortschritt.",
            "instruction": "Schauen Sie das Video, wiederholen Sie die Grammatik und bearbeiten Sie das Arbeitsbuch.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": "",
            "grammar_topic": ""
        }
    ]



# === C1 Schedule Template ===
def get_c1_schedule():
    return [
        {
            "day": 1,
            "topic": "C1 Welcome & Orientation",
            "chapter": "0.0",
            "goal": "Get familiar with the C1 curriculum and expectations.",
            "instruction": "Read the C1 orientation, join the forum, and write a short self-intro.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": ""
        },
        {
            "day": 2,
            "topic": "C1 Diagnostic Writing",
            "chapter": "0.1",
            "goal": "Write a sample essay for initial assessment.",
            "instruction": "Write and upload a short essay on the assigned topic.",
            "video": "",
            "grammarbook_link": "",
            "workbook_link": ""
        }
        # You can add more C1 lessons here in the future
    ]


# --- FORCE A MOCK LOGIN FOR TESTING ---
if "student_row" not in st.session_state:
    st.session_state["student_row"] = {
        "Name": "Test Student",
        "Level": "A1",
        "StudentCode": "demo001",
        "ClassName": "A1 Berlin Klasse",
    }

student_row = st.session_state.get("student_row", {})
student_level = student_row.get("Level", "A1").upper()

# --- Cache level schedules with TTL for periodic refresh ---
@st.cache_data(ttl=86400)
def load_level_schedules():
    return {
        "A1": get_a1_schedule(),
        "A2": get_a2_schedule(),
        "B1": get_b1_schedule(),
        "B2": get_b2_schedule(),
        "C1": get_c1_schedule(),
    }

# --- Helpers ---
def render_assignment_reminder():
    st.markdown(
        '''
        <div style="
            box-sizing: border-box;
            width: 100%;
            max-width: 600px;
            padding: 16px;
            background: #ffc107;
            color: #000;
            border-left: 6px solid #e0a800;
            margin: 16px auto;
            border-radius: 8px;
            font-size: 1.1rem;
            line-height: 1.4;
            text-align: center;
            overflow-wrap: break-word;
            word-wrap: break-word;
        ">
            ⬆️ <strong>Your Assignment:</strong><br>
            Complete the exercises in your <em>workbook</em> for this chapter.
        </div>
        ''',
        unsafe_allow_html=True
    )

def render_link(label, url):
    st.markdown(f"- [{label}]({url})")

@st.cache_data(ttl=86400)
def build_wa_message(name, code, level, day, chapter, answer):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    return (
        f"Learn Language Education Academy – Assignment Submission\n"
        f"Name: {name}\n"
        f"Code: {code}\n"
        f"Level: {level}\n"
        f"Day: {day}\n"
        f"Chapter: {chapter}\n"
        f"Date: {timestamp}\n"
        f"Answer: {answer if answer.strip() else '[See attached file/photo]'}"
    )

SLACK_DEBUG = (os.getenv("SLACK_DEBUG", "0") == "1")

def _slack_url() -> str:
    # 1) Render env var  2) optional fallback to st.secrets.slack.webhook_url
    url = (os.getenv("SLACK_WEBHOOK_URL") or "").strip()
    if not url:
        try:
            url = (st.secrets.get("slack", {}).get("webhook_url", "") if hasattr(st, "secrets") else "").strip()
        except Exception:
            url = ""
    return url

def notify_slack(text: str):
    """
    Returns (ok: bool, info: str). Uses one webhook for all events.
    Set SLACK_DEBUG=1 in Render to see failure details in-app (admins only).
    """
    url = _slack_url()
    if not url:
        return False, "missing_webhook"
    try:
        resp = requests.post(url, json={"text": text}, timeout=6)
        ok = 200 <= resp.status_code < 300
        return ok, f"status={resp.status_code}"
    except Exception as e:
        return False, str(e)
        
def highlight_terms(text, terms):
    if not text: return ""
    for term in terms:
        if not term.strip():
            continue
        pattern = re.compile(re.escape(term), re.IGNORECASE)
        text = pattern.sub(f"<span style='background:yellow;border-radius:0.23em;'>{term}</span>", text)
    return text

def filter_matches(lesson, terms):
    searchable = (
        str(lesson.get('topic', '')).lower() +
        str(lesson.get('chapter', '')).lower() +
        str(lesson.get('goal', '')).lower() +
        str(lesson.get('instruction', '')).lower() +
        str(lesson.get('grammar_topic', '')).lower() +
        str(lesson.get('day', '')).lower()
    )
    return any(term in searchable for term in terms)
    
def render_section(day_info, key, title, icon):
    content = day_info.get(key)
    if not content:
        return
    items = content if isinstance(content, list) else [content]
    st.markdown(f"#### {icon} {title}")
    for idx, part in enumerate(items):
        if len(items) > 1:
            st.markdown(f"###### {icon} Part {idx+1} of {len(items)}: Chapter {part.get('chapter','')}")
        if part.get('video'):
            st.video(part['video'])
        if part.get('grammarbook_link'):
            render_link("📘 Grammar Book (Notes)", part['grammarbook_link'])
            st.markdown(
                '<em>Further notice:</em> 📘 contains notes; 📒 is your workbook assignment.',
                unsafe_allow_html=True
            )
        if part.get('workbook_link'):
            render_link("📒 Workbook (Assignment)", part['workbook_link'])
            render_assignment_reminder()
        extras = part.get('extra_resources')
        if extras:
            for ex in (extras if isinstance(extras, list) else [extras]):
                render_link("🔗 Extra", ex)


def post_message(level, code, name, text, reply_to=None):
    posts_ref = db.collection("class_board").document(level).collection("posts")
    posts_ref.add({
        "student_code": code,
        "student_name": name,
        "text": text.strip(),
        "timestamp": datetime.utcnow(),
        "reply_to": reply_to,
    })

RESOURCE_LABELS = {
    'video': '🎥 Video',
    'grammarbook_link': '📘 Grammar',
    'workbook_link': '📒 Workbook',
    'extra_resources': '🔗 Extra'
}

# ---- Firestore Helpers ----
def load_notes_from_db(student_code):
    ref = db.collection("learning_notes").document(student_code)
    doc = ref.get()
    return doc.to_dict().get("notes", []) if doc.exists else []

def save_notes_to_db(student_code, notes):
    ref = db.collection("learning_notes").document(student_code)
    ref.set({"notes": notes}, merge=True)
    

if tab == "My Course":
    # === HANDLE ALL SWITCHING *BEFORE* ANY WIDGET ===
    # Jump flags set by buttons elsewhere
    if st.session_state.get("__go_classroom"):
        st.session_state["coursebook_subtab"] = "🧑‍🏫 Classroom"
        del st.session_state["__go_classroom"]
        st.rerun()

    if st.session_state.get("__go_notes"):
        st.session_state["coursebook_subtab"] = "📒 Learning Notes"
        del st.session_state["__go_notes"]
        st.rerun()

    # Backward-compat: older code may still set this
    if st.session_state.get("switch_to_notes"):
        st.session_state["coursebook_subtab"] = "📒 Learning Notes"
        del st.session_state["switch_to_notes"]
        st.rerun()

    # First run default
    if "coursebook_subtab" not in st.session_state:
        st.session_state["coursebook_subtab"] = "🧑‍🏫 Classroom"

    # Header (render once)
    st.markdown(
        '''
        <div style="
            padding: 16px;
            background: #007bff;
            color: #ffffff;
            border-radius: 8px;
            text-align: center;
            margin-bottom: 16px;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        ">
            <span style="font-size:1.8rem; font-weight:600;">📈 My Course</span>
        </div>
        ''',
        unsafe_allow_html=True
    )
    st.divider()

    # Subtabs (1: Classroom, 2: Course Book, 3: Learning Notes)
    cb_subtab = st.radio(
        "Select section:",
        ["🧑‍🏫 Classroom", "📘 Course Book", "📒 Learning Notes"],
        horizontal=True,
        key="coursebook_subtab"
    )


    # === COURSE BOOK SUBTAB ===
    if cb_subtab == "📘 Course Book":
        st.markdown(
            '''
            <div style="
                padding: 16px;
                background: #007bff;
                color: #ffffff;
                border-radius: 8px;
                text-align: center;
                margin-bottom: 16px;
                box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            ">
                <span style="font-size:1.8rem; font-weight:600;">📈 Course Book</span>
            </div>
            ''',
            unsafe_allow_html=True
        )
        st.divider()

        schedules = load_level_schedules()
        schedule = schedules.get(student_level, schedules.get("A1", []))

        query = st.text_input("🔍 Search for topic, chapter, grammar, day, or anything…")
        search_terms = [q for q in query.strip().lower().split() if q] if query else []

        if search_terms:
            matches = [(i, d) for i, d in enumerate(schedule) if filter_matches(d, search_terms)]
            if not matches:
                st.warning("No matching lessons. Try simpler terms or check spelling.")
                st.stop()

            labels = []
            for _, d in matches:
                title = highlight_terms(f"Day {d['day']}: {d['topic']}", search_terms)
                grammar = highlight_terms(d.get("grammar_topic", ""), search_terms)
                labels.append(
                    f"{title}  {'<span style=\"color:#007bff\">['+grammar+']</span>' if grammar else ''}"
                )

            # Bold header for lessons dropdown
            st.markdown(
                "<span style='font-weight:700; font-size:1rem;'>Lessons:</span>",
                unsafe_allow_html=True
            )
            sel = st.selectbox(
                "",  # label hidden
                list(range(len(matches))),
                format_func=lambda i: labels[i],
                key="course_search_sel"
            )
            idx = matches[sel][0]
        else:
            # Bold header for lesson/day dropdown
            st.markdown(
                "<span style='font-weight:700; font-size:1rem;'>Choose your lesson/day:</span>",
                unsafe_allow_html=True
            )
            idx = st.selectbox(
                "",  # label hidden
                range(len(schedule)),
                format_func=lambda i: f"Day {schedule[i]['day']} - {schedule[i]['topic']}"
            )

        st.divider()

        # Progress Bar
        total = len(schedule)
        done = idx + 1
        pct = int(done / total * 100) if total else 0
        st.progress(pct)
        st.markdown(f"**You’ve loaded {done} / {total} lessons ({pct}%)**")
        st.divider()

        # ===== COURSE BOOK INFO =====
        with st.expander("📚 Course Book & Study Recommendations", expanded=False):

            # Recommended time
            LEVEL_TIME = {"A1": 15, "A2": 25, "B1": 30, "B2": 40, "C1": 45}
            rec_time = LEVEL_TIME.get(student_level, 20)
            st.info(f"⏱️ **Recommended:** Invest about {rec_time} minutes to complete this lesson fully.")

            # Suggested end dates
            start_str = student_row.get("ContractStart", "")
            start_date = None
            for fmt in ("%m/%d/%Y", "%Y-%m-%d", "%d/%m/%Y"):
                try:
                    start_date = datetime.strptime(start_str, fmt).date()
                    break
                except:
                    continue

            if start_date:
                total = total  # assuming this variable is already defined earlier in your code
                # calculate weeks for different paces
                weeks_three = (total + 2) // 3
                weeks_two   = (total + 1) // 2
                weeks_one   = total

                end_three = start_date + timedelta(weeks=weeks_three)
                end_two   = start_date + timedelta(weeks=weeks_two)
                end_one   = start_date + timedelta(weeks=weeks_one)

                # spacer layout
                spacer, content = st.columns([3, 7])
                with content:
                    st.success(f"If you complete **three sessions per week**, you will finish by **{end_three.strftime('%A, %d %B %Y')}**.")
                    st.info(f"If you complete **two sessions per week**, you will finish by **{end_two.strftime('%A, %d %B %Y')}**.")
                    st.warning(f"If you complete **one session per week**, you will finish by **{end_one.strftime('%A, %d %B %Y')}**.")
            else:
                spacer, content = st.columns([3, 7])
                with content:
                    st.warning("❓ Start date missing or invalid. Please update your contract start date.")
#
        info = schedule[idx]
        # ---- Fix for highlight and header ----
        lesson_title = f"Day {info['day']}: {info['topic']}"
        highlighted_title = highlight_terms(lesson_title, search_terms)
        st.markdown(
            f"### {highlighted_title} (Chapter {info['chapter']})",
            unsafe_allow_html=True
        )
        st.divider()

        if info.get("grammar_topic"):
            st.markdown(
                f"**🔤 Grammar Focus:** {highlight_terms(info['grammar_topic'], search_terms)}",
                unsafe_allow_html=True
            )
        if info.get("goal"):
            st.markdown(f"**🎯 Goal:**  {info['goal']}")
        if info.get("instruction"):
            st.markdown(f"**📝 Instruction:**  {info['instruction']}")

        # ---- RENDER SECTION: lesen_hören, schreiben_sprechen, each with fallback YouTube link ----
        def render_section(day_info, key, title, icon):
            content = day_info.get(key)
            if not content:
                return
            items = content if isinstance(content, list) else [content]
            st.markdown(f"#### {icon} {title}")
            for idx, part in enumerate(items):
                if len(items) > 1:
                    st.markdown(
                        f"###### {icon} Part {idx+1} of {len(items)}: Chapter {part.get('chapter','')}"
                    )
                # --- Embed video and show link if available ---
                if part.get('video'):
                    st.video(part['video'])
                    st.markdown(f"[▶️ Watch on YouTube]({part['video']})")
                # --- Also support explicit youtube_link (if different from 'video') ---
                elif part.get('youtube_link'):
                    st.markdown(f"[▶️ Watch on YouTube]({part['youtube_link']})")
                if part.get('grammarbook_link'):
                    st.markdown(f"- [📘 Grammar Book (Notes)]({part['grammarbook_link']})")
                    st.markdown(
                        '<em>Further notice:</em> 📘 contains notes; 📒 is your workbook assignment.',
                        unsafe_allow_html=True
                    )
                if part.get('workbook_link'):
                    st.markdown(f"- [📒 Workbook (Assignment)]({part['workbook_link']})")
                    render_assignment_reminder()
                extras = part.get('extra_resources')
                if extras:
                    for ex in (extras if isinstance(extras, list) else [extras]):
                        st.markdown(f"- [🔗 Extra]({ex})")

        render_section(info, "lesen_hören", "Lesen & Hören", "📚")
        render_section(info, "schreiben_sprechen", "Schreiben & Sprechen", "📝")

        # ---- Show resource links for upper levels if needed ----
        if student_level in ["A2", "B1", "B2", "C1"]:
            for res, label in RESOURCE_LABELS.items():
                val = info.get(res)
                if val:
                    if res == "video":
                        st.video(val)
                        st.markdown(f"[▶️ Watch on YouTube]({val})")
                    else:
                        st.markdown(f"- [{label}]({val})", unsafe_allow_html=True)
            st.markdown(
                '<em>Further notice:</em> 📘 contains notes; 📒 is your workbook assignment.',
                unsafe_allow_html=True
            )

        # --- Translation Tools ---
        with st.expander("🌐 Translation Tools", expanded=False):
            st.markdown("---")
            st.markdown(
                '**Need translation?** '
                '[🌐 DeepL Translator](https://www.deepl.com/translator) &nbsp; | &nbsp; '
                '[🌐 Google Translate](https://translate.google.com)',
                unsafe_allow_html=True
            )
            st.caption("Copy any text from the course book and paste it into your preferred translator.")

        st.divider()

        # --- Video of the Day ---
        with st.expander("🎬 Video of the Day for Your Level", expanded=False):
            playlist_id = YOUTUBE_PLAYLIST_IDS.get(student_level)
            if playlist_id:
                video_list = fetch_youtube_playlist_videos(playlist_id, YOUTUBE_API_KEY)
                if video_list:
                    today_idx = date.today().toordinal()
                    pick = today_idx % len(video_list)
                    video = video_list[pick]
                    st.markdown(f"**{video['title']}**")
                    st.video(video['url'])
                else:
                    st.info("No videos found for your level’s playlist. Check back soon!")
            else:
                st.info("No playlist found for your level yet. Stay tuned!")

        st.divider()

        # === SUBMIT ASSIGNMENT (Render env secret + context banner + submit + lock; NO AUTO-REFRESH) ===
        st.markdown("### ✅ Submit Your Assignment")

        # Clear context banner so students see exactly where they are
        st.markdown(
            f"""
            <div style="box-sizing:border-box;padding:14px 16px;border-radius:10px;
                        background:#f0f9ff;border:1px solid #bae6fd;margin:6px 0 12px 0;">
              <div style="font-size:1.05rem;">
                📌 <b>You're on:</b> Level <b>{student_level}</b> • Day <b>{info['day']}</b> • Chapter <b>{info['chapter']}</b>
              </div>
              <div style="color:#0369a1;margin-top:4px;">
                Make sure this matches the assignment your tutor set. If not, change the lesson from the dropdown above.
              </div>
            </div>
            """,
            unsafe_allow_html=True
        )

        # Render env (Slack) — set on Render dashboard: SLACK_WEBHOOK_URL=...
        import os, requests
        from datetime import datetime

        def get_slack_webhook() -> str:
            return (os.getenv("SLACK_WEBHOOK_URL") or "").strip()

        # --- Draft persistence (save + load from Firestore) ---
        def save_draft_to_db(code, lesson_key, text):
            doc_ref = db.collection('draft_answers').document(code)
            doc_ref.set(
                {lesson_key: text, f"{lesson_key}__updated_at": datetime.utcnow()},
                merge=True
            )

        def load_draft_from_db(code, lesson_key) -> str:
            try:
                doc = db.collection('draft_answers').document(code).get()
                if doc.exists:
                    data = doc.to_dict() or {}
                    return data.get(lesson_key, "")
            except Exception:
                pass
            return ""

        code = student_row.get('StudentCode', 'demo001')
        lesson_key = f"draft_{info['chapter']}"     # unique per chapter
        chapter_name = f"{info['chapter']} – {info.get('topic', '')}"
        name = st.text_input("Name", value=student_row.get('Name', ''))

        # Persisted lock per lesson
        locked_key = f"{lesson_key}_locked"
        locked = st.session_state.get(locked_key, False)

        # One-time hydration from Firestore so text survives refresh/restart
        if not st.session_state.get(f"{lesson_key}__hydrated", False):
            existing = load_draft_from_db(code, lesson_key)
            if existing and not st.session_state.get(lesson_key):
                st.session_state[lesson_key] = existing
                st.info("💾 Loaded your saved draft.")
            st.session_state[f"{lesson_key}__hydrated"] = True

        # Answer Box (autosaves on change ONLY)
        st.subheader("✍️ Your Answer (Autosaves)")
        def autosave_draft():
            text = st.session_state.get(lesson_key, "")
            save_draft_to_db(code, lesson_key, text)
            st.session_state[f"{lesson_key}_saved"] = True
            st.session_state[f"{lesson_key}_saved_at"] = datetime.utcnow()

        st.text_area(
            "Type all your answers here",
            value=st.session_state.get(lesson_key, ""),
            height=500,
            key=lesson_key,
            on_change=autosave_draft,  # saves when the field loses focus or the widget updates
            disabled=locked,
            help="Draft autosaves when you click outside the box or change focus."
        )

        cols_save = st.columns([1,2])
        with cols_save[0]:
            if st.button("💾 Save Draft now", disabled=locked):
                autosave_draft()
                st.success("Draft saved.")
        with cols_save[1]:
            ts = st.session_state.get(f"{lesson_key}_saved_at")
            if ts:
                st.caption("Last saved: " + ts.strftime("%Y-%m-%d %H:%M") + " UTC")

        # Instructions
        with st.expander("📌 How to Submit", expanded=False):
            st.markdown(f"""
                1) Check you’re on the correct page: **Level {student_level} • Day {info['day']} • Chapter {info['chapter']}**.  
                2) Tick the two confirmations below.  
                3) Click **Confirm & Submit**.  
                4) Your box will lock (read-only).  
                _You’ll get an **email** when it’s marked. See **Results & Resources** for scores & feedback._
            """)

        # Slack notify helper (uses Render env only)
        def notify_slack_submission(webhook_url: str, *, student_name: str, student_code: str,
                                    level: str, day: int, chapter: str, receipt: str, preview: str):
            if not webhook_url:
                return
            text = (
                f"*New submission* • {student_name} ({student_code})\n"
                f"*Level:* {level}  •  *Day:* {day}\n"
                f"*Chapter:* {chapter}\n"
                f"*Ref:* `{receipt}`\n"
                f"*Preview:* {preview[:180]}{'…' if len(preview) > 180 else ''}"
            )
            try:
                requests.post(webhook_url, json={"text": text}, timeout=6)
            except Exception:
                pass  # don't block student

        # Firestore: create submission, return short ref for Slack (hidden from student)
        def submit_answer(code, name, level, day, chapter, lesson_key, answer):
            if not answer or not answer.strip():
                st.warning("Please type your answer before submitting.")
                return False, None
            posts_ref = db.collection("submissions").document(level).collection("posts")
            now = datetime.utcnow()
            payload = {
                "student_code": code,
                "student_name": name or "Student",
                "level": level,
                "day": day,
                "chapter": chapter,
                "lesson_key": lesson_key,
                "answer": answer.strip(),
                "status": "submitted",
                "created_at": now,
                "updated_at": now,
                "version": 1,
            }
            _, ref = posts_ref.add(payload)
            doc_id = ref.id
            short_ref = f"{doc_id[:8].upper()}-{day}"
            return True, short_ref

        # Two-step confirm + Submit / Save to Notes / Ask a Question
        col1, col2, col3 = st.columns(3)

        with col1:
            st.markdown("#### 🧾 Finalize")
            confirm_final = st.checkbox(
                f"I confirm this is my complete work for Level {student_level} • Day {info['day']} • Chapter {info['chapter']}.",
                key=f"confirm_final_{lesson_key}",
                disabled=locked
            )
            confirm_lock = st.checkbox(
                "I understand it will be locked after I submit.",
                key=f"confirm_lock_{lesson_key}",
                disabled=locked
            )
            can_submit = (confirm_final and confirm_lock and (not locked))

            if st.button("✅ Confirm & Submit", type="primary", disabled=not can_submit):
                ok, short_ref = submit_answer(
                    code=code,
                    name=name,
                    level=student_level,
                    day=info["day"],
                    chapter=chapter_name,
                    lesson_key=lesson_key,
                    answer=st.session_state.get(lesson_key, "")
                )
                if ok:
                    st.session_state[locked_key] = True
                    st.success("Submitted! Your work has been sent to your tutor.")
                    st.caption("You’ll be **emailed when it’s marked**. Check **Results & Resources** for your score and feedback.")

                    webhook = get_slack_webhook()
                    if webhook:
                        notify_slack_submission(
                            webhook_url=webhook,
                            student_name=name or "Student",
                            student_code=code,
                            level=student_level,
                            day=info["day"],
                            chapter=chapter_name,
                            receipt=short_ref,
                            preview=st.session_state.get(lesson_key, "")
                        )

        # --- Column 2: Ask the Teacher (jump to Classroom Q&A) ---
        with col2:
            st.markdown("#### ❓ Ask the Teacher")
            if st.button("Open Classroom Q&A", key=f"open_qna_{lesson_key}", disabled=locked):
                # set a jump flag; DON'T touch the radio key directly here
                st.session_state["__go_classroom"] = True
                st.rerun()

        # --- Column 3: Add Notes (just jump to Notes tab) ---
        with col3:
            st.markdown("#### 📝 Add Notes")
            if st.button("Open Notes", key=f"open_notes_{lesson_key}", disabled=locked):
                # set a jump flag; no prefills
                st.session_state["__go_notes"] = True
                st.rerun()

        st.divider()

        # Submission status (latest only; no receipt shown)
        def fetch_latest(level, code, lesson_key):
            posts_ref = db.collection("submissions").document(level).collection("posts")
            try:
                docs = posts_ref.where("student_code","==",code)\
                                .where("lesson_key","==",lesson_key)\
                                .order_by("updated_at", direction=firestore.Query.DESCENDING)\
                                .limit(1).stream()
                for d in docs:
                    return d.to_dict()
            except Exception:
                docs = posts_ref.where("student_code","==",code)\
                                .where("lesson_key","==",lesson_key)\
                                .stream()
                items = [d.to_dict() for d in docs]
                items.sort(key=lambda x: x.get("updated_at"), reverse=True)
                return items[0] if items else None
            return None

        latest = fetch_latest(student_level, code, lesson_key)
        if latest:
            ts = latest.get('updated_at')
            when = ts.strftime('%Y-%m-%d %H:%M') + " UTC" if ts else ""
            st.markdown(f"**Status:** `{latest.get('status','submitted')}`  {'·  **Updated:** ' + when if when else ''}")
            st.caption("You’ll receive an **email** when it’s marked. See **Results & Resources** for scores & feedback.")
        else:
            st.info("No submission yet. Complete the two confirmations and click **Confirm & Submit**.")

    if cb_subtab == "🧑‍🏫 Classroom":
        # --- Classroom banner (top of subtab) ---
        st.markdown(
            '''
            <div style="
                padding: 16px;
                background: #0ea5e9;
                color: #ffffff;
                border-radius: 8px;
                text-align: center;
                margin-bottom: 16px;
                box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            ">
                <span style="font-size:1.8rem; font-weight:600;">🧑‍🏫 Classroom</span>
            </div>
            ''',
            unsafe_allow_html=True
        )
        st.divider()

        # ---------- DB (Firestore) bootstrap ----------
        def _get_db():
            # Use existing global if present
            _existing = globals().get("db")
            if _existing is not None:
                return _existing
            # Try Firebase Admin SDK first (firestore.client())
            try:
                import firebase_admin
                from firebase_admin import firestore as fbfs
                if not firebase_admin._apps:
                    firebase_admin.initialize_app()
                return fbfs.client()
            except Exception:
                pass
            # Fallback to Google Cloud Firestore (firestore.Client())
            try:
                from google.cloud import firestore as gcf
                return gcf.Client()
            except Exception:
                st.error(
                    "Firestore client isn't configured. Provide Firebase Admin creds or set GOOGLE_APPLICATION_CREDENTIALS.",
                    icon="🛑",
                )
                raise

        db = _get_db()

        # helpers
        import math, os, requests
        try:
            import streamlit.components.v1 as components
        except Exception:
            components = None

        def _safe_str(v, default: str = "") -> str:
            if v is None:
                return default
            if isinstance(v, float):
                try:
                    if math.isnan(v):
                        return default
                except Exception:
                    pass
            s = str(v).strip()
            return "" if s.lower() in ("nan", "none") else s

        def _safe_upper(v, default: str = "") -> str:
            s = _safe_str(v, default)
            return s.upper() if s else default

        student_row   = st.session_state.get("student_row", {}) or {}
        student_code  = _safe_str(student_row.get("StudentCode"), "demo001")
        student_name  = _safe_str(student_row.get("Name"), "Student")
        student_level = _safe_upper(student_row.get("Level"), "A1")
        class_name    = _safe_str(student_row.get("ClassName")) or f"{student_level} General"

        ADMINS = set()
        try:
            ADMINS = set(st.secrets["roles"]["admins"])
        except Exception:
            pass
        IS_ADMIN = (student_code in ADMINS)

        # ---------- slack helper (use global notify_slack if present; else env/secrets) ----------
        def _notify_slack(text: str):
            try:
                fn = globals().get("notify_slack")
                if callable(fn):
                    try:
                        fn(text)
                        return
                    except Exception:
                        pass
                url = (os.getenv("SLACK_WEBHOOK_URL") or
                       (st.secrets.get("slack", {}).get("webhook_url", "") if hasattr(st, "secrets") else "")).strip()
                if url:
                    try:
                        requests.post(url, json={"text": text}, timeout=6)
                    except Exception:
                        pass
            except Exception:
                pass

        # ===================== ZOOM HEADER (official link + reminder to use calendar) =====================
        # ensure urllib alias exists
        try:
            _ = _urllib.quote
        except Exception:
            import urllib.parse as _urllib

        with st.container():
            st.markdown(
                """
                <div style="padding: 12px; background: #facc15; color: #000; border-radius: 8px;
                     font-size: 1rem; margin-bottom: 16px; text-align: left; font-weight: 600;">
                  📣 <b>Zoom Classroom (Official)</b><br>
                  This is the <u>official Zoom link</u> for your class. <span style="font-weight:500;">Add the calendar below to get notifications before each class.</span>
                </div>
                """,
                unsafe_allow_html=True,
            )

            ZOOM = {
                "link": "https://us06web.zoom.us/j/6886900916?pwd=bEdtR3RLQ2dGTytvYzNrMUV3eFJwUT09",
                "meeting_id": "688 690 0916",
                "passcode": "german",
            }
            # Allow secrets override
            try:
                zs = st.secrets.get("zoom", {})
                if zs.get("link"):       ZOOM["link"]       = zs["link"]
                if zs.get("meeting_id"): ZOOM["meeting_id"] = zs["meeting_id"]
                if zs.get("passcode"):   ZOOM["passcode"]   = zs["passcode"]
            except Exception:
                pass

            # Build iOS/Android deep-link (opens Zoom app directly)
            _mid_digits = ZOOM["meeting_id"].replace(" ", "")
            _pwd_enc = _urllib.quote(ZOOM["passcode"] or "")
            zoom_deeplink = f"zoommtg://zoom.us/join?action=join&confno={_mid_digits}&pwd={_pwd_enc}"

            z1, z2 = st.columns([3, 2])
            with z1:
                # Primary join button (browser)
                try:
                    st.link_button("➡️ Join Zoom Meeting (Browser)", ZOOM["link"], key="zoom_join_btn")
                except Exception:
                    st.markdown(f"[➡️ Join Zoom Meeting (Browser)]({ZOOM['link']})")

                # Secondary: open in Zoom app (mobile deep link)
                try:
                    st.link_button("📱 Open in Zoom App", zoom_deeplink, key="zoom_app_btn")
                except Exception:
                    st.markdown(f"[📱 Open in Zoom App]({zoom_deeplink})")

                st.write(f"**Meeting ID:** `{ZOOM['meeting_id']}`")
                st.write(f"**Passcode:** `{ZOOM['passcode']}`")

                # Copy helpers (mobile-friendly, safe escaping)
                _link_safe = ZOOM["link"].replace("'", "\\'")
                _id_safe   = ZOOM["meeting_id"].replace("'", "\\'")
                _pwd_safe  = ZOOM["passcode"].replace("'", "\\'")
                if components:
                    components.html(
                        f"""
                        <div style="display:flex;gap:8px;margin-top:8px;">
                          <button id="zCopyLink"
                                  style="padding:6px 10px;border-radius:8px;border:1px solid #cbd5e1;background:#f1f5f9;cursor:pointer;">
                            Copy Link
                          </button>
                          <button id="zCopyId"
                                  style="padding:6px 10px;border-radius:8px;border:1px solid #cbd5e1;background:#f1f5f9;cursor:pointer;">
                            Copy ID
                          </button>
                          <button id="zCopyPwd"
                                  style="padding:6px 10px;border-radius:8px;border:1px solid #cbd5e1;background:#f1f5f9;cursor:pointer;">
                            Copy Passcode
                          </button>
                        </div>
                        <script>
                          (function(){{
                            try {{
                              var link = '{_link_safe}', mid = '{_id_safe}', pwd = '{_pwd_safe}';
                              function wire(btnId, txt, label) {{
                                var b = document.getElementById(btnId);
                                if (!b) return;
                                b.addEventListener('click', function(){{
                                  navigator.clipboard.writeText(txt).then(function(){{
                                    b.innerText = '✓ Copied ' + label;
                                    setTimeout(function(){{ b.innerText = 'Copy ' + label; }}, 1500);
                                  }}).catch(function(){{}});
                                }});
                              }}
                              wire('zCopyLink', link, 'Link');
                              wire('zCopyId',   mid,  'ID');
                              wire('zCopyPwd',  pwd,  'Passcode');
                            }} catch(e) {{}}
                          }})();
                        </script>
                        """,
                        height=72,
                    )

            with z2:
                st.info(
                    f"You’re viewing: **{class_name}**  \n\n"
                    "✅ Use the **calendar below** to receive automatic class reminders.",
                    icon="📅",
                )

        st.divider()

       # ===================== CALENDAR TAB BANNER =====================
        with st.container():
            st.markdown(
                '''
                <div style="
                    padding: 12px;
                    background: #0ea5e9;
                    color: #ffffff;
                    border-radius: 8px;
                    text-align: center;
                    margin-bottom: 12px;
                    box-shadow: 0 2px 6px rgba(0,0,0,0.08);
                    font-weight: 600;
                ">
                    <span style="font-size:1.2rem;">📅 Calendar</span>
                    <div style="font-weight:500; font-size:0.98rem; margin-top:2px;">
                        You’re in the <u>Calendar</u> section — download the full course schedule or add reminders to your phone.
                    </div>
                </div>
                ''',
                unsafe_allow_html=True
            )
        st.divider()
#

        # ===================== CALENDAR QUICK ADD (no schedule/dictionary UI) =====================
        from datetime import datetime as _dt, timedelta as _td
        import re, uuid, json, io, requests
        import urllib.parse as _urllib

        # Try dateutil if available (for robust date parsing); fall back gracefully.
        try:
            from dateutil import parser as _dateparse
        except Exception:
            _dateparse = None

        def _load_group_schedules():
            # 1) global
            cfg = globals().get("GROUP_SCHEDULES")
            if isinstance(cfg, dict) and cfg:
                return cfg
            # 2) session_state
            cfg = st.session_state.get("GROUP_SCHEDULES")
            if isinstance(cfg, dict) and cfg:
                globals()["GROUP_SCHEDULES"] = cfg
                return cfg
            # 3) secrets
            try:
                raw = st.secrets.get("group_schedules", None)
                if raw:
                    cfg = json.loads(raw) if isinstance(raw, str) else raw
                    if isinstance(cfg, dict) and cfg:
                        st.session_state["GROUP_SCHEDULES"] = cfg
                        globals()["GROUP_SCHEDULES"] = cfg
                        return cfg
            except Exception:
                pass
            # 4) Firestore (optional)
            try:
                doc = db.collection("config").document("group_schedules").get()
                if doc and getattr(doc, "exists", False):
                    data = doc.to_dict() or {}
                    cfg = data.get("data", data)
                    if isinstance(cfg, dict) and cfg:
                        st.session_state["GROUP_SCHEDULES"] = cfg
                        globals()["GROUP_SCHEDULES"] = cfg
                        return cfg
            except Exception:
                pass
            # 5) BUILT-IN FALLBACK (kept private; we won't render it anywhere)
            return {
                "A1 Munich Klasse": {
                    "days": ["Monday", "Tuesday", "Wednesday"],
                    "time": "6:00pm–7:00pm",
                    "start_date": "2025-07-08",
                    "end_date": "2025-09-02",
                    "doc_url": "https://drive.google.com/file/d/1en_YG8up4C4r36v4r7E714ARcZyvNFD6/view?usp=sharing"
                },
                "A1 Berlin Klasse": {
                    "days": ["Thursday", "Friday", "Saturday"],
                    "time": "Thu/Fri: 6:00pm–7:00pm, Sat: 8:00am–9:00am",
                    "start_date": "2025-06-14",
                    "end_date": "2025-08-09",
                    "doc_url": "https://drive.google.com/file/d/1foK6MPoT_dc2sCxEhTJbtuK5ZzP-ERzt/view?usp=sharing"
                },
                "A1 Koln Klasse": {
                    "days": ["Thursday", "Friday", "Saturday"],
                    "time": "Thu/Fri: 6:00pm–7:00pm, Sat: 8:00am–9:00am",
                    "start_date": "2025-08-15",
                    "end_date": "2025-10-11",
                    "doc_url": "https://drive.google.com/file/d/1d1Ord557jGRn5NxYsmCJVmwUn1HtrqI3/view?usp=sharing"
                },
                "A2 Munich Klasse": {
                    "days": ["Monday", "Tuesday", "Wednesday"],
                    "time": "7:30pm–9:00pm",
                    "start_date": "2025-06-24",
                    "end_date": "2025-08-26",
                    "doc_url": "https://drive.google.com/file/d/1Zr3iN6hkAnuoEBvRELuSDlT7kHY8s2LP/view?usp=sharing"
                },
                "A2 Berlin Klasse": {
                    "days": ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"],
                    "time": "Mon–Wed: 11:00am–12:00pm, Thu/Fri: 11:00am–12:00pm, Wed: 2:00pm–3:00pm",
                    "start_date": "",
                    "end_date": "",
                    "doc_url": ""
                },
                "A2 Koln Klasse": {
                    "days": ["Wednesday", "Thursday", "Friday"],
                    "time": "11:00am–12:00pm",
                    "start_date": "2025-08-06",
                    "end_date": "2025-10-08",
                    "doc_url": "https://drive.google.com/file/d/19cptfdlmBDYe9o84b8ZCwujmxuMCKXAD/view?usp=sharing"
                },
                "B1 Munich Klasse": {
                    "days": ["Thursday", "Friday"],
                    "time": "7:30pm–9:00pm",
                    "start_date": "2025-08-07",
                    "end_date": "2025-11-07",
                    "doc_url": "https://drive.google.com/file/d/1CaLw9RO6H8JOr5HmwWOZA2O7T-bVByi7/view?usp=sharing"
                },
                "B2 Munich Klasse": {
                    "days": ["Friday", "Saturday"],
                    "time": "Fri: 2pm-3:30pm, Sat: 9:30am-10am",
                    "start_date": "2025-08-08",
                    "end_date": "2025-10-08",
                    "doc_url": "https://drive.google.com/file/d/1gn6vYBbRyHSvKgqvpj5rr8OfUOYRL09W/view?usp=sharing"
                },
            }

        # ---------- helpers to fetch & parse dates from schedule PDF (Drive) ----------
        def _gdrive_direct_download(url: str) -> bytes | None:
            if not url:
                return None
            m = re.search(r"/file/d/([A-Za-z0-9_-]{20,})/", url) or re.search(r"[?&]id=([A-Za-z0-9_-]{20,})", url)
            file_id = m.group(1) if m else None
            if not file_id:
                return None
            dl = f"https://drive.google.com/uc?export=download&id={file_id}"
            try:
                r = requests.get(dl, timeout=15)
                if r.status_code == 200 and r.content:
                    # If Google shows a confirmation page for large files, bail out (keep simple)
                    if b"uc-download-link" in r.content[:4000] and b"confirm" in r.content[:4000]:
                        return None
                    return r.content
            except Exception:
                pass
            return None

        def _extract_text_from_pdf(pdf_bytes: bytes) -> str:
            # Try pypdf first
            try:
                from pypdf import PdfReader
                t = []
                reader = PdfReader(io.BytesIO(pdf_bytes))
                for p in reader.pages:
                    try:
                        t.append(p.extract_text() or "")
                    except Exception:
                        t.append("")
                return "\n".join(t)
            except Exception:
                pass
            # Fallback: pdfminer (if available)
            try:
                from pdfminer.high_level import extract_text
                return extract_text(io.BytesIO(pdf_bytes)) or ""
            except Exception:
                return ""

        _DATE_PATTERNS = [
            r"\b(20\d{2}-\d{2}-\d{2})\b",  # 2025-08-15
            r"\b(\d{1,2}/\d{1,2}/20\d{2})\b",  # 08/15/2025 or 15/08/2025
            r"\b(\d{1,2}\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+20\d{2})\b",  # 15 Aug 2025
            r"\b((Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{1,2},\s*20\d{2})\b",  # Aug 15, 2025
        ]

        def _parse_any_date(raw: str):
            # Prefer dateutil if present
            if _dateparse:
                for dayfirst in (False, True):
                    try:
                        return _dateparse.parse(raw, dayfirst=dayfirst, fuzzy=True).date()
                    except Exception:
                        pass
            # Lightweight manual attempts
            for fmt in ("%Y-%m-%d", "%d %b %Y", "%b %d, %Y", "%m/%d/%Y", "%d/%m/%Y"):
                try:
                    return _dt.strptime(raw, fmt).date()
                except Exception:
                    pass
            return None

        def _find_dates_in_text(txt: str):
            found = []
            if not txt:
                return found
            for pat in _DATE_PATTERNS:
                for m in re.finditer(pat, txt, flags=re.IGNORECASE):
                    d = _parse_any_date(m.group(1))
                    if d:
                        found.append(d)
            # de-dup + sort
            uniq = []
            seen = set()
            for d in sorted(found):
                if d not in seen:
                    seen.add(d)
                    uniq.append(d)
            return uniq

        def infer_start_end_from_doc(doc_url: str):
            pdf_bytes = _gdrive_direct_download(doc_url)
            if not pdf_bytes:
                return None, None
            text = _extract_text_from_pdf(pdf_bytes)
            dates = _find_dates_in_text(text)
            if len(dates) >= 2:
                return dates[0], dates[-1]
            if len(dates) == 1:
                return dates[0], None
            return None, None

        GROUP_SCHEDULES = _load_group_schedules()

        # Pull class config quietly
        class_cfg   = GROUP_SCHEDULES.get(class_name, {})
        days        = class_cfg.get("days", [])
        time_str    = class_cfg.get("time", "")
        start_str   = class_cfg.get("start_date", "")
        end_str     = class_cfg.get("end_date", "")
        doc_url     = class_cfg.get("doc_url", "")

        # Parse dates
        start_date_obj = None
        end_date_obj   = None
        try:
            if start_str:
                start_date_obj = _dt.strptime(start_str, "%Y-%m-%d").date()
        except Exception:
            pass
        try:
            if end_str:
                end_date_obj = _dt.strptime(end_str, "%Y-%m-%d").date()
        except Exception:
            pass

        # If missing, try to infer from the schedule PDF
        _inferred_start = _inferred_end = False
        if (not start_date_obj or not end_date_obj) and doc_url:
            s, e = infer_start_end_from_doc(doc_url)
            if s and not start_date_obj:
                start_date_obj = s
                _inferred_start = True
            if e and not end_date_obj:
                end_date_obj = e
                _inferred_end = True

        if not (start_date_obj and end_date_obj and isinstance(time_str, str) and time_str.strip() and days):
            st.warning("This class doesn’t have a full calendar setup yet. Please contact the office.", icon="⚠️")
        else:
            # Tell students clearly the course period (and note if inferred)
            _note_bits = []
            if _inferred_start or _inferred_end:
                _note_bits.append("dates inferred from the schedule document")
            _note = f" ({', '.join(_note_bits)})" if _note_bits else ""
            st.info(
                f"**Course period:** {start_date_obj.strftime('%d %b %Y')} → {end_date_obj.strftime('%d %b %Y')}{_note}",
                icon="📅",
            )

            # ---------- helpers ----------
            _WKD_ORDER = ["MO","TU","WE","TH","FR","SA","SU"]
            _FULL_TO_CODE = {
                "monday":"MO","tuesday":"TU","wednesday":"WE","thursday":"TH","friday":"FR","saturday":"SA","sunday":"SU",
                "mon":"MO","tue":"TU","tues":"TU","wed":"WE","thu":"TH","thur":"TH","thurs":"TH","fri":"FR","sat":"SA","sun":"SU"
            }

            def _to_24h(h, m, ampm):
                h = int(h); m = int(m); ap = ampm.lower()
                if ap == "pm" and h != 12: h += 12
                if ap == "am" and h == 12: h = 0
                return h, m

            def _parse_time_component(s):
                s = s.strip().lower()
                m = re.match(r"^(\d{1,2})(?::(\d{2}))?\s*(am|pm)$", s)
                if not m: return None
                h = m.group(1); mm = m.group(2) or "0"; ap = m.group(3)
                return _to_24h(h, mm, ap)

            def _parse_time_range(rng):
                rng = rng.strip().lower().replace("–","-").replace("—","-")
                parts = [p.strip() for p in rng.split("-")]
                if len(parts) != 2: return None
                a = _parse_time_component(parts[0]); b = _parse_time_component(parts[1])
                if not a or not b: return None
                return a, b

            def _expand_day_token(tok):
                tok = tok.strip().lower().replace("–","-").replace("—","-")
                if "-" in tok:  # mon–wed
                    a, b = [t.strip() for t in tok.split("-", 1)]
                    a_code = _FULL_TO_CODE.get(a, ""); b_code = _FULL_TO_CODE.get(b, "")
                    if a_code and b_code:
                        ai = _WKD_ORDER.index(a_code); bi = _WKD_ORDER.index(b_code)
                        return _WKD_ORDER[ai:bi+1] if ai <= bi else _WKD_ORDER[ai:] + _WKD_ORDER[:bi+1]
                    return []
                c = _FULL_TO_CODE.get(tok, "")
                return [c] if c else []

            def _parse_time_blocks(time_str, days_list):
                if not (isinstance(time_str, str) and time_str.strip()):
                    return []
                s = time_str.strip()
                if ":" in s:  # grouped "Days: time"
                    blocks = []
                    groups = [g.strip() for g in s.split(",") if g.strip()]
                    for g in groups:
                        if ":" not in g:
                            continue
                        left, right = [x.strip() for x in g.split(":", 1)]
                        day_tokens = re.split(r"/", left)
                        codes = []
                        for tok in day_tokens:
                            codes.extend(_expand_day_token(tok))
                        tr = _parse_time_range(right)
                        if codes and tr:
                            (sh, sm), (eh, em) = tr
                            blocks.append({"byday": sorted(set(codes), key=_WKD_ORDER.index),
                                           "start": (sh, sm), "end": (eh, em)})
                    return blocks
                # single time for given days[]
                tr = _parse_time_range(s)
                if not tr: return []
                (sh, sm), (eh, em) = tr
                codes = []
                for d in (days_list or []):
                    c = _FULL_TO_CODE.get(str(d).lower().strip(), "")
                    if c: codes.append(c)
                codes = sorted(set(codes), key=_WKD_ORDER.index) or _WKD_ORDER[:]
                return [{"byday": codes, "start": (sh, sm), "end": (eh, em)}]

            def _next_on_or_after(d, weekday_index):  # Mon=0..Sun=6
                delta = (weekday_index - d.weekday()) % 7
                return d + _td(days=delta)

            # Build ICS (with 15-minute preset reminder + URL field)
            _blocks = _parse_time_blocks(time_str, days)
            _zl = (ZOOM or {}).get("link", ""); _zid = (ZOOM or {}).get("meeting_id", ""); _zpw = (ZOOM or {}).get("passcode", "")
            _details = f"Zoom link: {_zl}\\nMeeting ID: {_zid}\\nPasscode: {_zpw}"
            _dtstamp = _dt.utcnow().strftime("%Y%m%dT%H%M%SZ")
            _until = _dt(end_date_obj.year, end_date_obj.month, end_date_obj.day, 23, 59, 59).strftime("%Y%m%dT%H%M%SZ")
            _summary = f"{class_name} — Live German Class"

            _ics_lines = [
                "BEGIN:VCALENDAR","VERSION:2.0","PRODID:-//Falowen//Course Scheduler//EN",
                "CALSCALE:GREGORIAN","METHOD:PUBLISH",
            ]

            if not _blocks:
                _start_dt = _dt(start_date_obj.year, start_date_obj.month, start_date_obj.day, 18, 0)
                _end_dt   = _dt(start_date_obj.year, start_date_obj.month, start_date_obj.day, 19, 0)
                _ics_lines += [
                    "BEGIN:VEVENT",
                    f"UID:{uuid.uuid4()}@falowen",
                    f"DTSTAMP:{_dtstamp}",
                    f"DTSTART:{_start_dt.strftime('%Y%m%dT%H%M%SZ')}",
                    f"DTEND:{_end_dt.strftime('%Y%m%dT%H%M%SZ')}",
                    f"SUMMARY:{_summary}",
                    f"DESCRIPTION:{_details}",
                    f"URL:{_zl}",
                    "LOCATION:Zoom",
                    # preset alert 15 minutes before
                    "BEGIN:VALARM",
                    "ACTION:DISPLAY",
                    "DESCRIPTION:Class starts soon",
                    "TRIGGER:-PT15M",
                    "END:VALARM",
                    "END:VEVENT",
                ]
            else:
                for blk in _blocks:
                    byday_codes = blk["byday"]
                    sh, sm = blk["start"]; eh, em = blk["end"]
                    _wmap = {"MO":0,"TU":1,"WE":2,"TH":3,"FR":4,"SA":5,"SU":6}
                    first_dates = []
                    for code in byday_codes:
                        widx = _wmap[code]
                        first_dates.append(_next_on_or_after(start_date_obj, widx))
                    first_date = min(first_dates)
                    dt_start = _dt(first_date.year, first_date.month, first_date.day, sh, sm)
                    dt_end   = _dt(first_date.year, first_date.month, first_date.day, eh, em)
                    _ics_lines += [
                        "BEGIN:VEVENT",
                        f"UID:{uuid.uuid4()}@falowen",
                        f"DTSTAMP:{_dtstamp}",
                        f"DTSTART:{dt_start.strftime('%Y%m%dT%H%M%SZ')}",
                        f"DTEND:{dt_end.strftime('%Y%m%dT%H%M%SZ')}",
                        f"RRULE:FREQ=WEEKLY;BYDAY={','.join(byday_codes)};UNTIL={_until}",
                        f"SUMMARY:{_summary}",
                        f"DESCRIPTION:{_details}",
                        f"URL:{_zl}",
                        "LOCATION:Zoom",
                        # preset alert 15 minutes before
                        "BEGIN:VALARM",
                        "ACTION:DISPLAY",
                        "DESCRIPTION:Class starts soon",
                        "TRIGGER:-PT15M",
                        "END:VALARM",
                        "END:VEVENT",
                    ]

            _ics_lines.append("END:VCALENDAR")
            _course_ics = "\n".join(_ics_lines)

            # UI (full course download only; next-session button removed)
            c1, c2 = st.columns([1, 1])
            with c1:
                st.download_button(
                    "⬇️ Download full course (.ics)",
                    data=_course_ics,
                    file_name=f"{class_name.replace(' ', '_')}_course.ics",
                    mime="text/calendar",
                    key="dl_course_ics",
                )
            with c2:
                st.caption("Calendar created. Use the download button to import the full course.")
#


            # --- Phone app quick links (Android) — concise only ---
            # Build per-block Google Calendar repeating links from the schedule
            _gcal_repeat_links = []
            try:
                if _blocks:
                    _wmap = {"MO":0,"TU":1,"WE":2,"TH":3,"FR":4,"SA":5,"SU":6}
                    _code_to_pretty = {"MO":"Mon","TU":"Tue","WE":"Wed","TH":"Thu","FR":"Fri","SA":"Sat","SU":"Sun"}

                    def _fmt_time(h, m):
                        ap = "AM" if h < 12 else "PM"
                        hh = h if 1 <= h <= 12 else (12 if h % 12 == 0 else h % 12)
                        return f"{hh}:{m:02d}{ap}"

                    for blk in _blocks:
                        byday_codes = blk["byday"]
                        sh, sm = blk["start"]; eh, em = blk["end"]

                        # First occurrence on/after course start for this block
                        first_dates = []
                        for code in byday_codes:
                            widx = _wmap[code]
                            first_dates.append(_next_on_or_after(start_date_obj, widx))
                        first_date = min(first_dates)

                        _start_dt = _dt(first_date.year, first_date.month, first_date.day, sh, sm)
                        _end_dt   = _dt(first_date.year, first_date.month, first_date.day, eh, em)
                        _start_str = _start_dt.strftime("%Y%m%dT%H%M%SZ")
                        _end_str   = _end_dt.strftime("%Y%m%dT%H%M%SZ")

                        # RRULE weekly until course end
                        _rrule = f"RRULE:FREQ=WEEKLY;BYDAY={','.join(byday_codes)};UNTIL={_until}"

                        # Friendly label e.g. "Thu/Fri 6:00PM–7:00PM" or "Sat 8:00AM–9:00AM"
                        _days_pretty = "/".join(_code_to_pretty[c] for c in byday_codes)
                        _label = f"{_days_pretty} {_fmt_time(sh, sm)}–{_fmt_time(eh, em)}"

                        _recur_url = (
                            "https://calendar.google.com/calendar/render"
                            f"?action=TEMPLATE"
                            f"&text={_urllib.quote(_summary)}"
                            f"&dates={_start_str}/{_end_str}"
                            f"&details={_urllib.quote(_details)}"
                            f"&location={_urllib.quote('Zoom')}"
                            f"&ctz={_urllib.quote('Africa/Accra')}"
                            f"&recur={_urllib.quote(_rrule)}"
                            f"&sf=true"
                        )
                        _gcal_repeat_links.append((_label, _recur_url))
            except Exception:
                _gcal_repeat_links = []

            # Render ultra-compact Android help with per-block links
            if _gcal_repeat_links:
                _items = "".join(
                    f"<li style='margin:4px 0;'><a href='{url.replace('&','&amp;')}' target='_blank'>Tap here: {lbl}</a></li>"
                    for (lbl, url) in _gcal_repeat_links
                )
                _phone_links_ul = f"<ul style='margin:6px 0 0 18px;padding:0;'>{_items}</ul>"
            else:
                _phone_links_ul = (
                    "<div style='margin:6px 0 0 2px;color:#444;'>"
                    "No repeating blocks are set yet. Ask the office to add your class times."
                    "</div>"
                )


            st.markdown(
                f"""
                **Computer or iPhone:** Download the **.ics** above and install.  
                - **Computer (Google Calendar web):** Go to [calendar.google.com](https://calendar.google.com) → **Settings** → **Import & export** → **Import** (you’ll see **“Imported X of X events.”**).
                - **iPhone (Apple Calendar):** Download the `.ics`, open it, choose your notification preference, then **Done**.

                **Android (Google Calendar app):** The app **can’t import `.ics`**. So use the links below to add it on your phone (**with repeat**):
                {_phone_links_ul}
                <div style="margin:8px 0 0 2px;">
                </div>
                """,
                unsafe_allow_html=True,
            )

        # ===================== CLASS ROSTER =====================

        # Subtle banner above the expander to draw attention
        st.markdown(
            """
            <div style="
                padding:10px 12px;
                background:#f0f9ff;
                border:1px solid #bae6fd;
                border-radius:12px;
                margin: 6px 0 8px 0;
                display:flex;align-items:center;gap:8px;">
              <span style="font-size:1.05rem;">👥 <b>Class Members</b></span>
              <span style="font-size:.92rem;color:#055d87;">Tap below to open and view the list</span>
            </div>
            """,
            unsafe_allow_html=True,
        )

        # Light CSS to make *all* expanders stand out a bit more
        st.markdown(
            """
            <style>
              /* Make expander headers pop a little */
              div[data-testid="stExpander"] > details > summary {
                  background:#f0f9ff !important;
                  border:1px solid #bae6fd !important;
                  border-radius:12px !important;
                  padding:10px 12px !important;
              }
              div[data-testid="stExpander"] > details[open] > summary {
                  background:#e0f2fe !important;
                  border-color:#7dd3fc !important;
              }
            </style>
            """,
            unsafe_allow_html=True,
        )

        with st.expander("👥 Class Members", expanded=False):
            try:
                df_students = load_student_data()

                # Normalize required columns
                for col in ("ClassName", "Name", "Email", "Location"):
                    if col not in df_students.columns:
                        df_students[col] = ""
                    df_students[col] = df_students[col].fillna("").astype(str).str.strip()

                # Filter to this class
                same_class = df_students[df_students["ClassName"] == class_name].copy()

                # Tiny header line inside with class + count
                _n = len(same_class)
                st.markdown(
                    f"""
                    <div style="display:flex;justify-content:space-between;align-items:center;margin:4px 0 6px 0;">
                      <div style="font-weight:600;color:#0f172a;">{class_name}</div>
                      <span style="background:#0ea5e922;border:1px solid #0ea5e9;color:#0369a1;
                                   padding:3px 8px;border-radius:999px;font-size:.9rem;">
                        {_n} member{'' if _n==1 else 's'}
                      </span>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

                # Columns to display (no StudentCode)
                cols_show = [c for c in ["Name", "Email", "Location"] if c in same_class.columns]

                if not same_class.empty and cols_show:
                    st.dataframe(
                        same_class[cols_show].reset_index(drop=True),
                        use_container_width=True,
                        hide_index=True,
                    )
                else:
                    st.info("No members found for this class yet.")
            except Exception as e:
                st.warning(f"Couldn’t load the class roster right now. {e}")
#


          # ===================== ANNOUNCEMENTS (CSV) + REPLIES (FIRESTORE) =====================

        # Prefer cached helper if exists; else fallback to direct CSV
        try:
            df = fetch_announcements_csv()
        except Exception:
            df = pd.DataFrame()
        if df.empty:
            CSV_URL = "https://docs.google.com/spreadsheets/d/16gjj0krncWsDwMfMbhlxODPSJsI50fuHAzkF7Prrs1k/export?format=csv&gid=0"
            try:
                df = pd.read_csv(CSV_URL)
            except Exception:
                df = pd.DataFrame()

        # Helpers (links, parsing, ids)
        URL_RE = re.compile(r"(https?://[^\s]+)")

        # ---------- Announcement banner (with NEW count) ----------
        _new_badge_html = ""
        try:
            from datetime import datetime as _dt
            _today = _dt.today().date()
            _recent = 0
            if not df.empty and "Date" in df.columns:
                # Try dateutil if available from earlier; fall back to common formats
                def _parse_date_any(s: str):
                    s = str(s).strip()
                    if not s:
                        return None
                    if 'dateutil' in globals() and _dateparse:
                        try:
                            return _dateparse.parse(s).date()
                        except Exception:
                            pass
                    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y", "%d-%m-%Y"):
                        try:
                            return _dt.strptime(s, fmt).date()
                        except Exception:
                            continue
                    return None

                for v in df["Date"].astype(str).tolist():
                    d = _parse_date_any(v)
                    if d and (_today - d).days <= 7:
                        _recent += 1

            if _recent > 0:
                _new_badge_html = f"<span style='margin-left:8px;background:#16a34a;color:#fff;padding:2px 8px;border-radius:999px;font-size:0.8rem;'>NEW · {_recent}</span>"
        except Exception:
            pass

        with st.container():
            st.markdown(
                f'''
                <div style="
                    padding:12px;
                    background: linear-gradient(90deg,#0ea5e9,#22c55e);
                    color:#ffffff;
                    border-radius:8px;
                    margin-bottom:12px;
                    box-shadow:0 2px 6px rgba(0,0,0,0.08);
                    display:flex;align-items:center;justify-content:space-between;">
                    <div style="font-weight:700;font-size:1.15rem;">📢 Announcements {_new_badge_html}</div>
                    <div style="font-size:0.92rem;opacity:.9;">Latest class updates, deadlines & links</div>
                </div>
                ''',
                unsafe_allow_html=True
            )
#


        def _short_label_from_url(u: str) -> str:
            try:
                p = urllib.parse.urlparse(u)
                host = (p.netloc or "").replace("www.", "")
                path = (p.path or "").strip("/")
                label = host if not path else f"{host}/{path}"
                return label[:60] + ("…" if len(label) > 60 else "")
            except Exception:
                return u[:60] + ("…" if len(u) > 60 else "")

        def _guess_link_emoji_and_label(u: str):
            lu = u.lower()
            if "zoom.us" in lu: return "🎦", None
            if "youtu" in lu:   return "▶️", None
            if lu.endswith(".pdf"): return "📄", None
            if "drive.google" in lu: return "🟢", None
            if "deepl.com" in lu: return "🌐", None
            if "google.com" in lu: return "🔗", None
            return "🔗", None

        # Normalize CSV into canonical columns
        if not df.empty:
            df.columns = [str(c).strip() for c in df.columns]
            lower_map = {c.lower(): c for c in df.columns}

            def _col(name: str):
                return lower_map.get(name.lower())

            for logical in ("announcement", "class", "date", "pinned"):
                if _col(logical) is None:
                    df[logical] = ""

            rename_map = {}
            if _col("announcement"): rename_map[_col("announcement")] = "Announcement"
            if _col("class"):        rename_map[_col("class")]        = "Class"
            if _col("date"):         rename_map[_col("date")]         = "Date"
            if _col("pinned"):       rename_map[_col("pinned")]       = "Pinned"
            df = df.rename(columns=rename_map)

            for c in ("Announcement", "Class", "Date", "Pinned"):
                if c not in df.columns:
                    df[c] = ""

            # Optional Link/Links column
            link_key = lower_map.get("link") or lower_map.get("links")
            df["Links"] = [[] for _ in range(len(df))]
            if link_key:
                def _split_links(val):
                    s = str(val or "").strip()
                    if not s:
                        return []
                    parts = [p for chunk in s.split(",") for p in chunk.split()]
                    return [p.strip() for p in parts if p.strip().lower().startswith(("http://", "https://"))]
                df["Links"] = df[link_key].apply(_split_links)

            # Normalize pinned
            def _norm_pinned(v) -> bool:
                s = str(v).strip().lower()
                return s in {"true", "yes", "1"}
            df["Pinned"] = df["Pinned"].apply(_norm_pinned)

            # Parse dates
            def _parse_dt(x):
                for fmt in ("%Y-%m-%d %H:%M", "%Y/%m/%d %H:%M", "%d/%m/%Y %H:%M", "%Y-%m-%d", "%d/%m/%Y"):
                    try:
                        return datetime.strptime(str(x), fmt)
                    except Exception:
                        continue
                try:
                    return pd.to_datetime(x, errors="coerce")
                except Exception:
                    return pd.NaT
            df["__dt"] = df["Date"].apply(_parse_dt)

            # Append auto-detected links
            def _append_detected_links(row):
                txt = str(row.get("Announcement", "") or "")
                found = URL_RE.findall(txt)
                existing = list(row.get("Links", []) or [])
                merged, seen = [], set()
                for url in existing + found:
                    if url not in seen:
                        merged.append(url); seen.add(url)
                return merged
            df["Links"] = df.apply(_append_detected_links, axis=1)

            # Stable ID
            def _ann_id(row):
                try:
                    raw = f"{row.get('Class','')}|{row.get('Date','')}|{row.get('Announcement','')}".encode("utf-8")
                    return hashlib.sha1(raw).hexdigest()[:16]
                except Exception:
                    return str(uuid4()).replace("-", "")[:16]
            df["__id"] = df.apply(_ann_id, axis=1)

        # Firestore reply helpers (with IDs for edit/delete)
        def _ann_reply_coll(ann_id: str):
            return (db.collection("class_announcements")
                     .document(class_name)
                     .collection("replies")
                     .document(ann_id)
                     .collection("posts"))

        def _load_replies_with_ids(ann_id: str):
            try:
                docs = list(_ann_reply_coll(ann_id).order_by("timestamp").stream())
            except Exception:
                docs = list(_ann_reply_coll(ann_id).stream())
                docs.sort(key=lambda d: (d.to_dict() or {}).get("timestamp"))
            out = []
            for d in docs:
                x = d.to_dict() or {}
                x["__id"] = d.id
                out.append(x)
            return out

        def _update_reply_text(ann_id: str, reply_id: str, new_text: str):
            _ann_reply_coll(ann_id).document(reply_id).update({
                "text": new_text.strip(),
                "edited_at": datetime.utcnow(),
                "edited_by": student_name,
                "edited_by_code": student_code,
            })

        def _delete_reply(ann_id: str, reply_id: str):
            _ann_reply_coll(ann_id).document(reply_id).delete()

        # Controls + render
        if df.empty:
            st.info("No announcements yet.")
        else:
            c1, c2, c3 = st.columns([1, 2, 1])
            with c1:
                show_only_pinned = st.checkbox("Show only pinned", value=False, key="ann_only_pinned")
            with c2:
                search_term = st.text_input("Search announcements…", "", key="ann_search")
            with c3:
                if st.button("↻ Refresh", key="ann_refresh"):
                    try:
                        st.cache_data.clear()
                    except Exception:
                        pass
                    st.rerun()

            # Filter for this class
            df["__class_norm"] = (
                df["Class"].astype(str)
                .str.replace(r"\s+", " ", regex=True)
                .str.strip()
                .str.lower()
            )
            class_norm = re.sub(r"\s+", " ", class_name.strip().lower())
            view = df[df["__class_norm"] == class_norm].copy()

            if show_only_pinned:
                view = view[view["Pinned"] == True]
            if search_term.strip():
                q = search_term.lower()
                view = view[view["Announcement"].astype(str).str.lower().str.contains(q)]

            view.sort_values("__dt", ascending=False, inplace=True, na_position="last")
            pinned_df = view[view["Pinned"] == True]
            latest_df = view[view["Pinned"] == False]

            def render_announcement(row, is_pinned=False):
                # teacher card
                try:
                    ts_label = row.get("__dt").strftime("%d %b %H:%M")
                except Exception:
                    ts_label = ""
                st.markdown(
                    f"<div style='padding:10px 12px; background:{'#fff7ed' if is_pinned else '#f8fafc'}; "
                    f"border:1px solid #e5e7eb; border-radius:8px; margin:8px 0;'>"
                    f"{'📌 <b>Pinned</b> • ' if is_pinned else ''}"
                    f"<b>Teacher</b> <span style='color:#888;'>{ts_label} GMT</span><br>"
                    f"{row.get('Announcement','')}"
                    f"</div>",
                    unsafe_allow_html=True,
                )

                # links
                links = row.get("Links") or []
                if isinstance(links, str):
                    links = [links] if links.strip() else []
                if links:
                    st.markdown("**🔗 Links:**")
                    for u in links:
                        emoji, label = _guess_link_emoji_and_label(u)
                        label = label or _short_label_from_url(u)
                        st.markdown(f"- {emoji} [{label}]({u})")

                # replies
                ann_id = row.get("__id")
                replies = _load_replies_with_ids(ann_id)
                if replies:
                    for r in replies:
                        ts = r.get("timestamp")
                        when = ""
                        try:
                            when = ts.strftime("%d %b %H:%M") + " UTC"
                        except Exception:
                            pass
                        edited_badge = ""
                        if r.get("edited_at"):
                            try:
                                edited_badge = f" <span style='color:#aaa;'>(edited {r['edited_at'].strftime('%d %b %H:%M')} UTC)</span>"
                            except Exception:
                                edited_badge = " <span style='color:#aaa;'>(edited)</span>"

                        st.markdown(
                            f"<div style='margin-left:20px; color:#444;'>↳ <b>{r.get('student_name','')}</b> "
                            f"<span style='color:#bbb;'>{when}</span>{edited_badge}<br>"
                            f"{r.get('text','')}</div>",
                            unsafe_allow_html=True,
                        )

                        # edit/delete (own or admin)
                        can_edit = IS_ADMIN or (r.get("student_code") == student_code)
                        if can_edit:
                            c_ed, c_del = st.columns([1, 1])
                            with c_ed:
                                if st.button("✏️ Edit", key=f"ann_edit_reply_{ann_id}_{r['__id']}"):
                                    st.session_state[f"edit_mode_{ann_id}_{r['__id']}"] = True
                                    st.session_state[f"edit_text_{ann_id}_{r['__id']}"] = r.get("text", "")
                                    st.rerun()
                            with c_del:
                                if st.button("🗑️ Delete", key=f"ann_del_reply_{ann_id}_{r['__id']}"):
                                    _delete_reply(ann_id, r["__id"])
                                    _notify_slack(
                                        f"🗑️ *Announcement reply deleted* — {class_name}\n"
                                        f"*By:* {student_name} ({student_code})\n"
                                        f"*When:* {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC"
                                    )
                                    st.success("Reply deleted.")
                                    st.rerun()

                            # inline editor
                            if st.session_state.get(f"edit_mode_{ann_id}_{r['__id']}", False):
                                new_txt = st.text_area(
                                    "Edit reply",
                                    key=f"ann_editbox_{ann_id}_{r['__id']}",
                                    value=st.session_state.get(f"edit_text_{ann_id}_{r['__id']}", r.get("text", "")),
                                    height=100,
                                )
                                ec1, ec2 = st.columns([1, 1])
                                with ec1:
                                    if st.button("💾 Save", key=f"ann_save_reply_{ann_id}_{r['__id']}"):
                                        if new_txt.strip():
                                            _update_reply_text(ann_id, r["__id"], new_txt)
                                            _notify_slack(
                                                f"✏️ *Announcement reply edited* — {class_name}\n"
                                                f"*By:* {student_name} ({student_code})\n"
                                                f"*When:* {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC\n"
                                                f"*Preview:* {new_txt[:180]}{'…' if len(new_txt)>180 else ''}"
                                            )
                                            st.success("Reply updated.")
                                        st.session_state.pop(f"edit_mode_{ann_id}_{r['__id']}", None)
                                        st.session_state.pop(f"edit_text_{ann_id}_{r['__id']}", None)
                                        st.rerun()
                                with ec2:
                                    if st.button("❌ Cancel", key=f"ann_cancel_reply_{ann_id}_{r['__id']}"):
                                        st.session_state.pop(f"edit_mode_{ann_id}_{r['__id']}", None)
                                        st.session_state.pop(f"edit_text_{ann_id}_{r['__id']}", None)
                                        st.rerun()

                # new reply (single click -> rerun)
                with st.expander(f"Reply ({ann_id[:6]})", expanded=False):
                    ta_key = f"ann_reply_box_{ann_id}"
                    flag_key = f"__clear_{ta_key}"
                    if st.session_state.get(flag_key):
                        st.session_state.pop(flag_key, None)
                        st.session_state[flag_key] = True
                    reply_text = st.text_area(
                        f"Reply to {ann_id}",
                        key=ta_key,
                        height=90,
                        placeholder="Write your reply…"
                    )
                    if st.button("Send Reply", key=f"ann_send_reply_{ann_id}") and reply_text.strip():
                        payload = {
                            "student_code": student_code,
                            "student_name": student_name,
                            "text": reply_text.strip(),
                            "timestamp": datetime.utcnow(),
                        }
                        _ann_reply_coll(ann_id).add(payload)
                        _notify_slack(
                            f"💬 *New announcement reply* — {class_name}\n"
                            f"*By:* {student_name} ({student_code})\n"
                            f"*When:* {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC\n"
                            f"*Preview:* {payload['text'][:180]}{'…' if len(payload['text'])>180 else ''}"
                        )
                        st.session_state[flag_key] = True
                        st.success("Reply sent!")
                        st.rerun()

            # render all
            for _, row in pinned_df.iterrows():
                render_announcement(row, is_pinned=True)
            for _, row in latest_df.iterrows():
                render_announcement(row, is_pinned=False)

        st.divider()

        # ===================== CLASS Q&A (POST / REPLY + EDIT/DELETE) =====================

        # Firestore collection handle
        q_base = db.collection("class_qna").document(class_name).collection("questions")

        # --- Compute NEW (≤7 days) and UNANSWERED counts for badges ---
        _new7, _unans, _total = 0, 0, 0
        try:
            from datetime import datetime as _dt
            _now = _dt.utcnow()

            # Try ordered fetch first; fall back to basic stream
            try:
                _qdocs = list(q_base.order_by("created_at", direction="DESCENDING").limit(250).stream())
            except Exception:
                _qdocs = list(q_base.stream())

            def _to_datetime_any(v):
                if v is None:
                    return None
                # Firestore Timestamp object?
                try:
                    if hasattr(v, "to_datetime"):
                        return v.to_datetime()
                except Exception:
                    pass
                # Seconds/nanos style?
                try:
                    if hasattr(v, "seconds"):
                        return _dt.utcfromtimestamp(int(v.seconds))
                except Exception:
                    pass
                # String parse
                try:
                    if 'dateutil' in globals() and _dateparse:
                        return _dateparse.parse(str(v))
                except Exception:
                    pass
                for fmt in ("%Y-%m-%d", "%Y-%m-%d %H:%M:%S", "%d/%m/%Y", "%m/%d/%Y", "%d-%m-%Y"):
                    try:
                        return _dt.strptime(str(v), fmt)
                    except Exception:
                        continue
                return None

            for _doc in _qdocs:
                _d = (_doc.to_dict() or {})
                _total += 1

                # Replies count (supports several schema styles)
                _rc = 0
                if isinstance(_d.get("answers"), list):
                    _rc = len(_d["answers"])
                elif isinstance(_d.get("replies"), list):
                    _rc = len(_d["replies"])
                elif isinstance(_d.get("reply_count"), int):
                    _rc = int(_d["reply_count"])
                if _rc == 0:
                    _unans += 1

                # New in last 7 days?
                _created = _to_datetime_any(_d.get("created_at") or _d.get("ts") or _d.get("timestamp"))
                if _created and (_now - _created).days <= 7:
                    _new7 += 1
        except Exception:
            pass

        # --- Render banner with badges ---
        _badges = []
        if _new7 > 0:
            _badges.append(
                f"<span style='margin-left:8px;background:#16a34a;color:#fff;padding:2px 8px;"
                f"border-radius:999px;font-size:0.8rem;'>NEW · {_new7}</span>"
            )
        if _unans > 0:
            _badges.append(
                f"<span style='margin-left:8px;background:#f97316;color:#fff;padding:2px 8px;"
                f"border-radius:999px;font-size:0.8rem;'>UNANSWERED · {_unans}</span>"
            )
        _badge_html = "".join(_badges)

        with st.container():
            st.markdown(
                f'''
                <div style="
                    padding:12px;
                    background: linear-gradient(90deg,#6366f1,#0ea5e9);
                    color:#ffffff;
                    border-radius:8px;
                    margin-bottom:12px;
                    box-shadow:0 2px 6px rgba(0,0,0,0.08);
                    display:flex;align-items:center;justify-content:space-between;">
                    <div style="font-weight:700;font-size:1.15rem;">💬 Class Q&amp;A {_badge_html}</div>
                    <div style="font-size:0.92rem;opacity:.9;">
                        Ask a question • Help classmates with answers
                    </div>
                </div>
                ''',
                unsafe_allow_html=True
            )

        # (keep your formatter as-is)
        def _fmt_ts(ts):
            try:
                return ts.strftime("%d %b %H:%M")
            except Exception:
                return ""
#

        # Post a new question (single click -> rerun)
        with st.expander("➕ Ask a new question", expanded=False):
            # clear form values on next run if flagged
            if st.session_state.get("__clear_q_form"):
                st.session_state.pop("__clear_q_form", None)
                st.session_state["q_topic"] = ""
                st.session_state["q_text"] = ""
            topic = st.text_input("Topic (optional)", key="q_topic")
            new_q = st.text_area("Your question", key="q_text", height=80)
            if st.button("Post Question", key="qna_post_question") and new_q.strip():
                q_id = str(uuid4())[:8]
                payload = {
                    "question": new_q.strip(),
                    "asked_by_name": student_name,
                    "asked_by_code": student_code,
                    "timestamp": datetime.utcnow(),
                    "topic": (topic or "").strip(),
                }
                q_base.document(q_id).set(payload)
                preview = (payload["question"][:180] + "…") if len(payload["question"]) > 180 else payload["question"]
                topic_tag = f" • Topic: {payload['topic']}" if payload["topic"] else ""
                _notify_slack(
                    f"❓ *New class question* — {class_name}{topic_tag}\n"
                    f"*From:* {student_name} ({student_code})\n"
                    f"*When:* {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC\n"
                    f"*Q:* {preview}"
                )
                # clear and rerun
                st.session_state["__clear_q_form"] = True
                st.success("Question posted!")
                st.rerun()

        # Controls
        colsa, colsb, colsc = st.columns([2, 1, 1])
        with colsa:
            q_search = st.text_input("Search questions (text or topic)…", key="q_search")
        with colsb:
            show_latest = st.toggle("Newest first", value=True, key="q_show_latest")
        with colsc:
            if st.button("↻ Refresh", key="qna_refresh"):
                st.rerun()

        # Load questions (fresh each run)
        try:
            q_docs = list(q_base.order_by("timestamp", direction=firestore.Query.DESCENDING).stream())
            questions = [dict(d.to_dict() or {}, id=d.id) for d in q_docs]
        except Exception:
            q_docs = list(q_base.stream())
            questions = [dict(d.to_dict() or {}, id=d.id) for d in q_docs]
            questions.sort(key=lambda x: x.get("timestamp"), reverse=True)

        # Filter & order
        if q_search.strip():
            ql = q_search.lower()
            questions = [
                q for q in questions
                if ql in str(q.get("question", "")).lower() or ql in str(q.get("topic", "")).lower()
            ]
        if not show_latest:
            questions = list(reversed(questions))

        # Render questions
        if not questions:
            st.info("No questions yet.")
        else:
            for q in questions:
                q_id = q.get("id", "")
                ts = q.get("timestamp")
                ts_label = _fmt_ts(ts)

                topic_html = (
                    f"<div style='font-size:0.9em;color:#666;'>{q.get('topic','')}</div>"
                    if q.get("topic") else ""
                )
                st.markdown(
                    f"<div style='padding:10px;background:#f8fafc;border:1px solid #ddd;border-radius:6px;margin:6px 0;'>"
                    f"<b>{q.get('asked_by_name','')}</b>"
                    f"<span style='color:#aaa;'> • {ts_label}</span>"
                    f"{topic_html}"
                    f"{q.get('question','')}"
                    f"</div>",
                    unsafe_allow_html=True
                )

                # Edit/Delete controls for the question
                can_modify_q = (q.get("asked_by_code") == student_code) or IS_ADMIN
                if can_modify_q:
                    qc1, qc2, _ = st.columns([1, 1, 6])
                    with qc1:
                        if st.button("✏️ Edit", key=f"q_edit_btn_{q_id}"):
                            st.session_state[f"q_editing_{q_id}"] = True
                            st.session_state[f"q_edit_text_{q_id}"] = q.get("question", "")
                            st.session_state[f"q_edit_topic_{q_id}"] = q.get("topic", "")
                    with qc2:
                        if st.button("🗑️ Delete", key=f"q_del_btn_{q_id}"):
                            # delete replies first
                            try:
                                r_ref = q_base.document(q_id).collection("replies")
                                for rdoc in r_ref.stream():
                                    rdoc.reference.delete()
                            except Exception:
                                pass
                            q_base.document(q_id).delete()
                            _notify_slack(
                                f"🗑️ *Q&A question deleted* — {class_name}\n"
                                f"*By:* {student_name} ({student_code}) • QID: {q_id}\n"
                                f"*When:* {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC"
                            )
                            st.success("Question deleted.")
                            st.rerun()

                    # Inline edit form
                    if st.session_state.get(f"q_editing_{q_id}", False):
                        with st.form(f"q_edit_form_{q_id}"):
                            new_topic = st.text_input(
                                "Edit topic (optional)",
                                value=st.session_state.get(f"q_edit_topic_{q_id}", ""),
                                key=f"q_edit_topic_input_{q_id}"
                            )
                            new_text = st.text_area(
                                "Edit question",
                                value=st.session_state.get(f"q_edit_text_{q_id}", ""),
                                key=f"q_edit_text_input_{q_id}",
                                height=100
                            )
                            save_edit = st.form_submit_button("💾 Save")
                            cancel_edit = st.form_submit_button("❌ Cancel")
                        if save_edit and new_text.strip():
                            q_base.document(q_id).update({
                                "question": new_text.strip(),
                                "topic": (new_topic or "").strip(),
                                "edited_at": datetime.utcnow(),
                            })
                            _notify_slack(
                                f"✏️ *Q&A question edited* — {class_name}\n"
                                f"*By:* {student_name} ({student_code}) • QID: {q_id}\n"
                                f"*When:* {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC\n"
                                f"*New:* {(new_text[:180] + '…') if len(new_text) > 180 else new_text}"
                            )
                            st.session_state[f"q_editing_{q_id}"] = False
                            st.success("Question updated.")
                            st.rerun()
                        if cancel_edit:
                            st.session_state[f"q_editing_{q_id}"] = False
                            st.rerun()

                # Load replies
                r_ref = q_base.document(q_id).collection("replies")
                try:
                    replies_docs = list(r_ref.order_by("timestamp").stream())
                except Exception:
                    replies_docs = list(r_ref.stream())
                    replies_docs.sort(key=lambda r: (r.to_dict() or {}).get("timestamp"))

                if replies_docs:
                    for r in replies_docs:
                        rid = r.id
                        r_data = r.to_dict() or {}
                        r_label = _fmt_ts(r_data.get("timestamp"))
                        st.markdown(
                            f"<div style='margin-left:20px;color:#444;'>↳ <b>{r_data.get('replied_by_name','')}</b> "
                            f"<span style='color:#bbb;'>{r_label}</span><br>"
                            f"{r_data.get('reply_text','')}</div>",
                            unsafe_allow_html=True
                        )

                        # Edit/Delete for replies
                        can_modify_r = (r_data.get("replied_by_code") == student_code) or IS_ADMIN
                        if can_modify_r:
                            rc1, rc2, _ = st.columns([1, 1, 6])
                            with rc1:
                                if st.button("✏️ Edit", key=f"r_edit_btn_{q_id}_{rid}"):
                                    st.session_state[f"r_editing_{q_id}_{rid}"] = True
                                    st.session_state[f"r_edit_text_{q_id}_{rid}"] = r_data.get("reply_text", "")
                            with rc2:
                                if st.button("🗑️ Delete", key=f"r_del_btn_{q_id}_{rid}"):
                                    r.reference.delete()
                                    _notify_slack(
                                        f"🗑️ *Q&A reply deleted* — {class_name}\n"
                                        f"*By:* {student_name} ({student_code}) • QID: {q_id}\n"
                                        f"*When:* {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC"
                                    )
                                    st.success("Reply deleted.")
                                    st.rerun()

                            if st.session_state.get(f"r_editing_{q_id}_{rid}", False):
                                with st.form(f"r_edit_form_{q_id}_{rid}"):
                                    new_rtext = st.text_area(
                                        "Edit reply",
                                        value=st.session_state.get(f"r_edit_text_{q_id}_{rid}", ""),
                                        key=f"r_edit_text_input_{q_id}_{rid}",
                                        height=80
                                    )
                                    rsave = st.form_submit_button("💾 Save")
                                    rcancel = st.form_submit_button("❌ Cancel")
                                if rsave and new_rtext.strip():
                                    r.reference.update({
                                        "reply_text": new_rtext.strip(),
                                        "edited_at": datetime.utcnow(),
                                    })
                                    _notify_slack(
                                        f"✏️ *Q&A reply edited* — {class_name}\n"
                                        f"*By:* {student_name} ({student_code}) • QID: {q_id}\n"
                                        f"*When:* {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC\n"
                                        f"*New:* {(new_rtext[:180] + '…') if len(new_rtext) > 180 else new_rtext}"
                                    )
                                    st.session_state[f"r_editing_{q_id}_{rid}"] = False
                                    st.success("Reply updated.")
                                    st.rerun()
                                if rcancel:
                                    st.session_state[f"r_editing_{q_id}_{rid}"] = False
                                    st.rerun()

                # Reply form (anyone can answer) — single click -> rerun
                input_key = f"q_reply_box_{q_id}"
                clear_key = f"__clear_{input_key}"
                if st.session_state.get(clear_key):
                    st.session_state.pop(clear_key, None)
                    st.session_state[clear_key] = True
                reply_text = st.text_input(
                    f"Reply to Q{q_id}",
                    key=input_key,
                    placeholder="Write your reply…"
                )
                if st.button(f"Send Reply {q_id}", key=f"q_reply_btn_{q_id}") and reply_text.strip():
                    reply_payload = {
                        "reply_text": reply_text.strip(),
                        "replied_by_name": student_name,
                        "replied_by_code": student_code,
                        "timestamp": datetime.utcnow(),
                    }
                    r_ref = q_base.document(q_id).collection("replies")
                    r_ref.document(str(uuid4())[:8]).set(reply_payload)
                    prev = (reply_payload["reply_text"][:180] + "…") if len(reply_payload["reply_text"]) > 180 else reply_payload["reply_text"]
                    _notify_slack(
                        f"💬 *New Q&A reply* — {class_name}\n"
                        f"*By:* {student_name} ({student_code})  •  *QID:* {q_id}\n"
                        f"*When:* {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC\n"
                        f"*Reply:* {prev}"
                    )
                    st.session_state[clear_key] = True
                    st.success("Reply sent!")
                    st.rerun()
#


    # === LEARNING NOTES SUBTAB ===
    elif cb_subtab == "📒 Learning Notes":
        st.markdown("""
            <div style="padding: 14px; background: #8d4de8; color: #fff; border-radius: 8px; 
            text-align:center; font-size:1.5rem; font-weight:700; margin-bottom:16px; letter-spacing:.5px;">
            📒 My Learning Notes
            </div>
        """, unsafe_allow_html=True)

        student_code = st.session_state.get("student_code", "demo001")
        key_notes = f"notes_{student_code}"

        if key_notes not in st.session_state:
            st.session_state[key_notes] = load_notes_from_db(student_code)
        notes = st.session_state[key_notes]

        if st.session_state.get("switch_to_edit_note"):
            st.session_state["course_notes_radio"] = "➕ Add/Edit Note"
            del st.session_state["switch_to_edit_note"]
        elif st.session_state.get("switch_to_library"):
            st.session_state["course_notes_radio"] = "📚 My Notes Library"
            del st.session_state["switch_to_library"]

        notes_subtab = st.radio(
            "Notebook",
            ["➕ Add/Edit Note", "📚 My Notes Library"],
            horizontal=True,
            key="course_notes_radio"
        )

        if notes_subtab == "➕ Add/Edit Note":
            # >>>> New helper message for pre-filled note context <<<<
            editing = st.session_state.get("edit_note_idx", None) is not None
            if editing:
                idx = st.session_state["edit_note_idx"]
                title = st.session_state.get("edit_note_title", "")
                tag = st.session_state.get("edit_note_tag", "")
                text = st.session_state.get("edit_note_text", "")
            else:
                title, tag, text = "", "", ""

            if title and tag:
                st.info(f"You're adding a note for **{title}** ({tag}).")

            st.markdown("#### ✍️ Create a new note or update an old one")

            with st.form("note_form", clear_on_submit=not editing):
                new_title = st.text_input("Note Title", value=title, max_chars=50)
                new_tag = st.text_input("Category/Tag (optional)", value=tag, max_chars=20)
                new_text = st.text_area("Your Note", value=text, height=200, max_chars=3000)
                save_btn = st.form_submit_button("💾 Save Note")
                cancel_btn = editing and st.form_submit_button("❌ Cancel Edit")

            if save_btn:
                timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
                if not new_title.strip():
                    st.warning("Please enter a title.")
                    st.stop()
                note = {
                    "title": new_title.strip().title(),
                    "tag": new_tag.strip().title(),
                    "text": new_text.strip(),
                    "pinned": False,
                    "created": timestamp,
                    "updated": timestamp
                }
                if editing:
                    notes[idx] = note
                    for k in ["edit_note_idx", "edit_note_title", "edit_note_text", "edit_note_tag"]:
                        if k in st.session_state: del st.session_state[k]
                    st.success("Note updated!")
                else:
                    notes.insert(0, note)
                    st.success("Note added!")
                st.session_state[key_notes] = notes
                save_notes_to_db(student_code, notes)
                st.session_state["switch_to_library"] = True
                st.rerun()

            if cancel_btn:
                for k in ["edit_note_idx", "edit_note_title", "edit_note_text", "edit_note_tag"]:
                    if k in st.session_state: del st.session_state[k]
                st.session_state["switch_to_library"] = True
                st.rerun()

        elif notes_subtab == "📚 My Notes Library":
            st.markdown("#### 📚 All My Notes")

            if not notes:
                st.info("No notes yet. Add your first note in the ➕ tab!")
            else:
                search_term = st.text_input("🔎 Search your notes…", "")
                if search_term.strip():
                    filtered = []
                    st.markdown('<div style="height:10px"></div>', unsafe_allow_html=True)
                    for n in notes:
                        if (search_term.lower() in n.get("title","").lower() or 
                            search_term.lower() in n.get("tag","").lower() or 
                            search_term.lower() in n.get("text","").lower()):
                            filtered.append(n)
                    notes_to_show = filtered
                    if not filtered:
                        st.warning("No matching notes found!")
                else:
                    notes_to_show = notes

                # --- Download Buttons (TXT, PDF, DOCX) FOR ALL NOTES ---
                all_notes = []
                for n in notes_to_show:
                    note_text = f"Title: {n.get('title','')}\n"
                    if n.get('tag'):
                        note_text += f"Tag: {n['tag']}\n"
                    note_text += n.get('text','') + "\n"
                    note_text += f"Date: {n.get('updated', n.get('created',''))}\n"
                    note_text += "-"*32 + "\n"
                    all_notes.append(note_text)
                txt_data = "\n".join(all_notes)

                st.download_button(
                    label="⬇️ Download All Notes (TXT)",
                    data=txt_data.encode("utf-8"),
                    file_name=f"{student_code}_notes.txt",
                    mime="text/plain"
                )

                # --- PDF Download (all notes, Unicode/emoji ready!) ---
                class PDF(FPDF):
                    def header(self):
                        self.set_font('DejaVu', '', 16)
                        self.cell(0, 12, "My Learning Notes", align="C", ln=1)
                        self.ln(5)
                pdf = PDF()
                pdf.add_font('DejaVu', '', './font/DejaVuSans.ttf', uni=True)
                pdf.add_page()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.set_font("DejaVu", '', 13)
                pdf.cell(0, 10, "Table of Contents", ln=1)
                pdf.set_font("DejaVu", '', 11)
                for idx, note in enumerate(notes_to_show):
                    pdf.cell(0, 8, f"{idx+1}. {note.get('title','')} - {note.get('created', note.get('updated',''))}", ln=1)
                pdf.ln(5)
                for n in notes_to_show:
                    pdf.set_font("DejaVu", '', 13)
                    pdf.cell(0, 10, f"Title: {n.get('title','')}", ln=1)
                    pdf.set_font("DejaVu", '', 11)
                    if n.get("tag"):
                        pdf.cell(0, 8, f"Tag: {n['tag']}", ln=1)
                    pdf.set_font("DejaVu", '', 12)
                    for line in n.get('text','').split("\n"):
                        pdf.multi_cell(0, 7, line)
                    pdf.ln(1)
                    pdf.set_font("DejaVu", '', 11)
                    pdf.cell(0, 8, f"Date: {n.get('updated', n.get('created',''))}", ln=1)
                    pdf.ln(5)
                    pdf.set_font("DejaVu", '', 10)
                    pdf.cell(0, 4, '-' * 55, ln=1)
                    pdf.ln(8)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                    pdf.output(tmp_pdf.name)
                    tmp_pdf.seek(0)
                    pdf_bytes = tmp_pdf.read()
                os.remove(tmp_pdf.name)
                st.download_button(
                    label="⬇️ Download All Notes (PDF)",
                    data=pdf_bytes,
                    file_name=f"{student_code}_notes.pdf",
                    mime="application/pdf"
                )

                # --- DOCX Download (all notes) ---
                def export_notes_to_docx(notes, student_code="student"):
                    doc = Document()
                    doc.add_heading("My Learning Notes", 0)
                    doc.add_heading("Table of Contents", level=1)
                    for idx, note in enumerate(notes):
                        doc.add_paragraph(f"{idx+1}. {note.get('title', '(No Title)')} - {note.get('created', note.get('updated',''))}")
                    doc.add_page_break()
                    for note in notes:
                        doc.add_heading(note.get('title','(No Title)'), level=1)
                        if note.get("tag"):
                            doc.add_paragraph(f"Tag: {note.get('tag','')}")
                        doc.add_paragraph(note.get('text', ''))
                        doc.add_paragraph(f"Date: {note.get('created', note.get('updated',''))}")
                        doc.add_paragraph('-' * 40)
                        doc.add_paragraph("")
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
                        doc.save(f.name)
                        return f.name
                docx_path = export_notes_to_docx(notes_to_show, student_code)
                with open(docx_path, "rb") as f:
                    st.download_button(
                        label="⬇️ Download All Notes (DOCX)",
                        data=f.read(),
                        file_name=f"{student_code}_notes.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                os.remove(docx_path)

                st.markdown("---")
                pinned_notes = [n for n in notes_to_show if n.get("pinned")]
                other_notes = [n for n in notes_to_show if not n.get("pinned")]
                show_list = pinned_notes + other_notes
                for i, note in enumerate(show_list):
                    st.markdown(
                        f"<div style='padding:12px 0 6px 0; font-weight:600; color:#7c3aed; font-size:1.18rem;'>"
                        f"{'📌 ' if note.get('pinned') else ''}{note.get('title','(No Title)')}"
                        f"</div>", unsafe_allow_html=True)
                    if note.get("tag"):
                        st.caption(f"🏷️ Tag: {note['tag']}")
                    st.markdown(
                        f"<div style='margin-top:-5px; margin-bottom:6px; font-size:1.08rem; line-height:1.7;'>{note['text']}</div>",
                        unsafe_allow_html=True)
                    st.caption(f"🕒 {note.get('updated',note.get('created',''))}")

                    # --- Per-Note Download Buttons (TXT, PDF, DOCX) ---
                    download_cols = st.columns([1,1,1])
                    with download_cols[0]:
                        # TXT per note
                        txt_note = f"Title: {note.get('title','')}\n"
                        if note.get('tag'):
                            txt_note += f"Tag: {note['tag']}\n"
                        txt_note += note.get('text', '') + "\n"
                        txt_note += f"Date: {note.get('updated', note.get('created',''))}\n"
                        st.download_button(
                            label="⬇️ TXT",
                            data=txt_note.encode("utf-8"),
                            file_name=f"{student_code}_{note.get('title','note').replace(' ','_')}.txt",
                            mime="text/plain",
                            key=f"download_txt_{i}"
                        )
                    with download_cols[1]:
                        # PDF per note (Unicode/emoji ready!)
                        class SingleNotePDF(FPDF):
                            def header(self):
                                self.set_font('DejaVu', '', 13)
                                self.cell(0, 10, note.get('title','Note'), ln=True, align='C')
                                self.ln(2)
                        pdf_note = SingleNotePDF()
                        pdf_note.add_font('DejaVu', '', './font/DejaVuSans.ttf', uni=True)
                        pdf_note.add_page()
                        pdf_note.set_font("DejaVu", '', 12)
                        if note.get("tag"):
                            pdf_note.cell(0, 8, f"Tag: {note.get('tag','')}", ln=1)
                        for line in note.get('text','').split("\n"):
                            pdf_note.multi_cell(0, 7, line)
                        pdf_note.ln(1)
                        pdf_note.set_font("DejaVu", '', 11)
                        pdf_note.cell(0, 8, f"Date: {note.get('updated', note.get('created',''))}", ln=1)
                        pdf_bytes_single = pdf_note.output(dest="S").encode("latin1", "replace")
                        st.download_button(
                            label="⬇️ PDF",
                            data=pdf_bytes_single,
                            file_name=f"{student_code}_{note.get('title','note').replace(' ','_')}.pdf",
                            mime="application/pdf",
                            key=f"download_pdf_{i}"
                        )
                    with download_cols[2]:
                        # DOCX per note
                        doc_single = Document()
                        doc_single.add_heading(note.get('title','(No Title)'), level=1)
                        if note.get("tag"):
                            doc_single.add_paragraph(f"Tag: {note.get('tag','')}")
                        doc_single.add_paragraph(note.get('text', ''))
                        doc_single.add_paragraph(f"Date: {note.get('updated', note.get('created',''))}")
                        single_docx_io = io.BytesIO()
                        doc_single.save(single_docx_io)
                        st.download_button(
                            label="⬇️ DOCX",
                            data=single_docx_io.getvalue(),
                            file_name=f"{student_code}_{note.get('title','note').replace(' ','_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_docx_{i}"
                        )

                    cols = st.columns([1,1,1,1])
                    with cols[0]:
                        if st.button("✏️ Edit", key=f"edit_{i}"):
                            st.session_state["edit_note_idx"] = i
                            st.session_state["edit_note_title"] = note["title"]
                            st.session_state["edit_note_text"] = note["text"]
                            st.session_state["edit_note_tag"] = note.get("tag", "")
                            st.session_state["switch_to_edit_note"] = True
                            st.rerun()
                    with cols[1]:
                        if st.button("🗑️ Delete", key=f"del_{i}"):
                            notes.remove(note)
                            st.session_state[key_notes] = notes
                            save_notes_to_db(student_code, notes)
                            st.success("Note deleted.")
                            st.rerun()
                    with cols[2]:
                        if note.get("pinned"):
                            if st.button("📌 Unpin", key=f"unpin_{i}"):
                                note["pinned"] = False
                                st.session_state[key_notes] = notes
                                save_notes_to_db(student_code, notes)
                                st.rerun()
                        else:
                            if st.button("📍 Pin", key=f"pin_{i}"):
                                note["pinned"] = True
                                st.session_state[key_notes] = notes
                                save_notes_to_db(student_code, notes)
                                st.rerun()
                    with cols[3]:
                        st.caption("")



# =========================== MY RESULTS & RESOURCES ===========================
# Safe utilities (define only if missing to avoid duplicates)
if "html_stdlib" not in globals():
    import html as html_stdlib
if "urllib" not in globals():
    import urllib
if "linkify_html" not in globals():
    def linkify_html(text):
        """Escape HTML and convert URLs in plain text to anchor tags."""
        s = "" if text is None or (isinstance(text, float) and pd.isna(text)) else str(text)
        s = html_stdlib.escape(s)
        s = re.sub(r'(https?://[^\s<]+)', r'<a href="\1" target="_blank" rel="noopener">\1</a>', s)
        return s
if "_clean_link" not in globals():
    def _clean_link(val) -> str:
        """Return a clean string or '' if empty/NaN/common placeholders."""
        if val is None: return ""
        if isinstance(val, float) and pd.isna(val): return ""
        s = str(val).strip()
        return "" if s.lower() in {"", "nan", "none", "null", "0"} else s
if "_is_http_url" not in globals():
    def _is_http_url(s: str) -> bool:
        try:
            u = urllib.parse.urlparse(str(s))
            return u.scheme in ("http", "https") and bool(u.netloc)
        except Exception:
            return False

# Reuse the app’s schedules provider if available (no duplicate calls)
def _get_level_schedules():
    if "load_level_schedules" in globals() and callable(load_level_schedules):
        return load_level_schedules()
    # Fallback (won’t run if you’ve got load_level_schedules)
    def _safe(fn):
        try: return fn()
        except Exception: return []
    return {
        "A1": _safe(get_a1_schedule),
        "A2": _safe(get_a2_schedule),
        "B1": _safe(get_b1_schedule),
        "B2": _safe(get_b2_schedule),
        "C1": _safe(get_c1_schedule),
    }

# Plain/emoji score label once; reuse everywhere
if "score_label_fmt" not in globals():
    def score_label_fmt(score, *, plain=False):
        try:
            s = float(score)
        except Exception:
            return "" if not plain else "Needs Improvement"
        if s >= 90: return "Excellent 🌟" if not plain else "Excellent"
        if s >= 75: return "Good 👍"      if not plain else "Good"
        if s >= 60: return "Sufficient ✔️" if not plain else "Sufficient"
        return "Needs Improvement ❗" if not plain else "Needs Improvement"

# PDF text sanitizer defined up-front (header needs it)
if "clean_for_pdf" not in globals():
    import unicodedata as _ud
    def clean_for_pdf(text):
        if not isinstance(text, str):
            text = str(text)
        text = _ud.normalize('NFKD', text)
        text = ''.join(c if 32 <= ord(c) <= 255 else '?' for c in text)
        return text.replace('\n', ' ').replace('\r', ' ')

# Prefer secrets/env for sheet; fallback to constant
def _results_csv_url():
    try:
        u = (st.secrets.get("results", {}).get("csv_url", "") if hasattr(st, "secrets") else "").strip()
        if u: return u
    except Exception:
        pass
    return "https://docs.google.com/spreadsheets/d/1BRb8p3Rq0VpFCLSwL4eS9tSgXBo9hSWzfW_J_7W36NQ/gviz/tq?tqx=out:csv"

# Cached fetch of scores (robust columns)
@st.cache_data(ttl=600)
def fetch_scores(csv_url: str):
    resp = requests.get(csv_url, timeout=8)
    resp.raise_for_status()
    df = pd.read_csv(io.StringIO(resp.text), engine='python')
    # normalize columns
    df.columns = [str(c).strip().lower().replace("studentcode", "student_code") for c in df.columns]
    # a few friendly aliases
    aliases = {
        "assignment/chapter": "assignment",
        "chapter": "assignment",
        "score (%)": "score",
    }
    for src, dst in aliases.items():
        if src in df.columns and dst not in df.columns:
            df = df.rename(columns={src: dst})
    required = ["student_code", "name", "assignment", "score", "date", "level"]
    missing = [c for c in required if c not in df.columns]
    if missing:
        # Return empty with diagnostic columns so UI can error cleanly
        return pd.DataFrame(columns=required)
    df = df.dropna(subset=["student_code", "assignment", "score", "date", "level"])
    return df

# Tiny helpers for current user
def _get_current_student():
    row = st.session_state.get("student_row", {}) or {}
    code = (row.get("StudentCode") or st.session_state.get("student_code", "") or "").strip()
    name = (row.get("Name") or st.session_state.get("student_name", "") or "").strip()
    level = (row.get("Level") or "").strip().upper()
    return code, name, level

if tab == "My Results and Resources":
    # Header
    st.markdown(
        '''
        <div style="
            padding: 8px 12px;
            background: #17a2b8;
            color: #fff;
            border-radius: 6px;
            text-align: center;
            margin-bottom: 8px;
            font-size: 1.3rem;
        ">
            📊 My Results & Resources
        </div>
        ''',
        unsafe_allow_html=True
    )
    st.divider()

    # Live CSV URL (secrets/env-aware)
    GOOGLE_SHEET_CSV = _results_csv_url()

    # Utility: manual download link for PDF
    def _pdf_dl_link(pdf_bytes, filename="results.pdf"):
        import base64 as _b64
        b64 = _b64.b64encode(pdf_bytes).decode()
        return f'<a href="data:application/pdf;base64,{b64}" download="{filename}" style="font-size:1.1em;font-weight:600;color:#2563eb;">📥 Click here to download PDF (manual)</a>'

    # Refresh
    if st.button("🔄 Refresh for your latest results"):
        st.cache_data.clear()
        st.success("Cache cleared! Reloading…")
        st.rerun()

    # Load data
    df_scores = fetch_scores(GOOGLE_SHEET_CSV)
    required_cols = {"student_code", "name", "assignment", "score", "date", "level"}
    if not required_cols.issubset(df_scores.columns):
        st.error("Data format error. Please contact support.")
        st.write("Columns found:", df_scores.columns.tolist())
        st.stop()

    # Current student
    student_code, student_name, _ = _get_current_student()
    code_key = (student_code or "").lower().strip()

    st.header("📈 My Results and Resources Hub")
    st.markdown("View and download your assignment history. All results are private and only visible to you.")

    # Filter to user
    df_user = df_scores[df_scores.student_code.astype(str).str.lower().str.strip() == code_key]
    if df_user.empty:
        st.info("No results yet. Complete an assignment to see your scores!")
        st.stop()

    # Level selector
    df_user = df_user.copy()
    df_user["level"] = df_user["level"].astype(str).str.upper().str.strip()
    levels = sorted(df_user["level"].unique())
    level = st.selectbox("Select level:", levels)
    df_lvl = df_user[df_user.level == level].copy()

    # Metrics
    totals = {"A1": 18, "A2": 29, "B1": 28, "B2": 24, "C1": 24}
    total = int(totals.get(level, 0))
    completed = int(df_lvl["assignment"].nunique())
    df_lvl["score"] = pd.to_numeric(df_lvl["score"], errors="coerce")
    avg_score = float(df_lvl["score"].mean() or 0)
    best_score = float(df_lvl["score"].max() or 0)

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Total Assignments", total)
    c2.metric("Completed", completed)
    c3.metric("Average Score", f"{avg_score:.1f}")
    c4.metric("Best Score", f"{best_score:.0f}")

    # Detailed results
    st.markdown("---")
    st.info("🔎 **Scroll down and expand the box below to see your full assignment history and feedback!**")

    # Default display (available to PDF section below)
    df_display = (
        df_lvl.sort_values(["assignment", "score"], ascending=[True, False])
              .reset_index(drop=True)
    )
    # Ensure optional cols exist
    if "comments" not in df_display.columns: df_display["comments"] = ""
    if "link" not in df_display.columns: df_display["link"] = ""

    with st.expander("📋 SEE DETAILED RESULTS (ALL ASSIGNMENTS & FEEDBACK)", expanded=False):
        base_cols = ["assignment", "score", "date", "comments", "link"]
        for _, row in df_display[base_cols].iterrows():
            perf = score_label_fmt(row["score"])
            comment_html = linkify_html(row["comments"])
            ref_link = _clean_link(row.get("link"))
            show_ref = bool(ref_link) and _is_http_url(ref_link) and pd.notna(pd.to_numeric(row["score"], errors="coerce"))

            st.markdown(
                f"""
                <div style="margin-bottom: 18px;">
                    <span style="font-size:1.05em;font-weight:600;">{row['assignment']}</span><br>
                    Score: <b>{row['score']}</b> <span style='margin-left:12px;'>{perf}</span> | Date: {row['date']}<br>
                    <div style='margin:8px 0; padding:10px 14px; background:#f2f8fa; border-left:5px solid #007bff; border-radius:7px; color:#333; font-size:1em;'>
                        <b>Feedback:</b> {comment_html}
                    </div>
                </div>
                """,
                unsafe_allow_html=True
            )
            if show_ref:
                st.markdown(
                    f'🔍 <a href="{ref_link}" target="_blank" rel="noopener">View answer reference (Lesen & Hören)</a>',
                    unsafe_allow_html=True
                )
            st.divider()

    # Badges
    st.markdown("---")
    st.markdown("### 🏅 Badges & Trophies")
    with st.expander("What badges can you earn?", expanded=False):
        st.markdown(
            """
            - 🏆 **Completion Trophy**: Finish all assignments for your level.
            - 🥇 **Gold Badge**: Maintain an average score above 80.
            - 🥈 **Silver Badge**: Average score above 70.
            - 🥉 **Bronze Badge**: Average score above 60.
            - 🌟 **Star Performer**: Score 85 or higher on any assignment.
            """
        )

    badge_count = 0
    if completed >= total and total > 0:
        st.success("🏆 **Congratulations!** You have completed all assignments for this level!")
        badge_count += 1
    if avg_score >= 90:
        st.info("🥇 **Gold Badge:** Average score above 90!")
        badge_count += 1
    elif avg_score >= 75:
        st.info("🥈 **Silver Badge:** Average score above 75!")
        badge_count += 1
    elif avg_score >= 60:
        st.info("🥉 **Bronze Badge:** Average score above 60!")
        badge_count += 1
    if best_score >= 95:
        st.info("🌟 **Star Performer:** You scored 95 or above on an assignment!")
        badge_count += 1
    if badge_count == 0:
        st.warning("No badges yet. Complete more assignments to earn badges!")

    # Skipped assignments (use shared schedule cache)
    schedules_map = _get_level_schedules()
    schedule = schedules_map.get(level, [])
    def _extract_all_nums(chapter_str):
        parts = re.split(r'[_\s,;]+', str(chapter_str))
        nums = []
        for part in parts:
            m = re.search(r'\d+(?:\.\d+)?', part)
            if m: nums.append(float(m.group()))
        return nums

    completed_nums = set()
    for _, row in df_lvl.iterrows():
        for num in _extract_all_nums(row["assignment"]):
            completed_nums.add(num)
    last_num = max(completed_nums) if completed_nums else 0.0

    skipped_assignments = []
    for lesson in schedule:
        chapter_field = lesson.get("chapter", "")
        lesson_nums = _extract_all_nums(chapter_field)
        day = lesson.get("day", "")
        has_assignment = lesson.get("assignment", False)
        for chap_num in lesson_nums:
            if has_assignment and chap_num < last_num and chap_num not in completed_nums:
                skipped_assignments.append(f"Day {day}: Chapter {chapter_field} – {lesson.get('topic','')}")
                break
    if skipped_assignments:
        st.markdown(
            f"""
            <div style="
                background-color: #fff3cd;
                border-left: 6px solid #ffecb5;
                color: #7a6001;
                padding: 16px 18px;
                border-radius: 8px;
                margin: 12px 0;
                font-size: 1.05em;">
                <b>⚠️ You have skipped the following assignments.<br>
                Please complete them for full progress:</b><br>
                {"<br>".join(skipped_assignments)}
            </div>
            """,
            unsafe_allow_html=True
        )

    # Next assignment recommendation (skip Schreiben & Sprechen-only)
    def _is_recommendable(lesson):
        topic = str(lesson.get("topic", "")).lower()
        return not ("schreiben" in topic and "sprechen" in topic)
    def _extract_max_num(chapter):
        nums = re.findall(r'\d+(?:\.\d+)?', str(chapter))
        return max([float(n) for n in nums], default=None)

    completed_chapters = []
    for a in df_lvl["assignment"]:
        n = _extract_max_num(a)
        if n is not None: completed_chapters.append(n)
    last_num = max(completed_chapters) if completed_chapters else 0.0

    next_assignment = None
    for lesson in schedule:
        chap_num = _extract_max_num(lesson.get("chapter", ""))
        if not _is_recommendable(lesson):
            continue
        if chap_num and chap_num > last_num:
            next_assignment = lesson
            break
    if next_assignment:
        st.success(
            f"**Your next recommended assignment:**\n\n"
            f"**Day {next_assignment.get('day','?')}: {next_assignment.get('chapter','?')} – {next_assignment.get('topic','')}**\n\n"
            f"**Goal:** {next_assignment.get('goal','')}\n\n"
            f"**Instruction:** {next_assignment.get('instruction','')}"
        )
    else:
        st.info("🎉 Great Job!")

    # ======================= PDF SUMMARY DOWNLOAD =======================
    COL_ASSN_W, COL_SCORE_W, COL_DATE_W = 45, 18, 30
    PAGE_WIDTH, MARGIN = 210, 10
    FEEDBACK_W = PAGE_WIDTH - 2 * MARGIN - (COL_ASSN_W + COL_SCORE_W + COL_DATE_W)
    LOGO_URL = "https://i.imgur.com/iFiehrp.png"

    @st.cache_data(ttl=3600)
    def fetch_logo():
        try:
            r = requests.get(LOGO_URL, timeout=6)
            r.raise_for_status()
            import tempfile
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            tmp.write(r.content); tmp.flush()
            return tmp.name
        except Exception:
            return None

    from fpdf import FPDF
    class PDFReport(FPDF):
        def header(self):
            logo_path = fetch_logo()
            if logo_path:
                try:
                    self.image(logo_path, 10, 8, 30)
                    self.ln(20)
                except Exception:
                    self.ln(20)
            else:
                self.ln(28)
            self.set_font("Arial", 'B', 16)
            self.cell(0, 12, clean_for_pdf("Learn Language Education Academy"), ln=1, align='C')
            self.ln(3)
        def footer(self):
            self.set_y(-15)
            self.set_font("Arial", 'I', 9)
            self.set_text_color(120, 120, 120)
            footer_text = clean_for_pdf("Learn Language Education Academy — Results generated on ") + pd.Timestamp.now().strftime("%d.%m.%Y")
            self.cell(0, 8, footer_text, 0, 0, 'C')
            self.set_text_color(0, 0, 0)
            self.alias_nb_pages()

    if st.button("⬇️ Download PDF Summary"):
        pdf = PDFReport()
        pdf.add_page()

        # Student Info
        pdf.set_font("Arial", '', 12)
        # Find a name to show (prefer df_user)
        try:
            shown_name = df_user.name.iloc[0]
        except Exception:
            shown_name = student_name or "Student"
        pdf.cell(0, 8, clean_for_pdf(f"Name: {shown_name}"), ln=1)
        pdf.cell(0, 8, clean_for_pdf(f"Code: {code_key}     Level: {level}"), ln=1)
        pdf.cell(0, 8, clean_for_pdf(f"Date: {pd.Timestamp.now():%Y-%m-%d %H:%M}"), ln=1)
        pdf.ln(5)

        # Summary Metrics
        pdf.set_font("Arial", 'B', 13)
        pdf.cell(0, 10, clean_for_pdf("Summary Metrics"), ln=1)
        pdf.set_font("Arial", '', 11)
        pdf.cell(0, 8, clean_for_pdf(f"Total: {total}   Completed: {completed}   Avg: {avg_score:.1f}   Best: {best_score:.0f}"), ln=1)
        pdf.ln(6)

        # Table
        pdf.set_font("Arial", 'B', 11)
        pdf.set_fill_color(235, 235, 245)
        pdf.cell(COL_ASSN_W, 9, "Assignment", 1, 0, 'C', True)
        pdf.cell(COL_SCORE_W, 9, "Score", 1, 0, 'C', True)
        pdf.cell(COL_DATE_W, 9, "Date", 1, 0, 'C', True)
        pdf.cell(FEEDBACK_W, 9, "Feedback", 1, 1, 'C', True)

        pdf.set_font("Arial", '', 10)
        pdf.set_fill_color(249, 249, 249)
        row_fill = False

        for _, row in df_display.iterrows():
            assn = clean_for_pdf(str(row['assignment'])[:24])
            score_txt = clean_for_pdf(str(row['score']))
            date_txt = clean_for_pdf(str(row['date']))
            label = clean_for_pdf(score_label_fmt(row['score'], plain=True))
            pdf.cell(COL_ASSN_W, 8, assn, 1, 0, 'L', row_fill)
            pdf.cell(COL_SCORE_W, 8, score_txt, 1, 0, 'C', row_fill)
            pdf.cell(COL_DATE_W, 8, date_txt, 1, 0, 'C', row_fill)
            pdf.multi_cell(FEEDBACK_W, 8, label, 1, 'C', row_fill)
            row_fill = not row_fill

        pdf_bytes = pdf.output(dest='S').encode('latin1', 'replace')
        st.download_button(
            label="Download PDF",
            data=pdf_bytes,
            file_name=f"{code_key}_results_{level}.pdf",
            mime="application/pdf"
        )
        st.markdown(_pdf_dl_link(pdf_bytes, f"{code_key}_results_{level}.pdf"), unsafe_allow_html=True)
        st.info("If the button does not work, right-click the blue link above and choose 'Save link as...' to download your PDF.")

    # ======================= USEFUL RESOURCES =======================
    st.markdown("---")
    st.subheader("📚 Useful Resources")
    st.markdown(
        """
**1. [A1 Schreiben Practice Questions](https://drive.google.com/file/d/1X_PFF2AnBXSrGkqpfrArvAnEIhqdF6fv/view?usp=sharing)**  
Practice writing tasks and sample questions for A1.

**2. [A1 Exams Sprechen Guide](https://drive.google.com/file/d/1UWvbCCCcrW3_j9x7pOuWug6_Odvzcvaa/view?usp=sharing)**  
Step-by-step guide to the A1 speaking exam.

**3. [German Writing Rules](https://drive.google.com/file/d/1o7_ez3WSNgpgxU_nEtp6EO1PXDyi3K3b/view?usp=sharing)**  
Tips and grammar rules for better writing.

**4. [A2 Sprechen Guide](https://drive.google.com/file/d/1TZecDTjNwRYtZXpEeshbWnN8gCftryhI/view?usp=sharing)**  
A2-level speaking exam guide.

**5. [B1 Sprechen Guide](https://drive.google.com/file/d/1snk4mL_Q9-xTBXSRfgiZL_gYRI9tya8F/view?usp=sharing)**  
How to prepare for your B1 oral exam.
        """
    )


# ================================
# 5. EXAMS MODE & CUSTOM CHAT — uses your prompts + bubble UI + highlighting
# ================================

# —— keep Firestore `db` and OpenAI `client` from above (not redefined here) ——

# Optional: progress saver (kept from your code; safe if unused)
def save_exam_progress(student_code, progress_items):
    doc_ref = db.collection("exam_progress").document(student_code)
    doc = doc_ref.get()
    data = doc.to_dict() if doc.exists else {}
    all_progress = data.get("completed", [])
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    for item in progress_items:
        already = any(
            p["level"] == item["level"] and
            p["teil"] == item["teil"] and
            p["topic"] == item["topic"]
            for p in all_progress
        )
        if not already:
            all_progress.append({
                "level": item["level"],
                "teil": item["teil"],
                "topic": item["topic"],
                "date": now
            })
    doc_ref.set({"completed": all_progress}, merge=True)

# Simple back-step that returns to Stage 1 (used in buttons)
def back_step():
    for key in [
        "falowen_mode", "falowen_level", "falowen_teil",
        "falowen_exam_topic", "falowen_exam_keyword",
        "remaining_topics", "used_topics", "falowen_messages"
    ]:
        st.session_state.pop(key, None)
    st.session_state["_falowen_loaded"] = False
    st.session_state["falowen_stage"] = 1
    st.rerun()

# --- CONFIG (same doc, no duplicate db init) ---
exam_sheet_id = "1zaAT5NjRGKiITV7EpuSHvYMBHHENMs9Piw3pNcyQtho"
exam_sheet_name = "exam_topics"
exam_csv_url = f"https://docs.google.com/spreadsheets/d/{exam_sheet_id}/gviz/tq?tqx=out:csv&sheet={exam_sheet_name}"

@st.cache_data
def load_exam_topics():
    df = pd.read_csv(exam_csv_url)
    for col in ['Level', 'Teil', 'Topic/Prompt', 'Keyword/Subtopic']:
        if col not in df.columns:
            df[col] = ""
    # strip
    for c in df.columns:
        if df[c].dtype == "O":
            df[c] = df[c].astype(str).str.strip()
    return df

df_exam = load_exam_topics()

# ================= UI styles: bubbles + highlights (yours, restored) =================
bubble_user = (
    "background:#1976d2; color:#fff; border-radius:18px 18px 2px 18px;"
    "padding:10px 16px; margin:5px 0 5px auto; max-width:90vw; display:inline-block; font-size:1.12em;"
    "box-shadow:0 2px 8px rgba(0,0,0,0.09); word-break:break-word;"
)
bubble_assistant = (
    "background:#faf9e4; color:#2d2d2d; border-radius:18px 18px 18px 2px;"
    "padding:10px 16px; margin:5px auto 5px 0; max-width:90vw; display:inline-block; font-size:1.12em;"
    "box-shadow:0 2px 8px rgba(0,0,0,0.09); word-break:break-word;"
)
highlight_words = [
    "Fehler", "Tipp", "Achtung", "gut", "korrekt", "super", "nochmals",
    "Bitte", "Vergessen Sie nicht"
]

import re, random

def highlight_keywords(text, words, ignore_case=True):
    flags = re.IGNORECASE if ignore_case else 0
    for w in words:
        pattern = r'\b' + re.escape(w) + r'\b'
        text = re.sub(
            pattern,
            lambda m: f"<span style='background:#ffe082; color:#d84315; font-weight:bold;'>{m.group(0)}</span>",
            text,
            flags=flags,
        )
    return text

def clear_falowen_chat(student_code, mode, level, teil):
    """Deletes the saved chat for a particular student/mode/level/teil from Firestore."""
    chat_key = f"{mode}_{level}_{teil or 'custom'}"
    doc_ref = db.collection("falowen_chats").document(student_code)
    doc = doc_ref.get()
    if doc.exists:
        data = doc.to_dict()
        chats = data.get("chats", {})
        if chat_key in chats:
            del chats[chat_key]
            doc_ref.set({"chats": chats}, merge=True)

# ====== Quick links (kept) ======
lesen_links = {
    "A1": [("Goethe A1 Lesen (Lesen & Hören page)", "https://www.goethe.de/ins/mm/en/spr/prf/gzsd1/ueb.html")],
    "A2": [("Goethe A2 Lesen (Lesen & Hören page)", "https://www.goethe.de/ins/mm/en/spr/prf/gzsd2/ueb.html")],
    "B1": [("Goethe B1 Lesen (Lesen & Hören page)", "https://www.goethe.de/ins/mm/en/spr/prf/gzb1/ueb.html")],
    "B2": [("Goethe B2 Lesen (Lesen & Hören page)", "https://www.goethe.de/ins/mm/en/spr/prf/gzb2/ue9.html")],
    "C1": [("Goethe C1 Lesen (Lesen & Hören page)", "https://www.goethe.de/ins/be/en/spr/prf/gzc1/u24.html")],
}
hoeren_links = {
    "A1": [("Goethe A1 Hören (Lesen & Hören page)", "https://www.goethe.de/ins/mm/en/spr/prf/gzsd1/ueb.html")],
    "A2": [("Goethe A2 Hören (Lesen & Hören page)", "https://www.goethe.de/ins/mm/en/spr/prf/gzsd2/ueb.html")],
    "B1": [("Goethe B1 Hören (Lesen & Hören page)", "https://www.goethe.de/ins/mm/en/spr/prf/gzb1/ueb.html")],
    "B2": [("Goethe B2 Hören (Lesen & Hören page)", "https://www.goethe.de/ins/mm/en/spr/prf/gzb2/ue9.html")],
    "C1": [("Goethe C1 Hören (Lesen & Hören page)", "https://www.goethe.de/ins/be/en/spr/prf/gzc1/u24.html")],
}

# ================= PROMPT BUILDERS (yours, unchanged) =================
def build_a1_exam_intro():
    return (
        "**A1 – Teil 1: Basic Introduction**\n\n"
        "In the A1 exam's first part, you will be asked to introduce yourself. "
        "Typical information includes: your **Name, Land, Wohnort, Sprachen, Beruf, Hobby**.\n\n"
        "After your introduction, you will be asked 3 basic questions such as:\n"
        "- Haben Sie Geschwister?\n"
        "- Wie alt ist deine Mutter?\n"
        "- Bist du verheiratet?\n\n"
        "You might also be asked to spell your name (**Buchstabieren**). "
        "Please introduce yourself now using all the keywords above."
    )

def build_exam_instruction(level, teil):
    # ... (UNCHANGED: keep the long per-level/per-teil instructions you shared)
    # Paste your full original version here (omitted for brevity in this snippet)
    # — I kept your earlier long strings exactly as-is in my local version.
    # BEGIN exact content
    if level == "A1":
        if "Teil 1" in teil:
            return build_a1_exam_intro()
        elif "Teil 2" in teil:
            return (
                "**A1 – Teil 2: Question and Answer**\n\n"
                "You will get a topic and a keyword. Your job: ask a question using the keyword, "
                "then answer it yourself. Example: Thema: Geschäft – Keyword: schließen → "
                "Wann schließt das Geschäft?\nLet's try one. Type 'Yes' in the chatbox so we start?"
            )
        elif "Teil 3" in teil:
            return (
                "**A1 – Teil 3: Making a Request**\n\n"
                "You'll receive a prompt (e.g. 'Radio anmachen'). Write a polite request or imperative. "
                "Example: Können Sie bitte das Radio anmachen?\nReady?"
                "Type Yes in the chatbox so we start?"
            )
    if level == "A2":
        if "Teil 1" in teil:
            return (
                "**A2 – Teil 1: Fragen zu Schlüsselwörtern**\n\n"
                "You'll get a topic (e.g. 'Wohnort'). Ask a question, then answer it yourself. "
                "When you're ready, type 'Begin'."
            )
        elif "Teil 2" in teil:
            return (
                "**A2 – Teil 2: Über das Thema sprechen**\n\n"
                "Talk about the topic in 3–4 sentences. I'll correct and give tips. Start when ready."
            )
        elif "Teil 3" in teil:
            return (
                "**A2 – Teil 3: Gemeinsam planen**\n\n"
                "Let's plan something together. Respond and make suggestions. Start when ready."
            )
    if level == "B1":
        if "Teil 1" in teil:
            return (
                "**B1 – Teil 1: Gemeinsam planen**\n\n"
                "We'll plan an activity together (e.g., a trip or party). Give your ideas and answer questions."
            )
        elif "Teil 2" in teil:
            return (
                "**B1 – Teil 2: Präsentation**\n\n"
                "Give a short presentation on the topic (about 2 minutes). I'll ask follow-up questions."
            )
        elif "Teil 3" in teil:
            return (
                "**B1 – Teil 3: Feedback & Fragen stellen**\n\n"
                "Answer questions about your presentation. I'll give you feedback on your language and structure."
            )
    if level == "B2":
        if "Teil 1" in teil:
            return (
                "**B2 – Teil 1: Diskussion**\n\n"
                "We'll discuss a topic. Express your opinion and justify it."
            )
        elif "Teil 2" in teil:
            return (
                "**B2 – Teil 2: Präsentation**\n\n"
                "Present a topic in detail. I'll challenge your points and help you improve."
            )
        elif "Teil 3" in teil:
            return (
                "**B2 – Teil 3: Argumentation**\n\n"
                "Argue your perspective. I'll give feedback and counterpoints."
            )
    if level == "C1":
        if "Teil 1" in teil:
            return (
                "**C1 – Teil 1: Vortrag**\n\n"
                "Bitte halte einen kurzen Vortrag zum Thema. Ich werde anschließend Fragen stellen und deine Sprache bewerten."
            )
        elif "Teil 2" in teil:
            return (
                "**C1 – Teil 2: Diskussion**\n\n"
                "Diskutiere mit mir über das gewählte Thema. Ich werde kritische Nachfragen stellen."
            )
        elif "Teil 3" in teil:
            return (
                "**C1 – Teil 3: Bewertung**\n\n"
                "Bewerte deine eigene Präsentation. Was würdest du beim nächsten Mal besser machen?"
            )
    return ""
    # END exact content

def build_exam_system_prompt(level, teil):
    # ... (UNCHANGED: keep your detailed persona rules; exact copy from your code)
    # BEGIN exact content from your message (trimmed here for space)
    if level == "A1":
        if "Teil 1" in teil:
            return (
                "You are Herr Felix, a supportive A1 German examiner. "
                "Ask the student to introduce themselves using the keywords (Name, Land, Wohnort, Sprachen, Beruf, Hobby). "
                "Check if all info is given, correct any errors (explain in English), and give the right way to say things in German. "
                "1. Always explain errors and suggestion in English only. Only next question should be German. They are just A1 student "
                "After their intro, ask these three questions one by one: "
                "'Haben Sie Geschwister?', 'Wie alt ist deine Mutter?', 'Bist du verheiratet?'. "
                "Correct their answers (explain in English). At the end, mention they may be asked to spell their name ('Buchstabieren') and wish them luck."
                "Give them a score out of 25 and let them know if they passed or not"
            )
        elif "Teil 2" in teil:
            return (
                "You are Herr Felix, an A1 examiner. Randomly give the student a Thema and Keyword from the official list. "
                "Let them know you have 52 cards available and here to help them prepare for the exams. Let them know they can relax and continue another time when tired. Explain in English "
                "Tell them to ask a question with the keyword and answer it themselves, then correct their German (explain errors in English, show the correct version), and move to the next topic."
                 "1.After every input, let them know if they passed or not and explain why you said so"
            )
        elif "Teil 3" in teil:
            return (
                "You are Herr Felix, an A1 examiner. Give the student a prompt (e.g. 'Radio anmachen'). "
                "Let them know you have 20 cards available and you here to help them prepare for the exams. Let them know they can relax and continue another time when tired. Explain in English "
                "Ask them to write a polite request or imperative and answer themseves like their partners will do. Check if it's correct and polite, explain errors in English, and provide the right German version. Then give the next prompt."
                " They respond using Ja gerne or In ordnung. They can also answer using Ja, Ich kann and the question of the verb at the end (e.g 'Ich kann das Radio anmachen'). "
            )
    if level == "A2":
        if "Teil 1" in teil:
            return (
                "You are Herr Felix, a Goethe A2 examiner. Give a topic from the A2 list. "
                "Always let the student know that you are to help them pass their exams so they should sit for some minutes and be consistent. Teach them how to pass the exams."
                "1. After student input, let the student know you will ask just 3 questions and after give a score out of 25 marks "
                "2. Use phrases like your next recommended question to ask for the next question"
                "Ask the student to ask and answer a question on it. Always correct their German (explain errors in English), show the correct version, and encourage."
                "Ask one question at a time"
                "Pick 3 random keywords from the topic and ask the student 3 questions only per keyword. One question based on one keyword"
                "When student make mistakes and explaining, use English and simple German to explain the mistake and make correction"
                "After the third questions, mark the student out of 25 marks and tell the student whether they passed or not. Explain in English for them to understand"
            )
        elif "Teil 2" in teil:
            return (
                "You are Herr Felix, an A2 examiner. Give a topic. Student gives a short monologue. Correct errors (in English), give suggestions, and follow up with one question."
                "Always let the student know that you are to help them pass their exams so they should sit for some minutes and be consistent. Teach them how to pass the exams."
                "1. After student input, let the student know you will ask just 3 questions and after give a score out of 25 marks "
                "2. Use phrases like your next recommended question to ask for the next question"
                "Pick 3 random keywords from the topic and ask the student 3 questions only per keyword. One question based on one keyword"
                "When student make mistakes and explaining, use English and simple German to explain the mistake and make correction"
                "After the third questions, mark the student out of 25 marks and tell the student whether they passed or not. Explain in English for them understand"
            )
        elif "Teil 3" in teil:
            return (
                "You are Herr Felix, an A2 examiner. Plan something together (e.g., going to the cinema). Check student's suggestions, correct errors, and keep the conversation going."
                "Always let the student know that you are to help them pass their exams so they should sit for some minutes and be consistent. Teach them how to pass the exams."
                "Alert students to be able to plan something with you for you to agree with exact 5 prompts"
                "After the last prompt, mark the student out of 25 marks and tell the student whether they passed or not. Explain in English for them to understand"
            )
    if level == "B1":
        if "Teil 1" in teil:
            return (
                "You are Herr Felix, a Goethe B1 supportive examiner. You and the student plan an activity together. "
                "Always give feedback in both German and English, correct mistakes, suggest improvements, and keep it realistic."
                "Always let the student know that you are to help them pass their exams so they should sit for some minutes and be consistent. Teach them how to pass the exams."
                "1. Give short answers that encourages the student to also type back"
                "2. After student input, let the student know you will ask just 5 questions and after give a score out of 25 marks. Explain in English for them to understand. "
                "3. Ask only 5 questions and try and end the conversation"
                "4. Give score after every presentation whether they passed or not"
                "5. Use phrases like your next recommended question to ask for the next question"
            )
        elif "Teil 2" in teil:
            return (
                "You are Herr Felix, a Goethe B1 examiner. Student gives a presentation. Give constructive feedback in German and English, ask for more details, and highlight strengths and weaknesses."
                "Always let the student know that you are to help them pass their exams so they should sit for some minutes and be consistent. Teach them how to pass the exams."
                "1. After student input, let the student know you will ask just 3 questions and after give a score out of 25 marks. Explain in English for them to understand. "
                "2. Ask only 3 questions and one question at a time"
                "3. Dont make your reply too long and complicated but friendly"
                "4. After your third question, mark and give the student their scores"
                "5. Use phrases like your next recommended question to ask for the next question"
            )
        elif "Teil 3" in teil:
            return (
                "You are Herr Felix, a Goethe B1 examiner. Student answers questions about their presentation. "
                "Always let the student know that you are to help them pass their exams so they should sit for some minutes and be consistent. Teach them to pass the exams. Tell them to ask questions if they dont understand and ask for translations of words. You can help than they going to search for words "
                "Give exam-style feedback (in German and English), correct language, and motivate."
                "1. Ask only 3 questions and one question at a time"
                "2. Dont make your reply too long and complicated but friendly"
                "3. After your third question, mark and give the student their scores out of 25 marks. Explain in English for them to understand"
                "4. Use phrases like your next recommended question to ask for the next question"
            )
    if level == "B2":
        if "Teil 1" in teil:
            return (
                "You are Herr Felix, a B2 examiner. Discuss a topic with the student. Challenge their points. Correct errors (mostly in German, but use English if it's a big mistake), and always provide the correct form."
            )
        elif "Teil 2" in teil:
            return (
                "You are Herr Felix, a B2 examiner. Listen to the student's presentation. Give high-level feedback (mostly in German), ask probing questions, and always highlight advanced vocabulary and connectors."
            )
        elif "Teil 3" in teil:
            return (
                "You are Herr Felix, a B2 examiner. Argue your perspective. Give detailed, advanced corrections (mostly German, use English if truly needed). Encourage native-like answers."
            )
    if level == "C1":
        if "Teil 1" in teil or "Teil 2" in teil or "Teil 3" in teil:
            return (
                "Du bist Herr Felix, ein C1-Prüfer. Sprich nur Deutsch. "
                "Stelle herausfordernde Fragen, gib ausschließlich auf Deutsch Feedback, und fordere den Studenten zu komplexen Strukturen auf."
            )
    return ""

def build_custom_chat_prompt(level):
    # exact content from your message kept
    if level == "C1":
        return (
            "You are supportive German C1 Teacher. Speak both english and German "
            "Ask student one question at a time"
            "Suggest useful phrases student can use to begin their phrases"
            "Check if student is writing on C1 Level"
            "After correction, proceed to the next question using the phrase your next recommended question"
            "When there is error, correct for the student and teach them how to say it correctly"
            "Stay on one topic and always ask next question. After 5 intelligent questions only on a topic, give the student their performance and scores and suggestions to improve"
            "Help student progress from B2 to C1 with your support and guidance"
        )
    if level in ["A1", "A2", "B1", "B2"]:
        correction_lang = "in English" if level in ["A1", "A2"] else "half in English and half in German"
        return (
            f"You are Herr Felix, a supportive and innovative German teacher. "
            f"1. Congratulate the student in English for the topic and give interesting tips on the topic. Always let the student know how the session is going to go in English. It shouldnt just be questions but teach them also. The total number of questios,what they should expect,what they would achieve at the end of the session. Let them know they can ask questions or ask for translation if they dont understand anything. You are ready to always help "
            f"2. If student input looks like a letter question instead of a topic for discussion, then prompt them that you are trained to only help them with their speaking so they should rather paste their letter question in the ideas generator in the schreiben tab. "
            f"Promise them that if they answer all 6 questions, you use their own words to build a presentation of 60 words for them. They record it as mp3 or wav on their phones and upload at the Pronunciation & Speaking Checker tab under the Exams Mode & Custom Chat. They only have to be consistent "
            f"Pick 3 useful keywords related to the student's topic and use them as the focus for conversation. Give students ideas and how to build their points for the conversation in English. "
            f"For each keyword, ask the student up to 2 creative, diverse and interesting questions in German only based on student language level, one at a time, not all at once. Just askd the question and don't let student know this is the keyword you are using. "
            f"After each student answer, give feedback and a suggestion to extend their answer if it's too short. Feedback in English and suggestion in German. "
            f" Explain difficult words when level is A1,A2,B1,B2. "
            f"IMPORTANT: If a student asks 3 grammar questions in a row without trying to answer your conversation questions, respond warmly but firmly: remind them to check their course book using the search button for grammar explanations. Explain that reading their book will help them become more independent and confident as a learner. Kindly pause grammar explanations until they have checked the book and tried the conversation questions. Stay positive, but firm about using the resources. If they still have a specific question after reading, gladly help. "
            f"After keyword questions, continue with other random follow-up questions that reflect student selected level about the topic in German (until you reach 6 questions in total). "
            f"Never ask more than 2 questions about the same keyword. "
            f"After the student answers 6 questions, write a summary of their performance: what they did well, mistakes, and what to improve in English and end the chat with motivation and tips. "
            f"Also give them 60 words from their own words in a presentation form that they can use in class. Add your own points if their words and responses were small. Tell them to improve on it, record with phones as wav or mp3 and upload at Pronunciation & Speaking Checker for further assessment and learn to speak without reading "
            f"All feedback and corrections should be {correction_lang}. "
            f"Encourage the student and keep the chat motivating. "
        )
    return ""

# ================= SESSION DEFAULTS (reuse your falowen_* keys) =================
default_state = {
    "falowen_stage": 1,                  # 1: mode, 2: level, 3: part, 4: chat, 5: summary, 99: pron checker
    "falowen_mode": None,                # **RENAMED choices in UI below**
    "falowen_level": None,
    "falowen_teil": None,
    "falowen_messages": [],
    "falowen_turn_count": 0,
    "custom_topic_intro_done": False,
    "custom_chat_level": None,
    "falowen_exam_topic": None,
    "falowen_exam_keyword": None,
}
for key, val in default_state.items():
    if key not in st.session_state:
        st.session_state[key] = val

if tab == "Exams Mode & Custom Chat":
    st.markdown(
        '''
        <div style="padding: 8px 12px; background: #28a745; color: #fff; border-radius: 6px;
                    text-align: center; margin-bottom: 8px; font-size: 1.3rem;">
            🗣️ Exams Mode & Custom Chat
        </div>
        ''',
        unsafe_allow_html=True
    )
    st.divider()

    # ===== Login context (reuse app login; no duplicate UI here) =====
    if "student_code" not in st.session_state or not st.session_state["student_code"]:
        st.warning("Please log in on the main page to continue.")
        st.stop()
    code = st.session_state["student_code"]

    # ——— Step 1: Mode ———
    if st.session_state["falowen_stage"] == 1:
        st.subheader("Step 1: Choose Practice Mode")
        st.info(
            """
            - **Exams Mode**: Chat with an examiner (Sprechen) and quick links to official Lesen/Hören.
            - **Custom Chat**: Free conversation on your topic with feedback.
            - **Pronunciation & Speaking Checker**: Upload a short audio for scoring and tips.
            """,
            icon="ℹ️"
        )
        mode = st.radio(
            "How would you like to practice?",
            ["Exams Mode", "Custom Chat", "Pronunciation & Speaking Checker"],
            key="falowen_mode_center"
        )
        if st.button("Next ➡️", key="falowen_next_mode"):
            st.session_state["falowen_mode"] = mode
            st.session_state["falowen_stage"] = 99 if mode == "Pronunciation & Speaking Checker" else 2
            st.session_state["falowen_level"] = None
            st.session_state["falowen_teil"] = None
            st.session_state["falowen_messages"] = []
            st.session_state["custom_topic_intro_done"] = False
            st.rerun()

    # ——— Step 2: Level ———
    if st.session_state["falowen_stage"] == 2:
        st.subheader("Step 2: Choose Your Level")
        level = st.radio("Select your level:", ["A1","A2","B1","B2","C1"], key="falowen_level_center")

        col1, col2 = st.columns(2)
        with col1:
            if st.button("⬅️ Back", key="falowen_back1"):
                st.session_state["falowen_stage"] = 1
                st.session_state["falowen_messages"] = []
                st.session_state["_falowen_loaded"] = False
                st.rerun()
        with col2:
            if st.button("Next ➡️", key="falowen_next_level"):
                st.session_state["falowen_level"] = level
                st.session_state["falowen_stage"] = 3 if st.session_state["falowen_mode"] == "Exams Mode" else 4
                st.session_state["falowen_teil"] = None
                st.session_state["falowen_messages"] = []
                st.session_state["custom_topic_intro_done"] = False
                st.rerun()
        st.stop()

    # ——— Step 3: Exam Part or Lesen/Hören links ———
    if st.session_state["falowen_stage"] == 3:
        st.subheader("Step 3: Choose Exam Part")
        teil_options = {
            "A1": ["Teil 1 – Basic Introduction", "Teil 2 – Question and Answer", "Teil 3 – Making A Request",
                   "Lesen – Past Exam Reading", "Hören – Past Exam Listening"],
            "A2": ["Teil 1 – Fragen zu Schlüsselwörtern", "Teil 2 – Über das Thema sprechen", "Teil 3 – Gemeinsam planen",
                   "Lesen – Past Exam Reading", "Hören – Past Exam Listening"],
            "B1": ["Teil 1 – Gemeinsam planen (Dialogue)", "Teil 2 – Präsentation (Monologue)", "Teil 3 – Feedback & Fragen stellen",
                   "Lesen – Past Exam Reading", "Hören – Past Exam Listening"],
            "B2": ["Teil 1 – Diskussion", "Teil 2 – Präsentation", "Teil 3 – Argumentation",
                   "Lesen – Past Exam Reading", "Hören – Past Exam Listening"],
            "C1": ["Teil 1 – Vortrag", "Teil 2 – Diskussion", "Teil 3 – Bewertung",
                   "Lesen – Past Exam Reading", "Hören – Past Exam Listening"],
        }
        level = st.session_state["falowen_level"]
        teil = st.radio("Which exam part?", teil_options[level], key="falowen_teil_center")

        if "Lesen" in teil or "Hören" in teil:
            if "Lesen" in teil:
                st.markdown(
                    """
                    <div style="background:#e1f5fe;border-radius:10px;
                                padding:1.1em 1.4em;margin:1.2em 0;">
                      <span style="font-size:1.18em;color:#0277bd;">
                        <b>📖 Past Exam: Lesen (Reading)</b>
                      </span><br><br>
                    """,
                    unsafe_allow_html=True
                )
                for label, url in lesen_links.get(level, []):
                    st.markdown(
                        f'<a href="{url}" target="_blank" style="font-size:1.10em;color:#1976d2;font-weight:600">'
                        f'👉 {label}</a><br>',
                        unsafe_allow_html=True
                    )
                st.markdown("</div>", unsafe_allow_html=True)

            if "Hören" in teil:
                st.markdown(
                    """
                    <div style="background:#ede7f6;border-radius:10px;
                                padding:1.1em 1.4em;margin:1.2em 0;">
                      <span style="font-size:1.18em;color:#512da8;">
                        <b>🎧 Past Exam: Hören (Listening)</b>
                      </span><br><br>
                    """,
                    unsafe_allow_html=True
                )
                for label, url in hoeren_links.get(level, []):
                    st.markdown(
                        f'<a href="{url}" target="_blank" style="font-size:1.10em;color:#5e35b1;font-weight:600">'
                        f'👉 {label}</a><br>',
                        unsafe_allow_html=True
                    )
                st.markdown("</div>", unsafe_allow_html=True)

            if st.button("⬅️ Back", key="lesen_hoeren_back"):
                st.session_state["falowen_stage"] = 2
                st.session_state["falowen_messages"] = []
                st.rerun()

        else:
            # Topic picker (your format: "Topic/Prompt" + "Keyword/Subtopic")
            teil_number = teil.split()[1]  # e.g., "1"
            exam_topics = df_exam[(df_exam["Level"] == level) & (df_exam["Teil"] == f"Teil {teil_number}")].copy()

            topics_list = []
            if not exam_topics.empty:
                topic_vals   = exam_topics["Topic/Prompt"].astype(str).str.strip()
                keyword_vals = exam_topics["Keyword/Subtopic"].astype(str).str.strip()
                topics_list  = [
                    f"{t} – {k}" if k else t
                    for t, k in zip(topic_vals, keyword_vals) if t
                ]

            search = st.text_input("🔍 Search topic or keyword...", "")
            filtered = [t for t in topics_list if search.lower() in t.lower()] if search else topics_list

            if filtered:
                st.markdown("**Preview: Available Topics**")
                for t in filtered[:6]:
                    st.markdown(f"- {t}")
                if len(filtered) > 6:
                    with st.expander(f"See all {len(filtered)} topics"):
                        col1, col2 = st.columns(2)
                        for i, t in enumerate(filtered):
                            with (col1 if i % 2 == 0 else col2):
                                st.markdown(f"- {t}")

                choice = st.selectbox("Pick your topic (or choose random):", ["(random)"] + filtered, key="topic_picker")
                chosen = random.choice(filtered) if choice == "(random)" else choice

                if " – " in chosen:
                    topic, keyword = chosen.split(" – ", 1)
                else:
                    topic, keyword = chosen, None

                st.session_state["falowen_exam_topic"]   = topic
                st.session_state["falowen_exam_keyword"] = keyword
                st.success(f"**Your exam topic is:** {topic}" + (f" – {keyword}" if keyword else ""))

            else:
                st.info("No topics found. Try a different search.")

            col1, col2 = st.columns([1, 2])
            with col1:
                if st.button("⬅️ Back", key="falowen_back_part"):
                    st.session_state["falowen_stage"]    = 2
                    st.session_state["falowen_messages"] = []
                    st.rerun()
            with col2:
                if st.button("Start Practice", key="falowen_start_practice"):
                    st.session_state["falowen_teil"]            = teil
                    st.session_state["falowen_stage"]           = 4
                    st.session_state["falowen_messages"]        = []
                    st.session_state["custom_topic_intro_done"] = False
                    st.session_state["remaining_topics"]        = filtered.copy()
                    random.shuffle(st.session_state["remaining_topics"])
                    st.session_state["used_topics"]             = []
                    st.rerun()

    # ——— Step 4: Chat (Exam or Custom) ———
    if st.session_state.get("falowen_stage") == 4:
        level = st.session_state.get("falowen_level")
        teil  = st.session_state.get("falowen_teil")
        mode  = st.session_state.get("falowen_mode")
        is_exam = mode == "Exams Mode"
        student_code = st.session_state.get("student_code", "demo")

        # Load chat from db once
        if not st.session_state.get("_falowen_loaded"):
            # reuse same storage key format
            def _chat_key(mode, level, teil): return f"{mode}_{level}_{(teil or 'custom')}"
            doc = db.collection("falowen_chats").document(student_code).get()
            if doc.exists:
                chats = (doc.to_dict() or {}).get("chats", {})
                loaded = chats.get(_chat_key(mode, level, teil), [])
                if loaded:
                    st.session_state["falowen_messages"] = loaded
            st.session_state["_falowen_loaded"] = True

        # Initial instruction if chat is empty (uses YOUR builders)
        if not st.session_state["falowen_messages"]:
            instruction = build_exam_instruction(level, teil) if is_exam else (
                "Hallo! 👋 What would you like to talk about? Give me details of what you want so I can understand."
            )
            st.session_state["falowen_messages"].append({"role": "assistant", "content": instruction})
            # save
            try:
                doc = db.collection("falowen_chats").document(student_code)
                snap = doc.get()
                chats = snap.to_dict().get("chats", {}) if snap.exists else {}
                chats[f"{mode}_{level}_{(teil or 'custom')}"] = st.session_state["falowen_messages"]
                doc.set({"chats": chats}, merge=True)
            except Exception:
                pass

        # Build system prompt (YOUR detailed personas)
        if is_exam:
            if (not st.session_state.get("falowen_exam_topic")) and st.session_state.get("remaining_topics"):
                next_topic = st.session_state["remaining_topics"].pop(0)
                if " – " in next_topic:
                    topic, keyword = next_topic.split(" – ", 1)
                    st.session_state["falowen_exam_topic"] = topic
                    st.session_state["falowen_exam_keyword"] = keyword
                else:
                    st.session_state["falowen_exam_topic"] = next_topic
                    st.session_state["falowen_exam_keyword"] = None
                st.session_state["used_topics"].append(next_topic)
            base_prompt = build_exam_system_prompt(level, teil)
            topic = st.session_state.get("falowen_exam_topic")
            if topic:
                system_prompt = f"{base_prompt} Thema: {topic}."
                if st.session_state.get("falowen_exam_keyword"):
                    system_prompt += f" Keyword: {st.session_state['falowen_exam_keyword']}."
            else:
                system_prompt = base_prompt
        else:
            system_prompt = build_custom_chat_prompt(level)

        # Render chat (your bubble UI + highlights)
        for msg in st.session_state["falowen_messages"]:
            if msg["role"] == "assistant":
                with st.chat_message("assistant", avatar="🧑‍🏫"):
                    st.markdown(
                        "<span style='color:#cddc39;font-weight:bold'>🧑‍🏫 Herr Felix:</span><br>"
                        f"<div style='{bubble_assistant}'>{highlight_keywords(msg['content'], highlight_words)}</div>",
                        unsafe_allow_html=True
                    )
            else:
                with st.chat_message("user"):
                    st.markdown(
                        f"<div style='display:flex;justify-content:flex-end;'>"
                        f"<div style='{bubble_user}'>🗣️ {msg['content']}</div></div>",
                        unsafe_allow_html=True
                    )

        # Downloads
        if st.session_state["falowen_messages"]:
            from fpdf import FPDF
            def falowen_download_pdf(messages, filename):
                def safe_latin1(text): return text.encode("latin1","replace").decode("latin1")
                pdf = FPDF(); pdf.add_page(); pdf.set_font("Arial", size=12)
                for m in messages:
                    who = "Herr Felix" if m["role"]=="assistant" else "Student"
                    pdf.multi_cell(0, 8, safe_latin1(f"{who}: {m['content']}"))
                    pdf.ln(1)
                return pdf.output(dest='S').encode('latin1','replace')

            teil_str = str(teil) if teil else "chat"
            pdf_bytes = falowen_download_pdf(st.session_state["falowen_messages"], f"Falowen_Chat_{level}_{teil_str.replace(' ', '_')}")
            st.download_button(
                "⬇️ Download Chat as PDF",
                pdf_bytes,
                file_name=f"Falowen_Chat_{level}_{teil_str.replace(' ', '_')}.pdf",
                mime="application/pdf"
            )
            chat_as_text = "\n".join([f"{m['role'].capitalize()}: {m['content']}" for m in st.session_state["falowen_messages"]])
            st.download_button(
                "⬇️ Download Chat as TXT",
                chat_as_text.encode("utf-8"),
                file_name=f"Falowen_Chat_{level}_{teil_str.replace(' ', '_')}.txt",
                mime="text/plain"
            )

        # Actions
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🗑️ Delete All Chat History"):
                try:
                    db.collection("falowen_chats").document(student_code).delete()
                except Exception as e:
                    st.error(f"Could not delete chat history: {e}")
                else:
                    for key in [
                        "falowen_stage", "falowen_mode", "falowen_level", "falowen_teil",
                        "falowen_messages", "custom_topic_intro_done", "falowen_exam_topic",
                        "falowen_exam_keyword", "remaining_topics", "used_topics", "_falowen_loaded"
                    ]:
                        st.session_state.pop(key, None)
                    st.session_state["falowen_stage"] = 1
                    st.success("All chat history deleted.")
                    st.rerun()
        with col2:
            if st.button("⬅️ Back"):
                back_step()

        # Input + model call
        user_input = st.chat_input("Type your answer or message here...", key="falowen_user_input")
        if user_input:
            st.session_state["falowen_messages"].append({"role": "user", "content": user_input})
            try:
                if "inc_sprechen_usage" in globals():
                    inc_sprechen_usage(student_code)
            except Exception:
                pass

            with st.chat_message(
                "assistant",
                avatar="🧑‍🏫"
            ):
                with st.spinner("🧑‍🏫 Herr Felix is typing..."):
                    messages = [{"role": "system", "content": system_prompt}] + st.session_state["falowen_messages"]
                    try:
                        resp = client.chat.completions.create(
                            model="gpt-4o",
                            messages=messages,
                            temperature=0.15,
                            max_tokens=600
                        )
                        ai_reply = (resp.choices[0].message.content or "").strip()
                    except Exception as e:
                        ai_reply = f"Sorry, an error occurred: {e}"

                st.markdown(
                    "<span style='color:#cddc39;font-weight:bold'>🧑‍🏫 Herr Felix:</span><br>"
                    f"<div style='{bubble_assistant}'>{highlight_keywords(ai_reply, highlight_words)}</div>",
                    unsafe_allow_html=True
                )

            st.session_state["falowen_messages"].append({"role": "assistant", "content": ai_reply})

            # save thread
            try:
                key = f"{mode}_{level}_{(teil or 'custom')}"
                doc = db.collection("falowen_chats").document(student_code)
                snap = doc.get()
                chats = snap.to_dict().get("chats", {}) if snap.exists else {}
                chats[key] = st.session_state["falowen_messages"]
                doc.set({"chats": chats}, merge=True)
            except Exception:
                pass

        st.divider()
        if st.button("✅ End Session & Show Summary"):
            st.session_state["falowen_stage"] = 5
            st.rerun()

    # ——— Step 5: Summary ———
    if st.session_state.get("falowen_stage") == 5:
        st.success("🎉 Practice Session Complete!")
        st.markdown("#### Your Exam Summary")
        if st.session_state.get("falowen_messages"):
            for msg in st.session_state["falowen_messages"]:
                who = "👨‍🎓 You" if msg["role"] == "user" else "🧑‍🏫 Herr Felix"
                st.markdown(f"**{who}:** {msg['content']}")
            # downloads (same as above)
            from fpdf import FPDF
            def _pdf(messages, filename):
                def s(x): return (x or "").encode("latin1","replace").decode("latin1")
                pdf = FPDF(); pdf.add_page(); pdf.set_font("Arial", size=12)
                for m in messages:
                    who = "Herr Felix" if m["role"]=="assistant" else "Student"
                    pdf.multi_cell(0, 8, s(f"{who}: {m['content']}"))
                    pdf.ln(1)
                return pdf.output(dest='S').encode('latin1','replace')
            teil_str = (st.session_state.get("falowen_teil") or "chat").replace(" ","_")
            level_str = st.session_state.get("falowen_level") or ""
            pdf_bytes = _pdf(st.session_state["falowen_messages"], f"Falowen_Chat_{level_str}_{teil_str}")
            st.download_button("⬇️ Download Chat as PDF", pdf_bytes, file_name=f"Falowen_Chat_{level_str}_{teil_str}.pdf", mime="application/pdf")
            txt = "\n".join([f"{m['role'].capitalize()}: {m['content']}" for m in st.session_state["falowen_messages"]])
            st.download_button("⬇️ Download Chat as TXT", txt.encode("utf-8"), file_name=f"Falowen_Chat_{level_str}_{teil_str}.txt", mime="text/plain")

        if st.button("⬅️ Back"):
            back_step()

    # ——— Stage 99: Pronunciation & Speaking Checker
    if st.session_state.get("falowen_stage") == 99:
        import datetime as _dt
        from io import BytesIO

        today_str = _dt.date.today().isoformat()
        uploads_ref = db.collection("pron_uses").document(st.session_state["student_code"])
        doc = uploads_ref.get()
        data = doc.to_dict() if doc.exists else {}
        last_date = data.get("date")
        count = data.get("count", 0)

        if last_date != today_str:
            count = 0
        if count >= 3:
            st.warning("You’ve hit your daily upload limit (3). Try again tomorrow.")
            st.stop()

        st.subheader("🎤 Pronunciation & Speaking Checker")
        st.info(
            """
            Record or upload your speaking sample below (max 60 seconds).
            • Use your phone's voice recorder **or** visit [vocaroo.com](https://vocaroo.com) and download the recording file to your device.  
            • Then tap **Browse** and select the saved **.wav / .mp3 / .m4a** audio file.
            """
        )

        # Android help (non-blocking)
        with st.expander("Having trouble on Android? Tap for quick steps"):
            st.markdown(
                """
                1) Open this page in **Chrome** (not the in-app browser from WhatsApp/Telegram).  
                2) In your recorder app, **Save** the file first, then choose **Files → Audio/Recordings** (not Photos).  
                3) If it still fails, use **vocaroo.com** and download as **.mp3**.
                """
            )

        audio_file = st.file_uploader(
            "Upload your audio file (≤ 60 seconds, WAV/MP3/M4A).",
            type=["mp3", "wav", "m4a", "3gp", "aac", "ogg", "webm"],
            accept_multiple_files=False,
            key="pron_audio_uploader",
        )

        if audio_file:
            # Preflight checks
            max_mb = 24  # Whisper hard limit ~25MB; keep a little headroom
            size_ok = getattr(audio_file, "size", None)
            if size_ok is not None and size_ok > max_mb * 1024 * 1024:
                st.error(f"File is larger than {max_mb} MB. Please trim or export at a lower bitrate.")
                st.stop()

            allowed_types = {
                "audio/mpeg", "audio/mp3", "audio/wav", "audio/x-wav",
                "audio/x-m4a", "audio/m4a", "audio/mp4",
                "audio/3gpp", "video/3gpp",
                "audio/aac", "audio/x-aac",
                "audio/ogg", "audio/webm", "video/webm",
                "application/octet-stream",  # some Android pickers use this
            }
            allowed_exts = (".mp3", ".wav", ".m4a", ".3gp", ".aac", ".ogg", ".webm")

            file_type = (audio_file.type or "").lower()
            file_name = (audio_file.name or "speech").lower()

            # Some Android pickers give no extension; infer one if needed
            def _ensure_ext(name: str, mime: str) -> str:
                if name.endswith(allowed_exts):
                    return name
                if "wav" in mime:
                    return name + ".wav"
                if "mpeg" in mime or "mp3" in mime:
                    return name + ".mp3"
                if "m4a" in mime or "mp4" in mime or "aac" in mime:
                    return name + ".m4a"
                if "3gpp" in mime:
                    return name + ".3gp"
                if "ogg" in mime:
                    return name + ".ogg"
                if "webm" in mime:
                    return name + ".webm"
                return name + ".mp3"

            file_name = _ensure_ext(file_name, file_type)

            if not (file_type.startswith("audio/") or file_type in allowed_types or file_name.endswith(allowed_exts)):
                st.error("Please upload a supported audio file (.mp3, .wav, .m4a, .3gp, .aac, .ogg, .webm).")
                st.stop()

            # Show the inline player
            st.audio(audio_file)

            # Read clean bytes for processing
            try:
                audio_file.seek(0)
            except Exception:
                pass
            audio_bytes = audio_file.read()
            if not audio_bytes or len(audio_bytes) < 512:
                st.error("We received an empty or very short file. Please re-record and ensure the file is saved.")
                st.stop()

            # Try to normalize to 16k mono WAV (best for Whisper). Fall back if conversion not available.
            buf_for_whisper = None
            try:
                from pydub import AudioSegment  # requires ffmpeg on the server
                import tempfile, os

                # Persist to temp with extension so pydub can guess the format
                tmp_in = tempfile.NamedTemporaryFile(delete=False, suffix=file_name[file_name.rfind("."):])
                tmp_in.write(audio_bytes)
                tmp_in.flush(); tmp_in.close()

                seg = AudioSegment.from_file(tmp_in.name)
                seg = seg.set_channels(1).set_frame_rate(16000)
                out_io = BytesIO()
                seg.export(out_io, format="wav")
                out_io.seek(0)
                setattr(out_io, "name", "speech.wav")
                buf_for_whisper = out_io

                try:
                    os.unlink(tmp_in.name)
                except Exception:
                    pass

            except Exception:
                # If pydub/ffmpeg is not installed or conversion fails, send original bytes with a usable name
                raw_io = BytesIO(audio_bytes)
                raw_io.seek(0)
                setattr(raw_io, "name", file_name)
                buf_for_whisper = raw_io

            # Transcribe (German only)
            try:
                transcript_resp = client.audio.transcriptions.create(
                    file=buf_for_whisper,
                    model="whisper-1",
                    language="de",
                    temperature=0,
                    prompt="Dies ist deutsche Sprache. Bitte nur transkribieren (keine Übersetzung).",
                )
                transcript_text = transcript_resp.text or ""
            except Exception as e:
                st.error(f"Sorry, could not process audio: {e}")
                st.stop()

            if not transcript_text.strip():
                st.error("We couldn't detect speech in the file. Please speak louder or reduce background noise and try again.")
                st.stop()

            st.markdown(f"**Transcribed (German):**  \n> {transcript_text}")

            # Evaluate in English
            eval_prompt = (
                "You are an English-speaking tutor evaluating a **German** speaking sample.\n"
                f'The student said (in German): "{transcript_text}"\n\n'
                "Please provide scores **in English only**:\n"
                "• Rate Pronunciation, Grammar, and Fluency each from 0–100.\n"
                "• Give three concise, actionable tips for each category.\n"
                "• Do not translate the student's text; focus on evaluating it.\n\n"
                "Respond exactly in this format:\n"
                "Pronunciation: XX/100\nTips:\n1. …\n2. …\n3. …\n\n"
                "Grammar: XX/100\nTips:\n1. …\n2. …\n3. …\n\n"
                "Fluency: XX/100\nTips:\n1. …\n2. …\n3. …"
            )

            with st.spinner("Evaluating your sample..."):
                try:
                    eval_resp = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {
                                "role": "system",
                                "content": (
                                    "You are an English-speaking tutor evaluating German speech. "
                                    "Always answer in clear, concise English using the requested format."
                                ),
                            },
                            {"role": "user", "content": eval_prompt},
                        ],
                        temperature=0.2,
                    )
                    result_text = eval_resp.choices[0].message.content
                except Exception as e:
                    st.error(f"Evaluation error: {e}")
                    result_text = None

            if result_text:
                st.markdown(result_text)
                uploads_ref.set({"count": count + 1, "date": today_str})
                st.info("💡 Tip: Use **Custom Chat** first to build ideas, then record and upload here.")
                if st.button("🔄 Try Another"):
                    st.rerun()
            else:
                st.error("Could not get feedback. Please try again later.")

        if st.button("⬅️ Back to Start"):
            st.session_state["falowen_stage"] = 1
            st.rerun()
#



# =========================================
# End
# =========================================

# =========================================
# Vocab
# =========================================

# sentence_bank.py
SENTENCE_BANK = {
    "A1": [
        {
            "prompt_en": "I go jogging every morning.",
            "target_de": "Ich gehe jeden Morgen joggen.",
            "tokens": ["Ich", "gehe", "jeden", "Morgen", "joggen", "."],
            "distractors": ["oft", "im", "Park", "später"],
            "hint_en": "Verb in 2nd position; time can follow subject.",
            "grammar_tag": "Verb-2; TMP",
            "weight": 1
        },
        {
            "prompt_en": "Do you have siblings?",
            "target_de": "Hast du Geschwister?",
            "tokens": ["Hast", "du", "Geschwister", "?"],
            "distractors": ["die", "hast", "ist", "Wo"],
            "hint_en": "Yes/No question: verb first.",
            "grammar_tag": "Ja/Nein-Frage",
            "weight": 1
        },
        {
            "prompt_en": "We are going to the supermarket today.",
            "target_de": "Wir gehen heute zum Supermarkt.",
            "tokens": ["Wir", "gehen", "heute", "zum", "Supermarkt", "."],
            "distractors": ["ins", "gehen", "morgen"],
            "hint_en": "Verb 2nd, time after subject, place after time.",
            "grammar_tag": "TMP",
            "weight": 1
        },
        {
            "prompt_en": "My name is Anna.",
            "target_de": "Ich heiße Anna.",
            "tokens": ["Ich", "heiße", "Anna", "."],
            "distractors": ["bin", "Name", "habe"],
            "hint_en": "Introduce yourself with ‘heißen’.",
            "grammar_tag": "Vorstellung",
            "weight": 1
        },
        {
            "prompt_en": "We live in Berlin.",
            "target_de": "Wir wohnen in Berlin.",
            "tokens": ["Wir", "wohnen", "in", "Berlin", "."],
            "distractors": ["nach", "wohne", "im"],
            "hint_en": "‘wohnen’ + in + city.",
            "grammar_tag": "Präpositionen",
            "weight": 1
        },
        {
            "prompt_en": "I would like a coffee, please.",
            "target_de": "Ich möchte einen Kaffee, bitte.",
            "tokens": ["Ich", "möchte", "einen", "Kaffee", ",", "bitte", "."],
            "distractors": ["haben", "die", "mochte"],
            "hint_en": "möchte + Akkusativ.",
            "grammar_tag": "Bestellung",
            "weight": 1
        },
        {
            "prompt_en": "The bus arrives at 8 o'clock.",
            "target_de": "Der Bus kommt um acht Uhr an.",
            "tokens": ["Der", "Bus", "kommt", "um", "acht", "Uhr", "an", "."],
            "distractors": ["an", "fahren", "achtzehn"],
            "hint_en": "Separable verb ‘ankommen’.",
            "grammar_tag": "Trennbare Verben",
            "weight": 1
        },
        {
            "prompt_en": "Where is the toilet?",
            "target_de": "Wo ist die Toilette?",
            "tokens": ["Wo", "ist", "die", "Toilette", "?"],
            "distractors": ["wann", "wer", "woher"],
            "hint_en": "W-Question: verb in 2nd position.",
            "grammar_tag": "Fragen",
            "weight": 1
        },
        {
            "prompt_en": "I am learning German.",
            "target_de": "Ich lerne Deutsch.",
            "tokens": ["Ich", "lerne", "Deutsch", "."],
            "distractors": ["lernen", "lernst", "sprichst"],
            "hint_en": "Simple present tense, verb 2nd.",
            "grammar_tag": "Präsens",
            "weight": 1
        },
        {
            "prompt_en": "She works in a school.",
            "target_de": "Sie arbeitet in einer Schule.",
            "tokens": ["Sie", "arbeitet", "in", "einer", "Schule", "."],
            "distractors": ["im", "arbeiten", "ein"],
            "hint_en": "in + Dativ for location.",
            "grammar_tag": "Präpositionen + Dativ",
            "weight": 1
        },
        {
            "prompt_en": "What is your phone number?",
            "target_de": "Wie ist deine Telefonnummer?",
            "tokens": ["Wie", "ist", "deine", "Telefonnummer", "?"],
            "distractors": ["Wo", "Wann", "Wer"],
            "hint_en": "Use ‘Wie ist…?’ to ask for numbers.",
            "grammar_tag": "Fragen",
            "weight": 1
        },
        {
            "prompt_en": "I like pizza.",
            "target_de": "Ich mag Pizza.",
            "tokens": ["Ich", "mag", "Pizza", "."],
            "distractors": ["möchte", "liebe", "esse"],
            "hint_en": "Use ‘mögen’ to talk about likes.",
            "grammar_tag": "Modalverb mögen",
            "weight": 1
        },
        {
            "prompt_en": "Can you repeat that, please?",
            "target_de": "Kannst du das bitte wiederholen?",
            "tokens": ["Kannst", "du", "das", "bitte", "wiederholen", "?"],
            "distractors": ["kannst", "wiederhole", "du"],
            "hint_en": "Yes/No question: modal verb first.",
            "grammar_tag": "Modalverben; Frage",
            "weight": 1
        },
        {
            "prompt_en": "The bakery is next to the bank.",
            "target_de": "Die Bäckerei ist neben der Bank.",
            "tokens": ["Die", "Bäckerei", "ist", "neben", "der", "Bank", "."],
            "distractors": ["neben", "dem", "Bank"],
            "hint_en": "neben + Dativ (location).",
            "grammar_tag": "Wechselpräposition (Dativ)",
            "weight": 1
        },
        {
            "prompt_en": "I don’t understand.",
            "target_de": "Ich verstehe nicht.",
            "tokens": ["Ich", "verstehe", "nicht", "."],
            "distractors": ["kein", "keine", "nichts"],
            "hint_en": "Use ‘nicht’ to negate the verb.",
            "grammar_tag": "Negation",
            "weight": 1
        },
        {
            "prompt_en": "At what time does the class start?",
            "target_de": "Um wie viel Uhr beginnt der Kurs?",
            "tokens": ["Um", "wie", "viel", "Uhr", "beginnt", "der", "Kurs", "?"],
            "distractors": ["Wann", "beginnen", "Kurs"],
            "hint_en": "Asking for time with ‘Um wie viel Uhr…’.",
            "grammar_tag": "Fragen; Zeit",
            "weight": 1
        },
        {
            "prompt_en": "I’m sorry, I’m late.",
            "target_de": "Entschuldigung, ich bin spät.",
            "tokens": ["Entschuldigung", ",", "ich", "bin", "spät", "."],
            "distractors": ["später", "habe", "ist"],
            "hint_en": "Fixed apology phrase.",
            "grammar_tag": "Redemittel",
            "weight": 1
        },
        {
            "prompt_en": "We need two tickets.",
            "target_de": "Wir brauchen zwei Tickets.",
            "tokens": ["Wir", "brauchen", "zwei", "Tickets", "."],
            "distractors": ["brauche", "Ticket", "zweite"],
            "hint_en": "Plural nouns without article in general count.",
            "grammar_tag": "Akkusativ; Plural",
            "weight": 1
        },
        {
            "prompt_en": "He is from Spain.",
            "target_de": "Er kommt aus Spanien.",
            "tokens": ["Er", "kommt", "aus", "Spanien", "."],
            "distractors": ["von", "Spanischem", "Spanier"],
            "hint_en": "aus + Land for origin.",
            "grammar_tag": "Präpositionen",
            "weight": 1
        },
        {
            "prompt_en": "The window is open.",
            "target_de": "Das Fenster ist offen.",
            "tokens": ["Das", "Fenster", "ist", "offen", "."],
            "distractors": ["auf", "öffnen", "macht"],
            "hint_en": "Simple statement with ‘sein’.",
            "grammar_tag": "Präsens sein",
            "weight": 1
        }
    ],

    "A2": [
        {
            "prompt_en": "I am staying at home because I am sick.",
            "target_de": "Ich bleibe heute zu Hause, weil ich krank bin.",
            "tokens": ["Ich", "bleibe", "heute", "zu", "Hause", ",", "weil", "ich", "krank", "bin", "."],
            "distractors": ["deshalb", "werde", "morgen"],
            "hint_en": "‘weil’ sends the verb to the end.",
            "grammar_tag": "weil",
            "weight": 1
        },
        {
            "prompt_en": "Tomorrow I will visit my friend.",
            "target_de": "Morgen besuche ich meinen Freund.",
            "tokens": ["Morgen", "besuche", "ich", "meinen", "Freund", "."],
            "distractors": ["werde", "besuchen", "Freunde"],
            "hint_en": "Time first → inversion (verb before subject).",
            "grammar_tag": "Inversion",
            "weight": 1
        },
        {
            "prompt_en": "She is reading a book and drinking tea.",
            "target_de": "Sie liest ein Buch und trinkt Tee.",
            "tokens": ["Sie", "liest", "ein", "Buch", "und", "trinkt", "Tee", "."],
            "distractors": ["oder", "Bücher", "trinken"],
            "hint_en": "Coordinate clauses with ‘und’.",
            "grammar_tag": "Konjunktionen",
            "weight": 1
        },
        {
            "prompt_en": "He has to go to the doctor.",
            "target_de": "Er muss zum Arzt gehen.",
            "tokens": ["Er", "muss", "zum", "Arzt", "gehen", "."],
            "distractors": ["geht", "gehen", "ins"],
            "hint_en": "Modal verb + infinitive at the end.",
            "grammar_tag": "Modalverben",
            "weight": 1
        },
        {
            "prompt_en": "We are interested in the new film.",
            "target_de": "Wir interessieren uns für den neuen Film.",
            "tokens": ["Wir", "interessieren", "uns", "für", "den", "neuen", "Film", "."],
            "distractors": ["an", "im", "alte"],
            "hint_en": "sich interessieren für + Akkusativ.",
            "grammar_tag": "Reflexiv + Präposition",
            "weight": 1
        },
        {
            "prompt_en": "It’s raining, therefore we’re staying inside.",
            "target_de": "Es regnet, deshalb bleiben wir drinnen.",
            "tokens": ["Es", "regnet", ",", "deshalb", "bleiben", "wir", "drinnen", "."],
            "distractors": ["weil", "obwohl", "damit"],
            "hint_en": "‘deshalb’ = connector; main clause word order.",
            "grammar_tag": "Folge: deshalb",
            "weight": 1
        },
        {
            "prompt_en": "I’m trying to learn more German.",
            "target_de": "Ich versuche, mehr Deutsch zu lernen.",
            "tokens": ["Ich", "versuche", ",", "mehr", "Deutsch", "zu", "lernen", "."],
            "distractors": ["lernen", "zum", "Deutsch"],
            "hint_en": "zu + Infinitiv construction.",
            "grammar_tag": "zu-Infinitiv",
            "weight": 1
        },
        {
            "prompt_en": "When I have time, I cook.",
            "target_de": "Wenn ich Zeit habe, koche ich.",
            "tokens": ["Wenn", "ich", "Zeit", "habe", ",", "koche", "ich", "."],
            "distractors": ["Weil", "Dann", "habe"],
            "hint_en": "Subordinate clause first → inversion in main clause.",
            "grammar_tag": "Temporalsatz wenn",
            "weight": 1
        },
        {
            "prompt_en": "I have already finished my homework.",
            "target_de": "Ich habe meine Hausaufgaben schon fertig gemacht.",
            "tokens": ["Ich", "habe", "meine", "Hausaufgaben", "schon", "fertig", "gemacht", "."],
            "distractors": ["bin", "gemacht", "machen"],
            "hint_en": "Perfekt with ‘haben’.",
            "grammar_tag": "Perfekt",
            "weight": 1
        },
        {
            "prompt_en": "We moved to a bigger apartment.",
            "target_de": "Wir sind in eine größere Wohnung umgezogen.",
            "tokens": ["Wir", "sind", "in", "eine", "größere", "Wohnung", "umgezogen", "."],
            "distractors": ["haben", "umgezogen", "umziehen"],
            "hint_en": "Perfekt with ‘sein’ (movement change).",
            "grammar_tag": "Perfekt mit sein",
            "weight": 1
        },
        {
            "prompt_en": "First we eat, then we go for a walk.",
            "target_de": "Zuerst essen wir, dann gehen wir spazieren.",
            "tokens": ["Zuerst", "essen", "wir", ",", "dann", "gehen", "wir", "spazieren", "."],
            "distractors": ["weil", "obwohl", "spazierengehen"],
            "hint_en": "Sequencing with adverbs; verb 2nd each clause.",
            "grammar_tag": "Satzadverbien",
            "weight": 1
        },
        {
            "prompt_en": "I don’t have any time today.",
            "target_de": "Ich habe heute keine Zeit.",
            "tokens": ["Ich", "habe", "heute", "keine", "Zeit", "."],
            "distractors": ["nicht", "kein", "Zeiten"],
            "hint_en": "Use ‘kein/keine’ to negate nouns.",
            "grammar_tag": "Negation mit kein",
            "weight": 1
        },
        {
            "prompt_en": "We’re looking forward to the weekend.",
            "target_de": "Wir freuen uns auf das Wochenende.",
            "tokens": ["Wir", "freuen", "uns", "auf", "das", "Wochenende", "."],
            "distractors": ["für", "am", "im"],
            "hint_en": "sich freuen auf + Akkusativ.",
            "grammar_tag": "Reflexiv + Präp.",
            "weight": 1
        },
        {
            "prompt_en": "Could you help me, please?",
            "target_de": "Könnten Sie mir bitte helfen?",
            "tokens": ["Könnten", "Sie", "mir", "bitte", "helfen", "?"],
            "distractors": ["Kannst", "hilfst", "Hilfe"],
            "hint_en": "Polite request with Konjunktiv II of ‘können’.",
            "grammar_tag": "Höflichkeit",
            "weight": 1
        },
        {
            "prompt_en": "I have been living here for two years.",
            "target_de": "Ich wohne seit zwei Jahren hier.",
            "tokens": ["Ich", "wohne", "seit", "zwei", "Jahren", "hier", "."],
            "distractors": ["für", "vor", "Jahre"],
            "hint_en": "seit + Dativ for duration up to now.",
            "grammar_tag": "Zeitangabe seit",
            "weight": 1
        },
        {
            "prompt_en": "As soon as I finish work, I call you.",
            "target_de": "Sobald ich mit der Arbeit fertig bin, rufe ich dich an.",
            "tokens": ["Sobald", "ich", "mit", "der", "Arbeit", "fertig", "bin", ",", "rufe", "ich", "dich", "an", "."],
            "distractors": ["weil", "deshalb", "rufen"],
            "hint_en": "Subordinate clause first; separable verb ‘anrufen’.",
            "grammar_tag": "Temporalsatz sobald; trennbar",
            "weight": 1
        },
        {
            "prompt_en": "I don’t know if he is at home.",
            "target_de": "Ich weiß nicht, ob er zu Hause ist.",
            "tokens": ["Ich", "weiß", "nicht", ",", "ob", "er", "zu", "Hause", "ist", "."],
            "distractors": ["dass", "weil", "wann"],
            "hint_en": "Indirect yes/no question with ‘ob’.",
            "grammar_tag": "Nebensatz ob",
            "weight": 1
        },
        {
            "prompt_en": "My sister is taller than me.",
            "target_de": "Meine Schwester ist größer als ich.",
            "tokens": ["Meine", "Schwester", "ist", "größer", "als", "ich", "."],
            "distractors": ["wie", "groß", "am"],
            "hint_en": "Comparative with ‘als’.",
            "grammar_tag": "Komparativ",
            "weight": 1
        },
        {
            "prompt_en": "I need to pick up the package.",
            "target_de": "Ich muss das Paket abholen.",
            "tokens": ["Ich", "muss", "das", "Paket", "abholen", "."],
            "distractors": ["hole", "ab", "abgeholt"],
            "hint_en": "Modal + separable verb (infinitive at the end).",
            "grammar_tag": "Modal + trennbar",
            "weight": 1
        },
        {
            "prompt_en": "He likes playing football the most.",
            "target_de": "Am liebsten spielt er Fußball.",
            "tokens": ["Am", "liebsten", "spielt", "er", "Fußball", "."],
            "distractors": ["Lieblings", "am", "liebe"],
            "hint_en": "Superlative of ‘gern’: gern → lieber → am liebsten.",
            "grammar_tag": "Steigerung gern",
            "weight": 1
        }
    ],

    "B1": [
        {
            "prompt_en": "I know that you are coming tomorrow.",
            "target_de": "Ich weiß, dass du morgen kommst.",
            "tokens": ["Ich", "weiß", ",", "dass", "du", "morgen", "kommst", "."],
            "distractors": ["kommst", "dann", "sein"],
            "hint_en": "‘dass’ clause: verb at the end.",
            "grammar_tag": "dass",
            "weight": 1
        },
        {
            "prompt_en": "Although it was raining, we went out.",
            "target_de": "Obwohl es geregnet hat, sind wir ausgegangen.",
            "tokens": ["Obwohl", "es", "geregnet", "hat", ",", "sind", "wir", "ausgegangen", "."],
            "distractors": ["Weil", "Deshalb", "ob"],
            "hint_en": "Concessive clause with ‘obwohl’; Perfekt.",
            "grammar_tag": "Obwohl; Perfekt",
            "weight": 1
        },
        {
            "prompt_en": "Could you tell me where the station is?",
            "target_de": "Könnten Sie mir sagen, wo der Bahnhof ist?",
            "tokens": ["Könnten", "Sie", "mir", "sagen", ",", "wo", "der", "Bahnhof", "ist", "?"],
            "distractors": ["wann", "wer", "wie"],
            "hint_en": "Indirect question: verb at the end.",
            "grammar_tag": "Indirekte Frage",
            "weight": 1
        },
        {
            "prompt_en": "He said that he would come later.",
            "target_de": "Er hat gesagt, dass er später kommen würde.",
            "tokens": ["Er", "hat", "gesagt", ",", "dass", "er", "später", "kommen", "würde", "."],
            "distractors": ["wird", "kommt", "kam"],
            "hint_en": "Reported speech with ‘würde’.",
            "grammar_tag": "Indirekte Rede (würde)",
            "weight": 1
        },
        {
            "prompt_en": "If I had more time, I would travel more.",
            "target_de": "Wenn ich mehr Zeit hätte, würde ich mehr reisen.",
            "tokens": ["Wenn", "ich", "mehr", "Zeit", "hätte", ",", "würde", "ich", "mehr", "reisen", "."],
            "distractors": ["habe", "werde", "würden"],
            "hint_en": "Irrealis with Konjunktiv II.",
            "grammar_tag": "Konjunktiv II Konditional",
            "weight": 1
        },
        {
            "prompt_en": "The book that I am reading is exciting.",
            "target_de": "Das Buch, das ich lese, ist spannend.",
            "tokens": ["Das", "Buch", ",", "das", "ich", "lese", ",", "ist", "spannend", "."],
            "distractors": ["welche", "was", "dem"],
            "hint_en": "Relative clause with ‘das’.",
            "grammar_tag": "Relativsatz",
            "weight": 1
        },
        {
            "prompt_en": "I’m used to getting up early.",
            "target_de": "Ich bin daran gewöhnt, früh aufzustehen.",
            "tokens": ["Ich", "bin", "daran", "gewöhnt", ",", "früh", "aufzustehen", "."],
            "distractors": ["gewohnt", "aufstehen", "früher"],
            "hint_en": "Adjective + zu-Infinitiv; fixed phrase.",
            "grammar_tag": "zu-Infinitiv; Redemittel",
            "weight": 1
        },
        {
            "prompt_en": "The film was not as good as expected.",
            "target_de": "Der Film war nicht so gut, wie erwartet.",
            "tokens": ["Der", "Film", "war", "nicht", "so", "gut", ",", "wie", "erwartet", "."],
            "distractors": ["als", "besser", "am"],
            "hint_en": "so … wie for comparison of equality.",
            "grammar_tag": "Vergleich so…wie",
            "weight": 1
        },
        {
            "prompt_en": "While he was cooking, I set the table.",
            "target_de": "Während er kochte, deckte ich den Tisch.",
            "tokens": ["Während", "er", "kochte", ",", "deckte", "ich", "den", "Tisch", "."],
            "distractors": ["Wenn", "Als", "Nachdem"],
            "hint_en": "Temporal clause with ‘während’ (Präteritum).",
            "grammar_tag": "Temporalsatz während",
            "weight": 1
        },
        {
            "prompt_en": "After we arrived, we called our parents.",
            "target_de": "Nachdem wir angekommen waren, haben wir unsere Eltern angerufen.",
            "tokens": ["Nachdem", "wir", "angekommen", "waren", ",", "haben", "wir", "unsere", "Eltern", "angerufen", "."],
            "distractors": ["Nachdem", "ist", "rufen"],
            "hint_en": "Plusquamperfekt in the subordinate clause.",
            "grammar_tag": "Nachdem; Plusquamperfekt",
            "weight": 1
        },
        {
            "prompt_en": "You should do more sport.",
            "target_de": "Du solltest mehr Sport machen.",
            "tokens": ["Du", "solltest", "mehr", "Sport", "machen", "."],
            "distractors": ["sollst", "Sporten", "machst"],
            "hint_en": "Advice with Konjunktiv II of ‘sollen’.",
            "grammar_tag": "Ratschlag",
            "weight": 1
        },
        {
            "prompt_en": "The meeting was postponed because the boss was ill.",
            "target_de": "Die Besprechung wurde verschoben, weil der Chef krank war.",
            "tokens": ["Die", "Besprechung", "wurde", "verschoben", ",", "weil", "der", "Chef", "krank", "war", "."],
            "distractors": ["ist", "hat", "verschob"],
            "hint_en": "Passive in Präteritum + ‘weil’.",
            "grammar_tag": "Passiv Präteritum; weil",
            "weight": 1
        },
        {
            "prompt_en": "I’m looking for a job that offers flexibility.",
            "target_de": "Ich suche eine Stelle, die Flexibilität bietet.",
            "tokens": ["Ich", "suche", "eine", "Stelle", ",", "die", "Flexibilität", "bietet", "."],
            "distractors": ["welche", "bieten", "anbietet"],
            "hint_en": "Relative clause with ‘die’.",
            "grammar_tag": "Relativsatz",
            "weight": 1
        },
        {
            "prompt_en": "It depends on the weather.",
            "target_de": "Es hängt vom Wetter ab.",
            "tokens": ["Es", "hängt", "vom", "Wetter", "ab", "."],
            "distractors": ["von", "Wetter", "ist"],
            "hint_en": "Verb-preposition phrase with separable verb.",
            "grammar_tag": "Verb + Präp.; trennbar",
            "weight": 1
        },
        {
            "prompt_en": "As far as I know, the store is closed.",
            "target_de": "Soweit ich weiß, ist das Geschäft geschlossen.",
            "tokens": ["Soweit", "ich", "weiß", ",", "ist", "das", "Geschäft", "geschlossen", "."],
            "distractors": ["Sofern", "Soviel", "war"],
            "hint_en": "Fixed phrase ‘Soweit ich weiß’.",
            "grammar_tag": "Redemittel",
            "weight": 1
        },
        {
            "prompt_en": "He apologized for the mistake.",
            "target_de": "Er hat sich für den Fehler entschuldigt.",
            "tokens": ["Er", "hat", "sich", "für", "den", "Fehler", "entschuldigt", "."],
            "distractors": ["entschuldigte", "entschuldigen", "bei"],
            "hint_en": "sich entschuldigen für + Akk.",
            "grammar_tag": "Reflexiv + Präp.",
            "weight": 1
        },
        {
            "prompt_en": "If the train is late, we will take a taxi.",
            "target_de": "Falls der Zug verspätet ist, nehmen wir ein Taxi.",
            "tokens": ["Falls", "der", "Zug", "verspätet", "ist", ",", "nehmen", "wir", "ein", "Taxi", "."],
            "distractors": ["Wenn", "würden", "nahm"],
            "hint_en": "Conditional with ‘falls’.",
            "grammar_tag": "Konditionalsatz",
            "weight": 1
        },
        {
            "prompt_en": "I ended up buying the cheaper one.",
            "target_de": "Am Ende habe ich das günstigere gekauft.",
            "tokens": ["Am", "Ende", "habe", "ich", "das", "günstigere", "gekauft", "."],
            "distractors": ["Endlich", "gekauft", "kaufe"],
            "hint_en": "Idiomatic time adverb + Perfekt.",
            "grammar_tag": "Zeitangabe; Perfekt",
            "weight": 1
        },
        {
            "prompt_en": "The more I practice, the better I get.",
            "target_de": "Je mehr ich übe, desto besser werde ich.",
            "tokens": ["Je", "mehr", "ich", "übe", ",", "desto", "besser", "werde", "ich", "."],
            "distractors": ["umso", "je", "bester"],
            "hint_en": "Comparative correlative ‘je … desto’.",
            "grammar_tag": "Je…desto",
            "weight": 1
        },
        {
            "prompt_en": "I didn’t expect that.",
            "target_de": "Damit habe ich nicht gerechnet.",
            "tokens": ["Damit", "habe", "ich", "nicht", "gerechnet", "."],
            "distractors": ["Dafür", "Darauf", "rechnete"],
            "hint_en": "Fixed verb-preposition phrase.",
            "grammar_tag": "Redemittel; Verb + Präp.",
            "weight": 1
        }
    ],

    "B2": [
        {
            "prompt_en": "The car that I bought is red.",
            "target_de": "Das Auto, das ich gekauft habe, ist rot.",
            "tokens": ["Das", "Auto", ",", "das", "ich", "gekauft", "habe", ",", "ist", "rot", "."],
            "distractors": ["welche", "hatte", "mehr"],
            "hint_en": "Relative clause: verb at the end of the clause.",
            "grammar_tag": "Relativsatz",
            "weight": 1
        },
        {
            "prompt_en": "It is assumed that prices will rise.",
            "target_de": "Es wird angenommen, dass die Preise steigen werden.",
            "tokens": ["Es", "wird", "angenommen", ",", "dass", "die", "Preise", "steigen", "werden", "."],
            "distractors": ["steigen", "gestiegen", "wurden"],
            "hint_en": "Impersonal passive + ‘dass’.",
            "grammar_tag": "Passiv unpersönlich",
            "weight": 1
        },
        {
            "prompt_en": "Despite the rain, the concert took place.",
            "target_de": "Trotz des Regens hat das Konzert stattgefunden.",
            "tokens": ["Trotz", "des", "Regens", "hat", "das", "Konzert", "stattgefunden", "."],
            "distractors": ["Obwohl", "wegen", "stattfindet"],
            "hint_en": "Genitive with ‘trotz’.",
            "grammar_tag": "Präp. mit Genitiv",
            "weight": 1
        },
        {
            "prompt_en": "He explained the problem in a way that everyone understood it.",
            "target_de": "Er erklärte das Problem so, dass es alle verstanden.",
            "tokens": ["Er", "erklärte", "das", "Problem", "so", ",", "dass", "es", "alle", "verstanden", "."],
            "distractors": ["damit", "weil", "obwohl"],
            "hint_en": "Consecutive clause ‘so … dass’.",
            "grammar_tag": "Konsekutivsatz",
            "weight": 1
        },
        {
            "prompt_en": "If I had known that earlier, I would have reacted differently.",
            "target_de": "Hätte ich das früher gewusst, hätte ich anders reagiert.",
            "tokens": ["Hätte", "ich", "das", "früher", "gewusst", ",", "hätte", "ich", "anders", "reagiert", "."],
            "distractors": ["Wenn", "würde", "gewollt"],
            "hint_en": "Inversion with omitted ‘wenn’; Konjunktiv II Vergangenheit.",
            "grammar_tag": "Konditionalsatz; Konjunktiv II",
            "weight": 1
        },
        {
            "prompt_en": "The project was completed within the agreed time frame.",
            "target_de": "Das Projekt wurde innerhalb des vereinbarten Zeitrahmens abgeschlossen.",
            "tokens": ["Das", "Projekt", "wurde", "innerhalb", "des", "vereinbarten", "Zeitrahmens", "abgeschlossen", "."],
            "distractors": ["im", "zwischen", "Zeit"],
            "hint_en": "Nominal style + Genitive after preposition.",
            "grammar_tag": "Nominalstil; Genitiv",
            "weight": 1
        },
        {
            "prompt_en": "The article deals with the topic of climate change.",
            "target_de": "Der Artikel setzt sich mit dem Thema Klimawandel auseinander.",
            "tokens": ["Der", "Artikel", "setzt", "sich", "mit", "dem", "Thema", "Klimawandel", "auseinander", "."],
            "distractors": ["über", "an", "darüber"],
            "hint_en": "Fixed reflexive verb + Präposition.",
            "grammar_tag": "Verb + Präp.",
            "weight": 1
        },
        {
            "prompt_en": "He denied having made a mistake.",
            "target_de": "Er bestritt, einen Fehler gemacht zu haben.",
            "tokens": ["Er", "bestritt", ",", "einen", "Fehler", "gemacht", "zu", "haben", "."],
            "distractors": ["dass", "zu", "machen"],
            "hint_en": "zu-Infinitiv (Perfekt) after certain verbs.",
            "grammar_tag": "zu-Infinitiv Perfekt",
            "weight": 1
        },
        {
            "prompt_en": "The results, which were published yesterday, are surprising.",
            "target_de": "Die Ergebnisse, die gestern veröffentlicht wurden, sind überraschend.",
            "tokens": ["Die", "Ergebnisse", ",", "die", "gestern", "veröffentlicht", "wurden", ",", "sind", "überraschend", "."],
            "distractors": ["welche", "worden", "waren"],
            "hint_en": "Relative clause + passive.",
            "grammar_tag": "Relativsatz; Passiv",
            "weight": 1
        },
        {
            "prompt_en": "In contrast to last year, sales have increased.",
            "target_de": "Im Gegensatz zum letzten Jahr sind die Umsätze gestiegen.",
            "tokens": ["Im", "Gegensatz", "zum", "letzten", "Jahr", "sind", "die", "Umsätze", "gestiegen", "."],
            "distractors": ["Gegenteil", "zum", "wurden"],
            "hint_en": "Fixed prepositional phrase.",
            "grammar_tag": "Feste Wendung",
            "weight": 1
        },
        {
            "prompt_en": "It is questionable whether the plan will work.",
            "target_de": "Es ist fraglich, ob der Plan funktionieren wird.",
            "tokens": ["Es", "ist", "fraglich", ",", "ob", "der", "Plan", "funktionieren", "wird", "."],
            "distractors": ["dass", "wenn", "würde"],
            "hint_en": "‘ob’ clause expressing doubt.",
            "grammar_tag": "Indirekte Frage ob",
            "weight": 1
        },
        {
            "prompt_en": "The more complex the task, the more time we need.",
            "target_de": "Je komplexer die Aufgabe ist, desto mehr Zeit brauchen wir.",
            "tokens": ["Je", "komplexer", "die", "Aufgabe", "ist", ",", "desto", "mehr", "Zeit", "brauchen", "wir", "."],
            "distractors": ["umso", "je", "braucht"],
            "hint_en": "‘je … desto’ with adjective in comparative.",
            "grammar_tag": "Je…desto",
            "weight": 1
        },
        {
            "prompt_en": "Contrary to expectations, the meeting was short.",
            "target_de": "Entgegen den Erwartungen war die Besprechung kurz.",
            "tokens": ["Entgegen", "den", "Erwartungen", "war", "die", "Besprechung", "kurz", "."],
            "distractors": ["Gegen", "Entgegen", "Erwartung"],
            "hint_en": "Preposition ‘entgegen’ takes Dative (plural).",
            "grammar_tag": "Präp. Dativ",
            "weight": 1
        },
        {
            "prompt_en": "He acted as if nothing had happened.",
            "target_de": "Er verhielt sich, als ob nichts passiert wäre.",
            "tokens": ["Er", "verhielt", "sich", ",", "als", "ob", "nichts", "passiert", "wäre", "."],
            "distractors": ["war", "sei", "würde"],
            "hint_en": "‘als ob’ + Konjunktiv II (past).",
            "grammar_tag": "Vergleichssatz als ob",
            "weight": 1
        },
        {
            "prompt_en": "It was not until yesterday that I received the email.",
            "target_de": "Erst gestern habe ich die E-Mail bekommen.",
            "tokens": ["Erst", "gestern", "habe", "ich", "die", "E-Mail", "bekommen", "."],
            "distractors": ["Nur", "erst", "bekam"],
            "hint_en": "Focus with ‘erst’ + inversion.",
            "grammar_tag": "Fokus; Inversion",
            "weight": 1
        },
        {
            "prompt_en": "Given the circumstances, the decision is understandable.",
            "target_de": "Angesichts der Umstände ist die Entscheidung nachvollziehbar.",
            "tokens": ["Angesichts", " der", " Umstände", " ist", " die", " Entscheidung", " nachvollziehbar", "."],
            "distractors": ["Wegen", "Trotz", "Angesicht"],
            "hint_en": "Genitive preposition ‘angesichts’.",
            "grammar_tag": "Präp. Genitiv",
            "weight": 1
        },
        {
            "prompt_en": "He is considered a reliable employee.",
            "target_de": "Er gilt als zuverlässiger Mitarbeiter.",
            "tokens": ["Er", "gilt", "als", "zuverlässiger", "Mitarbeiter", "."],
            "distractors": ["giltet", "wie", "zuverlässig"],
            "hint_en": "Verb ‘gelten als’.",
            "grammar_tag": "Verb + als",
            "weight": 1
        },
        {
            "prompt_en": "We must ensure that all data is protected.",
            "target_de": "Wir müssen sicherstellen, dass alle Daten geschützt sind.",
            "tokens": ["Wir", "müssen", "sicherstellen", ",", "dass", "alle", "Daten", "geschützt", "sind", "."],
            "distractors": ["werden", "wurden", "schützen"],
            "hint_en": "Verb + ‘dass’-Satz.",
            "grammar_tag": "dass-Satz",
            "weight": 1
        },
        {
            "prompt_en": "Instead of complaining, we should look for solutions.",
            "target_de": "Anstatt zu jammern, sollten wir nach Lösungen suchen.",
            "tokens": ["Anstatt", "zu", "jammern", ",", "sollten", "wir", "nach", "Lösungen", "suchen", "."],
            "distractors": ["stattdessen", "für", "sucht"],
            "hint_en": "‘anstatt zu’ + Infinitiv.",
            "grammar_tag": "Infinitivgruppe",
            "weight": 1
        }
    ],

    "C1": [
        {
            "prompt_en": "Had he prepared better, the outcome would have been different.",
            "target_de": "Hätte er sich besser vorbereitet, wäre das Ergebnis anders ausgefallen.",
            "tokens": ["Hätte", "er", "sich", "besser", "vorbereitet", ",", "wäre", "das", "Ergebnis", "anders", "ausgefallen", "."],
            "distractors": ["Wenn", "hatte", "würde"],
            "hint_en": "Omitted ‘wenn’; Konjunktiv II Vergangenheit.",
            "grammar_tag": "Irrealis; Konjunktiv II",
            "weight": 1
        },
        {
            "prompt_en": "The measures, some of which are controversial, were approved.",
            "target_de": "Die Maßnahmen, von denen einige umstritten sind, wurden verabschiedet.",
            "tokens": ["Die", "Maßnahmen", ",", "von", "denen", "einige", "umstritten", "sind", ",", "wurden", "verabschiedet", "."],
            "distractors": ["die", "welche", "worden"],
            "hint_en": "Prepositional relative clause.",
            "grammar_tag": "Relativsatz mit Präp.",
            "weight": 1
        },
        {
            "prompt_en": "Considering the latest findings, a reassessment seems necessary.",
            "target_de": "In Anbetracht der neuesten Erkenntnisse erscheint eine Neubewertung notwendig.",
            "tokens": ["In", "Anbetracht", "der", "neuesten", "Erkenntnisse", "erscheint", "eine", "Neubewertung", "notwendig", "."],
            "distractors": ["Aufgrund", "Anbetracht", "scheint"],
            "hint_en": "Genitive prepositional phrase; formal register.",
            "grammar_tag": "Nominalstil; Genitiv",
            "weight": 1
        },
        {
            "prompt_en": "It is to be feared that the situation will escalate.",
            "target_de": "Es ist zu befürchten, dass sich die Lage zuspitzen wird.",
            "tokens": ["Es", "ist", "zu", "befürchten", ",", "dass", "sich", "die", "Lage", "zuspitzen", "wird", "."],
            "distractors": ["befürchtet", "zu", "zuspitzt"],
            "hint_en": "zu-Infinitiv + ‘dass’.",
            "grammar_tag": "zu-Infinitiv; dass",
            "weight": 1
        },
        {
            "prompt_en": "Contrary to what was assumed, the figures are inaccurate.",
            "target_de": "Entgegen der Annahme erweisen sich die Zahlen als ungenau.",
            "tokens": ["Entgegen", "der", "Annahme", "erweisen", "sich", "die", "Zahlen", "als", "ungenau", "."],
            "distractors": ["Gegen", "Annähme", "ungenaue"],
            "hint_en": "‘sich erweisen als’ + Prädikativ.",
            "grammar_tag": "Verb + als",
            "weight": 1
        },
        {
            "prompt_en": "Only by investing more can we maintain our competitiveness.",
            "target_de": "Nur durch höhere Investitionen können wir unsere Wettbewerbsfähigkeit erhalten.",
            "tokens": ["Nur", "durch", "höhere", "Investitionen", "können", "wir", "unsere", "Wettbewerbsfähigkeit", "erhalten", "."],
            "distractors": ["könnten", "erhält", "bei"],
            "hint_en": "Fronted adverbial → inversion.",
            "grammar_tag": "Inversion; Fokus",
            "weight": 1
        },
        {
            "prompt_en": "He failed to recognize the risks associated with the plan.",
            "target_de": "Er versäumte, die mit dem Plan verbundenen Risiken zu erkennen.",
            "tokens": ["Er", "versäumte", ",", "die", "mit", "dem", "Plan", "verbundenen", "Risiken", "zu", "erkennen", "."],
            "distractors": ["verbundene", "Risiko", "erkennt"],
            "hint_en": "Participle attribute + zu-Infinitiv.",
            "grammar_tag": "Partizipialattribut",
            "weight": 1
        },
        {
            "prompt_en": "As was to be expected, the negotiations dragged on.",
            "target_de": "Wie zu erwarten war, zogen sich die Verhandlungen in die Länge.",
            "tokens": ["Wie", "zu", "erwarten", "war", ",", "zogen", "sich", "die", "Verhandlungen", "in", "die", "Länge", "."],
            "distractors": ["Wie", "erwartet", "wurden"],
            "hint_en": "Fixed impersonal construction.",
            "grammar_tag": "Feste Wendung",
            "weight": 1
        },
        {
            "prompt_en": "Even if the proposal is revised, fundamental issues remain.",
            "target_de": "Selbst wenn der Vorschlag überarbeitet wird, bleiben grundlegende Probleme bestehen.",
            "tokens": ["Selbst", "wenn", "der", "Vorschlag", "überarbeitet", "wird", ",", "bleiben", "grundlegende", "Probleme", "bestehen", "."],
            "distractors": ["obwohl", "wären", "bleibt"],
            "hint_en": "Concessive conditional ‘selbst wenn’.",
            "grammar_tag": "Konzessivsatz",
            "weight": 1
        },
        {
            "prompt_en": "What is crucial is not the speed but the accuracy.",
            "target_de": "Entscheidend ist nicht die Geschwindigkeit, sondern die Genauigkeit.",
            "tokens": ["Entscheidend", "ist", "nicht", "die", "Geschwindigkeit", ",", "sondern", "die", "Genauigkeit", "."],
            "distractors": ["aber", "doch", "genau"],
            "hint_en": "Cleft-like emphasis; ‘sondern’ after negation.",
            "grammar_tag": "Fokus; sondern",
            "weight": 1
        },
        {
            "prompt_en": "He is said to have influenced the decision.",
            "target_de": "Er soll die Entscheidung beeinflusst haben.",
            "tokens": ["Er", "soll", "die", "Entscheidung", "beeinflusst", "haben", "."],
            "distractors": ["sollte", "hat", "wurde"],
            "hint_en": "Modalverb ‘sollen’ for report/rumor.",
            "grammar_tag": "Indirektheit",
            "weight": 1
        },
        {
            "prompt_en": "The more attention is paid to details, the fewer errors occur.",
            "target_de": "Je mehr auf Details geachtet wird, desto weniger Fehler treten auf.",
            "tokens": ["Je", "mehr", "auf", "Details", "geachtet", "wird", ",", "desto", "weniger", "Fehler", "treten", "auf", "."],
            "distractors": ["je", "weniger", "tritt"],
            "hint_en": "Impersonal passive + je/desto.",
            "grammar_tag": "Passiv; Je…desto",
            "weight": 1
        },
        {
            "prompt_en": "This is a development whose consequences are still unforeseeable.",
            "target_de": "Dies ist eine Entwicklung, deren Folgen noch unabsehbar sind.",
            "tokens": ["Dies", "ist", "eine", "Entwicklung", ",", "deren", "Folgen", "noch", "unabsehbar", "sind", "."],
            "distractors": ["deren", "welcher", "denen"],
            "hint_en": "Genitive relative pronoun ‘deren’.",
            "grammar_tag": "Relativpronomen Genitiv",
            "weight": 1
        },
        {
            "prompt_en": "Not only did the team miss the deadline, but costs also exploded.",
            "target_de": "Nicht nur verpasste das Team die Frist, sondern auch die Kosten explodierten.",
            "tokens": ["Nicht", "nur", "verpasste", "das", "Team", "die", "Frist", ",", "sondern", "auch", "die", "Kosten", "explodierten", "."],
            "distractors": ["aber", "sondern", "explodiert"],
            "hint_en": "‘Nicht nur … sondern auch’ with inversion.",
            "grammar_tag": "Korrelative Konjunktion",
            "weight": 1
        },
        {
            "prompt_en": "There is reason to assume that demand will decrease.",
            "target_de": "Es gibt Anlass zu der Annahme, dass die Nachfrage zurückgehen wird.",
            "tokens": ["Es", "gibt", "Anlass", "zu", "der", "Annahme", ",", "dass", "die", "Nachfrage", "zurückgehen", "wird", "."],
            "distractors": ["zum", "gehen", "würde"],
            "hint_en": "Nominal phrase + ‘dass’.",
            "grammar_tag": "Nominalstil",
            "weight": 1
        },
        {
            "prompt_en": "Far from being perfect, the plan nevertheless offers a basis for discussion.",
            "target_de": "Weit davon entfernt, perfekt zu sein, bietet der Plan dennoch eine Diskussionsgrundlage.",
            "tokens": ["Weit", "davon", "entfernt", ",", "perfekt", "zu", "sein", ",", "bietet", "der", "Plan", "dennoch", "eine", "Diskussionsgrundlage", "."],
            "distractors": ["obwohl", "perfekt", "ist"],
            "hint_en": "Participial preface + main clause.",
            "grammar_tag": "Partizipialkonstruktion",
            "weight": 1
        },
        {
            "prompt_en": "Whether the project will be funded remains to be seen.",
            "target_de": "Ob das Projekt finanziert wird, bleibt abzuwarten.",
            "tokens": ["Ob", "das", "Projekt", "finanziert", "wird", ",", "bleibt", "abzuwarten", "."],
            "distractors": ["dass", "zu", "abwarten"],
            "hint_en": "Impersonal construction with ‘bleibt abzuwarten’.",
            "grammar_tag": "Unpersönliche Form",
            "weight": 1
        },
        {
            "prompt_en": "It is precisely here that the difficulties arise.",
            "target_de": "Gerade hier ergeben sich die Schwierigkeiten.",
            "tokens": ["Gerade", "hier", "ergeben", "sich", "die", "Schwierigkeiten", "."],
            "distractors": ["ergeben", "gibt", "sich"],
            "hint_en": "Focus adverb ‘gerade’.",
            "grammar_tag": "Fokus",
            "weight": 1
        },
        {
            "prompt_en": "No sooner had we started than problems emerged.",
            "target_de": "Kaum hatten wir begonnen, traten schon Probleme auf.",
            "tokens": ["Kaum", "hatten", "wir", "begonnen", ",", "traten", "schon", "Probleme", "auf", "."],
            "distractors": ["Kaum", "beginnen", "aufgetreten"],
            "hint_en": "‘Kaum …, da/als’ pattern; here without ‘da’.",
            "grammar_tag": "Temporale Inversion",
            "weight": 1
        },
        {
            "prompt_en": "It remains unclear to what extent the rule applies.",
            "target_de": "Unklar bleibt, inwiefern die Regel gilt.",
            "tokens": ["Unklar", "bleibt", ",", "inwiefern", "die", "Regel", "gilt", "."],
            "distractors": ["wiefern", "obwohl", "giltet"],
            "hint_en": "Fronted predicate + indirect question.",
            "grammar_tag": "Inversion; Indirekte Frage",
            "weight": 1
        }
    ]
}

# =========================================
# Vocab
# =========================================

# (Removed duplicate SENTENCE_BANK redefinition — it was overwriting the full bank)

# If you initialize Firestore elsewhere, expose it here.
# This helper prevents NameError if db isn't ready.
def _get_db():
    try:
        return db  # provided by your app elsewhere
    except NameError:
        return None


# ================================
# HELPERS: Level loading (Google Sheet)
# ================================
@st.cache_data
def load_student_levels():
    """
    Load the roster with a 'Level' column.
    Expected columns (case-insensitive): student_code, level
    We normalize headers and try common alternatives for student code and level.
    """
    sheet_id = "12NXf5FeVHr7JJT47mRHh7Jp-TC1yhPS7ZG6nzZVTt1U"
    csv_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv"
    df = pd.read_csv(csv_url)
    # normalize headers
    df.columns = [c.strip().lower() for c in df.columns]

    # try to align student_code column
    code_col_candidates = ["student_code", "studentcode", "code", "student id", "id"]
    level_col_candidates = ["level", "klasse", "stufe"]
    code_col = next((c for c in code_col_candidates if c in df.columns), None)
    level_col = next((c for c in level_col_candidates if c in df.columns), None)

    if code_col is None or level_col is None:
        st.error(
            f"Roster is missing required columns. "
            f"Found: {list(df.columns)}; need one of {code_col_candidates} and one of {level_col_candidates}."
        )
        # still return something so callers don't crash
        df["__dummy_code__"] = "demo001"
        df["__dummy_level__"] = "A1"
        return df.rename(columns={"__dummy_code__": "student_code", "__dummy_level__": "level"})

    # rename to canonical names
    df = df.rename(columns={code_col: "student_code", level_col: "level"})
    return df

def get_student_level(student_code: str, default: str = "A1") -> str:
    """
    Return student's Level (A1..C1) from the roster for this student_code.
    Case-insensitive match on student_code.
    """
    try:
        df = load_student_levels()
        # ensure columns exist after normalization/rename
        if "student_code" not in df.columns or "level" not in df.columns:
            return default
        sc = str(student_code).strip().lower()
        row = df[df["student_code"].astype(str).str.strip().str.lower() == sc]
        if not row.empty:
            return str(row.iloc[0]["level"]).upper().strip()
        return default
    except Exception as e:
        st.warning(f"Could not load level from roster ({e}). Using default {default}.")
        return default


def vocab_attempt_exists(student_code: str, session_id: str) -> bool:
    """Check if an attempt with this session_id already exists for the student."""
    if not session_id:
        return False
    _db = _get_db()
    if _db is None:
        return False

    doc_ref = _db.collection("vocab_stats").document(student_code)
    doc = doc_ref.get()
    if not doc.exists:
        return False

    data = doc.to_dict() or {}
    history = data.get("history", [])
    return any(h.get("session_id") == session_id for h in history)


def save_vocab_attempt(student_code, level, total, correct, practiced_words, session_id=None):
    """
    Save one vocab practice attempt to Firestore.
    Duplicate-safe using session_id.
    """
    _db = _get_db()
    if _db is None:
        st.warning("Firestore not initialized; skipping stats save.")
        return

    if not session_id:
        session_id = str(uuid4())

    if vocab_attempt_exists(student_code, session_id):
        return

    doc_ref = _db.collection("vocab_stats").document(student_code)
    doc = doc_ref.get()
    data = doc.to_dict() if doc.exists else {}
    history = data.get("history", [])

    attempt = {
        "level": level,
        "total": int(total) if total is not None else 0,
        "correct": int(correct) if correct is not None else 0,
        "practiced_words": list(practiced_words or []),
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "session_id": session_id,
    }

    history.append(attempt)
    completed = {w for a in history for w in a.get("practiced_words", [])}

    doc_ref.set({
        "history":           history,
        "last_practiced":    attempt["timestamp"],
        "completed_words":   sorted(completed),
        "total_sessions":    len(history),
    }, merge=True)


def get_vocab_stats(student_code):
    """Load vocab practice stats from Firestore (or defaults)."""
    _db = _get_db()
    if _db is None:
        return {
            "history":           [],
            "last_practiced":    None,
            "completed_words":   [],
            "total_sessions":    0,
        }

    doc_ref = _db.collection("vocab_stats").document(student_code)
    doc = doc_ref.get()
    if doc.exists:
        data = doc.to_dict() or {}
        # Ensure we don't return "best"
        return {
            "history": data.get("history", []),
            "last_practiced": data.get("last_practiced"),
            "completed_words": data.get("completed_words", []),
            "total_sessions": data.get("total_sessions", 0),
        }

    return {
        "history":           [],
        "last_practiced":    None,
        "completed_words":   [],
        "total_sessions":    0,
    }


# ================================
# HELPERS: Writing (Sentence Builder) persistence
# ================================
def save_sentence_attempt(student_code, level, target_sentence, chosen_sentence, correct, tip):
    """Append a sentence-builder attempt to Firestore."""
    _db = _get_db()
    if _db is None:
        st.warning("Firestore not initialized; skipping sentence stats save.")
        return

    doc_ref = _db.collection("sentence_builder_stats").document(student_code)
    doc = doc_ref.get()
    data = doc.to_dict() if doc.exists else {}
    history = data.get("history", [])
    history.append({
        "level": level,
        "target": target_sentence,
        "chosen": chosen_sentence,
        "correct": bool(correct),
        "tip": tip,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
    })
    doc_ref.set({
        "history": history,
        "last_played": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "total_sessions": len(history),
    }, merge=True)

def get_sentence_progress(student_code: str, level: str):
    """
    Returns (correct_unique_count, total_items_for_level)
    based on history in 'sentence_builder_stats'.
    """
    _db = _get_db()
    correct_set = set()
    if _db is not None:
        ref = _db.collection("sentence_builder_stats").document(student_code)
        doc = ref.get()
        if doc.exists:
            data = doc.to_dict()
            for h in data.get("history", []):
                if h.get("level") == level and h.get("correct"):
                    correct_set.add(h.get("target"))
    total_items = len(SENTENCE_BANK.get(level, []))
    return len(correct_set), total_items


# ================================
# HELPERS: Vocab CSV (optional flashcards list)
# ================================
@st.cache_data
def load_vocab_lists():
    """
    Optional CSV for flashcards: columns Level, German, English
    """
    sheet_id = "1I1yAnqzSh3DPjwWRh9cdRSfzNSPsi7o4r5Taj9Y36NU"  # Vocab sheet
    csv_url  = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv&gid=0"
    try:
        df = pd.read_csv(csv_url)
    except Exception as e:
        st.error(f"Could not fetch vocab CSV: {e}")
        return {}
    df.columns = df.columns.str.strip()
    missing = [c for c in ("Level","German","English") if c not in df.columns]
    if missing:
        st.error(f"Missing column(s) in your vocab sheet: {missing}")
        return {}
    df = df[["Level","German","English"]].dropna()
    lists = {}
    for lvl, grp in df.groupby("Level"):
        lists[lvl] = list(zip(grp["German"], grp["English"]))
    return lists

VOCAB_LISTS = load_vocab_lists()


# ================================
# SMALL UI + CHECK HELPERS
# ================================
def render_message(role, msg):
    align   = "left"   if role=="assistant" else "right"
    bgcolor = "#FAFAFA" if role=="assistant" else "#D2F8D2"
    bordcol = "#CCCCCC"
    label   = "Herr Felix 👨‍🏫" if role=="assistant" else "You"
    style = (
        f"padding:14px; border-radius:12px; max-width:96vw; "
        f"margin:7px 0; text-align:{align}; background:{bgcolor}; "
        f"border:1px solid {bordcol}; font-size:1.12em; word-break:break-word;"
    )
    st.markdown(f"<div style='{style}'><b>{label}:</b> {msg}</div>", unsafe_allow_html=True)

def clean_text(text):
    return text.replace("the ", "").replace(",", "").replace(".", "").strip().lower()

def is_correct_answer(user_input, answer):
    import re
    # Clean both sides; accept comma/;/ slash-separated variants
    possible = [clean_text(a) for a in re.split(r"[,/;]", str(answer))]
    return clean_text(str(user_input)) in possible

def normalize_join(tokens):
    """Join tokens and fix spaces before punctuation for sentence builder."""
    s = " ".join(tokens)
    for p in [",", ".", "!", "?", ":", ";"]:
        s = s.replace(f" {p}", p)
    return s

# ---------- Dictionary helpers ----------
def _flatten_vocab_entries(vocab_lists: dict):
    rows = []
    for lvl, pairs in (vocab_lists or {}).items():
        for de, en in pairs:
            rows.append({
                "level":   str(lvl).upper().strip(),
                "german":  str(de).strip(),
                "english": str(en).strip()
            })
    return rows

@st.cache_data(show_spinner=False)
def _dict_tts_bytes_de(word: str, slow: bool = False):
    """Return MP3 bytes for a German word (cached)."""
    try:
        from gtts import gTTS
        import io
        t = gTTS(text=word, lang="de", slow=bool(slow))
        buf = io.BytesIO()
        t.write_to_fp(buf)
        buf.seek(0)
        return buf.read()
    except Exception:
        return None

# ================================
# TAB: Vocab Trainer (locked by Level)
# ================================
if tab == "Vocab Trainer":
    # --- Who is this? ---
    student_code = st.session_state.get("student_code", "demo001")

    # --- Lock the level from your Sheet ---
    student_level_locked = get_student_level(
        student_code,
        default=(st.session_state.get("student_level", "A1"))
    )
    if not student_level_locked:
        student_level_locked = "A1"

    # Header
    st.markdown(
        """
        <div style="
            padding:8px 12px; background:#6f42c1; color:#fff;
            border-radius:6px; text-align:center; margin-bottom:8px;
            font-size:1.3rem;">
        📚 Vocab Trainer
        </div>
        """,
        unsafe_allow_html=True
    )
    st.markdown(f"**Practicing Level:** `{student_level_locked}` (from your profile)")
    st.caption("Your level is loaded automatically from the school list. Ask your tutor if this looks wrong.")
    st.divider()

    subtab = st.radio(
        "Choose practice:",
        ["Sentence Builder", "Vocab Practice", "Dictionary"],
        horizontal=True,
        key="vocab_practice_subtab"
    )

    # ===========================
    # SUBTAB: Sentence Builder
    # ===========================
    if subtab == "Sentence Builder":
        student_level = student_level_locked
        st.info(f"✍️ You are practicing **Sentence Builder** at **{student_level}** (locked from your profile).")

        # --- Guide & Progress (collapsed) ---
        with st.expander("✍️ Sentence Builder — Guide & Progress", expanded=False):
            done_unique, total_items = get_sentence_progress(student_code, student_level)
            pct = int((done_unique / total_items) * 100) if total_items else 0
            st.progress(pct)
            st.caption(f"**Overall Progress:** {done_unique} / {total_items} unique sentences correct ({pct}%).")
            st.markdown(
                """
                <div style="padding:10px 14px; background:#7b2ff2; color:#fff; border-radius:8px; text-align:center;">
                  ✍️ <b>Sentence Builder</b> — Click the words in the correct order!
                </div>
                """,
                unsafe_allow_html=True
            )
            st.caption(
                "Tip: Click words to build the sentence. Clear to reset, Check to submit, "
                "Next for a new one."
            )
            st.markdown(
                "**What these numbers mean:**  \n"
                "- **Score** = Correct sentences *this session*.  \n"
                "- **Progress** (bar above) = Unique sentences you have *ever* solved at this level."
            )

        # ---- Session state defaults ----
        init_defaults = {
            "sb_round": 0,
            "sb_pool": None,
            "sb_pool_level": None,
            "sb_current": None,
            "sb_shuffled": [],
            "sb_selected_idx": [],
            "sb_score": 0,
            "sb_total": 0,
            "sb_feedback": "",
            "sb_correct": None,
        }
        for k, v in init_defaults.items():
            if k not in st.session_state:
                st.session_state[k] = v

        # ---- Init / Level change ----
        if (st.session_state.sb_pool is None) or (st.session_state.sb_pool_level != student_level):
            import random
            st.session_state.sb_pool_level = student_level
            st.session_state.sb_pool = SENTENCE_BANK.get(
                student_level, SENTENCE_BANK.get("A1", [])
            ).copy()
            random.shuffle(st.session_state.sb_pool)
            st.session_state.sb_round = 0
            st.session_state.sb_score = 0
            st.session_state.sb_total = 0
            st.session_state.sb_feedback = ""
            st.session_state.sb_correct = None
            st.session_state.sb_current = None
            st.session_state.sb_selected_idx = []
            st.session_state.sb_shuffled = []

        def new_sentence():
            import random
            # Refill pool if empty
            if not st.session_state.sb_pool:
                st.session_state.sb_pool = SENTENCE_BANK.get(
                    student_level, SENTENCE_BANK.get("A1", [])
                ).copy()
                random.shuffle(st.session_state.sb_pool)
            if st.session_state.sb_pool:
                st.session_state.sb_current = st.session_state.sb_pool.pop()
                words = st.session_state.sb_current.get("tokens", [])[:]
                random.shuffle(words)
                st.session_state.sb_shuffled = words
                st.session_state.sb_selected_idx = []
                st.session_state.sb_feedback = ""
                st.session_state.sb_correct = None
                st.session_state.sb_round += 1
            else:
                st.warning("No sentences available for this level.")

        if st.session_state.sb_current is None:
            new_sentence()

        # ---- Top metrics for session ----
        cols = st.columns([3, 2, 2])
        with cols[0]:
            st.session_state.setdefault("sb_target", 5)
            _ = st.number_input(
                "Number of sentences this session",
                min_value=1, max_value=20,
                key="sb_target"
            )
        target = int(st.session_state.sb_target)
        with cols[1]:
            st.metric("Score (this session)", f"{st.session_state.sb_score}")
        with cols[2]:
            st.metric("Progress (this session)", f"{st.session_state.sb_total}/{target}")

        st.divider()

        # --- English prompt panel ---
        cur = st.session_state.sb_current or {}
        prompt_en = cur.get("prompt_en", "")
        hint_en = cur.get("hint_en", "")
        grammar_tag = cur.get("grammar_tag", "")

        if prompt_en:
            st.markdown(
                f"""
                <div style="box-sizing:border-box; padding:12px 14px; margin:6px 0 14px 0;
                            background:#f0f9ff; border:1px solid #bae6fd; border-left:6px solid #0ea5e9;
                            border-radius:10px;">
                  <div style="font-size:1.05rem;">
                    🇬🇧 <b>Translate into German:</b> <span style="color:#0b4a6f">{prompt_en}</span>
                  </div>
                </div>
                """,
                unsafe_allow_html=True
            )
            with st.expander("💡 Need a nudge? (Hint)"):
                if hint_en:
                    st.markdown(f"**Hint:** {hint_en}")
                if grammar_tag:
                    st.caption(f"Grammar: {grammar_tag}")

        # ---- Word buttons ----
        st.markdown("#### 🧩 Click the words in order")
        if st.session_state.sb_shuffled:
            word_cols = st.columns(min(6, len(st.session_state.sb_shuffled)) or 1)
            for i, w in enumerate(st.session_state.sb_shuffled):
                selected = i in st.session_state.sb_selected_idx
                btn_label = f"✅ {w}" if selected else w
                col = word_cols[i % len(word_cols)]
                with col:
                    if st.button(btn_label, key=f"sb_word_{st.session_state.sb_round}_{i}", disabled=selected):
                        st.session_state.sb_selected_idx.append(i)
                        st.rerun()

        # ---- Preview ----
        chosen_tokens = [st.session_state.sb_shuffled[i] for i in st.session_state.sb_selected_idx]
        st.markdown("#### ✨ Your sentence")
        st.code(normalize_join(chosen_tokens) if chosen_tokens else "—", language="text")

        # ---- Actions ----
        a, b, c = st.columns(3)
        with a:
            if st.button("🧹 Clear"):
                st.session_state.sb_selected_idx = []
                st.session_state.sb_feedback = ""
                st.session_state.sb_correct = None
                st.rerun()
        with b:
            if st.button("✅ Check"):
                target_sentence = st.session_state.sb_current.get("target_de", "").strip()
                chosen_sentence = normalize_join(chosen_tokens).strip()
                correct = (chosen_sentence.lower() == target_sentence.lower())
                st.session_state.sb_correct = correct
                st.session_state.sb_total += 1
                if correct:
                    st.session_state.sb_score += 1
                    st.session_state.sb_feedback = "✅ **Correct!** Great job!"
                else:
                    tip = st.session_state.sb_current.get("hint_en", "")
                    st.session_state.sb_feedback = (
                        f"❌ **Not quite.**\n\n**Correct:** {target_sentence}\n\n*Tip:* {tip}"
                    )
                save_sentence_attempt(
                    student_code=student_code,
                    level=student_level,
                    target_sentence=target_sentence,
                    chosen_sentence=chosen_sentence,
                    correct=correct,
                    tip=st.session_state.sb_current.get("hint_en", ""),
                )
                st.rerun()
        with c:
            next_disabled = (st.session_state.sb_correct is None)
            if st.button("➡️ Next", disabled=next_disabled):
                if st.session_state.sb_total >= st.session_state.sb_target:
                    st.success(f"Session complete! Score: {st.session_state.sb_score}/{st.session_state.sb_total}")
                new_sentence()
                st.rerun()

        # ---- Feedback box ----
        if st.session_state.sb_feedback:
            if st.session_state.sb_correct:
                st.success(st.session_state.sb_feedback)
            else:
                st.info(st.session_state.sb_feedback)

    # ===========================
    # SUBTAB: Vocab Practice (flashcards)
    # ===========================
    elif subtab == "Vocab Practice":
        # init session vars
        defaults = {
            "vt_history": [],
            "vt_list": [],
            "vt_index": 0,
            "vt_score": 0,
            "vt_total": None,
            "vt_saved": False,
            "vt_session_id": None,
        }
        for k, v in defaults.items():
            st.session_state.setdefault(k, v)

        # --- Stats ---
        with st.expander("📝 Your Vocab Stats", expanded=False):
            stats = get_vocab_stats(student_code)
            st.markdown(f"- **Sessions:** {stats['total_sessions']}")
            st.markdown(f"- **Last Practiced:** {stats['last_practiced']}")
            st.markdown(f"- **Unique Words:** {len(stats['completed_words'])}")

            if st.checkbox("Show Last 5 Sessions"):
                for a in stats["history"][-5:][::-1]:
                    st.markdown(
                        f"- {a['timestamp']} | {a['correct']}/{a['total']} | {a['level']}<br>"
                        f"<span style='font-size:0.9em;'>Words: {', '.join(a['practiced_words'])}</span>",
                        unsafe_allow_html=True
                    )

        # lock level
        level = student_level_locked
        items = VOCAB_LISTS.get(level, [])
        # re-use stats outside the expander
        if 'stats' not in locals():
            stats = get_vocab_stats(student_code)
        completed = set(stats["completed_words"])
        not_done = [p for p in items if p[0] not in completed]
        st.info(f"{len(not_done)} words NOT yet done at {level}.")

        # reset button
        if st.button("🔁 Start New Practice", key="vt_reset"):
            for k in defaults:
                st.session_state[k] = defaults[k]
            st.rerun()

        mode = st.radio("Select words:", ["Only new words", "All words"], horizontal=True, key="vt_mode")
        session_vocab = (not_done if mode == "Only new words" else items).copy()

        # pick count and start
        if st.session_state.vt_total is None:
            maxc = len(session_vocab)
            if maxc == 0:
                st.success("🎉 All done! Switch to 'All words' to repeat.")
                st.stop()

            count = st.number_input("How many today?", 1, maxc, min(7, maxc), key="vt_count")
            if st.button("Start", key="vt_start"):
                import random
                from uuid import uuid4
                random.shuffle(session_vocab)
                st.session_state.vt_list = session_vocab[:count]
                st.session_state.vt_total = count
                st.session_state.vt_index = 0
                st.session_state.vt_score = 0
                st.session_state.vt_history = [("assistant", f"Hallo! Ich bin Herr Felix. Let's do {count} words!")]
                st.session_state.vt_saved = False
                st.session_state.vt_session_id = str(uuid4())
                st.rerun()

        # show chat/history
        if st.session_state.vt_history:
            st.markdown("### 🗨️ Practice Chat")
            for who, msg in st.session_state.vt_history:
                render_message(who, msg)

        # practice loop
        tot = st.session_state.vt_total
        idx = st.session_state.vt_index
        if isinstance(tot, int) and idx < tot:
            word, answer = st.session_state.vt_list[idx]

            # audio
            if st.button("🔊 Play & Download", key=f"tts_{idx}"):
                try:
                    from gtts import gTTS
                    import tempfile
                    t = gTTS(text=word, lang="de")
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as fp:
                        t.save(fp.name)
                        st.audio(fp.name, format="audio/mp3")
                        fp.seek(0)
                        blob = fp.read()
                    st.download_button(
                        f"⬇️ {word}.mp3",
                        data=blob,
                        file_name=f"{word}.mp3",
                        mime="audio/mp3",
                        key=f"tts_dl_{idx}"
                    )
                except Exception as e:
                    st.error(f"Could not generate audio (gTTS): {e}")

            # bigger, bolder, clearer input
            st.markdown(
                """
                <style>
                div[data-baseweb="input"] input {
                    font-size: 18px !important;
                    font-weight: 600 !important;
                    color: black !important;
                }
                </style>
                """,
                unsafe_allow_html=True
            )

            usr = st.text_input(
                f"{word} = ?",
                key=f"vt_input_{idx}",
                placeholder="Type your answer here..."
            )

            if usr and st.button("Check", key=f"vt_check_{idx}"):
                st.session_state.vt_history.append(("user", usr))
                if is_correct_answer(usr, answer):
                    st.session_state.vt_score += 1
                    fb = f"✅ Correct! '{word}' = '{answer}'"
                else:
                    fb = f"❌ Nope. '{word}' = '{answer}'"
                st.session_state.vt_history.append(("assistant", fb))
                st.session_state.vt_index += 1
                st.rerun()

        # done
        if isinstance(tot, int) and idx >= tot:
            score = st.session_state.vt_score
            words = [w for w, _ in (st.session_state.vt_list or [])]
            st.markdown(f"### 🏁 Done! You scored {score}/{tot}.")

            # Save exactly once per session, duplicate-safe
            if not st.session_state.get("vt_saved", False):
                if not st.session_state.get("vt_session_id"):
                    from uuid import uuid4
                    st.session_state.vt_session_id = str(uuid4())
                if not vocab_attempt_exists(student_code, st.session_state.vt_session_id):
                    save_vocab_attempt(
                        student_code=student_code,
                        level=level,
                        total=tot,
                        correct=score,
                        practiced_words=words,
                        session_id=st.session_state.vt_session_id
                    )
                st.session_state.vt_saved = True
                st.rerun()

            if st.button("Practice Again", key="vt_again"):
                for k in defaults:
                    st.session_state[k] = defaults[k]
                st.rerun()

    # ===========================
    # SUBTAB: Dictionary (simple, sticky search, mobile-friendly)
    # ===========================
    elif subtab == "Dictionary":
        import io, json, difflib

        # ---------- Helpers ----------
        def _fallback_df(levels):
            rows = []
            for lvl in levels:
                for de, en in VOCAB_LISTS.get(lvl, []):
                    rows.append({"Level": lvl, "German": de, "English": en, "Pronunciation": ""})
            return pd.DataFrame(rows) if rows else pd.DataFrame(columns=["Level","German","English","Pronunciation"])

        def _merge_sentence_bank(df, levels):
            extra = []
            for lvl in levels:
                for item in SENTENCE_BANK.get(lvl, []):
                    for tok in item.get("tokens", []):
                        t = str(tok).strip()
                        if not t or t in [",", ".", "!", "?", ":", ";"]:
                            continue
                        if not ((df["German"] == t) & (df["Level"] == lvl)).any():
                            extra.append({"Level": lvl, "German": t, "English": "", "Pronunciation": ""})
            if extra:
                df = pd.concat([df, pd.DataFrame(extra)], ignore_index=True)
                df = df.drop_duplicates(subset=["Level","German"]).reset_index(drop=True)
            return df

        def _tts_bytes_de(text: str) -> bytes:
            try:
                from gtts import gTTS
                buf = io.BytesIO()
                gTTS(text=text, lang="de").write_to_fp(buf)
                buf.seek(0)
                return buf.read()
            except Exception:
                return b""

        def _json_from_text(raw: str) -> dict:
            txt = (raw or "").strip()
            if txt.startswith("```"):
                txt = txt.strip("`")
                if "\n" in txt:
                    txt = txt.split("\n", 1)[1]
                if "```" in txt:
                    txt = txt.split("```", 1)[0]
            try:
                return json.loads(txt)
            except Exception:
                return {}

        def _enrich_word(german: str, english_hint: str, level: str):
            """Fill Pronunciation + English + 2 examples if missing (quiet background call)."""
            try:
                prompt = (
                    "You are a precise German lexicographer.\n"
                    f'Word: "{german}"\n'
                    f'Known English hint (may be empty): "{english_hint}"\n'
                    f'Level: {level}\n\n'
                    "Return compact JSON with keys: ipa, english, examples (2 items with keys de and en)."
                )
                resp = client.chat.completions.create(
                    model="gpt-4o",
                    temperature=0.2,
                    max_tokens=220,
                    messages=[
                        {"role": "system", "content": "Return strict JSON only."},
                        {"role": "user", "content": prompt},
                    ],
                )
                data = _json_from_text(resp.choices[0].message.content)
                ipa = str(data.get("ipa", "") or "")
                eng = str(data.get("english", "") or english_hint or "")
                exs = data.get("examples", []) or []
                clean = []
                for ex in exs[:2]:
                    clean.append({
                        "de": str(ex.get("de", "") or ""),
                        "en": str(ex.get("en", "") or "")
                    })
                return {"pron": ipa, "english": eng, "examples": clean}
            except Exception:
                return {"pron": "", "english": english_hint or "", "examples": []}

        # diacritic/umlaut normalization so "ae" finds "ä", etc.
        _map = {"ä":"ae","ö":"oe","ü":"ue","ß":"ss"}
        def _norm(s: str) -> str:
            s = (s or "").strip().lower()
            for k,v in _map.items():
                s = s.replace(k, v)
            return "".join(ch for ch in s if ch.isalnum() or ch.isspace())

        # ---------- Build data (CSV + Sentence Bank) ----------
        levels = [student_level_locked]
        df_dict = _fallback_df(levels)
        df_dict = _merge_sentence_bank(df_dict, levels)
        for c in ["Level","German","English","Pronunciation"]:
            if c not in df_dict.columns:
                df_dict[c] = ""
        df_dict["g_norm"] = df_dict["German"].astype(str).map(_norm)
        df_dict["e_norm"] = df_dict["English"].astype(str).map(_norm)
        df_dict = df_dict.sort_values(["German"]).reset_index(drop=True)

        # ---------- Mobile-friendly sticky search ----------
        st.markdown(
            """
            <style>
              .sticky-search { position: sticky; top: 0; z-index: 999; background: white; padding: 8px 0 10px 0; }
              input[type="text"] { font-size: 18px !important; }
              .chip { display:inline-block; padding:6px 10px; border-radius:999px; border:1px solid #e5e7eb; margin-right:6px; margin-bottom:6px; }
            </style>
            """,
            unsafe_allow_html=True
        )
        with st.container():
            st.markdown('<div class="sticky-search">', unsafe_allow_html=True)
            cols = st.columns([6, 3, 3])
            with cols[0]:
                q = st.text_input("🔎 Search (German or English)", key="dict_q", placeholder="e.g., Wochenende, bakery, spielen")
            with cols[1]:
                search_in = st.selectbox("Field", ["Both", "German", "English"], index=0, key="dict_field")
            with cols[2]:
                match_mode = st.selectbox("Match", ["Contains", "Starts with", "Exact"], index=0, key="dict_mode")
            st.markdown('</div>', unsafe_allow_html=True)

        # ---------- Filter (and add enrichment when empty) ----------
        df_view = df_dict.copy()
        suggestions = []
        top_row = None

        if q:
            qn = _norm(q)

            # masks by field
            g_contains = df_view["g_norm"].str.contains(qn, na=False) if search_in in ("Both","German") else pd.Series([False]*len(df_view))
            g_starts   = df_view["g_norm"].str.startswith(qn, na=False) if search_in in ("Both","German") else pd.Series([False]*len(df_view))
            g_exact    = df_view["g_norm"].eq(qn) if search_in in ("Both","German") else pd.Series([False]*len(df_view))

            e_contains = df_view["e_norm"].str.contains(qn, na=False) if search_in in ("Both","English") else pd.Series([False]*len(df_view))
            e_starts   = df_view["e_norm"].str.startswith(qn, na=False) if search_in in ("Both","English") else pd.Series([False]*len(df_view))
            e_exact    = df_view["e_norm"].eq(qn) if search_in in ("Both","English") else pd.Series([False]*len(df_view))

            if match_mode == "Contains":
                mask = g_contains | e_contains
            elif match_mode == "Starts with":
                mask = g_starts | e_starts
            else:
                mask = g_exact | e_exact

            if mask.any():
                df_view = df_view[mask].copy().reset_index(drop=True)
                # prefer exact > starts > contains
                exact_mask  = (g_exact | e_exact)
                starts_mask = (g_starts | e_starts)
                if exact_mask.any():
                    top_row = df_view[exact_mask].iloc[0]
                elif starts_mask.any():
                    top_row = df_view[starts_mask].iloc[0]
                else:
                    top_row = df_view.iloc[0]
            else:
                # no local match → show fuzzy suggestions + enrich the query so learners still get value
                vocab_all = df_view["German"].astype(str).unique().tolist()
                suggestions = difflib.get_close_matches(q, vocab_all, n=5, cutoff=0.72)

                enrich = _enrich_word(q, "", student_level_locked)
                new_row = {
                    "Level": student_level_locked,
                    "German": q.capitalize() if q.islower() else q,
                    "English": enrich.get("english", ""),
                    "Pronunciation": enrich.get("pron", ""),
                    "g_norm": _norm(q),
                    "e_norm": _norm(enrich.get("english","")),
                }
                df_view = pd.concat([df_view, pd.DataFrame([new_row])], ignore_index=True)
                top_row = pd.Series(new_row)
                st.session_state.setdefault("dict_cache", {})
                st.session_state["dict_cache"][(new_row["German"], student_level_locked)] = {
                    "pron": new_row["Pronunciation"],
                    "english": new_row["English"],
                    "examples": enrich.get("examples", []),
                }
        else:
            # no query → show first word (nice landing state) if available
            if not df_view.empty:
                top_row = df_view.iloc[0]

        # ---------- Details (always ABOVE) ----------
        if "dict_cache" not in st.session_state:
            st.session_state["dict_cache"] = {}

        if top_row is not None and len(top_row) > 0:
            de  = str(top_row["German"])
            en  = str(top_row.get("English", "") or "")
            lvl = str(top_row.get("Level", student_level_locked))
            pron = str(top_row.get("Pronunciation", "") or "")

            cache_key = (de, lvl)
            cached = st.session_state["dict_cache"].get(cache_key, {})

            # Backfill pronunciation/english/examples if missing
            if not pron or not cached.get("examples"):
                enrich = _enrich_word(de, en, lvl)
                if not pron and enrich.get("pron"):
                    pron = enrich["pron"]
                if not en and enrich.get("english"):
                    en = enrich["english"]
                if enrich.get("examples"):
                    cached["examples"] = enrich["examples"]
                st.session_state["dict_cache"][cache_key] = {
                    "pron": pron, "english": en, "examples": cached.get("examples", [])
                }

            examples = st.session_state["dict_cache"].get(cache_key, {}).get("examples", [])

            st.markdown(f"### {de}")
            if en:
                st.markdown(f"**Meaning:** {en}")
            if pron:
                st.caption(f"**Pronunciation:** /{pron}/")

            audio_bytes = _tts_bytes_de(de)
            c1, c2 = st.columns([1, 2])
            with c1:
                if st.button("🔊 Pronounce", key=f"say_{de}_{lvl}"):
                    if audio_bytes:
                        st.audio(audio_bytes, format="audio/mp3")
            with c2:
                if audio_bytes:
                    st.download_button(
                        "⬇️ Download MP3",
                        data=audio_bytes,
                        file_name=f"{de}.mp3",
                        mime="audio/mpeg",
                        key=f"dl_{de}_{lvl}"
                    )
                else:
                    st.caption("Audio currently unavailable.")

            with st.expander("📌 Examples", expanded=True):
                if examples:
                    for ex in examples[:2]:
                        de_ex = (ex.get("de", "") or "").strip()
                        en_ex = (ex.get("en", "") or "").strip()
                        if de_ex:
                            st.markdown(f"- **{de_ex}**")
                            if en_ex:
                                st.caption(f"  ↳ {en_ex}")
                else:
                    st.caption("No examples yet.")

        # ---------- Did you mean (chips) ----------
        if q and suggestions:
            st.markdown("**Did you mean:**")
            bcols = st.columns(min(5, len(suggestions)))
            for i, s in enumerate(suggestions[:5]):
                with bcols[i]:
                    if st.button(s, key=f"sugg_{i}"):
                        st.session_state["dict_q"] = s
                        st.rerun()

        # ---------- Scrollable table INSIDE an expander (clean page) ----------
        with st.expander(f"Browse all words at level {student_level_locked}", expanded=False):
            df_show = df_view[["German","English","Pronunciation"]].copy()
            st.dataframe(df_show, use_container_width=True, height=420)
#

                

# ===== BUBBLE FUNCTION FOR CHAT DISPLAY =====
def bubble(role, text):
    color = "#7b2ff2" if role == "assistant" else "#222"
    bg = "#ede3fa" if role == "assistant" else "#f6f8fb"
    name = "Herr Felix" if role == "assistant" else "You"
    return f"""
        <div style="background:{bg};color:{color};margin-bottom:8px;padding:13px 15px;
        border-radius:14px;max-width:98vw;font-size:1.09rem;">
            <b>{name}:</b><br>{text}
        </div>
    """



# ===== Schreiben =====

db = firestore.client()

# -- Feedback HTML Highlight Helper --
highlight_words = ["correct", "should", "mistake", "improve", "tip"]

def highlight_feedback(text: str) -> str:
    # 1) Highlight “[correct]…[/correct]” spans in green
    text = re.sub(
        r"\[correct\](.+?)\[/correct\]",
        r"<span style="
        r"'background-color:#d4edda;"
        r"color:#155724;"
        r"border-radius:4px;"
        r"padding:2px 6px;"
        r"margin:0 2px;"
        r"font-weight:600;'"
        r">\1</span>",
        text,
        flags=re.DOTALL
    )

    # 2) Highlight “[wrong]…[/wrong]” spans in red with strikethrough
    text = re.sub(
        r"\[wrong\](.+?)\[/wrong\]",
        r"<span style="
        r"'background-color:#f8d7da;"
        r"color:#721c24;"
        r"border-radius:4px;"
        r"padding:2px 6px;"
        r"margin:0 2px;"
        r"text-decoration:line-through;"
        r"font-weight:600;'"
        r">\1</span>",
        text,
        flags=re.DOTALL
    )

    # 3) Bold keywords
    def repl_kw(m):
        return f"<strong style='color:#d63384'>{m.group(1)}</strong>"
    pattern = r"\b(" + "|".join(map(re.escape, highlight_words)) + r")\b"
    text = re.sub(pattern, repl_kw, text, flags=re.IGNORECASE)

    # 4) Restyle the final breakdown block as a simple, transparent list
    def _format_breakdown(m):
        lines = [line.strip() for line in m.group(0).splitlines() if line.strip()]
        items = "".join(f"<li style='margin-bottom:4px'>{line}</li>" for line in lines)
        return (
            "<ul style='margin:8px 0 12px 1em;"
            "padding:0;"
            "list-style:disc inside;"
            "font-size:0.95em;'>"
            f"{items}"
            "</ul>"
        )

    text = re.sub(
        r"(Grammar:.*?\nVocabulary:.*?\nSpelling:.*?\nStructure:.*)",
        _format_breakdown,
        text,
        flags=re.DOTALL
    )

    return text

# -- Firestore-only: Usage Limit (Daily Mark My Letter) --
def get_schreiben_usage(student_code):
    today = str(date.today())
    doc = db.collection("schreiben_usage").document(f"{student_code}_{today}").get()
    return doc.to_dict().get("count", 0) if doc.exists else 0

def inc_schreiben_usage(student_code):
    today = str(date.today())
    doc_ref = db.collection("schreiben_usage").document(f"{student_code}_{today}")
    doc = doc_ref.get()
    if doc.exists:
        doc_ref.update({"count": firestore.Increment(1)})
    else:
        doc_ref.set({"student_code": student_code, "date": today, "count": 1})

# -- Firestore-only: Submission + Full letter (Saves for feedback & stats) --
def save_submission(student_code: str, score: int, passed: bool, timestamp, level: str, letter: str):
    payload = {
        "student_code": student_code,
        "score": score,
        "passed": passed,
        "date": firestore.SERVER_TIMESTAMP,  # Always use server time
        "level": level,
        "assignment": "Schreiben Trainer",
        "letter": letter,
    }
    db.collection("schreiben_submissions").add(payload)

# -- Firestore-only: Recalculate All Schreiben Stats (called after every submission) --
def update_schreiben_stats(student_code: str):
    """
    Recalculates stats for a student after every submission.
    """
    submissions = db.collection("schreiben_submissions").where(
        "student_code", "==", student_code
    ).stream()

    total = 0
    passed = 0
    scores = []
    last_letter = ""
    last_attempt = None

    for doc in submissions:
        data = doc.to_dict()
        total += 1
        score = data.get("score", 0)
        scores.append(score)
        if data.get("passed"):
            passed += 1
        last_letter = data.get("letter", "") or last_letter
        last_attempt = data.get("date", last_attempt)

    pass_rate = (passed / total * 100) if total > 0 else 0
    best_score = max(scores) if scores else 0
    average_score = sum(scores) / total if scores else 0

    stats_ref = db.collection("schreiben_stats").document(student_code)
    stats_ref.set({
        "total": total,
        "passed": passed,
        "pass_rate": pass_rate,
        "best_score": best_score,
        "average_score": average_score,
        "last_attempt": last_attempt,
        "last_letter": last_letter,
        "attempts": scores
    }, merge=True)

# -- Firestore-only: Fetch stats for display (for status panel etc) --
def get_schreiben_stats(student_code: str):
    stats_ref = db.collection("schreiben_stats").document(student_code)
    doc = stats_ref.get()
    if doc.exists:
        return doc.to_dict()
    else:
        return {
            "total": 0, "passed": 0, "average_score": 0, "best_score": 0,
            "pass_rate": 0, "last_attempt": None, "attempts": [], "last_letter": ""
        }

# -- Firestore-only: Usage Limit (Daily Letter Coach) --
def get_letter_coach_usage(student_code):
    today = str(date.today())
    doc = db.collection("letter_coach_usage").document(f"{student_code}_{today}").get()
    return doc.to_dict().get("count", 0) if doc.exists else 0

def inc_letter_coach_usage(student_code):
    today = str(date.today())
    doc_ref = db.collection("letter_coach_usage").document(f"{student_code}_{today}")
    doc = doc_ref.get()
    if doc.exists:
        doc_ref.update({"count": firestore.Increment(1)})
    else:
        doc_ref.set({"student_code": student_code, "date": today, "count": 1})

# -- Firestore: Save/load Letter Coach progress --
def save_letter_coach_progress(student_code, level, prompt, chat):
    doc_ref = db.collection("letter_coach_progress").document(student_code)
    doc_ref.set({
        "student_code": student_code,
        "level": level,
        "prompt": prompt,
        "chat": chat,
        "date": firestore.SERVER_TIMESTAMP
    })

def load_letter_coach_progress(student_code):
    doc = db.collection("letter_coach_progress").document(student_code).get()
    if doc.exists:
        data = doc.to_dict()
        return data.get("prompt", ""), data.get("chat", [])
    else:
        return "", []


# --- Helper: Get level from Google Sheet (public CSV) ---

SHEET_URL = "https://docs.google.com/spreadsheets/d/12NXf5FeVHr7JJT47mRHh7Jp-TC1yhPS7ZG6nzZVTt1U/export?format=csv"

@st.cache_data(ttl=300)
def load_sheet():
    return pd.read_csv(SHEET_URL)

def get_level_from_code(student_code):
    df = load_sheet()
    student_code = str(student_code).strip().lower()
    # Make sure 'StudentCode' column exists and is lowercase
    if "StudentCode" not in df.columns:
        df.columns = [c.strip() for c in df.columns]
    if "StudentCode" in df.columns:
        matches = df[df["StudentCode"].astype(str).str.strip().str.lower() == student_code]
        if not matches.empty:
            # Handles NaN, empty cells
            level = matches.iloc[0]["Level"]
            return str(level).strip().upper() if pd.notna(level) else "A1"
    return "A1"




#Maincode for me

if tab == "Schreiben Trainer":
    st.markdown(
        '''
        <div style="
            padding: 8px 12px;
            background: #d63384;
            color: #fff;
            border-radius: 6px;
            text-align: center;
            margin-bottom: 8px;
            font-size: 1.3rem;">
            ✍️ Schreiben Trainer (Writing Practice)
        </div>
        ''',
        unsafe_allow_html=True
    )

    st.info(
        """
        ✍️ **This section is for Writing (Schreiben) only.**
        - Practice your German letters, emails, and essays for A1–C1 exams.
        - **Want to prepare for class presentations, topic expansion, or practice Speaking, Reading (Lesen), or Listening (Hören)?**  
          👉 Go to **Exam Mode & Custom Chat** (tab above)!
        - **Tip:** Choose your exam level on the right before submitting your letter. Your writing will be checked and scored out of 25 marks, just like in the real exam.
        """,
        icon="✉️"
    )

    st.divider()

    # --- Writing stats summary with Firestore ---
    student_code = st.session_state.get("student_code", "demo")
    stats = get_schreiben_stats(student_code)
    if stats:
        total = stats.get("total", 0)
        passed = stats.get("passed", 0)
        pass_rate = stats.get("pass_rate", 0)

        # Milestone and title logic
        if total <= 2:
            writer_title = "🟡 Beginner Writer"
            milestone = "Write 3 letters to become a Rising Writer!"
        elif total <= 5 or pass_rate < 60:
            writer_title = "🟡 Rising Writer"
            milestone = "Achieve 60% pass rate and 6 letters to become a Confident Writer!"
        elif total <= 7 or (60 <= pass_rate < 80):
            writer_title = "🔵 Confident Writer"
            milestone = "Reach 8 attempts and 80% pass rate to become an Advanced Writer!"
        elif total >= 8 and pass_rate >= 80 and not (total >= 10 and pass_rate >= 95):
            writer_title = "🟢 Advanced Writer"
            milestone = "Reach 10 attempts and 95% pass rate to become a Master Writer!"
        elif total >= 10 and pass_rate >= 95:
            writer_title = "🏅 Master Writer!"
            milestone = "You've reached the highest milestone! Keep maintaining your skills 🎉"
        else:
            writer_title = "✏️ Active Writer"
            milestone = "Keep going to unlock your next milestone!"

        st.markdown(
            f"""
            <div style="background:#fff8e1;padding:18px 12px 14px 12px;border-radius:12px;margin-bottom:12px;
                        box-shadow:0 1px 6px #00000010;">
                <span style="font-weight:bold;font-size:1.25rem;color:#d63384;">{writer_title}</span><br>
                <span style="font-weight:bold;font-size:1.09rem;color:#444;">📊 Your Writing Stats</span><br>
                <span style="color:#202020;font-size:1.05rem;"><b>Total Attempts:</b> {total}</span><br>
                <span style="color:#202020;font-size:1.05rem;"><b>Passed:</b> {passed}</span><br>
                <span style="color:#202020;font-size:1.05rem;"><b>Pass Rate:</b> {pass_rate:.1f}%</span><br>
                <span style="color:#e65100;font-weight:bold;font-size:1.03rem;">{milestone}</span>
            </div>
            """,
            unsafe_allow_html=True
        )
    else:
        st.info("No writing stats found yet. Write your first letter to see progress!")

    # --- Update session states for new student (preserves drafts, etc) ---
    prev_student_code = st.session_state.get("prev_student_code", None)
    if student_code != prev_student_code:
        stats = get_schreiben_stats(student_code)
        st.session_state[f"{student_code}_schreiben_input"] = stats.get("last_letter", "")
        st.session_state[f"{student_code}_last_feedback"] = None
        st.session_state[f"{student_code}_last_user_letter"] = None
        st.session_state[f"{student_code}_delta_compare_feedback"] = None
        st.session_state[f"{student_code}_final_improved_letter"] = ""
        st.session_state[f"{student_code}_awaiting_correction"] = False
        st.session_state[f"{student_code}_improved_letter"] = ""
        st.session_state["prev_student_code"] = student_code

    # --- Sub-tabs for the Trainer ---
    sub_tab = st.radio(
        "Choose Mode",
        ["Mark My Letter", "Ideas Generator (Letter Coach)"],
        horizontal=True,
        key=f"schreiben_sub_tab_{student_code}"
    )

        # --- Level picker: Auto-detect from student code (manual override removed) ---
    if student_code:
        detected_level = get_level_from_code(student_code)
        # Only apply detected level when first seeing this student code
        if st.session_state.get("prev_student_code_for_level") != student_code:
            st.session_state["schreiben_level"] = detected_level
            st.session_state["prev_student_code_for_level"] = student_code
    else:
        detected_level = "A1"
        if "schreiben_level" not in st.session_state:
            st.session_state["schreiben_level"] = detected_level

    # Ensure current writing level variable reflects auto-detected one
    schreiben_level = st.session_state.get("schreiben_level", "A1")

    st.markdown(
        f"<span style='color:gray;font-size:0.97em;'>Auto-detected level from your code: <b>{detected_level}</b></span>",
        unsafe_allow_html=True
    )


    st.divider()

    # ----------- 1. MARK MY LETTER -----------
    if sub_tab == "Mark My Letter":
        MARK_LIMIT = 3
        daily_so_far = get_schreiben_usage(student_code)
        st.markdown(f"**Daily usage:** {daily_so_far} / {MARK_LIMIT}")

        user_letter = st.text_area(
            "Paste or type your German letter/essay here.",
            key=f"{student_code}_schreiben_input",
            value=st.session_state.get(f"{student_code}_schreiben_input", ""),
            disabled=(daily_so_far >= MARK_LIMIT),
            height=400,
            placeholder="Write your German letter here..."
        )

        # AUTOSAVE LOGIC (save every edit that's different from last_letter)
        if (
            user_letter.strip() and
            user_letter != get_schreiben_stats(student_code).get("last_letter", "")
        ):
            doc_ref = db.collection("schreiben_stats").document(student_code)
            doc = doc_ref.get()
            data = doc.to_dict() if doc.exists else {}
            data["last_letter"] = user_letter
            doc_ref.set(data, merge=True)

        # --- Word count and Goethe exam rules ---
        import re
        def get_level_requirements(level):
            reqs = {
                "A1": {"min": 25, "max": 40, "desc": "A1 formal/informal letters should be 25–40 words. Cover all bullet points."},
                "A2": {"min": 30, "max": 40, "desc": "A2 formal/informal letters should be 30–40 words. Cover all bullet points."},
                "B1": {"min": 80, "max": 150, "desc": "B1 letters/essays should be about 80–150 words, with all points covered and clear structure."},
                "B2": {"min": 150, "max": 250, "desc": "B2 essays are 180–220 words, opinion essays or reports, with good structure and connectors."},
                "C1": {"min": 230, "max": 350, "desc": "C1 essays are 230–250+ words. Use advanced structures and express opinions clearly."}
            }
            return reqs.get(level.upper(), reqs["A1"])

        def count_words(text):
            return len(re.findall(r'\b\w+\b', text))

        if user_letter.strip():
            words = re.findall(r'\b\w+\b', user_letter)
            chars = len(user_letter)
            st.info(f"**Word count:** {len(words)} &nbsp;|&nbsp; **Character count:** {chars}")

            # -- Apply Goethe writing rules here --
            requirements = get_level_requirements(detected_level)  # << USE AUTO-DETECTED LEVEL
            word_count = count_words(user_letter)
            min_wc = requirements["min"]
            max_wc = requirements["max"]

            if detected_level in ("A1", "A2"):
                if word_count < min_wc:
                    st.error(f"⚠️ Your letter is too short for {detected_level} ({word_count} words). {requirements['desc']}")
                    st.stop()
                elif word_count > max_wc:
                    st.warning(f"ℹ️ Your letter is a bit long for {detected_level} ({word_count} words). The exam expects {min_wc}-{max_wc} words.")
            else:
                if word_count < min_wc:
                    st.error(f"⚠️ Your essay is too short for {detected_level} ({word_count} words). {requirements['desc']}")
                    st.stop()
                elif word_count > max_wc + 40 and detected_level in ("B1", "B2"):
                    st.warning(f"ℹ️ Your essay is longer than the usual limit for {detected_level} ({word_count} words). Try to stay within the guidelines.")

        # --------- Reset correction states (do not indent inside above ifs)
        for k, v in [
            ("last_feedback", None),
            ("last_user_letter", None),
            ("delta_compare_feedback", None),
            ("improved_letter", ""),
            ("awaiting_correction", False),
            ("final_improved_letter", "")
        ]:
            session_key = f"{student_code}_{k}"
            if session_key not in st.session_state:
                st.session_state[session_key] = v

        # Namespaced correction state per student (reset on session)
        for k, v in [
            ("last_feedback", None),
            ("last_user_letter", None),
            ("delta_compare_feedback", None),
            ("improved_letter", ""),
            ("awaiting_correction", False),
            ("final_improved_letter", "")
        ]:
            session_key = f"{student_code}_{k}"
            if session_key not in st.session_state:
                st.session_state[session_key] = v

        submit_disabled = daily_so_far >= MARK_LIMIT or not user_letter.strip()
        feedback_btn = st.button(
            "Get Feedback",
            type="primary",
            disabled=submit_disabled,
            key=f"feedback_btn_{student_code}"
        )

        if feedback_btn:
            st.session_state[f"{student_code}_awaiting_correction"] = True
            ai_prompt = (
                f"You are Herr Felix, a supportive and innovative German letter writing trainer.\n"
                f"You help students prepare for A1, A2, B1, B2, and C1 German exam letters or essays.\n"
                f"The student has submitted a {schreiben_level} German letter or essay.\n"
                f"Your job is to mark, score, and explain feedback in a kind, step-by-step way.\n"
                f"Always answer in English.\n"
                f"1. Give a quick summary (one line) of how well the student did overall.\n"
                f"2. Then show a detailed breakdown of strengths and weaknesses in 4 areas:\n"
                f"   Grammar, Vocabulary, Spelling, Structure.\n"
                f"3. For each area, say what was good and what should improve.\n"
                f"4. Highlight every mistake with [wrong]...[/wrong] and every good example with [correct]...[/correct].\n"
                f"5. Give 2-3 improvement tips in bullet points.\n"
                f"6. At the end, give a realistic score out of 25 in the format: Score: X/25.\n"
                f"7. For A1 and A2, be strict about connectors, basic word order, modal verbs, and correct formal/informal greeting.\n"
                f"8. For B1+, mention exam criteria and what examiner wants.\n"
                f"9. Never write a new letter for the student, only mark what they submit.\n"
                f"10. When possible, point out specific lines or examples from their letter in your feedback.\n"
                f"11. When student score is 18 or above then they have passed. When score is less than 18, is a fail and they must try again before submitting to prevent low marks.\n"
                f"12. After completion, remind them to only copy their improved letter without your feedback, go to 'my course' on the app and submit together with their lesen and horen answers. They only share the letter and feedback with their teacher for evaluation only when they preparing for the exams\n"
                
            )

            with st.spinner("🧑‍🏫 Herr Felix is typing..."):
                try:
                    completion = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "system", "content": ai_prompt},
                            {"role": "user", "content": user_letter},
                        ],
                        temperature=0.6,
                    )
                    feedback = completion.choices[0].message.content
                    st.session_state[f"{student_code}_last_feedback"] = feedback
                    st.session_state[f"{student_code}_last_user_letter"] = user_letter
                    st.session_state[f"{student_code}_delta_compare_feedback"] = None
                except Exception as e:
                    st.error("AI feedback failed. Please check your OpenAI setup.")
                    feedback = None

            if feedback:
                inc_schreiben_usage(student_code)
                st.markdown("---")
                st.markdown("#### 📝 Feedback from Herr Felix")
                st.markdown(highlight_feedback(feedback), unsafe_allow_html=True)
                st.session_state[f"{student_code}_awaiting_correction"] = True

                # --- Save to Firestore ---
                score_match = re.search(r"Score[: ]+(\d+)", feedback)
                score = int(score_match.group(1)) if score_match else 0
                passed = score >= 17
                save_submission(
                    student_code=student_code,
                    score=score,
                    passed=passed,
                    timestamp=None,  # Not needed
                    level=schreiben_level,
                    letter=user_letter
                )
                update_schreiben_stats(student_code)



        # --- Improvement section: Compare, download, WhatsApp ---
        if st.session_state.get(f"{student_code}_last_feedback") and st.session_state.get(f"{student_code}_last_user_letter"):
            st.markdown("---")
            st.markdown("#### 📝 Feedback from Herr Felix (Reference)")
            st.markdown(
                highlight_feedback(st.session_state[f"{student_code}_last_feedback"]),
                unsafe_allow_html=True
            )
            st.markdown(
                """
                <div style="background:#e3f7da; border-left:7px solid #44c767; 
                color:#295327; padding:1.15em; margin-top:1em; border-radius:10px; font-size:1.09em;">
                    🔁 <b>Try to improve your letter!</b><br>
                    Paste your improved version below and click <b>Compare My Improvement</b>.<br>
                    The AI will highlight what’s better, what’s still not fixed, and give extra tips.<br>
                    <b>You can download or share the improved version & new feedback below.</b>
                </div>
                """, unsafe_allow_html=True
            )
            improved_letter = st.text_area(
                "Your improved version (try to fix the mistakes Herr Felix mentioned):",
                key=f"{student_code}_improved_letter",
                height=400,
                placeholder="Paste your improved letter here..."
            )
            compare_clicked = st.button("Compare My Improvement", key=f"compare_btn_{student_code}")

            if compare_clicked and improved_letter.strip():
                ai_compare_prompt = (
                    "You are Herr Felix, a supportive German writing coach. "
                    "A student first submitted this letter:\n\n"
                    f"{st.session_state[f'{student_code}_last_user_letter']}\n\n"
                    "Your feedback was:\n"
                    f"{st.session_state[f'{student_code}_last_feedback']}\n\n"
                    "Now the student has submitted an improved version below.\n"
                    "Compare both versions and:\n"
                    "- Tell the student exactly what they improved, and which mistakes were fixed.\n"
                    "- Point out if there are still errors left, with new tips for further improvement.\n"
                    "- Encourage the student. If the improvement is significant, say so.\n"
                    "1. If student dont improve after the third try, end the chat politely and tell the student to try again tomorrow. Dont continue to give the feedback after third try.\n"
                    "2. Always explain your feeback in English for them to understand. You can still highlight their german phrases. But your correction should be english\n"
                    "3. For A1 and A2 students, make sure a sentence is not more than 7 words."
                    "4. For A1 and A2 students, break their phrases down for them when they use relative clauses."
                    "5. For A1 and A2 students, only recommend connectors such as deshalb, weil, ich mochte wissen,und,oder."
                    "- Give a revised score out of 25 (Score: X/25)."
                )
                with st.spinner("👨‍🏫 Herr Felix is comparing your improvement..."):
                    try:
                        result = client.chat.completions.create(
                            model="gpt-4o",
                            messages=[
                                {"role": "system", "content": ai_compare_prompt},
                                {"role": "user", "content": improved_letter}
                            ],
                            temperature=0.5,
                        )
                        compare_feedback = result.choices[0].message.content
                        st.session_state[f"{student_code}_delta_compare_feedback"] = compare_feedback
                        st.session_state[f"{student_code}_final_improved_letter"] = improved_letter
                    except Exception as e:
                        st.session_state[f"{student_code}_delta_compare_feedback"] = f"Sorry, there was an error comparing your letters: {e}"

            if st.session_state.get(f"{student_code}_delta_compare_feedback"):
                st.markdown("---")
                st.markdown("### 📝 Improvement Feedback from Herr Felix")
                st.markdown(highlight_feedback(st.session_state[f"{student_code}_delta_compare_feedback"]), unsafe_allow_html=True)

                # PDF & WhatsApp buttons
                from fpdf import FPDF
                import urllib.parse, os

                def sanitize_text(text):
                    return text.encode('latin-1', errors='replace').decode('latin-1')

                # PDF
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                improved_letter = st.session_state.get(f"{student_code}_final_improved_letter", "")
                improved_feedback = st.session_state[f"{student_code}_delta_compare_feedback"]
                pdf.multi_cell(0, 10, f"Your Improved Letter:\n\n{sanitize_text(improved_letter)}\n\nFeedback from Herr Felix:\n\n{sanitize_text(improved_feedback)}")
                pdf_output = f"Feedback_{student_code}_{schreiben_level}_improved.pdf"
                pdf.output(pdf_output)
                with open(pdf_output, "rb") as f:
                    pdf_bytes = f.read()
                st.download_button(
                    "⬇️ Download Improved Version + Feedback (PDF)",
                    pdf_bytes,
                    file_name=pdf_output,
                    mime="application/pdf"
                )
                os.remove(pdf_output)

                # WhatsApp share
                wa_message = (
                    f"Hi, here is my IMPROVED German letter and AI feedback:\n\n"
                    f"{improved_letter}\n\n"
                    f"Feedback:\n{st.session_state[f'{student_code}_delta_compare_feedback']}"
                )
                wa_url = (
                    "https://api.whatsapp.com/send"
                    "?phone=233205706589"
                    f"&text={urllib.parse.quote(wa_message)}"
                )
                st.markdown(
                    f"[📲 Send Improved Letter & Feedback to Tutor on WhatsApp]({wa_url})",
                    unsafe_allow_html=True
                )


    if sub_tab == "Ideas Generator (Letter Coach)":
        import io

        # === NAMESPACED SESSION KEYS (per student) ===
        student_code = st.session_state.get("student_code", "demo")
        ns_prefix = f"{student_code}_letter_coach_"
        def ns(key): return ns_prefix + key

        # --- Reset per-student Letter Coach state on student change ---
        prev_letter_coach_code = st.session_state.get("prev_letter_coach_code", None)
        if student_code != prev_letter_coach_code:
            last_prompt, last_chat = load_letter_coach_progress(student_code)
            st.session_state[ns("prompt")] = last_prompt or ""
            st.session_state[ns("chat")] = last_chat or []
            st.session_state[ns("stage")] = 1 if last_chat else 0
            st.session_state["prev_letter_coach_code"] = student_code

        # --- Set per-student defaults if missing ---
        for k, default in [("prompt", ""), ("chat", []), ("stage", 0)]:
            if ns(k) not in st.session_state:
                st.session_state[ns(k)] = default


        LETTER_COACH_PROMPTS = {
            "A1": (
                "You are Herr Felix, a creative and supportive German letter-writing coach for A1 students. "
                "Always reply in English, never in German. "
                "When a student submits something, first congratulate them with ideas about how to go about the letter. "
                "Analyze if their message is a new prompt, a continuation, or a question. "
                "If it's a question, answer simply and encourage them to keep building their letter step by step. "
                "If it's a continuation, review their writing so far and guide them to the next step. "
                "    1. Always give students short ideas, structure and tips and phrases on how to build their points for the conversation in English and simple German. Don't overfeed students, help them but let them think by themselves also. "
                "    2. For conjunctions, only suggest 'weil', 'deshalb', 'ich möchte wissen, ob' and 'ich möchte wissen, wann'. Don't recommend 'da', 'dass' and relative clauses. "
                "    3. For requests, teach them how to use 'Könnten Sie...' and how it ends with a main verb to make a request when necessary. "
                "    4. For formal/informal letter: guide them to use 'Ich schreibe Ihnen/dir...', and show how to use 'weil' with 'ich' and end with only 'möchte' to prevent mistakes. Be strict with this. "
                "    5. Always check that the student statement is not too long or complicated. For example, if they use two conjunctions, warn them and break it down for them. "
                "    6. Warn students if their statement per input is too long or complicated. When student statement has more than 7 or 8 words, break it down for them with full stops and simple conjunctions. "
                "    7. Always add your ideas after student submits their sentence if necessary. "
                "    8. Make sure the complete letter is between 25 and 35 words. "
                "    9. When the letter is about cancelling appointments, teach students how they can use reasons connected to weather and health to cancel appointments. Teach them how to use 'absagen' to cancel appointments. "
                "    10. For enquiries or registrations, teach students how to use 'Anfrage stellen' for the Ich schreibe. "
                "    11. When the letter is about registrations like a course, teach students how they can use 'anfangen', 'beginnen'. "
                "    12. Asking for price, teach them how to use 'wie viel kostet...' and how they should ask for price always when it is about enquiries. "
                "    13. Teach them to use 'Es tut mir leid.' to say sorry. "
                "    14. Always remind students to use 'Ich schreibe Ihnen/dir, weil ich ... möchte.' for their reasons. "
                "Always make grammar correction or suggest a better phrase when necessary. "
                "If it's a continuation, review their writing so far and guide them to the next step. "
                "If it's a new prompt, give a brief, simple overview (in English) of how to build their letter (greeting, introduction, reason, request, closing), with short examples for each. "
                "For the introduction, always remind the student to use: 'Ich schreibe Ihnen, weil ich ...' for formal letters or 'Ich schreibe dir, weil ich ...' for informal letters. "
                "For the main request, always recommend ending the sentence with 'möchte' or another basic modal verb, as this is the easiest and most correct way at A1 (e.g., 'Ich möchte einen Termin machen.'). "
                "After your overview or advice, always use the phrase 'Your next recommended step:' and ask for only the next part—first the greeting (wait for it), then only the introduction (wait for it), then reason, then request, then closing—one after the other, never more than one at a time. "
                "After each student reply, check their answer, give gentle feedback, and then again state 'Your next recommended step:' and prompt for just the next section. "
                "Only help with basic connectors ('und', 'aber', 'weil', 'deshalb', 'ich möchte wissen'). Never write the full letter yourself—coach one part at a time. "
                "The chat session should last for about 10 student replies. If the student is not done by then, gently remind them: 'Most letters can be completed in about 10 steps. Please try to finish soon.' "
                "If after 14 student replies, the letter is still not finished, end the session with: 'We have reached the end of this coaching session. Please copy your letter below so far and paste it into the “Mark My Letter” tool for full AI feedback and a score.' "
                "Throughout, your questions must be progressive, one at a time, and always guide the student logically through the structure."
            ),
            "A2": (
                "You are Herr Felix, a creative and supportive German letter-writing coach for A2 students. "
                "Always reply in English, never in German. "
                "Congratulate the student on their first submission with ideas about how to go about the letter. Analyze whether it is a prompt, a continuation, or a question. "
                "    1. Always give students short ideas, structure and tips and phrases on how to build their points for the conversation in English and simple German. Don't overfeed students; help them but let them think by themselves also. "
                "    2. For structure, require their letter to use clear sequencing with 'Zuerst' (for the first paragraph), 'Dann' or 'Außerdem' (for the body/second idea), and 'Zum Schluss' (for closing/last idea). "
                "       - Always recommend 'Zuerst' instead of 'Erstens' for A2 letters, as it is simpler and more natural for personal or exam letters. "
                "    3. For connectors, use 'und', 'aber', 'weil', 'denn', 'deshalb', 'ich mochte wissen, ob', 'ich mochte wissen, wann', 'ich mochte wissen, wo', and encourage linking words for clarity. Recommend one at a time in a statement to prevent mistakes. When a student use two or more conjucntion in one statement less than 7 words, simplify for them to use just once to prevent errors"
                "    4. After every reply, give a tip or phrase, but never write the full letter for them. "
                "    5. Remind them not to write sentences longer than 7–8 words; break long sentences into short, clear ideas. "
                "    6. Letter should be between 30 and 40 words. "
                "    7. For cancellations, suggest health/weather reasons ('Ich bin krank.', 'Es regnet stark.') and use 'absagen' (e.g., 'Ich schreibe Ihnen, weil ich absagen möchte.'). "
                "    8. For enquiries/registrations, show 'Anfrage stellen' (e.g., 'Ich schreibe Ihnen, weil ich eine Anfrage stellen möchte.') and include asking for price: 'Wie viel kostet...?'. "
                "    9. For appointments, recommend 'vereinbaren' ('Ich möchte einen neuen Termin vereinbaren.'). "
                "    10. To say sorry, use: 'Es tut mir leid.' "
                "    11. Always correct grammar and suggest improved phrases when needed. "
                "    12. At each step, say 'Your next recommended step:' and ask for only the next section (first greeting, then introduction, then body using 'Zuerst', 'Außerdem', then final point 'Zum Schluss', then polite closing phrase 'Ich freue mich'). "
                "    13. The session should be complete in about 10 student replies; if not, remind them to finish soon. After 14, end and tell the student to copy their letter below and paste into 'Mark My Letter' for feedback. "
                "    14. Throughout, do not write the whole letter—guide only one part at a time."
                
            ),
            "B1": (
                "You are Herr Felix, a supportive German letter/essay coach for B1 students. "
                "Always reply in English, never in German. "
                "Congratulate the student with ideas about how to go about the letter, analyze the type of submission, and determine whether it is a formal letter, informal letter, or opinion essay. "
                "If you are not sure, politely ask the student what type of writing they need help with. "
                f"1. Always give students short ideas,structure and tips and phrases on how to build their points for the conversation in English and simple German. Dont overfeed students, help them but let them think by themselves also "
                f"2. Always check to be sure their letters are organized with paragraphs using sequences and sentence starters "
                f"3. Always add your ideas after student submmit their sentence if necessary "
                f"4. Always be sure that students complete formal letter is between 40 to 50 words,informal letter and opinion essay between 80 to 90 words "
                f"5. When giving ideas for sentences, just give 2 to 3 words and tell student to continue from there. Let the student also think and dont over feed them. "
                "For a formal letter, give a brief overview of the structure (greeting, introduction, main reason/request, closing), with useful examples. "
                "Always make grammar correction or suggest a better phrase when necessary. "
                "For an informal letter, outline the friendly structure (greeting, introduction, reason, personal info, closing), with simple examples. "
                "For an opinion essay, provide a short overview: introduction (with phrases like 'Heutzutage ist ... ein wichtiges Thema.' or 'Ich bin der Meinung, dass...'), main points (advantages, disadvantages, opinion), connectors, and closing. "
                "After your overview, always use the phrase 'Your next recommended step:' and ask for only one section at a time—greeting, then introduction, then main points, then closing—never more than one at a time. "
                "After each answer, provide feedback, then again prompt with 'Your next recommended step:'. "
                "Encourage the use of appropriate connectors ('außerdem', 'trotzdem', 'weil', 'deshalb'). "
                "If the student is still writing after 10 turns, encourage them to finish. At 14, end the chat, reminding them to copy their letter below and paste their draft in 'Mark My Letter' for feedback."
            ),
            "B2": (
                "You are Herr Felix, a supportive German writing coach for B2 students. "
                "Always reply in English, never in German. "
                "Congratulate the student with ideas about how to go about the letter, analyze the type of input, and determine if it is a formal letter, informal letter, or an opinion/argumentative essay. "
                "If you are not sure, politely ask the student what type of writing they need help with. "
                f"1. Always give students short ideas,structure and tips and phrases on how to build their points for the conversation in English and simple German. Dont overfeed students, help them but let them think by themselves also "
                f"2. Always check to be sure their letters are organized with paragraphs using sequences and sentence starters "
                f"3. Always add your ideas after student submmit their sentence if necessary "
                f"4. Always be sure that students complete formal letter is between 100 to 150 words and opinion essay is 150 to 170 words "
                f"5. When giving ideas for sentences, just give 2 to 3 words and tell student to continue from there. Let the student also think and dont over feed them. "
                "Always make grammar correction or suggest a better phrase when necessary. "
                "For a formal letter, briefly outline the advanced structure: greeting, introduction, clear argument/reason, supporting details, closing—with examples. "
                "For an informal letter, outline a friendly but organized structure: greeting, personal introduction, main point/reason, examples, closing. "
                "For an opinion or argumentative essay, outline: introduction (with a strong thesis), arguments (with connectors and examples), counterarguments, connectors, conclusion, closing. "
                "After your overview or advice, always use the phrase 'Your next recommended step:' and ask for only one section at a time. "
                "After each student reply, give feedback, then use 'Your next recommended step:' again. "
                "Suggest and model advanced connectors ('denn', 'dennoch', 'außerdem', 'jedoch', 'zum Beispiel', 'einerseits...andererseits'). "
                "If the student is still writing after 10 turns, gently encourage finishing; after 14, end the chat and ask the student to copy their letter below and paste their draft in 'Mark My Letter' for feedback."
            ),
            "C1": (
                "You are Herr Felix, an advanced and supportive German writing coach for C1 students. "
                "Always reply in English, and in German when neccessary. If the German is difficult, explain it to the student "
                "Congratulate the student with ideas about how to go about the letter, analyze the type of input, and determine if it is a formal letter, informal letter, or an academic/opinion essay. "
                f"1. Always give students short ideas,structure and tips and phrases on how to build their points for the conversation in English and simple German. Dont overfeed students, help them but let them think by themselves also "
                f"2. Always check to be sure their letters are organized with paragraphs using sequence and sentence starters "
                f"3. Always add your ideas after student submmit their sentence if necessary "
                f"4. Always be sure that students complete formal letter is between 120 to 150 words and opinion essay is 230 to 250 words "
                f"5. When giving ideas for sentences, just give 2 to 3 words and tell student to continue from there. Let the student also think and dont over feed them. "
                "If you are not sure, politely ask the student what type of writing they need help with. "
                "For a formal letter, give a precise overview: greeting, sophisticated introduction, detailed argument, supporting evidence, closing, with nuanced examples. "
                "Always make grammar correction or suggest a better phrase when necessary. "
                "For an informal letter, outline a nuanced and expressive structure: greeting, detailed introduction, main point/reason, personal opinion, nuanced closing. "
                "For academic or opinion essays, provide a clear outline: introduction (with a strong thesis and background), well-structured arguments, counterpoints, advanced connectors, conclusion, and closing—with C1-level examples. "
                "After your overview or advice, always use the phrase 'Your next recommended step:' and ask for only one section at a time. "
                "After each answer, provide feedback, then again prompt with 'Your next recommended step:'. "
                "Model and suggest advanced connectors ('nicht nur... sondern auch', 'obwohl', 'dennoch', 'folglich', 'somit'). "
                "If the student is still writing after 10 turns, gently encourage finishing; after 14, end the chat and ask the student to  paste their draft in 'Mark My Letter' for feedback and a score."
            ),
        }

        def reset_letter_coach():
            for k in [
                "letter_coach_stage", "letter_coach_chat", "letter_coach_prompt",
                "letter_coach_type", "selected_letter_lines", "letter_coach_uploaded"
            ]:
                st.session_state[k] = 0 if k == "letter_coach_stage" else []
            st.session_state["letter_coach_uploaded"] = False

        def bubble(role, text):
            if role == "assistant":
                return f"""<div style='background: #f4eafd; color: #7b2ff2; border-radius: 16px 16px 16px 3px; margin-bottom: 8px; margin-right: 80px; box-shadow: 0 2px 8px rgba(123,47,242,0.08); padding: 13px 18px; text-align: left; max-width: 88vw; font-size: 1.12rem;'><b>👨‍🏫 Herr Felix:</b><br>{text}</div>"""
            return f"""<div style='background: #eaf4ff; color: #1a237e; border-radius: 16px 16px 3px 16px; margin-bottom: 8px; margin-left: 80px; box-shadow: 0 2px 8px rgba(26,35,126,0.07); padding: 13px 18px; text-align: right; max-width: 88vw; font-size: 1.12rem;'><b>🙋 You:</b><br>{text}</div>"""

        # --- General Instructions for Students (Minimal Welcome + Subline) ---
        st.markdown(
            """
            <div style="
                background: linear-gradient(97deg, #f4eafd 75%, #ffe0f5 100%);
                border-radius: 12px;
                border: 1px solid #e6d3fa;
                box-shadow: 0 2px 8px #e5e1fa22;
                padding: 0.75em 1em 0.72em 1em;
                margin-bottom: 1.1em;
                margin-top: 0.1em;
                color: #4b2976;
                font-size: 1.03rem;
                font-family: 'Segoe UI', 'Roboto', 'Arial', sans-serif;
                text-align: center;
                ">
                <span style="font-size:1.19em; vertical-align:middle;">✉️</span>
                <span style="font-size:1.05em; font-weight: 500; margin-left:0.24em;">
                    Welcome to <span style="color:#7b2ff2;">Letter Coach</span>
                </span>
                <div style="color:#b48be6; font-size:0.97em; margin-top:0.35em;">
                    Get started below 👇
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

        IDEAS_LIMIT = 14
        ideas_so_far = get_letter_coach_usage(student_code)
        st.markdown(f"**Daily usage:** {ideas_so_far} / {IDEAS_LIMIT}")
        if ideas_so_far >= IDEAS_LIMIT:
            st.warning("You have reached today's letter coach limit. Please come back tomorrow.")
            st.stop()

        # --- Stage 0: Prompt input ---
        if st.session_state[ns("stage")] == 0:
            st.markdown("### ✏️ Enter your exam prompt or draft to start coaching")
            with st.form(ns("prompt_form"), clear_on_submit=True):
                prompt = st.text_area(
                    "",
                    value=st.session_state[ns("prompt")],
                    height=120,
                    placeholder="e.g., Schreiben Sie eine formelle E-Mail an Ihre Nachbarin ..."
                )
                send = st.form_submit_button("✉️ Start Letter Coach")

            if prompt:
                word_count = len(prompt.split())
                char_count = len(prompt)
                st.markdown(
                    f"<div style='color:#7b2ff2; font-size:0.97em; margin-bottom:0.18em;'>"
                    f"Words: <b>{word_count}</b> &nbsp;|&nbsp; Characters: <b>{char_count}</b>"
                    "</div>",
                    unsafe_allow_html=True
                )

            if send and prompt:
                st.session_state[ns("prompt")] = prompt
                student_level = st.session_state.get("schreiben_level", "A1")
                system_prompt = LETTER_COACH_PROMPTS[student_level].format(prompt=prompt)
                chat_history = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ]
                try:
                    resp = client.chat.completions.create(
                        model="gpt-4o",
                        messages=chat_history,
                        temperature=0.22,
                        max_tokens=380
                    )
                    ai_reply = resp.choices[0].message.content
                except Exception as e:
                    ai_reply = "Sorry, there was an error generating a response. Please try again."
                chat_history.append({"role": "assistant", "content": ai_reply})

                st.session_state[ns("chat")] = chat_history
                st.session_state[ns("stage")] = 1
                inc_letter_coach_usage(student_code)
                save_letter_coach_progress(
                    student_code,
                    student_level,
                    st.session_state[ns("prompt")],
                    st.session_state[ns("chat")],
                )
                st.rerun()

            if prompt:
                st.markdown("---")
                st.markdown(f"📝 **Letter/Essay Prompt or Draft:**\n\n{prompt}")

        # --- Stage 1: Coaching Chat ---
        elif st.session_state[ns("stage")] == 1:
            st.markdown("---")
            st.markdown(f"📝 **Letter/Essay Prompt:**\n\n{st.session_state[ns('prompt')]}")
            chat_history = st.session_state[ns("chat")]
            for msg in chat_history[1:]:
                st.markdown(bubble(msg["role"], msg["content"]), unsafe_allow_html=True)
            num_student_turns = sum(1 for msg in chat_history[1:] if msg["role"] == "user")
            if num_student_turns == 10:
                st.info("🔔 You have written 10 steps. Most students finish in 7–10 turns. Try to complete your letter soon!")
            elif num_student_turns == 12:
                st.warning(
                    "⏰ You have reached 12 writing turns. "
                    "Usually, your letter should be complete by now. "
                    "If you want feedback, click **END SUMMARY** or download your letter as TXT. "
                    "You can always start a new session for more practice."
                )
            elif num_student_turns > 12:
                st.warning(
                    f"🚦 You are now at {num_student_turns} turns. "
                    "Long letters are okay, but usually a good letter is finished in 7–12 turns. "
                    "Try to wrap up, click **END SUMMARY** or download your letter as TXT."
                )

            with st.form(ns("letter_coach_chat_form"), clear_on_submit=True):
                user_input = st.text_area(
                    "",
                    value="",
                    key=ns("user_input"),
                    height=400,
                    placeholder="Type your reply, ask about a section, or paste your draft here..."
                )
                send = st.form_submit_button("Send")
            if send and user_input.strip():
                chat_history.append({"role": "user", "content": user_input})
                student_level = st.session_state.get("schreiben_level", "A1")
                system_prompt = LETTER_COACH_PROMPTS[student_level].format(prompt=st.session_state[ns("prompt")])
                with st.spinner("👨‍🏫 Herr Felix is typing..."):
                    resp = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[{"role": "system", "content": system_prompt}] + chat_history[1:] + [{"role": "user", "content": user_input}],
                        temperature=0.22,
                        max_tokens=380
                    )
                    ai_reply = resp.choices[0].message.content
                chat_history.append({"role": "assistant", "content": ai_reply})
                st.session_state[ns("chat")] = chat_history
                save_letter_coach_progress(
                    student_code,
                    student_level,
                    st.session_state[ns("prompt")],
                    st.session_state[ns("chat")],
                )
                st.rerun()

            # ----- LIVE AUTO-UPDATING LETTER DRAFT, Download + Copy -----
            import streamlit.components.v1 as components

            user_msgs = [
                msg["content"]
                for msg in st.session_state[ns("chat")][1:]
                if msg.get("role") == "user"
            ]

            st.markdown("""
                **📝 Your Letter Draft**
                - Tick the lines you want to include in your letter draft.
                - You can untick any part you want to leave out.
                - Only ticked lines will appear in your downloadable draft below.
            """)

            # Store selection in session state (keeps selection per student)
            if ns("selected_letter_lines") not in st.session_state or \
                len(st.session_state[ns("selected_letter_lines")]) != len(user_msgs):
                st.session_state[ns("selected_letter_lines")] = [True] * len(user_msgs)

            selected_lines = []
            for i, line in enumerate(user_msgs):
                st.session_state[ns("selected_letter_lines")][i] = st.checkbox(
                    line,
                    value=st.session_state[ns("selected_letter_lines")][i],
                    key=ns(f"letter_line_{i}")
                )
                if st.session_state[ns("selected_letter_lines")][i]:
                    selected_lines.append(line)

            letter_draft = "\n".join(selected_lines)

            # --- Live word/character count for the letter draft ---
            draft_word_count = len(letter_draft.split())
            draft_char_count = len(letter_draft)
            st.markdown(
                f"<div style='color:#7b2ff2; font-size:0.97em; margin-bottom:0.18em;'>"
                f"Words: <b>{draft_word_count}</b> &nbsp;|&nbsp; Characters: <b>{draft_char_count}</b>"
                "</div>",
                unsafe_allow_html=True
            )

            # --- Modern, soft header (copy/download) ---
            st.markdown(
                """
                <div style="
                    background:#23272b;
                    color:#eee;
                    border-radius:10px;
                    padding:0.72em 1.04em;
                    margin-bottom:0.4em;
                    font-size:1.07em;
                    font-weight:400;
                    border:1px solid #343a40;
                    box-shadow:0 2px 10px #0002;
                    text-align:left;
                ">
                    <span style="font-size:1.12em; color:#ffe082;">📝 Your Letter So Far</span><br>
                    <span style="font-size:1.00em; color:#b0b0b0;">copy often or download below to prevent data loss</span>
                </div>
                """,
                unsafe_allow_html=True
            )

            # --- Mobile-friendly copy/download box ---
            components.html(f"""
                <textarea id="letterBox_{student_code}" readonly rows="6" style="
                    width: 100%;
                    border-radius: 12px;
                    background: #f9fbe7;
                    border: 1.7px solid #ffe082;
                    color: #222;
                    font-size: 1.12em;
                    font-family: 'Fira Mono', 'Consolas', monospace;
                    padding: 1em 0.7em;
                    box-shadow: 0 2px 8px #ffe08266;
                    margin-bottom: 0.5em;
                    resize: none;
                    overflow:auto;
                " onclick="this.select()">{letter_draft}</textarea>
                <button onclick="navigator.clipboard.writeText(document.getElementById('letterBox_{student_code}').value)" 
                    style="
                        background:#ffc107;
                        color:#3e2723;
                        font-size:1.08em;
                        font-weight:bold;
                        padding:0.48em 1.12em;
                        margin-top:0.4em;
                        border:none;
                        border-radius:7px;
                        cursor:pointer;
                        box-shadow:0 2px 8px #ffe08255;
                        width:100%;
                        max-width:320px;
                        display:block;
                        margin-left:auto;
                        margin-right:auto;
                    ">
                    📋 Copy Text
                </button>
                <style>
                    @media (max-width: 480px) {{
                        #letterBox_{student_code} {{
                            font-size: 1.16em !important;
                            min-width: 93vw !important;
                        }}
                    }}
                </style>
            """, height=175)

            st.markdown("""
                <div style="
                    background:#ffe082;
                    padding:0.9em 1.2em;
                    border-radius:10px;
                    margin:0.4em 0 1.2em 0;
                    color:#543c0b;
                    font-weight:600;
                    border-left:6px solid #ffc107;
                    font-size:1.08em;">
                    📋 <span>On phone, tap in the box above to select all for copy.<br>
                    Or just tap <b>Copy Text</b>.<br>
                    To download, use the button below.</span>
                </div>
            """, unsafe_allow_html=True)

            st.download_button(
                "⬇️ Download Letter as TXT",
                letter_draft.encode("utf-8"),
                file_name="my_letter.txt"
            )

            if st.button("Start New Letter Coach"):
                st.session_state[ns("chat")] = []
                st.session_state[ns("prompt")] = ""
                st.session_state[ns("selected_letter_lines")] = []
                st.session_state[ns("stage")] = 0
                save_letter_coach_progress(
                    student_code,
                    st.session_state.get("schreiben_level", "A1"),
                    "",
                    [],
                )
                st.rerun()


# Inject PWA/meta link tags AFTER the hero (zero-height iframe)
    _inject_meta_tags()

    # Inject SEO head tags AFTER the hero (using components.html)
    components.html("""
    <script>
      document.title = "Falowen – Learn German with Learn Language Education Academy";
      const desc = "Falowen is the German learning companion from Learn Language Education Academy. Join live classes or self-study with A1–C1 courses, recorded lectures, and real progress tracking.";
      let m = document.querySelector('meta[name="description"]');
      if (!m) { m = document.createElement('meta'); m.name = "description"; document.head.appendChild(m); }
      m.setAttribute("content", desc);
      const canonicalHref = window.location.origin + "/";
      let link = document.querySelector('link[rel="canonical"]');
      if (!link) { link = document.createElement('link'); link.rel = "canonical"; document.head.appendChild(link); }
      link.href = canonicalHref;
      function setOG(p, v){ let t=document.querySelector(`meta[property="${p}"]`);
        if(!t){ t=document.createElement('meta'); t.setAttribute('property', p); document.head.appendChild(t); }
        t.setAttribute('content', v);
      }
      setOG("og:title", "Falowen – Learn German with Learn Language Education Academy");
      setOG("og:description", desc);
      setOG("og:type", "website");
      setOG("og:url", canonicalHref);
      const ld = {"@context":"https://schema.org","@type":"WebSite","name":"Falowen","alternateName":"Falowen by Learn Language Education Academy","url": canonicalHref};
      const s = document.createElement('script'); s.type = "application/ld+json"; s.text = JSON.stringify(ld); document.head.appendChild(s);
    </script>
    """, height=0)












